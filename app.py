"""
Power BI Documentation Web Application
Flask backend for generating Power BI documentation on-demand
"""

from flask import Flask, render_template, request, jsonify, send_file, redirect, url_for, session, flash
from datetime import datetime, timezone, timedelta
from functools import wraps
import os
import time
import threading
from dotenv import load_dotenv
import msal
import uuid
import requests
import json

# Import your existing modules
from powerbi_connector import PowerBIConnector
from ai_generator import AIDocGenerator
from document_creator import PowerBIDocumentCreator
from config import Config

# Import visual metadata extractor for deep search
from visual_metadata_extractor import VisualMetadataExtractor
import asyncio

# Load environment variables BEFORE catalog import (catalog_config reads env at import time)
load_dotenv(override=True)

# Precomputed tenant catalog (SharePoint / local) — fast path; live APIs remain fallback
try:
    from catalog_service import catalog_service
    CATALOG_AVAILABLE = True
    print("✅ Catalog service loaded (SharePoint/local fast path enabled when configured)")
except Exception as _catalog_import_err:
    catalog_service = None
    CATALOG_AVAILABLE = False
    print(f"⚠️ Catalog service not available: {_catalog_import_err}")

# Shared exclude: platform usage metrics + [App] shells (Catalog / Home / Decomm / etc.)
try:
    from catalog_service.thin_packs import is_excluded_report_name as _is_excluded_report_name
except Exception:
    def _is_excluded_report_name(name):  # type: ignore
        n = (name or "").strip()
        if not n:
            return False
        if n.startswith("[App]"):
            return True
        return n.casefold() in {
            "usage metrics report",
            "report usage metrics report",
            "dashboard usage metrics report",
        }

# Initialize Flask app
app = Flask(__name__)
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'powerbi-doc-generator-secret-key-2024')

# Session configuration
# IMPORTANT: Client-side signed cookies overflow (~4KB) once we store Power BI
# + Copilot JWTs + the MSAL token_cache. Browsers then drop the cookie → endless
# login loop. Use server-side filesystem sessions (cookie only holds a small id).
SESSION_MAX_HOURS = int(os.getenv('SESSION_MAX_HOURS', '12'))
app.config['SESSION_PERMANENT'] = False
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(hours=SESSION_MAX_HOURS)
app.config['SESSION_REFRESH_EACH_REQUEST'] = False

_on_azure = bool(os.getenv('WEBSITE_HOSTNAME'))
# Secure cookies only work over HTTPS. On localhost (http://) they are dropped by
# the browser → OAuth state missing on /getAToken → endless login loop.
# Force insecure cookies for local dev unless you explicitly serve local HTTPS.
_secure_env = (os.getenv('SESSION_COOKIE_SECURE') or '').strip().lower()
if _on_azure:
    app.config['SESSION_COOKIE_SECURE'] = True
elif _secure_env in ('1', 'true', 'yes', 'on'):
    app.config['SESSION_COOKIE_SECURE'] = True
    print("⚠️ SESSION_COOKIE_SECURE=true on non-Azure — only use with local HTTPS")
else:
    app.config['SESSION_COOKIE_SECURE'] = False
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
app.config['SESSION_COOKIE_NAME'] = 'pbi_session'

# Server-side session store (Flask-Session).
# Local: prefer %TEMP% (or FLASK_SESSION_DIR). Sessions under OneDrive paths often
# fail to read back the OAuth "state" → endless /login loop.
def _pick_session_dir() -> str:
    env = (os.getenv('FLASK_SESSION_DIR') or '').strip()
    candidates = []
    if env:
        candidates.append(env)
    if _on_azure:
        candidates.append('/home/data/flask_sessions')
    # Local non-synced dirs first
    candidates.append(os.path.join(os.environ.get('TEMP') or os.environ.get('TMP') or '/tmp', 'pbi_cc_flask_sessions'))
    candidates.append(os.path.join(os.getcwd(), 'data', 'flask_sessions'))
    last_err = None
    for cand in candidates:
        try:
            os.makedirs(cand, exist_ok=True)
            probe = os.path.join(cand, '.write_probe')
            with open(probe, 'w', encoding='utf-8') as f:
                f.write('ok')
            os.remove(probe)
            return cand
        except Exception as exc:
            last_err = exc
            continue
    raise RuntimeError(f"No writable Flask session directory (last error: {last_err})")


_sess_dir = _pick_session_dir()

# Default filesystem. Optional Redis when SESSION_REDIS_URL is set (recommended on Azure
# multi-instance / to survive worker recycles without losing mid-login OAuth state).
# Existing successful-login behavior is unchanged when Redis is unset.
app.config['SESSION_FILE_DIR'] = _sess_dir
# Old default 500 pruned server session files aggressively and could delete the
# mid-login OAuth "state" file between /login and /getAToken under load.
try:
    _sess_threshold = int(os.getenv('SESSION_FILE_THRESHOLD', '10000'))
except ValueError:
    _sess_threshold = 10000
app.config['SESSION_FILE_THRESHOLD'] = max(500, _sess_threshold)
app.config['SESSION_USE_SIGNER'] = True
app.config['SESSION_KEY_PREFIX'] = 'pbi_cc:'
# Ensure Flask always saves session after login/callback
app.config['SESSION_PERMANENT'] = False

_session_backend = 'unset'
_redis_url = (os.getenv('SESSION_REDIS_URL') or os.getenv('REDIS_URL') or '').strip()
if _redis_url:
    try:
        import redis as _redis_mod
        app.config['SESSION_TYPE'] = 'redis'
        app.config['SESSION_REDIS'] = _redis_mod.from_url(_redis_url)
        from flask_session import Session as _FlaskSession
        _FlaskSession(app)
        _session_backend = f'redis:{_redis_url.split("@")[-1] if "@" in _redis_url else "configured"}'
    except Exception as _redis_err:
        print(f"⚠️ SESSION_REDIS_URL set but Redis session init failed ({_redis_err}); using filesystem")
        _redis_url = ''

if not _redis_url:
    app.config['SESSION_TYPE'] = 'filesystem'
    try:
        from flask_session import Session as _FlaskSession
        _FlaskSession(app)
        _session_backend = f'filesystem:{_sess_dir} (threshold={app.config["SESSION_FILE_THRESHOLD"]})'
    except Exception as _sess_err:
        _session_backend = f'cookie-fallback ({_sess_err})'
        print(f"⚠️ Flask-Session unavailable, cookie sessions may overflow: {_sess_err}")

print(f"Session configuration:")
print(f"   SECRET_KEY: {'Set from environment' if os.getenv('SECRET_KEY') else 'Using default (set SECRET_KEY in production!)'}")
print(f"   Session backend: {_session_backend}")
print(f"   Absolute max age: {SESSION_MAX_HOURS}h from login")
print(f"   Cookie secure: {app.config['SESSION_COOKIE_SECURE']}")
print(f"   Cookie SameSite: {app.config['SESSION_COOKIE_SAMESITE']}")
print(f"   REDIRECT will log after auth constants load")

# ============================================================================
# CACHE CONTROL - Prevent stale content after deployments
# ============================================================================
@app.after_request
def add_cache_control_headers(response):
    """
    Prevent browser/CDN caching issues that cause stale content after deployments.

    Issue: Docker layer caching + browser caching = users seeing old code
    Solution: Force no-cache for dynamic content (HTML, JSON)
    """
    # Don't cache API/HTML by default — but keep explicit private caches
    # (e.g. /api/catalog/impact/tables sets private max-age for fast revisits).
    ct = response.content_type or ''
    existing_cc = (response.headers.get('Cache-Control') or '').lower()
    allow_private = 'private' in existing_cc and 'max-age' in existing_cc
    if ('application/json' in ct or 'text/html' in ct) and not allow_private:
        response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
        response.headers['Pragma'] = 'no-cache'
        response.headers['Expires'] = '0'
    return response

# ⚡ PERFORMANCE OPTIMIZATION: Enhanced caching system
# User-specific workspace caching
workspaces_cache = {}  # Format: {user_id: {'data': [...], 'timestamp': float}}
workspace_cache = {}   # Format: {workspace_id_user_id: {'data': {...}, 'timestamp': float}}
scanner_cache = {}     # Format: {workspace_id_dataset_id: {'data': {...}, 'timestamp': float}}
reports_cache = {}     # Format: {workspace_id_user_id_folder: {'data': [...], 'timestamp': float}}

# Cache durations (in seconds)
WORKSPACES_CACHE_DURATION = 300  # 5 minutes - workspaces list rarely changes
CACHE_DURATION = 300  # 5 minutes - general cache
SCANNER_CACHE_DURATION = 1800  # 30 minutes - scanner data is expensive to fetch

print("✅ Usage cache will persist in memory (cleared only on explicit refresh or restart)")

# Azure AD SSO Configuration
TENANT_ID = os.getenv('TENANT_ID')
CLIENT_ID = os.getenv('CLIENT_ID')
CLIENT_SECRET = os.getenv('CLIENT_SECRET')
AUTHORITY = f"https://login.microsoftonline.com/{TENANT_ID}"
REDIRECT_PATH = "/getAToken"  # Must match redirect URI in Azure AD app registration



# Auto-detect environment and set appropriate redirect URI.
# Localhost must NOT keep a production https://… REDIRECT_URI from .env —
# that causes Azure AD to bounce to the wrong host or drop the local session.
_env_redirect = (os.getenv('REDIRECT_URI') or '').strip()
if os.getenv('WEBSITE_HOSTNAME'):
    # Azure App Service — prefer explicit env, else hostname
    if _env_redirect and 'localhost' not in _env_redirect.lower():
        REDIRECT_URI = _env_redirect
    else:
        REDIRECT_URI = f"https://{os.getenv('WEBSITE_HOSTNAME')}{REDIRECT_PATH}"
    print(f"🌐 Running on Azure App Service: {REDIRECT_URI}")
else:
    # Local dev — always http://localhost:5000 unless env is already localhost
    if _env_redirect and 'localhost' in _env_redirect.lower():
        REDIRECT_URI = _env_redirect
    else:
        if _env_redirect:
            print(
                f"⚠️ Ignoring REDIRECT_URI={_env_redirect!r} on local run "
                f"(use http://localhost:5000{REDIRECT_PATH})"
            )
        REDIRECT_URI = f'http://localhost:5000{REDIRECT_PATH}'
    print(f"💻 Running locally: {REDIRECT_URI}")

# Scopes for user-delegated permissions
# IMPORTANT: Use .default scope to get all consented permissions for Power BI API
# This ensures we get a token for Power BI API, not Microsoft Graph
SCOPE = ["https://analysis.windows.net/powerbi/api/.default"]

# Fabric API scope - needed for getDefinition and other Fabric-specific endpoints
FABRIC_SCOPE = ["https://api.fabric.microsoft.com/.default"]

# Optional: If you also need Graph API access, request it separately
GRAPH_SCOPE = ["User.Read"]

# Base MSAL app (no per-request cache). Prefer _msal_for_request() when
# acquiring/refreshing tokens so the session token_cache is actually used.
msal_app = msal.ConfidentialClientApplication(
    CLIENT_ID,
    authority=AUTHORITY,
    client_credential=CLIENT_SECRET,
)


def _load_cache():
    """Load MSAL token cache from the Flask session."""
    cache = msal.SerializableTokenCache()
    if session.get("token_cache"):
        cache.deserialize(session["token_cache"])
    return cache


def _save_cache(cache):
    """Persist MSAL token cache back into the Flask session."""
    if cache is not None and cache.has_state_changed:
        session["token_cache"] = cache.serialize()
        session.modified = True


def _msal_for_request(cache=None):
    """Confidential client bound to this request's serialized token cache.

    Without binding the cache, acquire_token_silent / get_accounts see an empty
    in-memory cache on every worker, so Fabric step-up / token refresh fails.
    """
    if cache is None:
        cache = _load_cache()
    app_cca = msal.ConfidentialClientApplication(
        CLIENT_ID,
        authority=AUTHORITY,
        client_credential=CLIENT_SECRET,
        token_cache=cache,
    )
    return app_cca, cache




# Initialize Power BI connector
config = Config()
powerbi = PowerBIConnector()
# Authentication will happen automatically when needed via _get_headers()

# Initialize AI generator
ai_generator = AIDocGenerator(
    openai_api_key=config.OPENAI_API_KEY
)

# Cache for workspaces and reports (to avoid repeated API calls)
# IMPORTANT: These are keyed by user_id to prevent cross-user data leakage
workspaces_cache = {}  # Key: user_id, Value: {'data': [], 'timestamp': 0}
reports_cache = {}  # Key: {workspace_id}_{user_id}, Value: {'data': [], 'timestamp': 0}
scanner_cache = {}  # Key: workspace_id, Value: {'data': {...}, 'timestamp': 0}
usage_cache = {}  # Key: workspace_id, Value: {'data': {...}, 'timestamp': datetime}
# Workspace folder tree (Fabric API) — shared across users, short TTL
workspace_folders_cache = {}  # Key: workspace_id, Value: {'folders': [], 'timestamp': float}
CACHE_DURATION = 300  # 5 minutes
SCANNER_CACHE_DURATION = 600  # 10 minutes for scanner results (expensive operation)
WORKSPACE_FOLDERS_CACHE_DURATION = 600  # 10 minutes — folder names change rarely

# Progress tracking for document generation
generation_progress = {}  # Key: job_id, Value: {'progress': 0-100, 'status': 'message', 'file_path': 'path/to/file.docx', 'complete': False, 'error': None}

# Persistent file-based cache directory for usage metrics
USAGE_CACHE_DIR = os.path.join(os.getcwd(), 'data', 'usage_cache')
os.makedirs(USAGE_CACHE_DIR, exist_ok=True)

def clear_user_cache(user_id):
    """Clear all cached data for a specific user"""
    global workspaces_cache, reports_cache

    # Clear workspace cache for this user
    cache_key = f"workspaces_{user_id}"
    if cache_key in workspaces_cache:
        del workspaces_cache[cache_key]
        print(f"🗑️ Cleared workspace cache for user: {user_id}")

    # Clear report caches for this user
    keys_to_delete = [k for k in reports_cache.keys() if k.endswith(f"_{user_id}")]
    for key in keys_to_delete:
        del reports_cache[key]
    if keys_to_delete:
        print(f"🗑️ Cleared {len(keys_to_delete)} report cache entries for user: {user_id}")

def clear_all_caches():
    """Clear all cached data - used on server startup or manual refresh"""
    global workspaces_cache, reports_cache, scanner_cache, workspace_folders_cache
    workspaces_cache = {}
    reports_cache = {}
    scanner_cache = {}
    workspace_folders_cache = {}
    print("🗑️ Cleared ALL workspace, report, scanner, and folder caches")


def _session_expired() -> bool:
    """True when absolute session age exceeds SESSION_MAX_HOURS (default 12)."""
    if 'user' not in session:
        return True
    started = session.get('login_at')
    if not started:
        # Legacy sessions without stamp — force re-login once
        return True
    try:
        from datetime import datetime, timezone
        if isinstance(started, (int, float)):
            started_ts = float(started)
        else:
            s = str(started)
            if s.endswith('Z'):
                s = s[:-1] + '+00:00'
            started_ts = datetime.fromisoformat(s).timestamp()
        age_sec = time.time() - started_ts
        return age_sec > (SESSION_MAX_HOURS * 3600)
    except Exception:
        return True


def login_required(f):
    """Decorator to require login; enforces browser-close cookie + 12h absolute max."""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user' not in session or _session_expired():
            if 'user' in session:
                # Absolute age exceeded — hard clear
                session.clear()
            if request.path.startswith('/api/'):
                return jsonify({
                    'success': False,
                    'error': 'Not authenticated or session expired',
                    'redirect': url_for('login')
                }), 401
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function


@app.route('/debug/env')
def debug_env():
    """Debug endpoint to check environment configuration"""
    return jsonify({
        'WEBSITE_HOSTNAME': os.getenv('WEBSITE_HOSTNAME'),
        'REDIRECT_URI_ENV': os.getenv('REDIRECT_URI'),
        'REDIRECT_URI_ACTUAL': REDIRECT_URI,
        'REDIRECT_PATH': REDIRECT_PATH,
        'SECRET_KEY_SET': bool(os.getenv('SECRET_KEY')),
        'CLIENT_ID': CLIENT_ID,
        'TENANT_ID': TENANT_ID,
        'AUTHORITY': AUTHORITY,
        'session_cookie_secure': app.config.get('SESSION_COOKIE_SECURE'),
        'session_cookie_samesite': app.config.get('SESSION_COOKIE_SAMESITE'),
        'request_scheme': request.scheme,
        'request_host': request.host,
        'full_url': request.url
    })


@app.route('/debug/session-test')
def session_test():
    """Test if sessions are working"""
    import uuid

    # Try to set a test value in session
    test_value = str(uuid.uuid4())
    session['test_key'] = test_value

    # Check if we can read it back
    retrieved = session.get('test_key')

    return jsonify({
        'session_working': retrieved == test_value,
        'test_value_set': test_value,
        'test_value_retrieved': retrieved,
        'session_keys': list(session.keys()),
        'cookie_name': app.config.get('SESSION_COOKIE_NAME'),
        'cookies_in_request': list(request.cookies.keys())
    })


def _oauth_redirect_uri() -> str:
    """Must match the browser host (localhost vs 127.0.0.1) and Entra registration."""
    if os.getenv('WEBSITE_HOSTNAME'):
        return REDIRECT_URI
    # Prefer URI stored on /login so code exchange matches authorize request
    stored = session.get('oauth_redirect_uri')
    if stored:
        return stored
    root = (request.url_root or 'http://localhost:5000/').rstrip('/')
    return f"{root}{REDIRECT_PATH}"


@app.route('/login')
def login():
    """Initiate Azure AD SSO login"""
    # Drop prior auth data but keep the filesystem session identity stable so the
    # browser cookie still maps to the same server-side file after AAD returns.
    for _k in (
        'user', 'access_token', 'fabric_access_token',
        'token_cache', 'login_at', 'next', 'test_key',
    ):
        session.pop(_k, None)

    # Generate a unique state token to prevent CSRF attacks
    session["state"] = str(uuid.uuid4())
    # Use the host the user opened (localhost ≠ 127.0.0.1 cookies)
    redirect_uri = REDIRECT_URI
    if not os.getenv('WEBSITE_HOSTNAME'):
        root = (request.url_root or 'http://localhost:5000/').rstrip('/')
        redirect_uri = f"{root}{REDIRECT_PATH}"
    session["oauth_redirect_uri"] = redirect_uri
    session.modified = True
    session.permanent = True  # keep cookie across the AAD round-trip

    print("\n🔐 LOGIN INITIATED:")
    print(f"   Redirect URI: {redirect_uri}")
    print(f"   Config REDIRECT_URI: {REDIRECT_URI}")
    print(f"   Request host: {request.host}")
    print(f"   State token: {session['state']}")
    print(f"   Cookie secure: {app.config.get('SESSION_COOKIE_SECURE')}")
    print(f"   Session dir: {app.config.get('SESSION_FILE_DIR')}")
    print(f"   Session keys: {list(session.keys())}")
    print(f"   Cookies in: {list(request.cookies.keys())}")

    # Build authorization URL with prompt=select_account to force account selection
    auth_url = msal_app.get_authorization_request_url(
        SCOPE,
        state=session["state"],
        redirect_uri=redirect_uri,
        prompt="select_account"  # Force user to select account (no auto-login)
    )

    print(f"   Auth URL: {auth_url[:120]}...")

    return redirect(auth_url)


@app.route(REDIRECT_PATH)
def authorized():
    """Handle the redirect from Azure AD after authentication"""

    redirect_uri = _oauth_redirect_uri()

    # Debug logging for production troubleshooting
    print(f"\n🔍 SSO CALLBACK RECEIVED:")
    print(f"   Request URL: {request.url}")
    print(f"   Request host: {request.host}")
    print(f"   Request state: {request.args.get('state')}")
    print(f"   Session state: {session.get('state')}")
    print(f"   Session keys: {list(session.keys())}")
    print(f"   oauth_redirect_uri: {redirect_uri}")
    print(f"   Cookies: {list(request.cookies.keys())}")
    print(f"   Has code: {bool(request.args.get('code'))}")
    print(f"   Has error: {bool(request.args.get('error'))}")

    # Verify state. On mismatch: clear orphan cookie/session once, then offer a clean
# /login — never leave users refreshing a one-time getAToken?code=… URL.
# Successful path below is unchanged (code exchange / session fill).
    if request.args.get('state') != session.get("state"):
        print(f"❌ STATE MISMATCH - Session may not be persisting!")
        print(f"   SESSION_COOKIE_SECURE={app.config.get('SESSION_COOKIE_SECURE')}")
        print(f"   REDIRECT_URI config={REDIRECT_URI}")
        print(f"   request_state={request.args.get('state')!r} session_state={session.get('state')!r}")
        print(f"   cookies={list(request.cookies.keys())} session_keys={list(session.keys())}")
        print(f"   session_backend={_session_backend}")

        # Drop broken server session + oauth leftovers so the next /login is clean.
        try:
            session.clear()
            session.modified = True
        except Exception as _clr_err:
            print(f"   session.clear failed: {_clr_err}")

        _cookie_name = app.config.get('SESSION_COOKIE_NAME') or 'pbi_session'
        _retry_cookie = 'pbi_oauth_retry'
        _is_azure = bool(os.getenv('WEBSITE_HOSTNAME'))
        _secure = bool(app.config.get('SESSION_COOKIE_SECURE'))
        _samesite = app.config.get('SESSION_COOKIE_SAMESITE') or 'Lax'
        # One auto-retry via /login (avoids stuck getAToken bookmark loops).
        # Tracked with a short-lived cookie so the Entra round-trip still counts
        # as "already retried" if state is missing again.
        _already_retried = (request.cookies.get(_retry_cookie) or '').strip() == '1'
        if not _already_retried:
            print("   → one-shot clean redirect to /login (cleared orphan session)")
            resp = redirect(url_for('login'))
            resp.set_cookie(
                _cookie_name,
                '',
                expires=0,
                max_age=0,
                path='/',
                secure=_secure,
                httponly=True,
                samesite=_samesite,
            )
            resp.set_cookie(
                _retry_cookie,
                '1',
                max_age=180,
                path='/',
                secure=_secure,
                httponly=True,
                samesite=_samesite,
            )
            return resp

        _help_prod = f"""
<p style="background:#eef6ff;border:1px solid #bcd;padding:12px;border-radius:8px">
  <b>Production tip:</b> Close this tab. Open
  <a href="{url_for('login')}"><code>/login</code></a> only
  (do not refresh the long <code>getAToken?code=…</code> URL).
  Ensure App Setting <code>SECRET_KEY</code> is set and stable across restarts.
  Optional: set <code>SESSION_REDIS_URL</code> so OAuth state survives worker recycles.
</p>"""
        _help_local = """
<ol>
<li>Always open <b>http://localhost:5000</b> (not 127.0.0.1) unless both URIs are in Entra.</li>
<li>Entra redirect URIs must include <code>http://localhost:5000/getAToken</code>.</li>
<li>.env: <code>SESSION_COOKIE_SECURE=false</code>, stable <code>SECRET_KEY</code>.</li>
<li>Use a fresh Incognito window, then open <a href="/login">/login</a>.</li>
<li><a href="/debug/session-test">/debug/session-test</a> — <code>session_working</code> must stay true.</li>
</ol>"""
        html = f"""<!doctype html><html><body style="font-family:Segoe UI,sans-serif;max-width:740px;margin:40px auto;padding:0 16px;line-height:1.45">
<h2>Sign-in session lost (OAuth state mismatch)</h2>
<p>Azure AD returned, but the server session no longer held the CSRF <code>state</code>
(orphan cookie, recycled worker, pruned session file, or stale <code>getAToken</code> URL).</p>
<pre style="background:#f4f4f4;padding:12px;overflow:auto">request_state = {request.args.get('state')!r}
session_state = {session.get('state')!r}
cookies = {list(request.cookies.keys())!r}
host = {request.host!r}
redirect_uri = {redirect_uri!r}
cookie_secure = {app.config.get('SESSION_COOKIE_SECURE')!r}
session_dir = {app.config.get('SESSION_FILE_DIR')!r}
session_backend = {_session_backend!r}
</pre>
{_help_prod if _is_azure else _help_local}
<p><a href="{url_for('login')}" style="display:inline-block;margin-top:8px;padding:10px 16px;background:#0b5fff;color:#fff;text-decoration:none;border-radius:6px">Try sign-in again</a>
 · <a href="/debug/env">/debug/env</a>
 · <a href="/debug/session-test">/debug/session-test</a></p>
</body></html>"""
        resp = app.make_response((html, 400))
        resp.set_cookie(
            _cookie_name,
            '',
            expires=0,
            max_age=0,
            path='/',
            secure=_secure,
            httponly=True,
            samesite=_samesite,
        )
        # Allow a future mismatch to auto-retry once again after user acts.
        resp.set_cookie(
            _retry_cookie,
            '',
            expires=0,
            max_age=0,
            path='/',
            secure=_secure,
            httponly=True,
            samesite=_samesite,
        )
        return resp

    # Check for errors from Azure AD
    if "error" in request.args:
        error_description = request.args.get("error_description", "Unknown error")
        flash(f'Authentication failed: {error_description}', 'error')
        return redirect(url_for("login"))

    # Exchange authorization code for Power BI access token
    if request.args.get('code'):
        # Drop any leftover step-up markers from older builds
        session.pop("oauth_purpose", None)
        session.pop("oauth_next", None)
        # Bind MSAL to the session token_cache so refresh tokens persist.
        cca, cache = _msal_for_request()

        print("\n🔐 ACQUIRING TOKEN WITH SCOPES:")
        for scope in SCOPE:
            print(f"   - {scope}")
        print(f"   Using redirect_uri={redirect_uri}")

        result = cca.acquire_token_by_authorization_code(
            request.args['code'],
            scopes=SCOPE,
            redirect_uri=redirect_uri,
        )

        if "error" in result:
            print(f"\n❌ TOKEN ACQUISITION FAILED:")
            print(f"   Error: {result.get('error')}")
            print(f"   Description: {result.get('error_description')}")
            print(f"   Correlation ID: {result.get('correlation_id')}")
            flash(f'Authentication failed: {result.get("error_description")}', 'error')
            return redirect(url_for("login"))

        print("\n✅ TOKEN ACQUIRED SUCCESSFULLY")
        print(f"   Scopes in result: {result.get('scope', 'N/A')}")
        print(f"   Token type: {result.get('token_type', 'N/A')}")
        print(f"   Expires in: {result.get('expires_in', 'N/A')} seconds")

        # Replace identity in-place — do NOT session.clear() (drops FS session
        # continuity / cookie mapping and can loop login on localhost).
        old_user_id = (session.get('user') or {}).get('oid')
        keep_keys = {"state", "oauth_redirect_uri"}
        for _k in list(session.keys()):
            if _k not in keep_keys:
                session.pop(_k, None)
        session.pop('state', None)

        # Non-permanent cookie → expires when browser is fully closed.
        # Absolute 12h bound enforced via login_at + login_required.
        session.permanent = False
        # Keep only fields we use — full id_token_claims can be large
        _claims = result.get("id_token_claims") or {}
        session["user"] = {
            "oid": _claims.get("oid"),
            "name": _claims.get("name"),
            "preferred_username": _claims.get("preferred_username")
                or _claims.get("upn")
                or _claims.get("email"),
            "email": _claims.get("email") or _claims.get("preferred_username"),
            "tid": _claims.get("tid"),
        }
        session["access_token"] = result.get("access_token")  # Power BI token
        session["login_at"] = datetime.now(timezone.utc).isoformat()
        session.modified = True
        _save_cache(cache)
        print(f"   Session after login keys: {list(session.keys())}")
        print(f"   User: {session['user'].get('preferred_username')}")
        if old_user_id and old_user_id != session['user'].get('oid'):
            print(f"   Previous oid was {old_user_id} (replaced)")

        # Get the new user ID and clear their old cache if any
        new_user_id = session.get('user', {}).get('oid')
        if new_user_id:
            clear_user_cache(new_user_id)
            print(f"\n👤 NEW LOGIN: {session.get('user', {}).get('name')} ({session.get('user', {}).get('preferred_username')})")
            print(f"   User ID: {new_user_id}")

        # Also clear old user cache if different user
        if old_user_id and old_user_id != new_user_id:
            clear_user_cache(old_user_id)

        # Optional one-shot post-login landing (scroll-to-enter). Not auto-zoom.
        session['show_login_landing'] = True
        resp = redirect(url_for("index"))
        # Clear one-shot OAuth retry marker after a successful sign-in.
        resp.set_cookie(
            'pbi_oauth_retry',
            '',
            expires=0,
            max_age=0,
            path='/',
            secure=bool(app.config.get('SESSION_COOKIE_SECURE')),
            httponly=True,
            samesite=app.config.get('SESSION_COOKIE_SAMESITE') or 'Lax',
        )
        return resp

    flash('No authorization code received', 'error')
    return redirect(url_for("login"))


@app.route('/logout')
def logout():
    """Logout and clear session"""
    session.clear()

    # Redirect to Azure AD logout endpoint
    logout_url = f"{AUTHORITY}/oauth2/v2.0/logout?post_logout_redirect_uri={url_for('login', _external=True)}"
    return redirect(logout_url)


@app.route('/api/debug/token')
@login_required
def debug_token():
    """Debug endpoint to check user's token and scopes"""
    try:
        import jwt
        token = get_user_powerbi_token()

        # Decode token without verification (just to inspect)
        decoded = jwt.decode(token, options={"verify_signature": False})

        return jsonify({
            'success': True,
            'user': session.get('user', {}),
            'scopes': decoded.get('scp', 'No scopes found'),
            'roles': decoded.get('roles', 'No roles found'),
            'aud': decoded.get('aud', 'No audience found'),
            'app_displayname': decoded.get('app_displayname', 'No app name found')
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        })


def _jwt_seconds_left(token):
    """Return seconds until JWT exp, or None if unreadable."""
    try:
        import base64
        import json
        import time as _t
        parts = (token or "").split(".")
        if len(parts) != 3:
            return None
        payload = parts[1] + "=" * (4 - len(parts[1]) % 4)
        data = json.loads(base64.b64decode(payload))
        exp = float(data.get("exp") or 0)
        return exp - _t.time()
    except Exception:
        return None


def get_user_powerbi_token():
    """Get the user's Power BI access token from session or refresh it if needed.

    Prefer the in-session JWT when it still has >5 minutes left so workspace
    loads do not hit MSAL/AAD on every folder-meta request.
    """
    # 1) Reuse valid session token (fast path — no AAD round-trip)
    if "access_token" in session:
        left = _jwt_seconds_left(session["access_token"])
        if left is not None and left > 300:
            # Quiet reuse — avoid log spam on every workspace select
            return session["access_token"]
        if left is not None and left <= 300:
            print(f"⚠️ Session Power BI token expiring in {int(left)}s — refreshing…")
            session.pop("access_token", None)

    # 2) MSAL silent refresh only when needed (session-bound cache)
    cca, cache = _msal_for_request()
    accounts = cca.get_accounts(
        username=session.get("user", {}).get("preferred_username")
    )
    if not accounts:
        accounts = cca.get_accounts()

    if accounts:
        print("🔄 Acquiring Power BI token silently (MSAL)…")
        result = cca.acquire_token_silent(SCOPE, account=accounts[0])
        if result and "access_token" in result and "error" not in result:
            print("✅ Power BI token acquired/refreshed")
            session["access_token"] = result["access_token"]
            _save_cache(cache)
            return result["access_token"]
        if result and "error" in result:
            print(f"⚠️ Token refresh error: {result.get('error_description', result.get('error'))}")

    print("❌ No valid Power BI token — user needs to re-authenticate")
    return None


def get_user_fabric_token():
    """
    Fabric API token (audience api.fabric.microsoft.com).

    Reuse session JWT when still valid (>5 min) so Report Catalog workspace
    loads do not call MSAL on every request.
    """
    # 1) Reuse valid Fabric token in session
    if "fabric_access_token" in session:
        left = _jwt_seconds_left(session["fabric_access_token"])
        if left is not None and left > 300:
            return session["fabric_access_token"]
        if left is not None and left <= 300:
            session.pop("fabric_access_token", None)

    # 2) Silent acquire with Fabric scope (session-bound MSAL cache)
    cca, cache = _msal_for_request()
    accounts = cca.get_accounts(
        username=session.get("user", {}).get("preferred_username")
    )
    if not accounts:
        accounts = cca.get_accounts()

    if accounts:
        result = cca.acquire_token_silent(FABRIC_SCOPE, account=accounts[0])
        if result and "access_token" in result and "error" not in result:
            print("✅ Fabric token acquired (silent)")
            session["fabric_access_token"] = result["access_token"]
            _save_cache(cache)
            return result["access_token"]

    # 3) OBO from Power BI token (often fails if assertion audience mismatch — keep as fallback)
    pbi_token = get_user_powerbi_token()
    if pbi_token:
        try:
            result = cca.acquire_token_on_behalf_of(
                user_assertion=pbi_token,
                scopes=FABRIC_SCOPE
            )
            if result and "access_token" in result and "error" not in result:
                print("✅ Fabric token acquired (OBO)")
                session["fabric_access_token"] = result["access_token"]
                _save_cache(cache)
                return result["access_token"]
            else:
                error = result.get('error', 'unknown') if result else 'no result'
                error_desc = result.get('error_description', '') if result else ''
                print(f"⚠️ Fabric OBO failed: {error} - {str(error_desc)[:120]}")
        except Exception as e:
            print(f"⚠️ Fabric OBO error: {e}")

    print("⚠️ Could not acquire Fabric API token")
    return None


def get_user_powerbi_headers():
    """Get HTTP headers for Power BI API requests using user's delegated token"""
    token = get_user_powerbi_token()

    if not token:
        raise Exception("User not authenticated or token expired. Please log in again.")

    # Debug: Decode token to check scopes (for troubleshooting)
    try:
        import base64
        import json
        # JWT tokens have 3 parts separated by dots
        parts = token.split('.')
        if len(parts) == 3:
            # Decode the payload (second part)
            # Add padding if needed
            payload = parts[1]
            payload += '=' * (4 - len(payload) % 4)
            decoded = base64.b64decode(payload)
            token_data = json.loads(decoded)

            print("\n🔍 TOKEN DEBUG INFO:")
            print(f"   Token Audience (aud): {token_data.get('aud', 'N/A')}")
            print(f"   Token Scopes (scp): {token_data.get('scp', 'N/A')}")
            print(f"   Token Roles (roles): {token_data.get('roles', 'N/A')}")
            print(f"   Token Issuer: {token_data.get('iss', 'N/A')}")
            print(f"   User UPN: {token_data.get('upn', 'N/A')}")
            print()
    except Exception as e:
        print(f"⚠️  Could not decode token for debugging: {e}")

    return {
        'Authorization': f'Bearer {token}',
        'Content-Type': 'application/json'
    }


@app.route('/health')
def health():
    """Health check endpoint for Docker and monitoring"""
    catalog_status = None
    if CATALOG_AVAILABLE and catalog_service is not None:
        try:
            catalog_status = catalog_service.status()
        except Exception as exc:
            catalog_status = {'error': str(exc)}
    return jsonify({
        'status': 'healthy',
        'service': 'powerbi-documentation',
        # Prefer CI-provided build ID if present; fallback to short git hash placeholder
        'version': os.getenv('BUILD_ID', os.getenv('COMMIT_SHA', 'unknown')),
        'timestamp': datetime.now(timezone.utc).isoformat(),
        'cache_fix': 'enabled',  # Indicates cache headers are active
        'catalog': catalog_status,
    }), 200


@app.route('/api/auth-status')
@login_required
def auth_status():
    """Test endpoint to check authentication status and token validity"""
    user_info = session.get('user', {})
    has_token = 'access_token' in session

    token_info = "No token"
    if has_token:
        token = session.get('access_token', '')
        token_info = f"Token exists (length: {len(token)})"

    return jsonify({
        'authenticated': True,
        'user_name': user_info.get('name', 'Unknown'),
        'user_email': user_info.get('preferred_username', 'Unknown'),
        'user_id': user_info.get('oid', 'Unknown'),
        'has_access_token': has_token,
        'token_info': token_info,
        'session_keys': list(session.keys())
    })


@app.route('/api/test-sso')
@login_required
def test_sso():
    """Test endpoint to verify if SSO-based workspace access is working"""
    import requests

    result = {
        'user_info': {
            'name': session.get('user', {}).get('name', 'Unknown'),
            'email': session.get('user', {}).get('preferred_username', 'Unknown'),
            'id': session.get('user', {}).get('oid', 'Unknown')
        },
        'has_session_token': 'access_token' in session,
        'test_results': {}
    }

    # Test 1: Try to get workspaces with user token
    try:
        headers = get_user_powerbi_headers()
        response = requests.get('https://api.powerbi.com/v1.0/myorg/groups', headers=headers)

        if response.status_code == 200:
            workspaces = response.json().get('value', [])
            result['test_results']['user_token'] = {
                'status': 'SUCCESS',
                'status_code': 200,
                'workspace_count': len(workspaces),
                'workspaces': [{'id': w['id'], 'name': w['name']} for w in workspaces[:10]]  # First 10
            }
        else:
            result['test_results']['user_token'] = {
                'status': 'FAILED',
                'status_code': response.status_code,
                'error': response.text[:500]
            }
    except Exception as e:
        result['test_results']['user_token'] = {
            'status': 'ERROR',
            'error': str(e)
        }

    # Test 2: Try to get workspaces with service principal (for comparison)
    try:
        sp_workspaces = powerbi.get_workspaces()
        result['test_results']['service_principal'] = {
            'status': 'SUCCESS',
            'workspace_count': len(sp_workspaces),
            'workspaces': [{'id': w['id'], 'name': w['name']} for w in sp_workspaces[:10]]  # First 10
        }
    except Exception as e:
        result['test_results']['service_principal'] = {
            'status': 'ERROR',
            'error': str(e)
        }

    # Determine which method is being used
    user_count = result['test_results'].get('user_token', {}).get('workspace_count', 0)
    sp_count = result['test_results'].get('service_principal', {}).get('workspace_count', 0)

    result['conclusion'] = {
        'sso_working': result['test_results'].get('user_token', {}).get('status') == 'SUCCESS',
        'user_workspace_count': user_count,
        'service_principal_workspace_count': sp_count,
        'recommendation': ''
    }

    if result['conclusion']['sso_working']:
        result['conclusion']['recommendation'] = '✅ SSO is working! User-delegated token successfully fetched workspaces.'
    else:
        result['conclusion']['recommendation'] = '❌ SSO is NOT working. User-delegated token failed. Check Azure AD permissions.'

    return jsonify(result)


@app.route('/')
@login_required
def index():
    """Home page - Dashboard Overview"""
    # Optional scroll-to-enter landing after SSO (user scrolls / clicks — no auto zoom)
    show_login_landing = bool(session.pop('show_login_landing', False))
    return render_template('home.html', show_login_landing=show_login_landing)


@app.route('/documentation')
@login_required
def documentation():
    """Documentation page - Report documentation generation"""
    can_archive = False
    try:
        from features.report_archive_service import user_can_archive
        email = session.get('user', {}).get('preferred_username') or ''
        can_archive = user_can_archive(email)
    except Exception:
        can_archive = False
    return render_template('index.html', can_archive_reports=can_archive)


@app.route('/semantic-models')
@login_required
def semantic_models_page():
    """Semantic Models page - Analyze and health-check semantic models"""
    return render_template('semantic_models.html')


@app.route('/impact')
@login_required
def impact_explorer_page():
    """
    Table impact explorer (EDW → report blast radius).
    Serves embedded Impact UI backed by precomputed catalog JSON.
    """
    return render_template('impact.html')


# =============================================================================
# CATALOG FAST PATH APIs (precomputed SharePoint/local metadata)
# Live Power BI APIs remain the fallback when catalog is unavailable.
# =============================================================================

def _user_allowed_workspace_ids():
    """
    Workspace IDs the current SSO user can access (from live /groups, cached).
    Returns None if undetermined (caller may skip filtering only for admin tools).
    Returns set() on hard failure after auth.
    """
    try:
        user_id = session.get('user', {}).get('oid', 'unknown')
        cache_key = f"workspaces_{user_id}"
        current_time = time.time()
        if cache_key in workspaces_cache and workspaces_cache[cache_key].get('data') and \
           (current_time - workspaces_cache[cache_key].get('timestamp', 0)) < CACHE_DURATION:
            return {w.get('id') for w in workspaces_cache[cache_key]['data'] if w.get('id')}

        # Lightweight live fetch (same as get_workspaces core)
        headers = get_user_powerbi_headers()
        resp = requests.get("https://api.powerbi.com/v1.0/myorg/groups", headers=headers, timeout=30)
        if resp.status_code != 200:
            print(f"⚠️ allowed workspaces fetch HTTP {resp.status_code}")
            return set()
        workspaces = resp.json().get('value', [])
        workspaces_cache[cache_key] = {'data': workspaces, 'timestamp': current_time}
        return {w.get('id') for w in workspaces if w.get('id')}
    except Exception as exc:
        print(f"⚠️ _user_allowed_workspace_ids: {exc}")
        return set()


@app.route('/api/catalog/status')
@login_required
def api_catalog_status():
    """Catalog availability, mode, freshness — for UI banners and ops."""
    if not CATALOG_AVAILABLE or catalog_service is None:
        return jsonify({
            'success': True,
            'enabled': False,
            'mode': 'off',
            'message': 'Catalog service not loaded; using live APIs only.',
        })
    try:
        st = catalog_service.status()
        st['success'] = True
        return jsonify(st)
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


def _perform_catalog_refresh() -> dict:
    """
    Shared body for interactive and machine catalog refresh.
    Force re-download of catalog artifacts from SharePoint into server memory/disk mirror.
    Also clears per-user /api/reports response cache so Report Catalog picks up new ops.
    """
    if not CATALOG_AVAILABLE or catalog_service is None:
        return {'success': False, 'error': 'Catalog not available', '_http': 503}

    catalog_service.invalidate()
    # Prefer thin packs first (fast UI), then heavy files for report/model APIs
    home = catalog_service.get_json('ui_home_index.json', force_refresh=True)
    tables = catalog_service.get_json('ui_impact_tables.json', force_refresh=True)
    try:
        impact_reports = catalog_service.get_json('ui_impact_reports.json', force_refresh=True)
    except Exception:
        impact_reports = None
    summary = catalog_service.get_summary(force_refresh=True)
    cat = catalog_service.get_workspace_catalog(force_refresh=True)
    impact = catalog_service.get_impact_index(force_refresh=True)
    # Rebuild report→sources pack if SharePoint has old pack without it
    if impact and (not impact_reports or not isinstance((impact_reports or {}).get('rows'), list)):
        try:
            catalog_service._ensure_thin_impact_pack(impact)
            impact_reports = catalog_service.get_json('ui_impact_reports.json')
        except Exception as exc:
            print(f"⚠️ thin impact reports pack rebuild: {exc}")
    # Ops snapshot used for last refresh / views — reload so Catalog is not stuck on old ops
    try:
        catalog_service.get_json('refresh_snapshot.json', force_refresh=True)
    except Exception:
        pass
    try:
        catalog_service.get_json('usage_snapshot.json', force_refresh=True)
    except Exception:
        pass

    # Ensure Home KPI detail lists + report directory exist (older SP packs may lack them)
    home_has_details = isinstance(home, dict) and isinstance(home.get('detailLists'), dict)
    report_dir = None
    try:
        report_dir = catalog_service.get_json('ui_report_directory.json')
    except Exception:
        report_dir = None
    need_home_rebuild = cat and (
        not home_has_details
        or not report_dir
        or not isinstance((report_dir or {}).get('rows'), list)
    )
    if need_home_rebuild:
        try:
            catalog_service._ensure_thin_home_pack(cat)
            home = catalog_service.get_json('ui_home_index.json')
            home_has_details = isinstance(home, dict) and isinstance(home.get('detailLists'), dict)
            report_dir = catalog_service.get_json('ui_report_directory.json')
        except Exception as exc:
            print(f"⚠️ thin home/report-directory pack rebuild after refresh: {exc}")

    # Drop in-process /api/reports shells + folder tree so next load is fresh
    cleared_reports = 0
    try:
        global reports_cache, workspace_folders_cache
        cleared_reports = len(reports_cache)
        reports_cache = {}
        workspace_folders_cache = {}
    except Exception:
        pass

    ops_at = None
    if isinstance(cat, dict):
        ops_at = cat.get('opsEnrichedAt') or cat.get('generatedAt')
    return {
        'success': True,
        'ui_home_index': bool(home),
        'ui_home_detailLists': home_has_details,
        'ui_report_directory': bool(report_dir and (report_dir.get('rows') is not None)),
        'ui_impact_tables': bool(tables),
        'ui_impact_reports': bool(impact_reports),
        'workspace_catalog': bool(cat),
        'impact_index': bool(impact),
        'summary': bool(summary),
        'opsEnrichedAt': ops_at,
        'generatedAt': (cat or {}).get('generatedAt') if isinstance(cat, dict) else None,
        'clearedReportsCacheEntries': cleared_reports,
        'status': catalog_service.status(),
    }


@app.route('/api/catalog/refresh', methods=['POST'])
@login_required
def api_catalog_refresh():
    """
    Force re-download of catalog artifacts from SharePoint into server memory/disk mirror.
    Also clears per-user /api/reports response cache so Report Catalog picks up new ops.
    """
    try:
        body = _perform_catalog_refresh()
        code = int(body.pop('_http', 200)) if isinstance(body, dict) else 200
        return jsonify(body), code
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/catalog/refresh-internal', methods=['POST'])
def api_catalog_refresh_internal():
    """
    Machine-callable catalog refresh after ops/fresh SharePoint publish.
    Auth: header X-Catalog-Refresh-Key (or Authorization: Bearer) must match
    App Setting CATALOG_REFRESH_SECRET. No user SSO session required.
    """
    secret = (os.getenv('CATALOG_REFRESH_SECRET') or '').strip()
    if not secret:
        return jsonify({
            'success': False,
            'error': 'CATALOG_REFRESH_SECRET is not configured on the app',
        }), 503

    provided = (
        (request.headers.get('X-Catalog-Refresh-Key') or '').strip()
        or (request.headers.get('X-Api-Key') or '').strip()
    )
    auth_h = (request.headers.get('Authorization') or '').strip()
    if auth_h.lower().startswith('bearer '):
        provided = provided or auth_h[7:].strip()
    if not provided or provided != secret:
        return jsonify({'success': False, 'error': 'Unauthorized'}), 401

    try:
        body = _perform_catalog_refresh()
        code = int(body.pop('_http', 200)) if isinstance(body, dict) else 200
        body['triggeredBy'] = 'internal'
        print(f"🔄 catalog refresh-internal OK opsEnrichedAt={body.get('opsEnrichedAt')}")
        return jsonify(body), code
    except Exception as e:
        print(f"❌ catalog refresh-internal failed: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/catalog/data/<path:name>')
@login_required
def api_catalog_data(name):
    """
    Serve *small* catalog JSON only (summary, ops_summary, ui packs).

    Large files (workspace_catalog, impact_index) are server-side only —
    use /api/catalog/impact/* and workspace report APIs instead.
    Browser never downloads SharePoint blobs.
    """
    if not CATALOG_AVAILABLE or catalog_service is None:
        return jsonify({'error': 'Catalog not available'}), 503
    safe = os.path.basename(name)
    if not safe.endswith('.json') or '..' in safe:
        return jsonify({'error': 'Only .json filenames allowed'}), 400
    # Hard block large artifacts from leaving the server
    try:
        from catalog_service import catalog_config as _ccfg
        blocked = getattr(_ccfg, 'BROWSER_BLOCKED_CATALOG_FILES', set())
        allowed = getattr(_ccfg, 'BROWSER_ALLOWED_CATALOG_FILES', None)
    except Exception:
        blocked = {
            'workspace_catalog.json', 'impact_index.json',
            'inventory.json', 'refresh_snapshot.json',
        }
        allowed = None
    if safe in blocked or (allowed is not None and safe not in allowed):
        return jsonify({
            'error': f'{safe} is server-side only and cannot be downloaded by the browser',
            'hint': (
                'Use thin APIs: /api/home-summary, /api/catalog/impact/tables, '
                '/api/catalog/impact/lookup, /api/catalog/impact/table?key=..., '
                '/api/reports?workspace_id=...'
            ),
        }), 403
    force = request.args.get('refresh') in ('1', 'true', 'yes')
    try:
        data = catalog_service.get_json(safe, force_refresh=force)
        if data is None:
            return jsonify({
                'error': f'File not available: {safe}',
                'hint': 'Run: python run_catalog_extract.py --fresh -v  (publishes to SharePoint latest/)',
                'status': catalog_service.status(),
            }), 404
        resp = jsonify(data)
        resp.headers['X-Data-Source'] = 'server-cache'
        return resp
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/catalog/impact/tables')
@login_required
def api_catalog_impact_tables():
    """
    Thin table list for Impact Explorer (no nested datasets).
    Strips bulky searchText field — browser builds search locally.
    Enables short browser/proxy cache so revisits are near-instant.
    """
    if not CATALOG_AVAILABLE or catalog_service is None:
        return jsonify({'success': False, 'error': 'Catalog not available'}), 503
    try:
        force = request.args.get('refresh') in ('1', 'true', 'yes')
        rows_in = catalog_service.impact_table_rows(force_refresh=force)
        # Compact rows for wire size (searchText alone is multi-MB)
        rows = []
        for r in rows_in or []:
            rows.append({
                'k': r.get('tableKey'),
                't': r.get('table'),
                'st': r.get('sourceType') or 'Unknown',
                'sv': r.get('server') or '',
                'db': r.get('database') or '',
                'sc': r.get('schema') or '',
                'mn': r.get('modelTableNames') or [],
                'rc': int(r.get('reportCount') or 0),
                'dc': int(r.get('datasetCount') or 0),
                'wc': int(r.get('workspaceCount') or 0),
            })
        summary = catalog_service.get_summary() or {}
        pack = catalog_service.get_json('ui_impact_tables.json') or {}
        payload = {
            'success': True,
            'v': 2,  # compact schema version
            'count': len(rows),
            'rows': rows,
            'generatedAt': pack.get('generatedAt') or summary.get('generatedAt'),
            'stats': summary.get('stats') or {},
            'source': 'server-thin',
        }
        resp = jsonify(payload)
        # Same user/session: allow brief cache so switching back to Impact is fast
        if not force:
            resp.headers['Cache-Control'] = 'private, max-age=120'
        else:
            resp.headers['Cache-Control'] = 'no-store'
        resp.headers['X-Data-Source'] = 'server-thin'
        return resp
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/catalog/impact/table')
@login_required
def api_catalog_impact_table_detail():
    """One table's full impact (drawer) — server reads index, browser gets one object."""
    if not CATALOG_AVAILABLE or catalog_service is None:
        return jsonify({'success': False, 'error': 'Catalog not available'}), 503
    key = (request.args.get('key') or request.args.get('table') or '').strip()
    if not key:
        return jsonify({'success': False, 'error': 'key (tableKey) required'}), 400
    try:
        entry = catalog_service.impact_table_detail(key)
        if not entry:
            return jsonify({'success': False, 'error': f'No impact entry for {key}'}), 404

        # Snapshot tenant-wide summary before ACL (grid uses this; drawer may be smaller)
        entry = dict(entry)
        tenant_summary = dict(entry.get('impactSummary') or {})
        entry['tenantImpactSummary'] = tenant_summary

        allowed = _user_allowed_workspace_ids()
        acl_applied = False
        if allowed is not None and len(allowed) > 0:
            acl_applied = True
            filtered_datasets = []
            for d in entry.get('datasets') or []:
                if not isinstance(d, dict):
                    continue
                d2 = dict(d)
                # Keep reports in workspaces the user can open
                reps = []
                for r in d2.get('reports') or []:
                    if not isinstance(r, dict):
                        continue
                    rid_ws = r.get('workspaceId') or d2.get('workspaceId') or ''
                    if rid_ws in allowed:
                        reps.append(r)
                d2['reports'] = reps
                # Keep dataset if its home workspace is allowed OR any remaining report is
                ds_ws = d2.get('workspaceId') or ''
                if ds_ws in allowed or reps:
                    # If dataset home is outside ACL but reports inside, still show
                    filtered_datasets.append(d2)
            entry['datasets'] = filtered_datasets

            report_ids = set()
            workspace_ids = set()
            for d in filtered_datasets:
                if d.get('workspaceId'):
                    workspace_ids.add(d['workspaceId'])
                for r in d.get('reports') or []:
                    if r.get('reportId'):
                        report_ids.add(r['reportId'])
                    if r.get('workspaceId'):
                        workspace_ids.add(r['workspaceId'])
            entry['impactSummary'] = {
                **tenant_summary,
                'datasetCount': len(filtered_datasets),
                'reportCount': len(report_ids),
                'workspaceCount': len(workspace_ids),
            }

        entry['aclApplied'] = acl_applied
        return jsonify({'success': True, 'table': entry})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/catalog/reports/search')
@login_required
def api_catalog_reports_search():
    """
    Reverse search: report name → workspace(s).
    Thin catalog directory (not full workspace_catalog in browser).
    ?q=wanek&limit=40
    """
    if not CATALOG_AVAILABLE or catalog_service is None:
        return jsonify({'success': False, 'error': 'Catalog not available'}), 503
    try:
        q = (request.args.get('q') or request.args.get('query') or '').strip()
        try:
            limit = int(request.args.get('limit') or 50)
        except Exception:
            limit = 50
        allowed = _user_allowed_workspace_ids()
        result = catalog_service.search_reports(
            query=q,
            allowed_workspace_ids=allowed if allowed is not None else None,
            limit=limit,
        )
        # Compact wire format
        rows = []
        for r in result.get('rows') or []:
            rows.append({
                'id': r.get('reportId'),
                'n': r.get('reportName') or '',
                'wid': r.get('workspaceId') or '',
                'wn': r.get('workspaceName') or '',
                'did': r.get('datasetId') or '',
            })
        payload = {
            'success': True,
            'query': result.get('query') or q,
            'count': len(rows),
            'total': result.get('total'),
            'capped': bool(result.get('capped')),
            'rows': rows,
            'source': 'ui_report_directory',
        }
        resp = jsonify(payload)
        resp.headers['Cache-Control'] = 'private, max-age=60'
        return resp
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/catalog/impact/reports')
@login_required
def api_catalog_impact_reports():
    """
    Thin report list for Impact Explorer «Report sources» tab (report → source counts).
    No nested source lists on the wire — use /impact/report for drawer detail.
    """
    if not CATALOG_AVAILABLE or catalog_service is None:
        return jsonify({'success': False, 'error': 'Catalog not available'}), 503
    try:
        force = request.args.get('refresh') in ('1', 'true', 'yes')
        allowed = _user_allowed_workspace_ids()
        rows_in = catalog_service.impact_report_rows(
            force_refresh=force,
            allowed_workspace_ids=allowed if allowed is not None else None,
        )
        rows = []
        for r in rows_in or []:
            rows.append({
                'id': r.get('reportId'),
                'n': r.get('reportName') or '',
                'wid': r.get('workspaceId') or '',
                'wn': r.get('workspaceName') or '',
                'rt': r.get('reportType') or '',
                'tc': int(r.get('tableCount') or 0),
                'dc': int(r.get('datasetCount') or 0),
                'st': r.get('sourceTypes') or [],
            })
        pack = catalog_service.get_json('ui_impact_reports.json') or {}
        payload = {
            'success': True,
            'v': 1,
            'count': len(rows),
            'rows': rows,
            'generatedAt': pack.get('generatedAt'),
            'source': 'server-thin',
        }
        resp = jsonify(payload)
        if not force:
            resp.headers['Cache-Control'] = 'private, max-age=120'
        else:
            resp.headers['Cache-Control'] = 'no-store'
        resp.headers['X-Data-Source'] = 'server-thin'
        return resp
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/catalog/impact/report')
@login_required
def api_catalog_impact_report_detail():
    """
    All sources for one report (SQL / Excel / SharePoint / model tables / …).
    Drawer payload for Report sources tab.
    """
    if not CATALOG_AVAILABLE or catalog_service is None:
        return jsonify({'success': False, 'error': 'Catalog not available'}), 503
    report_id = (request.args.get('report_id') or request.args.get('id') or '').strip()
    if not report_id:
        return jsonify({'success': False, 'error': 'report_id required'}), 400
    try:
        allowed = _user_allowed_workspace_ids()
        detail = catalog_service.impact_report_detail(
            report_id,
            allowed_workspace_ids=allowed if allowed is not None else None,
        )
        if not detail:
            return jsonify({'success': False, 'error': f'No sources for report {report_id}'}), 404
        return jsonify({'success': True, 'report': detail})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/catalog/impact/lookup')
@login_required
def api_catalog_impact_lookup():
    """Lookup table → datasets → reports blast radius from impact index (server-side)."""
    if not CATALOG_AVAILABLE or catalog_service is None:
        return jsonify({'success': False, 'error': 'Catalog not available'}), 503
    table = (request.args.get('table') or '').strip()
    if not table:
        return jsonify({'success': False, 'error': 'table query param required'}), 400
    try:
        hits = catalog_service.lookup_table(table)
        allowed = _user_allowed_workspace_ids()
        if allowed is not None and len(allowed) > 0:
            filtered = []
            for entry in hits:
                datasets = [d for d in (entry.get('datasets') or []) if d.get('workspaceId') in allowed]
                if datasets:
                    e2 = dict(entry)
                    e2['datasets'] = datasets
                    filtered.append(e2)
            hits = filtered
        return jsonify({'success': True, 'table': table, 'count': len(hits), 'results': hits})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/catalog/impact/model-details')
@login_required
def api_catalog_impact_model_details():
    """
    Thin semantic-model payload for Impact Explorer popup.
    Catalog-only (no full workspace_catalog in browser). Optional focus_table
    highlights the impact table the user came from.
    """
    if not CATALOG_AVAILABLE or catalog_service is None:
        return jsonify({'success': False, 'error': 'Catalog not available'}), 503
    dataset_id = (request.args.get('dataset_id') or '').strip()
    workspace_id = (request.args.get('workspace_id') or '').strip()
    focus_table = (request.args.get('focus_table') or request.args.get('table') or '').strip()
    model_table = (request.args.get('model_table') or '').strip()
    report_name = (request.args.get('report_name') or '').strip()
    report_id = (request.args.get('report_id') or '').strip()
    if not dataset_id:
        return jsonify({'success': False, 'error': 'dataset_id required'}), 400
    try:
        allowed = _user_allowed_workspace_ids()
        if allowed is not None and len(allowed) > 0 and workspace_id and workspace_id not in allowed:
            return jsonify({'success': False, 'error': 'Access denied for workspace'}), 403

        details = catalog_service.impact_model_details(
            dataset_id=dataset_id,
            workspace_id=workspace_id,
            focus_table=focus_table,
            model_table_name=model_table,
        )
        if not details:
            return jsonify({
                'success': False,
                'error': 'Dataset not found in catalog. Run a fresh extract if this model is new.',
            }), 404

        # ACL: if no workspace_id was passed, still enforce when we resolved one
        ws_resolved = details.get('workspaceId') or workspace_id
        if allowed is not None and len(allowed) > 0 and ws_resolved and ws_resolved not in allowed:
            return jsonify({'success': False, 'error': 'Access denied for workspace'}), 403

        payload = {
            'success': True,
            'source': 'catalog',
            'reportName': report_name or None,
            'reportId': report_id or None,
            **details,
        }
        resp = jsonify(payload)
        resp.headers['Cache-Control'] = 'private, max-age=60'
        resp.headers['X-Data-Source'] = 'catalog-impact-model'
        return resp
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/home-summary')
@login_required
def api_home_summary():
    """
    Fast Home dashboard: one response for all accessible workspaces.
    Prefers precomputed SharePoint/local catalog (report/inactive/orphaned counts).
    Falls back to workspace list only if catalog unavailable (UI may still use live paths).
    """
    try:
        allowed = _user_allowed_workspace_ids()
        # allowed may be empty set on failure — still try catalog unfiltered only if None
        if CATALOG_AVAILABLE and catalog_service is not None and catalog_service.is_available():
            summary = catalog_service.build_home_summary(
                allowed_workspace_ids=allowed if allowed is not None else None,
                inactive_days=30,
            )
            if summary:
                print(
                    f"⚡ HOME SUMMARY from catalog: ws={summary.get('workspaceCount')} "
                    f"reports={summary.get('totalReports')} "
                    f"inactive={summary.get('inactiveReports')} "
                    f"orphaned={summary.get('orphanedReports')} "
                    f"zeroViews={summary.get('zeroViewsReports')} "
                    f"opsEnrichedAt={summary.get('opsEnrichedAt')}"
                )
                return jsonify(summary)

        # Catalog miss — return minimal payload so UI can fall back
        return jsonify({
            'success': True,
            'source': 'none',
            'fallback': True,
            'message': 'Catalog not available; use live home loaders',
            'workspaceCount': len(allowed or []),
            'totalReports': None,
            'inactiveReports': None,
            'orphanedReports': None,
            'zeroViewsReports': None,
            'workspaces': [],
        })
    except Exception as e:
        print(f"❌ home-summary error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/home-summary/details')
@login_required
def api_home_summary_details():
    """
    Row lists for Home KPI tabs.
    ?metric=workspaces|reports|inactive|orphaned|zero_views
    """
    try:
        metric = (request.args.get('metric') or 'workspaces').strip().lower()
        try:
            limit = int(request.args.get('limit') or 5000)
        except Exception:
            limit = 5000
        limit = max(1, min(limit, 20000))
        allowed = _user_allowed_workspace_ids()
        if not CATALOG_AVAILABLE or catalog_service is None or not catalog_service.is_available():
            return jsonify({
                'success': False,
                'error': 'Catalog not available',
                'metric': metric,
                'rows': [],
            }), 503
        details = catalog_service.build_home_details(
            metric=metric,
            allowed_workspace_ids=allowed if allowed is not None else None,
            inactive_days=30,
            limit=limit,
        )
        if not details:
            return jsonify({'success': False, 'error': 'No details', 'metric': metric, 'rows': []}), 404
        if details.get('success') is False:
            return jsonify(details), 400
        return jsonify(details)
    except Exception as e:
        print(f"❌ home-summary/details error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/catalog/impact/top')
@login_required
def api_catalog_impact_top():
    """Top blast-radius tables from precomputed impact index."""
    if not CATALOG_AVAILABLE or catalog_service is None:
        return jsonify({'success': False, 'error': 'Catalog not available'}), 503
    try:
        n = min(int(request.args.get('n', 50)), 500)
        rows = catalog_service.impact_top(n=n)
        return jsonify({'success': True, 'count': len(rows), 'tables': rows,
                        'generatedAt': (catalog_service.get_impact_index() or {}).get('generatedAt')})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


def _resolve_person_identity(*values):
    """
    Normalize creator/modifier identity from Power BI API variants.

    Handles:
      - plain UPN/email/display strings (users AND team/DL mailboxes)
      - nested user objects (userPrincipalName, emailAddress, displayName, …)
      - skips empty / Unknown / bare GUIDs (keeps GUID only as last resort)

    Returns '' when nothing usable is found (UI maps that to N/A).
    """
    def _looks_like_guid(s: str) -> bool:
        s = (s or '').strip()
        if len(s) != 36 or s.count('-') != 4:
            return False
        hex_parts = s.replace('-', '')
        return len(hex_parts) == 32 and all(c in '0123456789abcdefABCDEF' for c in hex_parts)

    def _from_dict(d):
        if not isinstance(d, dict):
            return ''
        # Nested shapes sometimes appear from admin / Graph-ish payloads
        for nest_key in ('user', 'principal', 'identity', 'account'):
            nested = d.get(nest_key)
            if isinstance(nested, dict):
                got = _from_dict(nested)
                if got:
                    return got
        for key in (
            'userPrincipalName', 'emailAddress', 'email', 'mail',
            'principalName', 'upn', 'displayName', 'name',
            'identifier', 'id', 'objectId',
        ):
            raw = d.get(key)
            if raw is None:
                continue
            s = str(raw).strip()
            if not s:
                continue
            if s.lower() in {'unknown', 'n/a', 'none', 'null', '-', '—', 'undefined'}:
                continue
            if _looks_like_guid(s) and '@' not in s:
                continue  # prefer a human label later; GUID last-resort below
            return s
        # Last resort: any non-empty GUID id so row isn't blank
        for key in ('id', 'objectId', 'identifier'):
            raw = d.get(key)
            if raw and _looks_like_guid(str(raw)):
                return str(raw).strip()
        return ''

    guid_fallback = ''
    for value in values:
        if value is None:
            continue
        if isinstance(value, dict):
            got = _from_dict(value)
            if got:
                return got
            continue
        s = str(value).strip()
        if not s:
            continue
        if s.lower() in {'unknown', 'n/a', 'none', 'null', '-', '—', 'undefined'}:
            continue
        if _looks_like_guid(s) and '@' not in s:
            if not guid_fallback:
                guid_fallback = s
            continue
        return s
    return guid_fallback


def _pick_person(*values):
    """First non-empty resolved person/team identity."""
    for v in values:
        got = _resolve_person_identity(v)
        if got:
            return got
    return ''


def _pick_datetime(*values):
    """First non-empty datetime-like string from API variants."""
    for v in values:
        if v is None:
            continue
        s = str(v).strip()
        if s and s.lower() not in {'unknown', 'n/a', 'none', 'null', '-', '—'}:
            return s
    return ''


@app.route('/api/reports-metadata/<workspace_id>')
@login_required
def get_reports_metadata(workspace_id):
    """
    Creator / modifier / dates for reports in a workspace.

    Multi-source (does not break existing UI contract):
      1) Admin Scanner (createdBy / modifiedBy when tenant returns them)
      2) Groups REST /reports overlay — often has people + dates when Scanner
         omits team mailbox / DL / group principals
      3) Catalog fill for remaining gaps (after extract preserves owner fields)

    Response shape unchanged:
      { success, workspace_id, reports: [{ report_id, report_name, created_by,
        created_date_time, modified_by, modified_date_time, created_by_id,
        modified_by_id }], cached }
    Empty identity fields are '' (UI still shows N/A).
    """
    try:
        from scanner_connector import PowerBIScanner
        import time

        # Check cache first (5-minute TTL)
        cache_key = f"metadata_{workspace_id}"
        if cache_key in reports_cache:
            cached_data, timestamp = reports_cache[cache_key]
            if time.time() - timestamp < 300:  # 5 minutes
                print(f"📦 Using cached metadata for workspace {workspace_id}")
                return jsonify({
                    'success': True,
                    'workspace_id': workspace_id,
                    'reports': cached_data,
                    'cached': True
                })

        t0 = time.time()
        print(f"\n🔍 Fetching report metadata (REST → catalog → optional scanner) for workspace: {workspace_id}")
        by_id = {}

        def _ensure_row(rid, name=''):
            row = by_id.get(rid)
            if not row:
                row = {
                    'report_id': rid,
                    'report_name': name or '',
                    'created_by': '',
                    'created_date_time': '',
                    'modified_by': '',
                    'modified_date_time': '',
                    'created_by_id': '',
                    'modified_by_id': '',
                }
                by_id[rid] = row
            elif name and not row.get('report_name'):
                row['report_name'] = name
            return row

        def _merge_people(row, created_vals, modified_vals, created_dt_vals, modified_dt_vals,
                          created_id_vals=(), modified_id_vals=()):
            # Prefer first usable value already on row, then new sources
            row['created_by'] = _pick_person(row.get('created_by'), *created_vals)
            row['modified_by'] = _pick_person(row.get('modified_by'), *modified_vals)
            row['created_date_time'] = _pick_datetime(row.get('created_date_time'), *created_dt_vals)
            row['modified_date_time'] = _pick_datetime(row.get('modified_date_time'), *modified_dt_vals)
            if not row.get('created_by_id'):
                for v in created_id_vals:
                    if v:
                        row['created_by_id'] = str(v)
                        break
            if not row.get('modified_by_id'):
                for v in modified_id_vals:
                    if v:
                        row['modified_by_id'] = str(v)
                        break

        def _people_coverage():
            if not by_id:
                return 0.0
            hit = sum(
                1 for r in by_id.values()
                if r.get('created_by') or r.get('modified_by') or r.get('modified_date_time')
            )
            return hit / max(1, len(by_id))

        # ---- 1) Groups REST FIRST (fast; usually has createdBy/modifiedBy + dates) ----
        # Old order ran Admin Scanner first → multi-minute silent hang while UI spinners spin.
        rest_count = 0
        rest_filled = 0
        try:
            print(f"   🌐 REST /groups/.../reports starting… t+{time.time()-t0:.1f}s")
            headers = get_user_powerbi_headers()
            url = f"https://api.powerbi.com/v1.0/myorg/groups/{workspace_id}/reports"
            resp = requests.get(url, headers=headers, timeout=45)
            if resp.status_code == 200:
                for report in resp.json().get('value') or []:
                    rid = report.get('id')
                    if not rid:
                        continue
                    name = report.get('name') or ''
                    if _is_excluded_report_name(name):
                        continue
                    row = _ensure_row(rid, name)
                    before_c, before_m = row.get('created_by'), row.get('modified_by')
                    _merge_people(
                        row,
                        created_vals=(
                            report.get('createdBy'),
                            report.get('createdByUser'),
                            report.get('createdByUserPrincipalName'),
                        ),
                        modified_vals=(
                            report.get('modifiedBy'),
                            report.get('modifiedByUser'),
                            report.get('modifiedByUserPrincipalName'),
                        ),
                        created_dt_vals=(
                            report.get('createdDateTime'),
                            report.get('createdDate'),
                        ),
                        modified_dt_vals=(
                            report.get('modifiedDateTime'),
                            report.get('modifiedDate'),
                        ),
                        created_id_vals=(report.get('createdById'),),
                        modified_id_vals=(report.get('modifiedById'),),
                    )
                    rest_count += 1
                    if (row.get('created_by') and not before_c) or (row.get('modified_by') and not before_m):
                        rest_filled += 1
                print(
                    f"   🌐 REST done rows={rest_count} people_filled={rest_filled} "
                    f"coverage={_people_coverage():.0%} t+{time.time()-t0:.1f}s"
                )
            else:
                print(f"   ⚠️ REST reports HTTP {resp.status_code}: {resp.text[:200]}")
        except Exception as e:
            print(f"   ⚠️ REST metadata overlay failed: {e}")

        # ---- 2) Catalog fill (fast when SharePoint/cache warm) ----
        catalog_filled = 0
        if CATALOG_AVAILABLE and catalog_service is not None:
            try:
                print(f"   📦 Catalog gap-fill starting… t+{time.time()-t0:.1f}s")
                pack = catalog_service.get_workspace_reports(workspace_id)
                if pack:
                    for report in pack.get('reports') or []:
                        rid = report.get('id')
                        if not rid:
                            continue
                        row = by_id.get(rid) or _ensure_row(rid, report.get('name') or '')
                        before = (
                            row.get('created_by'), row.get('modified_by'),
                            row.get('created_date_time'), row.get('modified_date_time'),
                        )
                        _merge_people(
                            row,
                            created_vals=(report.get('createdBy'), report.get('created_by')),
                            modified_vals=(report.get('modifiedBy'), report.get('modified_by')),
                            created_dt_vals=(
                                report.get('createdDateTime'),
                                report.get('created_date_time'),
                                report.get('created_date'),
                            ),
                            modified_dt_vals=(
                                report.get('modifiedDateTime'),
                                report.get('modified_date_time'),
                                report.get('modified_date'),
                            ),
                            created_id_vals=(report.get('createdById'), report.get('created_by_id')),
                            modified_id_vals=(report.get('modifiedById'), report.get('modified_by_id')),
                        )
                        after = (
                            row.get('created_by'), row.get('modified_by'),
                            row.get('created_date_time'), row.get('modified_date_time'),
                        )
                        if after != before and any(after):
                            catalog_filled += 1
                print(
                    f"   📦 Catalog fill updates={catalog_filled} "
                    f"coverage={_people_coverage():.0%} t+{time.time()-t0:.1f}s"
                )
            except Exception as e:
                print(f"   ⚠️ Catalog metadata fill failed: {e}")

        # ---- 3) Scanner ONLY if coverage still poor (slow Admin scan — was causing 3–4 min UI hang) ----
        # Default skip when REST already covered most rows. Force with ?force_scanner=1
        force_scanner = str(request.args.get('force_scanner', '')).lower() in ('1', 'true', 'yes')
        coverage = _people_coverage()
        scanner_count = 0
        need_scanner = force_scanner or (not by_id) or (coverage < 0.35 and rest_count == 0)
        if need_scanner:
            try:
                print(
                    f"   🛰️ Scanner fallback (coverage={coverage:.0%} force={force_scanner}) "
                    f"t+{time.time()-t0:.1f}s — may take a while…"
                )
                scanner = PowerBIScanner()
                scan_result = scanner.run_scan(workspace_id=workspace_id) or {}
                for workspace in scan_result.get('workspaces') or []:
                    if workspace.get('id') != workspace_id:
                        continue
                    for report in workspace.get('reports') or []:
                        rid = report.get('id')
                        if not rid:
                            continue
                        row = _ensure_row(rid, report.get('name') or '')
                        _merge_people(
                            row,
                            created_vals=(
                                report.get('createdBy'),
                                report.get('createdByUser'),
                                report.get('createdByUserPrincipalName'),
                            ),
                            modified_vals=(
                                report.get('modifiedBy'),
                                report.get('modifiedByUser'),
                                report.get('modifiedByUserPrincipalName'),
                            ),
                            created_dt_vals=(
                                report.get('createdDateTime'),
                                report.get('createdDate'),
                            ),
                            modified_dt_vals=(
                                report.get('modifiedDateTime'),
                                report.get('modifiedDate'),
                            ),
                            created_id_vals=(report.get('createdById'),),
                            modified_id_vals=(report.get('modifiedById'),),
                        )
                        scanner_count += 1
                    break
                print(f"   🛰️ Scanner rows={scanner_count} t+{time.time()-t0:.1f}s")
            except Exception as e:
                print(f"   ⚠️ Scanner metadata failed: {e}")
        else:
            print(
                f"   ⏭️ Skipping Scanner (coverage={coverage:.0%} rest_rows={rest_count}) "
                f"— UI stays fast; pass force_scanner=1 if needed"
            )

        print(f"   ⏱️ metadata sources done in {time.time()-t0:.1f}s rows={len(by_id)}")

        if not by_id:
            return jsonify({
                'success': False,
                'error': 'No report metadata from Scanner, REST, or catalog'
            }), 500

        # Stable list; keep empty string for missing people (UI → N/A)
        reports_metadata = sorted(
            by_id.values(),
            key=lambda r: (r.get('report_name') or '').lower(),
        )
        with_people = sum(
            1 for r in reports_metadata
            if r.get('created_by') or r.get('modified_by')
        )
        print(
            f"✅ Metadata for {len(reports_metadata)} reports "
            f"({with_people} with created/modified identity)"
        )

        # Cache the results
        reports_cache[cache_key] = (reports_metadata, time.time())

        return jsonify({
            'success': True,
            'workspace_id': workspace_id,
            'reports': reports_metadata,
            'cached': False
        })

    except Exception as e:
        print(f"❌ Error fetching report metadata: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/lineage')
@login_required
def lineage():
    """Report Lineage page - View dataset lineage for reports"""
    return render_template('lineage.html')


@app.route('/api/dataset/lineage')
@login_required
def get_dataset_lineage():
    """
    API endpoint to get semantic model lineage

    Query Parameters:
        workspace_id: Power BI workspace GUID
        dataset_id: Dataset GUID

    Returns:
        JSON with dataset lineage including tables, M expressions, DAX measures
    """
    try:
        from features.semantic_model_lineage import SemanticModelLineage
        from scanner_connector import PowerBIScanner

        workspace_id = request.args.get('workspace_id')
        dataset_id = request.args.get('dataset_id')

        if not workspace_id or not dataset_id:
            return jsonify({
                'success': False,
                'error': 'workspace_id and dataset_id are required'
            }), 400

        # Get user token from session
        user_token = session.get('access_token')
        if not user_token:
            return jsonify({
                'success': False,
                'error': 'Not authenticated'
            }), 401

        # Initialize Scanner API with service principal token
        scanner_service = PowerBIScanner()
        scanner_service.access_token = scanner_service.get_access_token()

        # Create analyzer and get lineage
        analyzer = SemanticModelLineage(scanner_service)
        result = analyzer.get_dataset_lineage(workspace_id, dataset_id)

        return jsonify(result)

    except Exception as e:
        print(f"❌ Error in dataset lineage: {str(e)}")
        import traceback
        traceback.print_exc()

        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/orphaned-reports')
@login_required
def orphaned_reports_page():
    """Unowned Reports page — reports with no active owner signal."""
    return render_template('orphaned_reports.html')


def _decomm_dashboard_config():
    """
    Overview-tab Power BI report settings (Estate Decommissioning Programme).
    All optional — Detail tab never depends on these.
    """
    # Defaults match the published programme report in Fabric Admin Governance.
    # Override via env without code change.
    workspace_id = (
        os.getenv('DECOMM_DASHBOARD_WORKSPACE_ID')
        or '943b84dc-4f4f-8434-74396e772de1'
    ).strip()
    report_id = (
        os.getenv('DECOMM_DASHBOARD_REPORT_ID')
        or '5e4eb1ff-26b8-47d5-84ab-cecadcbb0c3f'
    ).strip()
    dataset_id = (os.getenv('DECOMM_DASHBOARD_DATASET_ID') or '').strip()
    service_url = (os.getenv('DECOMM_DASHBOARD_SERVICE_URL') or '').strip()
    if not service_url and workspace_id and report_id:
        service_url = (
            f'https://app.powerbi.com/groups/{workspace_id}/reports/{report_id}'
            f'/Overview?experience=power-bi'
        )
    embed_enabled = (os.getenv('DECOMM_DASHBOARD_EMBED') or 'true').strip().lower() in (
        '1', 'true', 'yes', 'on',
    )
    return {
        'workspaceId': workspace_id or None,
        'reportId': report_id or None,
        'datasetId': dataset_id or None,
        'serviceUrl': service_url or None,
        'embedEnabled': embed_enabled and bool(workspace_id and report_id),
        'title': os.getenv('DECOMM_DASHBOARD_TITLE')
        or 'Power BI Estate Decommissioning Programme',
    }


@app.route('/decommissioned-reports')
@login_required
def decommissioned_reports_page():
    """Reports archived to SharePoint Report Decommission Activity (file inventory)."""
    return render_template(
        'decommissioned_reports.html',
        decomm_dashboard=_decomm_dashboard_config(),
    )


@app.route('/api/decommissioned-reports/dashboard-config')
@login_required
def api_decommissioned_dashboard_config():
    """Public (auth) config for Overview tab — no secrets."""
    cfg = _decomm_dashboard_config()
    # Include feed path hints (no secrets) for UI help text
    try:
        from features.decommission_inventory import _dataset_feed_folder
        feed_folder = _dataset_feed_folder()
    except Exception:
        feed_folder = None
    return jsonify({
        'success': True,
        **cfg,
        'datasetFeedFolder': feed_folder,
        'datasetFeedFile': 'Decommissioned_Inventory_Latest.xlsx',
        'hasDatasetId': bool(cfg.get('datasetId')),
    })


@app.route('/api/decommissioned-reports/publish-dataset-feed', methods=['POST'])
@login_required
def api_decommissioned_publish_dataset_feed():
    """
    Scan SharePoint archive → overwrite fixed inventory Excel/CSV for Power BI.
    Optional body/query: refresh_dataset=1 to trigger dataset refresh after publish
    (requires DECOMM_DASHBOARD_DATASET_ID + user permission to refresh).
    Does not change Detail inventory API or archive folders.
    """
    try:
        from features.decommission_inventory import (
            publish_decommission_dataset_feed,
            trigger_decommission_dataset_refresh,
        )

        body = request.get_json(silent=True) or {}
        refresh_flag = (
            str(request.args.get('refresh_dataset') or body.get('refresh_dataset') or '')
            .strip()
            .lower()
        )
        do_refresh = refresh_flag in ('1', 'true', 'yes', 'on')

        force = str(
            request.args.get('refresh') or body.get('force_inventory_refresh') or '1'
        ).lower() in ('1', 'true', 'yes', 'on')

        result = publish_decommission_dataset_feed(force_inventory_refresh=force)
        if not result.get('success'):
            return jsonify(result), 502

        out = {'success': True, 'publish': result, 'datasetRefresh': None}
        if do_refresh:
            cfg = _decomm_dashboard_config()
            token = get_user_powerbi_token()
            dataset_id = cfg.get('datasetId') or ''
            # Resolve dataset from report if env not set (same report as Overview embed)
            if not dataset_id and token and cfg.get('workspaceId') and cfg.get('reportId'):
                try:
                    meta_url = (
                        f"https://api.powerbi.com/v1.0/myorg/groups/"
                        f"{cfg['workspaceId']}/reports/{cfg['reportId']}"
                    )
                    mr = requests.get(
                        meta_url,
                        headers={'Authorization': f'Bearer {token}'},
                        timeout=45,
                    )
                    if mr.ok:
                        dataset_id = (mr.json() or {}).get('datasetId') or ''
                except Exception as resolve_err:
                    print(f"   ⚠️ decomm dataset id resolve: {resolve_err}")
            refresh = trigger_decommission_dataset_refresh(
                access_token=token or '',
                workspace_id=cfg.get('workspaceId'),
                dataset_id=dataset_id or None,
            )
            out['datasetRefresh'] = refresh
            # Publish OK even if refresh skipped/failed — feed is the critical path
            if refresh.get('skipped'):
                out['warning'] = refresh.get('error')
            elif not refresh.get('success'):
                out['warning'] = refresh.get('error') or 'Dataset refresh failed'
        return jsonify(out)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/decommissioned-reports/dashboard-embed-token')
@login_required
def api_decommissioned_dashboard_embed_token():
    """
    Embed token for Overview tab using the signed-in user's Power BI token
    (GenerateToken against the programme report). Falls back gracefully so
    Detail tab and open-in-service still work if embed is blocked.
    """
    try:
        cfg = _decomm_dashboard_config()
        if not cfg.get('embedEnabled'):
            return jsonify({
                'success': False,
                'error': 'Dashboard embed is disabled (DECOMM_DASHBOARD_EMBED).',
                'serviceUrl': cfg.get('serviceUrl'),
            }), 400
        workspace_id = cfg.get('workspaceId')
        report_id = cfg.get('reportId')
        if not workspace_id or not report_id:
            return jsonify({
                'success': False,
                'error': 'DECOMM_DASHBOARD_WORKSPACE_ID / REPORT_ID not configured.',
                'serviceUrl': cfg.get('serviceUrl'),
            }), 400

        user_token = get_user_powerbi_token()
        if not user_token:
            return jsonify({
                'success': False,
                'error': 'No Power BI user token in session. Sign out and sign in again.',
                'serviceUrl': cfg.get('serviceUrl'),
            }), 401

        headers = {
            'Authorization': f'Bearer {user_token}',
            'Content-Type': 'application/json',
        }
        # Report meta (for embedUrl + dataset if not set)
        meta_url = (
            f'https://api.powerbi.com/v1.0/myorg/groups/{workspace_id}/reports/{report_id}'
        )
        meta_res = requests.get(meta_url, headers=headers, timeout=45)
        if meta_res.status_code == 401:
            return jsonify({
                'success': False,
                'error': 'Power BI token expired or unauthorized for this report. Re-login.',
                'serviceUrl': cfg.get('serviceUrl'),
            }), 401
        if meta_res.status_code == 404:
            return jsonify({
                'success': False,
                'error': 'Report not found or you do not have access in that workspace.',
                'serviceUrl': cfg.get('serviceUrl'),
            }), 404
        if not meta_res.ok:
            return jsonify({
                'success': False,
                'error': f'Failed to load report metadata (HTTP {meta_res.status_code}).',
                'detail': (meta_res.text or '')[:300],
                'serviceUrl': cfg.get('serviceUrl'),
            }), 502

        meta = meta_res.json() or {}
        embed_url = meta.get('embedUrl') or ''
        dataset_id = cfg.get('datasetId') or meta.get('datasetId') or ''

        token_body = {'accessLevel': 'View'}
        if dataset_id:
            token_body['datasetId'] = dataset_id
            token_body['allowSaveAs'] = False

        token_url = (
            f'https://api.powerbi.com/v1.0/myorg/groups/{workspace_id}'
            f'/reports/{report_id}/GenerateToken'
        )
        tok_res = requests.post(token_url, headers=headers, json=token_body, timeout=45)
        if not tok_res.ok:
            # Common: GenerateToken needs workspace permission / capacity
            err_txt = (tok_res.text or '')[:400]
            return jsonify({
                'success': False,
                'error': (
                    f'GenerateToken failed (HTTP {tok_res.status_code}). '
                    'You can still open the report in Power BI service.'
                ),
                'detail': err_txt,
                'serviceUrl': cfg.get('serviceUrl'),
                'embedUrl': embed_url or None,
                'reportId': report_id,
            }), 502

        tok = tok_res.json() or {}
        access_token = tok.get('token')
        if not access_token:
            return jsonify({
                'success': False,
                'error': 'GenerateToken returned no token.',
                'serviceUrl': cfg.get('serviceUrl'),
            }), 502

        return jsonify({
            'success': True,
            'token': access_token,
            'tokenId': tok.get('tokenId'),
            'expiration': tok.get('expiration'),
            'embedUrl': embed_url,
            'reportId': report_id,
            'workspaceId': workspace_id,
            'datasetId': dataset_id or None,
            'serviceUrl': cfg.get('serviceUrl'),
            'title': cfg.get('title'),
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        cfg = _decomm_dashboard_config()
        return jsonify({
            'success': False,
            'error': str(e),
            'serviceUrl': cfg.get('serviceUrl'),
        }), 500


@app.route('/api/decommissioned-reports')
@login_required
def api_decommissioned_reports():
    """
    List decommissioned reports from SharePoint archive tree.
    Query: workspace (optional name filter), q (search), refresh=1 to bypass cache.
    """
    try:
        from features.decommission_inventory import build_decommission_inventory

        force = request.args.get('refresh', '').lower() in ('1', 'true', 'yes')
        payload = build_decommission_inventory(force_refresh=force)
        if not payload.get('success'):
            return jsonify(payload), 502

        ws_filter = (request.args.get('workspace') or '').strip()
        q = (request.args.get('q') or '').strip().lower()

        rows = list(payload.get('rows') or [])
        if ws_filter:
            wsl = ws_filter.lower()
            rows = [r for r in rows if (r.get('workspaceName') or '').lower() == wsl]
        if q:
            def _match(r):
                blob = " ".join([
                    str(r.get('reportName') or ''),
                    str(r.get('workspaceName') or ''),
                    str(r.get('folderName') or ''),
                    str(r.get('batchFolder') or ''),
                    str(r.get('fileName') or ''),
                ]).lower()
                return q in blob
            rows = [r for r in rows if _match(r)]

        # Rebuild workspace groups for filtered report rows
        by_ws = {}
        for r in rows:
            wn = r.get('workspaceName') or 'Unknown'
            by_ws.setdefault(wn, []).append(r)

        # Keep empty SharePoint workspace folders (0 reports) from inventory.
        # Text search (q) only matches report rows — hide empties while searching.
        # Workspace dropdown filter still shows an empty folder when selected.
        if not q:
            for w in (payload.get('workspaces') or []):
                wn = (w.get('workspaceName') or 'Unknown')
                if ws_filter and wn.lower() != ws_filter.lower():
                    continue
                if wn not in by_ws:
                    by_ws[wn] = list(w.get('reports') or [])

        workspaces = [
            {
                'workspaceName': wn,
                'reportCount': len(rs),
                'reports': rs,
                'isEmpty': len(rs) == 0,
            }
            for wn, rs in sorted(by_ws.items(), key=lambda x: x[0].lower())
        ]

        empty_n = sum(1 for w in workspaces if not w.get('reportCount'))
        return jsonify({
            **payload,
            'rows': rows,
            'workspaces': workspaces,
            'totalReports': len(rows),
            'workspaceCount': len(workspaces),
            'emptyWorkspaceCount': empty_n,
            'filter': {'workspace': ws_filter or None, 'q': q or None},
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e), 'rows': [], 'workspaces': []}), 500


@app.route('/api/decommissioned-reports/export')
@login_required
def export_decommissioned_reports():
    """
    Excel export of SharePoint decommissioned-report inventory.
    Same columns as the Decommissioned Reports UI table.
    Honors optional workspace + q filters (same as list API).
    """
    try:
        import io
        from datetime import datetime, timezone

        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font, PatternFill
        from flask import send_file

        from features.decommission_inventory import build_decommission_inventory

        force = request.args.get('refresh', '').lower() in ('1', 'true', 'yes')
        payload = build_decommission_inventory(force_refresh=force)
        if not payload.get('success'):
            return jsonify({
                'success': False,
                'error': payload.get('error') or 'Failed to load decommission inventory',
            }), 502

        ws_filter = (request.args.get('workspace') or '').strip()
        q = (request.args.get('q') or '').strip().lower()

        rows = list(payload.get('rows') or [])
        if ws_filter:
            wsl = ws_filter.lower()
            rows = [r for r in rows if (r.get('workspaceName') or '').lower() == wsl]
        if q:
            def _match(r):
                blob = " ".join([
                    str(r.get('reportName') or ''),
                    str(r.get('workspaceName') or ''),
                    str(r.get('folderName') or ''),
                    str(r.get('batchFolder') or ''),
                    str(r.get('fileName') or ''),
                ]).lower()
                return q in blob
            rows = [r for r in rows if _match(r)]

        # Workspace A-Z, within each workspace newest decommissioned first
        from collections import defaultdict
        by_ws = defaultdict(list)
        for r in rows:
            by_ws[r.get('workspaceName') or 'Unknown'].append(r)
        ordered = []
        for wn in sorted(by_ws.keys(), key=lambda x: x.lower()):
            grp = by_ws[wn]
            grp.sort(
                key=lambda r: r.get('decommissionedAt') or r.get('lastModifiedAt') or '',
                reverse=True,
            )
            ordered.extend(grp)
        rows = ordered

        wb = Workbook()
        ws_out = wb.active
        ws_out.title = 'Decommissioned Reports'

        header_fill = PatternFill(start_color='2B6CB0', end_color='2B6CB0', fill_type='solid')
        header_font = Font(color='FFFFFF', bold=True, size=11)
        center_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
        left_align = Alignment(horizontal='left', vertical='center', wrap_text=True)

        headers = [
            '#',
            'Report',
            'File Name',
            'Workspace',
            'Folder',
            'Type',
            'Decommissioned',
            'Batch',
            'Size',
            'Size (bytes)',
            'File URL',
            'SharePoint Path',
        ]
        ws_out.append(headers)
        for col_num, _h in enumerate(headers, 1):
            cell = ws_out.cell(row=1, column=col_num)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = center_align

        ws_out.auto_filter.ref = f'A1:L1'
        ws_out.freeze_panes = 'A2'

        for i, r in enumerate(rows, 1):
            ws_out.append([
                i,
                r.get('reportName') or '',
                r.get('fileName') or '',
                r.get('workspaceName') or '',
                r.get('folderName') or '—',
                r.get('fileType') or '',
                r.get('decommissionedAtDisplay') or r.get('decommissionedAt') or '—',
                r.get('batchFolder') or '—',
                r.get('sizeDisplay') or '—',
                r.get('sizeBytes') if r.get('sizeBytes') is not None else '',
                r.get('webUrl') or '',
                r.get('sharePointPath') or '',
            ])
            rn = ws_out.max_row
            for c in range(1, 13):
                cell = ws_out.cell(row=rn, column=c)
                cell.alignment = center_align if c in (1, 6, 9) else left_align
            # Hyperlink file URL when present
            url = r.get('webUrl') or ''
            if url:
                link_cell = ws_out.cell(row=rn, column=11)
                try:
                    link_cell.hyperlink = url
                    link_cell.font = Font(color='0563C1', underline='single')
                except Exception:
                    pass

        if not rows:
            ws_out.append([
                '', 'No decommissioned report files found', '', '', '', '', '', '', '', '', '', ''
            ])

        widths = {
            'A': 6, 'B': 36, 'C': 36, 'D': 28, 'E': 22, 'F': 10,
            'G': 22, 'H': 40, 'I': 12, 'J': 14, 'K': 40, 'L': 50,
        }
        for letter, w in widths.items():
            ws_out.column_dimensions[letter].width = w

        # Meta sheet
        ws_meta = wb.create_sheet('Summary')
        ws_meta.append(['Decommissioned Reports export'])
        ws_meta.append(['Generated (UTC)', datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M')])
        ws_meta.append(['Source', 'SharePoint archive inventory'])
        ws_meta.append(['Base path', payload.get('basePath') or ''])
        ws_meta.append(['Reports in export', len(rows)])
        ws_meta.append(['Batch folders (all)', payload.get('batchCount') or 0])
        ws_meta.append(['Workspace filter', ws_filter or '(all)'])
        ws_meta.append(['Search filter', q or '(none)'])
        ws_meta.append([])
        ws_meta.append(['Note', payload.get('note') or ''])
        ws_meta.column_dimensions['A'].width = 28
        ws_meta.column_dimensions['B'].width = 80
        ws_meta.cell(row=1, column=1).font = Font(bold=True, size=13, color='FFFFFF')
        ws_meta.cell(row=1, column=1).fill = header_fill

        output = io.BytesIO()
        wb.save(output)
        output.seek(0)

        stamp = datetime.now(timezone.utc).strftime('%Y%m%d')
        safe_ws = ''.join(
            c if c.isalnum() or c in ('-', '_') else '_' for c in (ws_filter or 'All')
        )[:40]
        filename = f'Decommissioned_Reports_{safe_ws}_{stamp}.xlsx'
        print(f"📁 decommissioned export rows={len(rows)} file={filename}")

        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=filename,
        )
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/similarity-analysis')
@login_required
def similarity_analysis_page():
    """Similarity Analysis page - Discover similar reports, visuals, measures, and tables"""
    return render_template('similarity_analysis.html')




@app.route('/api/reports/similarity-analysis/<workspace_id>')
@login_required
def analyze_report_similarity(workspace_id):
    """Analyze similarity between reports in a workspace or across all workspaces to identify potential duplicates"""
    try:
        from scanner_connector import PowerBIScanner
        from difflib import SequenceMatcher
        import requests

        # Get threshold from query parameter (default 0.3 for 30%)
        threshold = float(request.args.get('threshold', 0.3))
        threshold_percent = int(threshold * 100)

        # Check if this is a cross-workspace analysis
        is_global = workspace_id.lower() == 'all' or workspace_id.lower() == 'global'

        if is_global:
            print(f"\n🌍 Starting GLOBAL similarity analysis across all workspaces (threshold: {threshold_percent}%)...")
        else:
            print(f"\n🔍 Starting similarity analysis for workspace: {workspace_id} (threshold: {threshold_percent}%)...")

        # Get user token
        user_token = session.get('access_token')
        if not user_token:
            return jsonify({'success': False, 'error': 'Not authenticated'}), 401

        # Initialize Scanner API
        scanner = PowerBIScanner()
        scanner.access_token = scanner.get_access_token()  # Use service principal token for Scanner API

        # Get list of workspaces to scan
        workspaces_to_scan = []
        if is_global:
            print("   📊 Fetching all accessible workspaces...")
            # Get all workspaces user has access to
            headers = {'Authorization': f'Bearer {user_token}', 'Content-Type': 'application/json'}
            ws_response = requests.get('https://api.powerbi.com/v1.0/myorg/groups', headers=headers)
            if ws_response.status_code == 200:
                all_workspaces = ws_response.json().get('value', [])
                workspaces_to_scan = [ws['id'] for ws in all_workspaces]
                print(f"   ✅ Found {len(workspaces_to_scan)} workspaces to analyze")
            else:
                return jsonify({'success': False, 'error': 'Failed to fetch workspaces'}), 500
        else:
            workspaces_to_scan = [workspace_id]

        # Run scan to get all reports with metadata
        print(f"   📊 Running Scanner API scan for {len(workspaces_to_scan)} workspace(s)...")

        # For global scan, aggregate data from multiple workspace scans
        all_scan_data = {'workspaces': []}

        if is_global:
            # Limit to first 10 workspaces for performance (can be adjusted)
            workspaces_to_scan = workspaces_to_scan[:10]
            print(f"   ⚠️  Limiting global scan to first {len(workspaces_to_scan)} workspaces for performance")

            for ws_id in workspaces_to_scan:
                try:
                    ws_scan = scanner.run_scan(workspace_id=ws_id)
                    if ws_scan and 'workspaces' in ws_scan:
                        all_scan_data['workspaces'].extend(ws_scan['workspaces'])
                        print(f"      ✅ Scanned workspace {ws_id}")
                except Exception as e:
                    print(f"      ⚠️  Failed to scan workspace {ws_id}: {e}")
                    continue
            scan_data = all_scan_data
        else:
            scan_data = scanner.run_scan(workspace_id=workspace_id)

        if not scan_data or "workspaces" not in scan_data:
            return jsonify({'success': False, 'error': 'Failed to scan workspace'}), 500

        # Extract reports and their metadata (ONLY from the scanned workspace)
        reports_data = []
        processed_report_ids = set()  # Track to avoid duplicates

        for ws in scan_data["workspaces"]:
            ws_id = ws.get("id")
            ws_name = ws.get('name', 'Unknown')

            # For single workspace analysis, only process the requested workspace
            if not is_global and ws_id != workspace_id:
                print(f"   ⚠️ Skipping workspace {ws_name} (not the requested workspace)")
                continue

            # Exclude App workspaces
            workspace_type = ws.get("type", "").lower()
            if workspace_type == "app" or "app" in ws_name.lower():
                print(f"   🚫 Skipping App workspace: {ws_name}")
                continue

            print(f"   ✓ Processing workspace: {ws_name} ({ws_id})")
            print(f"      Workspace type: {ws.get('type', 'Unknown')}")

            # Build dataset lookup for this workspace
            datasets_lookup = {}
            for dataset in ws.get("datasets", []):
                datasets_lookup[dataset.get("id")] = dataset

            # Process each report in THIS workspace only
            for report in ws.get("reports", []):
                report_id = report.get('id')
                report_name = report.get('name', '')

                # CRITICAL: Exclude App reports (they have [App] prefix or appId field)
                if '[App]' in report_name or report.get('appId'):
                    print(f"   🚫 Skipping App report: {report_name}")
                    continue

                # Skip if already processed (avoid duplicates)
                if report_id in processed_report_ids:
                    print(f"   ⚠️ Skipping duplicate report: {report_name}")
                    continue

                processed_report_ids.add(report_id)
                dataset_id = report.get("datasetId")
                dataset = datasets_lookup.get(dataset_id, {})

                print(f"   ✓ Including report: {report_name}")

                report_meta = {
                    'id': report_id,
                    'name': report_name,
                    'workspace_id': ws_id,
                    'workspace_name': ws_name,
                    'datasetId': dataset_id,
                    'tables': dataset.get('tables', []),
                    'expressions': dataset.get('expressions', []),
                    'measures': [],
                    'pages': [],
                    'visuals': []  # Will be populated with visual metadata
                }

                # Extract DAX measures from tables
                for table in dataset.get('tables', []):
                    for measure in table.get('measures', []):
                        report_meta['measures'].append({
                            'name': measure.get('name'),
                            'expression': measure.get('expression', '')
                        })

                reports_data.append(report_meta)

        print(f"   📊 Total reports found in workspace: {len(reports_data)}")

        # Get page/visual metadata using regular API
        print(f"   📄 Enriching with page/visual data for {len(reports_data)} reports...")
        headers = {'Authorization': f'Bearer {user_token}', 'Content-Type': 'application/json'}

        for report in reports_data:
            try:
                # Get workspace ID for API calls (use the report's workspace_id from metadata)
                report_ws_id = report.get('workspace_id', workspace_id)

                # Get pages
                pages_url = f"https://api.powerbi.com/v1.0/myorg/groups/{report_ws_id}/reports/{report['id']}/pages"
                pages_response = requests.get(pages_url, headers=headers)
                if pages_response.status_code == 200:
                    pages = pages_response.json().get('value', [])
                    report['pages'] = pages

                    # Extract visual metadata from scanner data if available
                    # Scanner API provides detailed visual information in the report sections
                    for page_data in scan_data.get('workspaces', []):
                        for report_data in page_data.get('reports', []):
                            if report_data.get('id') == report['id']:
                                # Extract visuals from pages in scanner data
                                for page in report_data.get('pages', []):
                                    for visual in page.get('visuals', []):
                                        visual_info = {
                                            'visual_type': visual.get('visualType', ''),
                                            'title': visual.get('title', ''),
                                            'fields': []
                                        }
                                        # Extract fields/measures used in the visual
                                        if 'config' in visual:
                                            # Parse visual config to extract field references
                                            config_str = str(visual.get('config', ''))
                                            # Simple extraction - look for field names in config
                                            # This is a simplified approach; full parsing would need JSON parsing
                                            visual_info['fields'] = []  # Placeholder

                                        report['visuals'].append(visual_info)
            except Exception as e:
                print(f"      ⚠️ Could not get pages for {report['name']}: {e}")

        # Compare all reports pairwise (ONLY workspace reports, not Apps)
        print(f"   🔄 Comparing {len(reports_data)} workspace reports...")

        # DEBUG: Show what data we have for each report
        for idx, r in enumerate(reports_data):
            print(f"      📊 Report {idx+1}: {r['name']}")
            print(f"         Tables: {len(r.get('tables', []))}")
            print(f"         Measures: {len(r.get('measures', []))}")
            print(f"         Pages: {len(r.get('pages', []))}")
            if r.get('tables'):
                table_names = [t.get('name') for t in r['tables'][:3]]
                print(f"         Table names: {table_names}")

        comparisons = []

        for i in range(len(reports_data)):
            for j in range(i + 1, len(reports_data)):
                report_a = reports_data[i]
                report_b = reports_data[j]

                # FINAL SAFETY CHECK: Ensure neither report is an App report
                if '[App]' in report_a['name'] or '[App]' in report_b['name']:
                    print(f"   🚫 Skipping comparison with App report: {report_a['name']} vs {report_b['name']}")
                    continue

                # Calculate similarity scores
                similarity = calculate_report_similarity(report_a, report_b)

                # Use dynamic threshold (convert percentage to decimal if needed)
                threshold_value = threshold * 100 if threshold <= 1 else threshold
                if similarity['overall_score'] >= threshold_value:
                    print(f"   ✓ Found similar pair: {report_a['name']} vs {report_b['name']} ({similarity['overall_score']}%)")

                    # Include workspace info for cross-workspace analysis
                    comp_data = {
                        'report_a': {
                            'id': report_a['id'],
                            'name': report_a['name'],
                            'workspace_id': report_a.get('workspace_id'),
                            'workspace_name': report_a.get('workspace_name')
                        },
                        'report_b': {
                            'id': report_b['id'],
                            'name': report_b['name'],
                            'workspace_id': report_b.get('workspace_id'),
                            'workspace_name': report_b.get('workspace_name')
                        },
                        'is_cross_workspace': report_a.get('workspace_id') != report_b.get('workspace_id'),
                        'similarity_score': similarity['overall_score'],
                        'dax_similarity': similarity['scores']['dax_similarity'],
                        'table_similarity': similarity['scores']['table_similarity'],
                        'page_similarity': similarity['scores']['page_similarity'],
                        'visual_similarity': similarity['scores']['visual_similarity'],
                        'identical_measures': similarity['details']['identical_measures'],
                        'logic_matched_measures': similarity['details']['logic_matched_measures'],
                        'similar_measures': similarity['details']['similar_measures'],
                        'unique_measures_a': similarity['details']['unique_to_a'],
                        'unique_measures_b': similarity['details']['unique_to_b'],
                        'identical_tables': similarity['details']['identical_tables'],
                        'unique_tables_a': similarity['details']['unique_tables_a'],
                        'unique_tables_b': similarity['details']['unique_tables_b'],
                        'identical_pages': similarity['details']['identical_pages'],
                        'unique_pages_a': similarity['details']['unique_pages_a'],
                        'unique_pages_b': similarity['details']['unique_pages_b'],
                        'identical_visuals': similarity['details']['identical_visuals'],
                        'similar_visuals': similarity['details']['similar_visuals']
                    }
                    comparisons.append(comp_data)

        # Sort by similarity score (highest first)
        comparisons.sort(key=lambda x: x['similarity_score'], reverse=True)

        print(f"   ✅ Found {len(comparisons)} report pairs with >= 70% similarity")

        return jsonify({
            'success': True,
            'workspace_id': workspace_id,
            'total_reports': len(reports_data),
            'similar_pairs': len(comparisons),
            'comparisons': comparisons
        })

    except Exception as e:
        print(f"❌ Error in similarity analysis: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/similarity-analysis/export', methods=['GET', 'POST'])
@login_required
def export_similarity_analysis():
    """Export similarity analysis results to Excel or CSV with full details"""
    try:
        import io
        from datetime import datetime
        import json

        # Support both GET (old way) and POST (new way with results data)
        if request.method == 'POST':
            workspace_id = request.form.get('workspace_id')
            export_format = request.form.get('format', 'excel')
            results_json = request.form.get('results')

            if not results_json:
                return jsonify({'success': False, 'error': 'results data is required'}), 400

            print(f"\n📊 Exporting pre-analyzed similarity results")
            print(f"   Workspace: {workspace_id}")
            print(f"   Format: {export_format}")

            # Parse the results from frontend
            comparisons = json.loads(results_json)
            print(f"   ✅ Received {len(comparisons)} comparison(s) from frontend")

        else:
            # GET request - old behavior (re-analyze)
            workspace_id = request.args.get('workspace_id')
            export_format = request.args.get('format', 'excel')

            if not workspace_id:
                return jsonify({'success': False, 'error': 'workspace_id is required'}), 400

            print(f"\n📊 Exporting similarity analysis for workspace: {workspace_id}")
            print(f"   Format: {export_format}")
            print(f"   ⚠️ WARNING: Re-analyzing (slow for large datasets)")

            # Get user token
            user_token = session.get('access_token')
            if not user_token:
                return jsonify({'success': False, 'error': 'Not authenticated'}), 401

            # Run the same similarity analysis
            from scanner_connector import PowerBIScanner
            scanner = PowerBIScanner()
            scanner.access_token = scanner.get_access_token()

            # Check if this is a cross-workspace analysis
            is_global = workspace_id.upper() == 'ALL' or workspace_id.upper() == 'GLOBAL'

            # Get list of workspaces to scan
            workspaces_to_scan = []
            if is_global:
                print("   🌍 Cross-workspace export - fetching all accessible workspaces...")
                # Get all workspaces user has access to
                headers = {'Authorization': f'Bearer {user_token}', 'Content-Type': 'application/json'}
                ws_response = requests.get('https://api.powerbi.com/v1.0/myorg/groups', headers=headers)
                if ws_response.status_code == 200:
                    all_workspaces = ws_response.json().get('value', [])
                    workspaces_to_scan = [ws['id'] for ws in all_workspaces[:10]]  # Limit to 10 for performance
                    print(f"   ✅ Will export from {len(workspaces_to_scan)} workspaces")
                else:
                    return jsonify({'success': False, 'error': 'Failed to fetch workspaces'}), 500
            else:
                workspaces_to_scan = [workspace_id]

            # Scan all workspaces
            all_scan_data = {'workspaces': []}
            for ws_id in workspaces_to_scan:
                try:
                    ws_scan = scanner.run_scan(workspace_id=ws_id)
                    if ws_scan and 'workspaces' in ws_scan:
                        all_scan_data['workspaces'].extend(ws_scan['workspaces'])
                        print(f"      ✅ Scanned workspace {ws_id}")
                except Exception as e:
                    print(f"      ⚠️ Failed to scan workspace {ws_id}: {e}")
                    continue

            scan_data = all_scan_data

            if not scan_data or "workspaces" not in scan_data:
                return jsonify({'success': False, 'error': 'Failed to scan workspace'}), 500

            # Extract reports
            reports_data = []
            processed_report_ids = set()

            for ws in scan_data["workspaces"]:
                ws_id = ws.get("id")
                ws_name = ws.get('name', 'Unknown')

                # For single workspace, only process requested workspace
                if not is_global and ws_id != workspace_id:
                    continue

                workspace_type = ws.get("type", "").lower()
                if workspace_type == "app" or "app" in ws_name.lower():
                    continue

                datasets_lookup = {}
                for dataset in ws.get("datasets", []):
                    datasets_lookup[dataset.get("id")] = dataset

                for report in ws.get("reports", []):
                    report_id = report.get('id')
                    report_name = report.get('name', '')

                    if '[App]' in report_name or report.get('appId'):
                        continue

                    if report_id in processed_report_ids:
                        continue

                    processed_report_ids.add(report_id)
                    dataset_id = report.get("datasetId")
                    dataset = datasets_lookup.get(dataset_id, {})

                    report_meta = {
                        'id': report_id,
                        'name': report_name,
                        'workspace_id': ws_id,
                        'workspace_name': ws_name,
                        'datasetId': dataset_id,
                        'tables': dataset.get('tables', []),
                        'measures': [],
                        'pages': []
                    }

                    for table in dataset.get('tables', []):
                        for measure in table.get('measures', []):
                            report_meta['measures'].append({
                                'name': measure.get('name'),
                                'expression': measure.get('expression', '')
                            })

                    reports_data.append(report_meta)

            # Get page/visual metadata
            headers = {'Authorization': f'Bearer {user_token}', 'Content-Type': 'application/json'}
            for report in reports_data:
                try:
                    report_ws_id = report.get('workspace_id', workspace_id)
                    pages_url = f"https://api.powerbi.com/v1.0/myorg/groups/{report_ws_id}/reports/{report['id']}/pages"
                    pages_response = requests.get(pages_url, headers=headers)
                    if pages_response.status_code == 200:
                        report['pages'] = pages_response.json().get('value', [])
                except Exception as e:
                    pass

            # Run comparisons
            print(f"\n📊 Starting similarity comparisons for {len(reports_data)} reports...")
            print(f"   Total comparisons to process: {len(reports_data) * (len(reports_data) - 1) // 2}")

            comparisons = []
            total_comparisons = len(reports_data) * (len(reports_data) - 1) // 2
            comparison_count = 0

            for i in range(len(reports_data)):
                for j in range(i + 1, len(reports_data)):
                    comparison_count += 1

                    # Progress update every 100 comparisons
                    if comparison_count % 100 == 0:
                        print(f"   Progress: {comparison_count}/{total_comparisons} comparisons ({int(comparison_count/total_comparisons*100)}%)")

                report_a = reports_data[i]
                report_b = reports_data[j]

                if '[App]' in report_a['name'] or '[App]' in report_b['name']:
                    continue

                similarity = calculate_report_similarity(report_a, report_b)

                # Use dynamic threshold (convert percentage to decimal if needed)
                threshold_value = threshold * 100 if threshold <= 1 else threshold
                if similarity['overall_score'] >= threshold_value:
                    # Check if cross-workspace
                    is_cross_ws = report_a.get('workspace_id') != report_b.get('workspace_id')

                    comparisons.append({
                        'report_a_name': report_a['name'],
                        'report_a_workspace': report_a.get('workspace_name', ''),
                        'report_b_name': report_b['name'],
                        'report_b_workspace': report_b.get('workspace_name', ''),
                        'is_cross_workspace': 'Yes' if is_cross_ws else 'No',
                        'similarity_score': similarity['overall_score'],
                        'dax_similarity': similarity['scores']['dax_similarity'],
                        'table_similarity': similarity['scores']['table_similarity'],
                        'page_similarity': similarity['scores']['page_similarity'],
                        'visual_similarity': similarity['scores'].get('visual_similarity', 0),
                        'identical_measures_count': len(similarity['details']['identical_measures']),
                        'identical_measures': ', '.join(similarity['details']['identical_measures']),
                        'logic_matched_measures_count': len(similarity['details'].get('logic_matched_measures', [])),
                        'logic_matched_measures': '; '.join([f"{m.get('measure_a')} ⟷ {m.get('measure_b')}" for m in similarity['details'].get('logic_matched_measures', [])]),
                        'identical_tables_count': len(similarity['details']['identical_tables']),
                        'identical_tables': ', '.join(similarity['details']['identical_tables']),
                        'identical_pages_count': len(similarity['details']['identical_pages']),
                        'identical_pages': ', '.join(similarity['details']['identical_pages']),
                        'identical_visuals_count': len(similarity['details'].get('identical_visuals', [])),
                        'similar_visuals_count': len(similarity['details'].get('similar_visuals', [])),
                        'unique_measures_a_count': len(similarity['details']['unique_to_a']),
                        'unique_measures_a': ', '.join(similarity['details']['unique_to_a']),
                        'unique_measures_b_count': len(similarity['details']['unique_to_b']),
                        'unique_measures_b': ', '.join(similarity['details']['unique_to_b']),
                        'unique_tables_a_count': len(similarity['details']['unique_tables_a']),
                        'unique_tables_a': ', '.join(similarity['details']['unique_tables_a']),
                        'unique_tables_b_count': len(similarity['details']['unique_tables_b']),
                        'unique_tables_b': ', '.join(similarity['details']['unique_tables_b']),
                        'unique_pages_a_count': len(similarity['details']['unique_pages_a']),
                        'unique_pages_a': ', '.join(similarity['details']['unique_pages_a']),
                        'unique_pages_b_count': len(similarity['details']['unique_pages_b']),
                        'unique_pages_b': ', '.join(similarity['details']['unique_pages_b'])
                    })

            print(f"\n✅ Comparison complete! Found {len(comparisons)} similar pairs (>= 70% similarity)")
            print(f"   Now generating export file...")

        # Enhanced sorting: First by Report A name (A-Z), then by Similarity Score (Descending)
        # Support both frontend format (report_a.name) and backend format (report_a_name)
        def get_report_a_name(comp):
            if 'report_a_name' in comp:
                return comp['report_a_name'].lower()
            elif 'report_a' in comp and isinstance(comp['report_a'], dict):
                return comp['report_a'].get('name', '').lower()
            return ''

        def get_similarity_score(comp):
            return comp.get('similarity_score', 0)

        comparisons.sort(key=lambda x: (get_report_a_name(x), -get_similarity_score(x)))

        # Normalize comparisons to flat structure for export (handle both frontend and backend formats)
        normalized_comparisons = []
        for comp in comparisons:
            normalized = {}

            # Handle both frontend format (nested) and backend format (flat)
            if 'report_a' in comp and isinstance(comp['report_a'], dict):
                # Frontend format
                normalized['report_a_name'] = comp['report_a'].get('name', '')
                normalized['report_a_workspace'] = comp['report_a'].get('workspace_name', '')
            else:
                # Backend format
                normalized['report_a_name'] = comp.get('report_a_name', '')
                normalized['report_a_workspace'] = comp.get('report_a_workspace', '')

            if 'report_b' in comp and isinstance(comp['report_b'], dict):
                # Frontend format
                normalized['report_b_name'] = comp['report_b'].get('name', '')
                normalized['report_b_workspace'] = comp['report_b'].get('workspace_name', '')
            else:
                # Backend format
                normalized['report_b_name'] = comp.get('report_b_name', '')
                normalized['report_b_workspace'] = comp.get('report_b_workspace', '')

            # Copy all other fields
            normalized['similarity_score'] = comp.get('similarity_score', 0)
            normalized['is_cross_workspace'] = comp.get('is_cross_workspace', 'No')
            normalized['dax_similarity'] = comp.get('scores', {}).get('dax_similarity', 0) if 'scores' in comp else comp.get('dax_similarity', 0)
            normalized['table_similarity'] = comp.get('scores', {}).get('table_similarity', 0) if 'scores' in comp else comp.get('table_similarity', 0)
            normalized['page_similarity'] = comp.get('scores', {}).get('page_similarity', 0) if 'scores' in comp else comp.get('page_similarity', 0)
            normalized['visual_similarity'] = comp.get('scores', {}).get('visual_similarity', 0) if 'scores' in comp else comp.get('visual_similarity', 0)

            # Handle details
            details = comp.get('details', {})
            normalized['identical_measures_count'] = len(details.get('identical_measures', []))
            normalized['identical_measures'] = ', '.join(details.get('identical_measures', []))
            normalized['logic_matched_measures_count'] = len(details.get('logic_matched_measures', []))

            # Logic matched measures formatting
            logic_matches = details.get('logic_matched_measures', [])
            if logic_matches:
                normalized['logic_matched_measures'] = '; '.join([f"{m.get('measure_a')} ⟷ {m.get('measure_b')}" for m in logic_matches])
            else:
                normalized['logic_matched_measures'] = ''

            normalized['identical_tables_count'] = len(details.get('identical_tables', []))
            normalized['identical_tables'] = ', '.join(details.get('identical_tables', []))
            normalized['identical_pages_count'] = len(details.get('identical_pages', []))
            normalized['identical_pages'] = ', '.join(details.get('identical_pages', []))
            normalized['identical_visuals_count'] = len(details.get('identical_visuals', []))
            normalized['similar_visuals_count'] = len(details.get('similar_visuals', []))
            normalized['unique_measures_a_count'] = len(details.get('unique_to_a', []))
            normalized['unique_measures_a'] = ', '.join(details.get('unique_to_a', []))
            normalized['unique_measures_b_count'] = len(details.get('unique_to_b', []))
            normalized['unique_measures_b'] = ', '.join(details.get('unique_to_b', []))
            normalized['unique_tables_a_count'] = len(details.get('unique_tables_a', []))
            normalized['unique_tables_a'] = ', '.join(details.get('unique_tables_a', []))
            normalized['unique_tables_b_count'] = len(details.get('unique_tables_b', []))
            normalized['unique_tables_b'] = ', '.join(details.get('unique_tables_b', []))
            normalized['unique_pages_a_count'] = len(details.get('unique_pages_a', []))
            normalized['unique_pages_a'] = ', '.join(details.get('unique_pages_a', []))
            normalized['unique_pages_b_count'] = len(details.get('unique_pages_b', []))
            normalized['unique_pages_b'] = ', '.join(details.get('unique_pages_b', []))

            normalized_comparisons.append(normalized)

        comparisons = normalized_comparisons

        # Prepare export metadata
        total_reports = len(comparisons) if request.method == 'POST' else len(reports_data)

        if export_format == 'excel':
            # Export to Excel with professional formatting
            try:
                import pandas as pd
                from openpyxl import load_workbook
                from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
                from openpyxl.utils import get_column_letter

                output = io.BytesIO()

                with pd.ExcelWriter(output, engine='openpyxl') as writer:
                    # ===== SHEET 1: Summary =====
                    summary_data = {
                        'Metric': [
                            'Report Title',
                            'Analysis Date',
                            'Analysis Time',
                            'Workspace ID',
                            'Total Reports Analyzed',
                            'Similar Pairs Found (≥70%)',
                            'Highest Similarity',
                            'Lowest Similarity'
                        ],
                        'Value': [
                            'Power BI Similarity Analysis',
                            datetime.now().strftime('%Y-%m-%d'),
                            datetime.now().strftime('%H:%M:%S'),
                            workspace_id,
                            total_reports,
                            len(comparisons),
                            max([c['similarity_score'] for c in comparisons]) if comparisons else 0,
                            min([c['similarity_score'] for c in comparisons]) if comparisons else 0
                        ]
                    }
                    pd.DataFrame(summary_data).to_excel(writer, sheet_name='Summary', index=False)

                    # ===== SHEET 2: Comparison Overview =====
                    overview_data = []
                    for comp in comparisons:
                        overview_data.append({
                            'Report A': comp['report_a_name'],
                            'Workspace A': comp.get('report_a_workspace', ''),
                            'Report B': comp['report_b_name'],
                            'Workspace B': comp.get('report_b_workspace', ''),
                            'Cross-Workspace': comp.get('is_cross_workspace', 'No'),
                            'Overall Similarity %': comp['similarity_score'],
                            'DAX Logic %': comp['dax_similarity'],
                            'Schema %': comp['table_similarity'],
                            'Pages %': comp['page_similarity'],
                            'Visuals %': comp.get('visual_similarity', 0),
                            'Common Measures': comp['identical_measures_count'],
                            'Logic-Matched Measures': comp.get('logic_matched_measures_count', 0),
                            'Common Tables': comp['identical_tables_count'],
                            'Common Pages': comp['identical_pages_count'],
                            'Identical Visuals': comp.get('identical_visuals_count', 0),
                            'Similar Visuals': comp.get('similar_visuals_count', 0)
                        })

                    df_overview = pd.DataFrame(overview_data)
                    df_overview.to_excel(writer, sheet_name='Comparison Overview', index=False)

                    # ===== SHEET 3: Detailed Breakdown =====
                    detailed_data = []
                    for comp in comparisons:
                        detailed_data.append({
                            'Report A': comp['report_a_name'],
                            'Workspace A': comp.get('report_a_workspace', ''),
                            'Report B': comp['report_b_name'],
                            'Workspace B': comp.get('report_b_workspace', ''),
                            'Cross-Workspace': comp.get('is_cross_workspace', 'No'),
                            'Overall Similarity %': comp['similarity_score'],
                            'DAX Logic %': comp['dax_similarity'],
                            'Schema %': comp['table_similarity'],
                            'Pages %': comp['page_similarity'],
                            'Visuals %': comp.get('visual_similarity', 0),
                            'Common Tables': comp['identical_tables'],
                            'Missing in A (Tables)': comp['unique_tables_b'],
                            'Missing in B (Tables)': comp['unique_tables_a'],
                            'Common Measures (Identical Name & DAX)': comp['identical_measures'],
                            'Logic-Matched Measures (Same DAX, Different Name)': comp.get('logic_matched_measures', ''),
                            'Missing in A (Measures)': comp['unique_measures_b'],
                            'Missing in B (Measures)': comp['unique_measures_a'],
                            'Common Pages': comp['identical_pages'],
                            'Missing in A (Pages)': comp['unique_pages_b'],
                            'Missing in B (Pages)': comp['unique_pages_a'],
                            'Identical Visuals': comp.get('identical_visuals_count', 0),
                            'Similar Visuals (70%+ Overlap)': comp.get('similar_visuals_count', 0)
                        })

                    df_detailed = pd.DataFrame(detailed_data)
                    df_detailed.to_excel(writer, sheet_name='Detailed Breakdown', index=False)

                # Now apply formatting to the workbook
                wb = load_workbook(output)

                # Define styles
                header_fill = PatternFill(start_color='667EEA', end_color='667EEA', fill_type='solid')
                header_font = Font(bold=True, color='FFFFFF', size=11)
                header_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                thin_border = Border(
                    left=Side(style='thin', color='E2E8F0'),
                    right=Side(style='thin', color='E2E8F0'),
                    top=Side(style='thin', color='E2E8F0'),
                    bottom=Side(style='thin', color='E2E8F0')
                )

                # Format each sheet
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]

                    # Format header row
                    for cell in ws[1]:
                        cell.fill = header_fill
                        cell.font = header_font
                        cell.alignment = header_alignment
                        cell.border = thin_border

                    # Auto-adjust column widths
                    for column in ws.columns:
                        max_length = 0
                        column_letter = get_column_letter(column[0].column)

                        for cell in column:
                            try:
                                cell_length = len(str(cell.value))
                                if cell_length > max_length:
                                    max_length = cell_length
                            except:
                                pass

                        # Set width with limits
                        adjusted_width = min(max(max_length + 2, 12), 60)
                        ws.column_dimensions[column_letter].width = adjusted_width

                    # Freeze top row
                    ws.freeze_panes = 'A2'

                    # Enable auto-filter on header row
                    if ws.max_row > 1:
                        ws.auto_filter.ref = ws.dimensions

                    # Apply borders to all cells
                    for row in ws.iter_rows(min_row=1, max_row=ws.max_row, min_col=1, max_col=ws.max_column):
                        for cell in row:
                            cell.border = thin_border
                            if cell.row > 1:  # Data rows
                                cell.alignment = Alignment(vertical='top', wrap_text=True)

                # Special formatting for Summary sheet
                if 'Summary' in wb.sheetnames:
                    ws_summary = wb['Summary']
                    ws_summary['A1'].font = Font(bold=True, size=12)
                    ws_summary.column_dimensions['A'].width = 30
                    ws_summary.column_dimensions['B'].width = 40

                # Save formatted workbook
                output = io.BytesIO()
                wb.save(output)
                output.seek(0)

                filename = f"similarity_analysis_{workspace_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"

                print(f"\n📥 Excel export ready!")
                print(f"   Filename: {filename}")
                print(f"   File size: {output.getbuffer().nbytes} bytes")
                print(f"   Sending file to client...")

                return send_file(
                    output,
                    mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                    as_attachment=True,
                    download_name=filename
                )

            except Exception as e:
                print(f"❌ Error creating Excel: {str(e)}")
                import traceback
                traceback.print_exc()
                # Fallback to CSV
                export_format = 'csv'

        if export_format == 'csv':  # CSV format
            import csv

            output = io.StringIO()
            writer = csv.writer(output)

            # Write summary
            writer.writerow(['Power BI Similarity Analysis Report'])
            writer.writerow(['Generated:', datetime.now().strftime('%Y-%m-%d %H:%M:%S')])
            writer.writerow(['Workspace ID:', workspace_id])
            writer.writerow(['Total Reports:', len(reports_data)])
            writer.writerow(['Similar Pairs:', len(comparisons)])
            writer.writerow([])

            # Write headers with enhanced columns
            writer.writerow([
                'Report A', 'Workspace A', 'Report B', 'Workspace B', 'Cross-Workspace',
                'Overall %', 'DAX Logic %', 'Schema %', 'Pages %', 'Visuals %',
                'Common Measures', 'Logic-Matched Measures', 'Common Tables', 'Common Pages',
                'Identical Visuals', 'Similar Visuals',
                'Unique Tables (A)', 'Unique Tables (B)',
                'Unique Measures (A)', 'Unique Measures (B)',
                'Unique Pages (A)', 'Unique Pages (B)'
            ])

            # Write data
            for comp in comparisons:
                writer.writerow([
                    comp['report_a_name'],
                    comp.get('report_a_workspace', ''),
                    comp['report_b_name'],
                    comp.get('report_b_workspace', ''),
                    comp.get('is_cross_workspace', 'No'),
                    comp['similarity_score'],
                    comp['dax_similarity'],
                    comp['table_similarity'],
                    comp['page_similarity'],
                    comp.get('visual_similarity', 0),
                    comp['identical_measures'],
                    comp.get('logic_matched_measures', ''),
                    comp['identical_tables'],
                    comp['identical_pages'],
                    comp.get('identical_visuals_count', 0),
                    comp.get('similar_visuals_count', 0),
                    comp['unique_tables_a'],
                    comp['unique_tables_b'],
                    comp['unique_measures_a'],
                    comp['unique_measures_b'],
                    comp['unique_pages_a'],
                    comp['unique_pages_b']
                ])

            output.seek(0)
            filename = f"similarity_analysis_{workspace_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"

            csv_bytes = output.getvalue().encode('utf-8')
            print(f"\n📄 CSV export ready!")
            print(f"   Filename: {filename}")
            print(f"   File size: {len(csv_bytes)} bytes")
            print(f"   Sending file to client...")

            return send_file(
                io.BytesIO(csv_bytes),
                mimetype='text/csv',
                as_attachment=True,
                download_name=filename
            )

    except Exception as e:
        print(f"❌ Error in export: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500


def calculate_report_similarity(report_a, report_b):
    """
    Calculate similarity between two reports based on DAX, structure, and visuals.
    Enhanced with logic-based DAX comparison and visual-level analysis.
    """
    from difflib import SequenceMatcher
    import re

    scores = {
        'dax_similarity': 0,
        'table_similarity': 0,
        'page_similarity': 0,
        'visual_similarity': 0,
        'overall_score': 0
    }

    details = {
        'identical_measures': [],
        'logic_matched_measures': [],  # NEW: Measures with same logic but different names
        'similar_measures': [],
        'unique_to_a': [],
        'unique_to_b': [],
        'identical_tables': [],
        'unique_tables_a': [],
        'unique_tables_b': [],
        'identical_pages': [],
        'unique_pages_a': [],
        'unique_pages_b': [],
        'identical_visuals': [],  # NEW: Visuals with same type and fields
        'similar_visuals': []  # NEW: Visuals with partial field overlap
    }

    # Helper function to normalize DAX expressions for logic comparison
    def normalize_dax(expression):
        """Normalize DAX expression by removing whitespace and converting to lowercase"""
        if not expression:
            return ""
        # Remove all whitespace
        normalized = re.sub(r'\s+', '', str(expression))
        # Convert to lowercase for case-insensitive comparison
        return normalized.lower()

    # 1. ENHANCED: Compare DAX Measures with Logic-Based Matching
    measures_a = {m['name']: m['expression'] for m in report_a.get('measures', [])}
    measures_b = {m['name']: m['expression'] for m in report_b.get('measures', [])}

    # Create lookup of normalized expressions to find logic matches
    expr_to_names_a = {}  # {normalized_expr: [measure_names]}
    expr_to_names_b = {}

    for name, expr in measures_a.items():
        norm_expr = normalize_dax(expr)
        if norm_expr not in expr_to_names_a:
            expr_to_names_a[norm_expr] = []
        expr_to_names_a[norm_expr].append(name)

    for name, expr in measures_b.items():
        norm_expr = normalize_dax(expr)
        if norm_expr not in expr_to_names_b:
            expr_to_names_b[norm_expr] = []
        expr_to_names_b[norm_expr].append(name)

    all_measure_names = set(measures_a.keys()) | set(measures_b.keys())
    if all_measure_names:
        identical_count = 0
        logic_matched_count = 0
        similar_count = 0
        matched_names = set()  # Track which measures have been matched

        # First pass: Exact name matches
        for name in list(all_measure_names):
            if name in measures_a and name in measures_b:
                expr_a = measures_a[name]
                expr_b = measures_b[name]

                if expr_a == expr_b:
                    identical_count += 1
                    details['identical_measures'].append(name)
                    matched_names.add(name)
                else:
                    # Calculate expression similarity
                    ratio = SequenceMatcher(None, str(expr_a), str(expr_b)).ratio()
                    if ratio >= 0.8:  # 80% similar
                        similar_count += 1
                        details['similar_measures'].append({'name': name, 'similarity': round(ratio * 100)})
                        matched_names.add(name)

        # Second pass: Logic-based matching (different names, same logic)
        for norm_expr_a, names_a in expr_to_names_a.items():
            if norm_expr_a in expr_to_names_b:
                names_b = expr_to_names_b[norm_expr_a]
                # Found measures with identical logic but potentially different names
                for name_a in names_a:
                    for name_b in names_b:
                        if name_a != name_b and name_a not in matched_names and name_b not in matched_names:
                            logic_matched_count += 1
                            details['logic_matched_measures'].append({
                                'measure_a': name_a,
                                'measure_b': name_b,
                                'expression': measures_a[name_a]
                            })
                            matched_names.add(name_a)
                            matched_names.add(name_b)
                            break  # Only match each measure once

        # Identify truly unique measures
        for name in all_measure_names:
            if name not in matched_names:
                if name in measures_a:
                    details['unique_to_a'].append(name)
                else:
                    details['unique_to_b'].append(name)

        # Calculate DAX similarity score (logic matches count as identical)
        total_matches = identical_count + logic_matched_count + similar_count * 0.5
        scores['dax_similarity'] = round((total_matches / len(all_measure_names)) * 100) if all_measure_names else 0

    # 2. Compare Tables/Data Model
    tables_a = set([t.get('name') for t in report_a.get('tables', [])])
    tables_b = set([t.get('name') for t in report_b.get('tables', [])])

    if tables_a or tables_b:
        common_tables = tables_a & tables_b
        details['identical_tables'] = list(common_tables)
        details['unique_tables_a'] = list(tables_a - tables_b)
        details['unique_tables_b'] = list(tables_b - tables_a)

        all_tables = tables_a | tables_b
        scores['table_similarity'] = round((len(common_tables) / len(all_tables)) * 100) if all_tables else 0

    # 3. Compare Page Structure
    pages_a = set([p.get('displayName', p.get('name', '')) for p in report_a.get('pages', [])])
    pages_b = set([p.get('displayName', p.get('name', '')) for p in report_b.get('pages', [])])

    if pages_a or pages_b:
        common_pages = pages_a & pages_b
        details['identical_pages'] = list(common_pages)
        details['unique_pages_a'] = list(pages_a - pages_b)
        details['unique_pages_b'] = list(pages_b - pages_a)

        all_pages = pages_a | pages_b
        scores['page_similarity'] = round((len(common_pages) / len(all_pages)) * 100) if all_pages else 0

    # 4. NEW: Compare Visual-Level Similarity
    visuals_a = report_a.get('visuals', [])
    visuals_b = report_b.get('visuals', [])

    if visuals_a or visuals_b:
        visual_matches = 0
        visual_partial_matches = 0

        for vis_a in visuals_a:
            vis_a_type = vis_a.get('visual_type', '')
            vis_a_title = vis_a.get('title', '')
            vis_a_fields = set(vis_a.get('fields', []))

            for vis_b in visuals_b:
                vis_b_type = vis_b.get('visual_type', '')
                vis_b_title = vis_b.get('title', '')
                vis_b_fields = set(vis_b.get('fields', []))

                # Check if visuals are identical (same type, title, and fields)
                if vis_a_type == vis_b_type and vis_a_title == vis_b_title and vis_a_fields == vis_b_fields:
                    visual_matches += 1
                    details['identical_visuals'].append({
                        'type': vis_a_type,
                        'title': vis_a_title,
                        'fields': list(vis_a_fields)
                    })
                    break
                # Check for partial matches (same type and overlapping fields)
                elif vis_a_type == vis_b_type and vis_a_fields and vis_b_fields:
                    common_fields = vis_a_fields & vis_b_fields
                    if len(common_fields) / max(len(vis_a_fields), len(vis_b_fields)) >= 0.7:  # 70% field overlap
                        visual_partial_matches += 1
                        details['similar_visuals'].append({
                            'type': vis_a_type,
                            'title_a': vis_a_title,
                            'title_b': vis_b_title,
                            'common_fields': list(common_fields),
                            'overlap_ratio': round(len(common_fields) / max(len(vis_a_fields), len(vis_b_fields)), 2)
                        })
                        break

        total_visuals = max(len(visuals_a), len(visuals_b))
        if total_visuals > 0:
            scores['visual_similarity'] = round(((visual_matches + visual_partial_matches * 0.5) / total_visuals) * 100)
        else:
            scores['visual_similarity'] = 0
    else:
        scores['visual_similarity'] = 0

    # Calculate overall similarity (weighted average)
    # DAX = 35%, Tables = 25%, Pages = 20%, Visuals = 20%
    scores['overall_score'] = round(
        scores['dax_similarity'] * 0.35 +
        scores['table_similarity'] * 0.25 +
        scores['page_similarity'] * 0.20 +
        scores['visual_similarity'] * 0.20
    )

    return {
        'scores': scores,
        'details': details,
        'overall_score': scores['overall_score']
    }


@app.route('/api/workspace/table-search/<workspace_id>')
@login_required
def search_workspace_tables(workspace_id):
    """Search for tables across workspace(s) - supports single workspace or 'ALL' for cross-workspace search"""
    try:
        from scanner_connector import PowerBIScanner
        import requests

        search_term = request.args.get('query', '').strip().lower()
        folder_id = request.args.get('folder_id', '').strip()  # NEW: Get folder filter (ignored for ALL)

        if not search_term:
            return jsonify({'success': False, 'error': 'Search query is required'}), 400

        # Check if cross-workspace search
        is_global = workspace_id.upper() == 'ALL' or workspace_id.upper() == 'GLOBAL'

        print(f"\n🔍 ===============================================")
        if is_global:
            print(f"🔍 TABLE SEARCH (CROSS-WORKSPACE): '{search_term}'")
        else:
            print(f"🔍 TABLE SEARCH: '{search_term}' in workspace {workspace_id}")
            if folder_id:
                print(f"🔍 FOLDER FILTER: {folder_id}")
        print(f"🔍 ===============================================")

        # Get user token for workspace enumeration and report lookups
        user_token = session.get('access_token')
        if not user_token:
            return jsonify({'success': False, 'error': 'Not authenticated'}), 401

        # Initialize Scanner API
        scanner = PowerBIScanner()
        scanner.access_token = scanner.get_access_token()

        # Get list of workspaces to scan
        workspaces_to_scan = []
        workspace_names = {}  # Map workspace_id -> workspace_name

        if is_global:
            print("   🌍 Cross-workspace search - fetching all accessible workspaces...")
            headers = {'Authorization': f'Bearer {user_token}', 'Content-Type': 'application/json'}
            ws_response = requests.get('https://api.powerbi.com/v1.0/myorg/groups', headers=headers)
            if ws_response.status_code == 200:
                all_workspaces = ws_response.json().get('value', [])
                workspaces_to_scan = [ws['id'] for ws in all_workspaces]
                workspace_names = {ws['id']: ws['name'] for ws in all_workspaces}
                print(f"   ✅ Will search across {len(workspaces_to_scan)} workspaces")
            else:
                return jsonify({'success': False, 'error': 'Failed to fetch workspaces'}), 500
        else:
            workspaces_to_scan = [workspace_id]

        # Scan all workspaces
        all_scan_data = {'workspaces': []}
        for ws_id in workspaces_to_scan:
            try:
                print(f"   📊 Scanning workspace {workspace_names.get(ws_id, ws_id)[:30]}...")
                ws_scan = scanner.run_scan(workspace_id=ws_id)
                if ws_scan and 'workspaces' in ws_scan:
                    all_scan_data['workspaces'].extend(ws_scan['workspaces'])
            except Exception as e:
                print(f"      ⚠️ Failed to scan workspace {ws_id}: {e}")
                continue

        scan_data = all_scan_data

        if not scan_data or "workspaces" not in scan_data:
            return jsonify({'success': False, 'error': 'Failed to scan workspaces'}), 500

        # NEW: Get folder hierarchy if folder filter is applied
        allowed_folder_ids = set()
        if folder_id:
            print(f"   🗂️  Folder filter requested: {folder_id}")

            if folder_id != '__ROOT__':
                # Fetch folder hierarchy from Fabric API
                try:
                    fabric_token = get_fabric_token()
                    fabric_url = f'https://api.fabric.microsoft.com/v1/workspaces/{workspace_id}/items'
                    fabric_headers = {
                        'Authorization': f'Bearer {fabric_token}',
                        'Content-Type': 'application/json'
                    }

                    print(f"   📁 Fetching folders from Fabric API...")
                    fabric_response = requests.get(fabric_url, headers=fabric_headers)

                    if fabric_response.status_code == 200:
                        items = fabric_response.json().get('value', [])
                        folder_map = {}

                        # Build folder hierarchy map
                        for item in items:
                            # Only process items that are folders (have displayName but are Reports)
                            if item.get('type') == 'Report':
                                current_folder_id = item.get('id')
                                current_parent_folder_id = item.get('parentFolderId')

                                if current_parent_folder_id:  # This report is in a folder
                                    # Track the folder if we haven't seen it yet
                                    if current_parent_folder_id not in folder_map:
                                        folder_map[current_parent_folder_id] = {
                                            'id': current_parent_folder_id,
                                            'parentFolderId': None  # Will be filled in later
                                        }

                        # Now get folder parent relationships by checking folder items
                        # (Folders themselves don't appear as items, only reports with parentFolderId)
                        # So we need to get unique folder structure from report placements

                        # Get all child folders recursively
                        allowed_folder_ids = get_all_child_folders_recursive(folder_id, folder_map)
                        allowed_folder_ids.add(folder_id)  # Include the selected folder itself

                        print(f"   📂 Folder filter active: {len(allowed_folder_ids)} folder(s) in tree (including children)")
                    else:
                        # Fallback: just use the exact folder ID without hierarchy
                        allowed_folder_ids = {folder_id}
                        print(f"   ⚠️ Could not fetch folder hierarchy, filtering by exact folder ID only")
                except Exception as e:
                    # Fallback: just use the exact folder ID without hierarchy
                    allowed_folder_ids = {folder_id}
                    print(f"   ⚠️ Error fetching folders: {e}, filtering by exact folder ID only")

        # Search through all tables
        matching_tables = []
        total_tables_scanned = 0

        for ws in scan_data["workspaces"]:
            # Extract workspace ID and name from each workspace
            ws_id = ws.get("id")
            ws_name = ws.get("name", "Unknown Workspace")

            for dataset in ws.get("datasets", []):
                dataset_id = dataset.get("id")
                dataset_name = dataset.get("name", "Unknown Dataset")

                # Get data sources for this dataset
                datasources = dataset.get("datasources", [])

                print(f"\n   📁 Dataset: {dataset_name} ({dataset_id})")

                for table in dataset.get("tables", []):
                    total_tables_scanned += 1
                    table_display_name = table.get("name", "")

                    print(f"      🔎 Scanning table: '{table_display_name}'")

                    # Extract native source table name from source expression or partition
                    native_source_name = None
                    source_expression = None

                    import re

                    # Try to extract from table source property (M query)
                    source = table.get("source", [])
                    if source and len(source) > 0:
                        # Source is typically an array of partition expressions
                        for src in source:
                            expr = src.get("expression", "")
                            if expr:
                                source_expression = expr
                                # DEBUG: Print first 300 characters of expression
                                print(f"         📝 M Query Expression (first 300 chars): {expr[:300]}...")

                                # ENHANCED REGEX PATTERNS for native table name extraction
                                # IMPORTANT: Check SQL patterns FIRST before generic patterns

                                # Pattern 1: SQL Database with curly braces - {[Schema="xxx",Item="yyy"]}
                                sql_curly_match = re.search(r'\{?\[Schema="([^"]+)",\s*Item="([^"]+)"\]\}?', expr)
                                if sql_curly_match:
                                    schema = sql_curly_match.group(1)
                                    item = sql_curly_match.group(2)
                                    # Check if schema contains database name with brackets
                                    if '[' in schema and ']' in schema:
                                        # Extract from format like [PowerBI_QTIL]
                                        schema = re.sub(r'[\[\]]', '', schema)
                                    native_source_name = f"[{schema}].[{item}]"
                                    print(f"      ✓ Extracted (Pattern 1 - SQL): {native_source_name} from display name '{table_display_name}'")
                                    break

                                # NEW Pattern: Extract from SQL Query parameter - FROM [schema].[table]
                                # Matches: Sql.Database(..., [Query="SELECT ... FROM [schema].[table] ..."])
                                sql_query_match = re.search(r'\[Query="[^"]*FROM\s+\[?([^\s\.\]]+)\]?\.\[?([^\s\.\]]+)\]?', expr, re.IGNORECASE)
                                if sql_query_match:
                                    schema = sql_query_match.group(1)
                                    item = sql_query_match.group(2)
                                    native_source_name = f"[{schema}].[{item}]"
                                    print(f"      ✓ Extracted (Pattern 1b - SQL Query FROM): {native_source_name} from display name '{table_display_name}'")
                                    break

                                # Pattern 2: SharePoint/Excel - [Name="xxx"] (Check BEFORE generic patterns)
                                name_match = re.search(r'\[Name="([^"]+)"\]', expr)
                                if name_match:
                                    native_source_name = name_match.group(1)
                                    print(f"      ✓ Extracted (Pattern 2 - SharePoint/Excel): {native_source_name} from display name '{table_display_name}'")
                                    break

                                # Pattern 3: Direct reference with brackets - #"[Schema].[TableName]"
                                bracket_match = re.search(r'#"\[([^\]]+)\]\.\[([^\]]+)\]"', expr)
                                if bracket_match:
                                    schema = bracket_match.group(1)
                                    item = bracket_match.group(2)
                                    native_source_name = f"[{schema}].[{item}]"
                                    print(f"      ✓ Extracted (Pattern 3 - Bracketed): {native_source_name} from display name '{table_display_name}'")
                                    break

                                # Pattern 4: Simple schema.table format - #"dbo.TableName"
                                # But NOT transformation steps like "Renamed Columns"
                                simple_match = re.search(r'#"([^"\s]+)\.([^"\s]+)"', expr)
                                if simple_match:
                                    schema = simple_match.group(1)
                                    item = simple_match.group(2)
                                    # Skip if it looks like a transformation step (contains spaces or common step words)
                                    if ' ' not in schema and ' ' not in item:
                                        native_source_name = f"{schema}.{item}"
                                        print(f"      ✓ Extracted (Pattern 4 - Schema.Table): {native_source_name} from display name '{table_display_name}'")
                                        break

                                # Pattern 5: SKIP - Don't extract generic transformation step names
                                # Old Pattern 4 was too broad and caught "Renamed Columns", "Changed Type", etc.
                                # We only want actual source table names, not transformation steps

                    # Also check partitions array (alternative location)
                    if not native_source_name:
                        partitions = table.get("partitions", [])
                        for partition in partitions:
                            source_expr = partition.get("source", {})
                            if isinstance(source_expr, dict):
                                expr = source_expr.get("expression", "")
                                if expr:
                                    source_expression = expr
                                    # DEBUG: Print first 300 characters of partition expression
                                    print(f"         📝 Partition M Query (first 300 chars): {expr[:300]}...")

                                    # Apply same enhanced patterns (SQL patterns FIRST)
                                    sql_curly_match = re.search(r'\{?\[Schema="([^"]+)",\s*Item="([^"]+)"\]\}?', expr)
                                    if sql_curly_match:
                                        schema = sql_curly_match.group(1)
                                        item = sql_curly_match.group(2)
                                        if '[' in schema and ']' in schema:
                                            schema = re.sub(r'[\[\]]', '', schema)
                                        native_source_name = f"[{schema}].[{item}]"
                                        print(f"      ✓ Extracted (Partition Pattern 1 - SQL): {native_source_name} from display name '{table_display_name}'")
                                        break

                                    # NEW: SQL Query FROM clause extraction
                                    sql_query_match = re.search(r'\[Query="[^"]*FROM\s+\[?([^\s\.\]]+)\]?\.\[?([^\s\.\]]+)\]?', expr, re.IGNORECASE)
                                    if sql_query_match:
                                        schema = sql_query_match.group(1)
                                        item = sql_query_match.group(2)
                                        native_source_name = f"[{schema}].[{item}]"
                                        print(f"      ✓ Extracted (Partition Pattern 1b - SQL Query FROM): {native_source_name} from display name '{table_display_name}'")
                                        break

                                    name_match = re.search(r'\[Name="([^"]+)"\]', expr)
                                    if name_match:
                                        native_source_name = name_match.group(1)
                                        print(f"      ✓ Extracted (Partition Pattern 2 - SharePoint/Excel): {native_source_name} from display name '{table_display_name}'")
                                        break

                                    bracket_match = re.search(r'#"\[([^\]]+)\]\.\[([^\]]+)\]"', expr)
                                    if bracket_match:
                                        schema = bracket_match.group(1)
                                        item = bracket_match.group(2)
                                        native_source_name = f"[{schema}].[{item}]"
                                        print(f"      ✓ Extracted (Partition Pattern 3 - Bracketed): {native_source_name} from display name '{table_display_name}'")
                                        break

                                    simple_match = re.search(r'#"([^"\s]+)\.([^"\s]+)"', expr)
                                    if simple_match:
                                        schema = simple_match.group(1)
                                        item = simple_match.group(2)
                                        if ' ' not in schema and ' ' not in item:
                                            native_source_name = f"{schema}.{item}"
                                            print(f"      ✓ Extracted (Partition Pattern 4 - Schema.Table): {native_source_name} from display name '{table_display_name}'")
                                            break

                    # Fallback: use display name if no native source found
                    if not native_source_name:
                        native_source_name = table_display_name
                        print(f"      ⚠️ No native source found, using display name: {table_display_name}")

                    # ENHANCED FUZZY MATCHING: Search in table names AND column names
                    # Normalize by removing brackets, underscores, and special characters for better matching
                    def normalize_for_search(text):
                        """Remove brackets, dots, underscores for fuzzy matching"""
                        # Remove brackets []
                        text = re.sub(r'[\[\]]', '', text)
                        # Replace dots and underscores with spaces for word matching
                        text = re.sub(r'[._]', ' ', text)
                        return text.lower()

                    native_normalized = normalize_for_search(native_source_name)
                    display_normalized = normalize_for_search(table_display_name)

                    # Search in both exact and normalized forms (TABLE NAME)
                    search_in_native = (search_term in native_source_name.lower() or
                                       search_term in native_normalized)
                    search_in_display = (search_term in table_display_name.lower() or
                                        search_term in display_normalized)

                    # NEW: Search in COLUMN NAMES
                    matched_columns = []
                    columns = table.get("columns", [])

                    for col in columns:
                        col_name = col.get('name', '')
                        col_source = col.get('sourceColumn', col_name)

                        # Normalize column names for search
                        col_name_normalized = normalize_for_search(col_name)
                        col_source_normalized = normalize_for_search(col_source)

                        # Check if search term matches column name or source column
                        if (search_term in col_name.lower() or
                            search_term in col_name_normalized or
                            search_term in col_source.lower() or
                            search_term in col_source_normalized):
                            matched_columns.append({
                                'name': col_name,
                                'source_name': col_source,
                                'dataType': col.get('dataType', 'Unknown')
                            })

                    # Table is a match if:
                    # 1. Table name matches (native or display)
                    # 2. OR any column name matches
                    is_match = search_in_native or search_in_display or len(matched_columns) > 0

                    # Debug logging
                    if is_match:
                        print(f"         ✅ MATCH FOUND!")
                        print(f"            Native: {native_source_name} (normalized: {native_normalized})")
                        print(f"            Display: {table_display_name} (normalized: {display_normalized})")
                        print(f"            Match in native: {search_in_native}, Match in display: {search_in_display}")
                        if len(matched_columns) > 0:
                            print(f"            Match in columns: {len(matched_columns)} column(s) - {[c['name'] for c in matched_columns]}")

                    if is_match:
                        # Build complete column info (already retrieved columns above for matching)
                        column_info = []

                        for col in columns:
                            # Also extract native column source if available
                            col_display_name = col.get('name', '')
                            col_source_expr = col.get('sourceColumn', col_display_name)

                            column_info.append({
                                'name': col_display_name,
                                'source_name': col_source_expr,
                                'dataType': col.get('dataType', 'Unknown'),
                                'isHidden': col.get('isHidden', False)
                            })

                        # Get measures in this table
                        measures = table.get("measures", [])
                        measure_info = [m.get('name', '') for m in measures]

                        # Find which reports use this dataset
                        report_names = []
                        for report in ws.get("reports", []):
                            if report.get("datasetId") == dataset_id:
                                # NEW: Filter by folder if folder_id is specified
                                if folder_id:
                                    report_folder_id = report.get('folderObjectId') or report.get('folderId')

                                    # Handle ROOT folder filter
                                    if folder_id == '__ROOT__':
                                        if report_folder_id:  # Skip reports WITH folders
                                            continue
                                    else:
                                        # Only include reports in allowed folders
                                        if not report_folder_id or report_folder_id not in allowed_folder_ids:
                                            continue

                                # Skip [App] shells + platform usage metrics reports
                                if _is_excluded_report_name(report.get('name')):
                                    continue

                                report_names.append({
                                    'id': report.get('id'),
                                    'name': report.get('name', 'Unknown Report')
                                })

                        # Get datasource information
                        datasource_info = []
                        for ds in datasources:
                            connection_details = ds.get("connectionDetails", {})
                            datasource_info.append({
                                'datasourceType': ds.get("datasourceType", "Unknown"),
                                'server': connection_details.get("server", "N/A"),
                                'database': connection_details.get("database", "N/A"),
                                'url': connection_details.get("url", "N/A")
                            })

                        # FALLBACK: Extract datasource from M query expression if not in datasources array
                        if not datasource_info or (len(datasource_info) == 1 and datasource_info[0]['server'] == 'N/A'):
                            if source_expression:
                                # Extract from Sql.Database("server", "database", ...)
                                sql_db_match = re.search(r'Sql\.Database\("([^"]+)",\s*"([^"]+)"', source_expression, re.IGNORECASE)
                                if sql_db_match:
                                    server = sql_db_match.group(1)
                                    database = sql_db_match.group(2)
                                    datasource_info = [{
                                        'datasourceType': 'Sql',
                                        'server': server,
                                        'database': database,
                                        'url': 'N/A'
                                    }]
                                    print(f"         📡 Extracted datasource from M query: {server} / {database}")

                        matching_tables.append({
                            'workspace_id': ws_id,
                            'workspace_name': workspace_names.get(ws_id, ws_name),
                            'table_name': table_display_name,
                            'native_source_name': native_source_name,
                            'source_expression': source_expression,
                            'has_different_names': native_source_name != table_display_name,
                            'dataset_id': dataset_id,
                            'dataset_name': dataset_name,
                            'reports': report_names,
                            'column_count': len(columns),
                            'columns': column_info,
                            'measure_count': len(measures),
                            'measures': measure_info,
                            'datasources': datasource_info,
                            'is_hidden': table.get('isHidden', False),
                            'matched_columns': matched_columns,  # NEW: Include matched columns
                            'match_type': 'column' if len(matched_columns) > 0 and not (search_in_native or search_in_display) else 'table'
                        })

        # Sort by native source name (prioritize native name)
        matching_tables.sort(key=lambda x: x['native_source_name'])

        print(f"\n🔍 ===============================================")
        print(f"🔍 SEARCH SUMMARY")
        print(f"🔍 ===============================================")
        print(f"   📊 Total tables scanned: {total_tables_scanned}")
        print(f"   ✅ Matching tables found: {len(matching_tables)}")
        if len(matching_tables) > 0:
            print(f"   📋 Results:")
            for tbl in matching_tables:
                print(f"      - {tbl['native_source_name']} (display: {tbl['table_name']})")
        print(f"🔍 ===============================================\n")

        return jsonify({
            'success': True,
            'workspace_id': workspace_id,
            'search_term': search_term,
            'result_count': len(matching_tables),
            'workspaces_scanned': len(workspaces_to_scan) if is_global else None,
            'tables': matching_tables
        })

    except Exception as e:
        print(f"❌ Error in table search: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/workspace/table-columns/<workspace_id>/<dataset_name>/<table_name>')
@login_required
def get_table_columns(workspace_id, dataset_name, table_name):
    """Get column details for a specific table in a dataset"""
    try:
        from scanner_connector import PowerBIScanner

        print(f"\n📋 FETCHING COLUMNS FOR TABLE: {table_name}")
        print(f"   Workspace: {workspace_id}")
        print(f"   Dataset: {dataset_name}")

        # Initialize Scanner API
        scanner = PowerBIScanner()
        scanner.access_token = scanner.get_access_token()

        # Scan the workspace to get dataset metadata
        scan_result = scanner.run_scan(workspace_id=workspace_id)

        if not scan_result or 'workspaces' not in scan_result:
            return jsonify({'success': False, 'error': 'Failed to scan workspace'}), 500

        # Find the target dataset in the first workspace
        target_dataset = None
        if scan_result['workspaces']:
            workspace_data = scan_result['workspaces'][0]
            for dataset in workspace_data.get('datasets', []):
                if dataset.get('name') == dataset_name:
                    target_dataset = dataset
                    break

        if not target_dataset:
            return jsonify({'success': False, 'error': f'Dataset "{dataset_name}" not found'}), 404

        # Find the target table
        target_table = None
        if 'tables' in target_dataset:
            for table in target_dataset['tables']:
                if table.get('name') == table_name:
                    target_table = table
                    break

        if not target_table:
            return jsonify({'success': False, 'error': f'Table "{table_name}" not found in dataset'}), 404

        # Extract column information
        columns = []
        if 'columns' in target_table:
            for col in target_table['columns']:
                columns.append({
                    'name': col.get('name', 'Unknown'),
                    'dataType': col.get('dataType', 'Unknown'),
                    'expression': col.get('expression', None),
                    'isHidden': col.get('isHidden', False),
                    'sortByColumn': col.get('sortByColumn', None)
                })

        print(f"   ✅ Found {len(columns)} columns")

        return jsonify({
            'success': True,
            'table_name': table_name,
            'dataset_name': dataset_name,
            'column_count': len(columns),
            'columns': columns
        })

    except Exception as e:
        print(f"❌ Error fetching table columns: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/workspace/table-search/export', methods=['POST'])
@login_required
def export_table_search():
    """Export table dictionary search results to Excel"""
    try:
        import pandas as pd
        import io
        from datetime import datetime

        data = request.get_json()
        workspace_id = data.get('workspace_id')
        search_term = data.get('search_term')
        tables = data.get('results', [])

        if not tables:
            return jsonify({'success': False, 'error': 'No results to export'}), 400

        print(f"\n📥 Exporting table search results: {len(tables)} table(s)")

        # Create Excel file
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            # Prepare data for export
            export_data = []
            for table in tables:
                base_row = {
                    'Workspace': table.get('workspace_name', 'N/A'),
                    'Table Name': table.get('table_name'),
                    'Native Source Name': table.get('native_source_name'),
                    'Dataset': table.get('dataset_name'),
                    'Column Count': table.get('column_count', 0),
                    'Measure Count': table.get('measure_count', 0),
                    'Reports Using This Table': '; '.join([r['name'] for r in table.get('reports', [])]),
                    'Columns': '; '.join([c['name'] for c in table.get('columns', [])]),
                    'Measures': '; '.join(table.get('measures', [])),
                    'Data Sources': '; '.join([f"{ds.get('datasourceType')} - {ds.get('server')}/{ds.get('database')}"
                                               for ds in table.get('datasources', [])])
                }
                export_data.append(base_row)

            # Create DataFrame and write to Excel
            df = pd.DataFrame(export_data)
            df.to_excel(writer, sheet_name='Table Dictionary', index=False)

            # Auto-adjust column widths
            worksheet = writer.sheets['Table Dictionary']
            for idx, col in enumerate(df.columns):
                max_length = max(df[col].astype(str).apply(len).max(), len(col)) + 2
                worksheet.column_dimensions[chr(65 + idx)].width = min(max_length, 50)

        output.seek(0)

        # Generate filename
        filename = f"table_dictionary_{search_term}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"

        print(f"   ✅ Export complete: {filename}")

        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        print(f"❌ Error exporting table search: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/workspace/deep-content-search/export', methods=['POST'])
@login_required
def export_deep_search():
    """Export deep content search results to Excel"""
    try:
        import pandas as pd
        import io
        from datetime import datetime

        data = request.get_json()
        workspace_id = data.get('workspace_id')
        search_term = data.get('search_term')
        matches = data.get('results', [])

        if not matches:
            return jsonify({'success': False, 'error': 'No results to export'}), 400

        print(f"\n📥 Exporting deep search results: {len(matches)} match(es)")

        # Create Excel file
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            # Prepare data for export
            export_data = []
            for match in matches:
                row = {
                    'Report Name': match.get('report_name'),
                    'Dataset': match.get('dataset_name'),
                    'Page': match.get('page_name'),
                    'Visual Title': match.get('visual_title'),
                    'Visual Type': match.get('visual_type'),
                    'Table Name': match.get('table_name', ''),
                    'Table Native Name': match.get('table_native_name', ''),
                    'Column Name': match.get('column_name', ''),
                    'Column Source Name': match.get('column_source_name', ''),
                    'Column Data Type': match.get('column_dataType', ''),
                    'Is Alias': 'Yes' if match.get('is_alias') else 'No',
                    'Match Type': match.get('match_type', 'Field'),
                    'Matched Fields': ', '.join(match.get('matched_fields', [])) if match.get('matched_fields') else '',
                    'Data Source': match.get('datasource', 'N/A'),
                    'Workspace': match.get('workspace_name', '') if 'workspace_name' in match else ''
                }
                export_data.append(row)

            # Create DataFrame and write to Excel
            df = pd.DataFrame(export_data)
            df.to_excel(writer, sheet_name='Deep Search Results', index=False)

            # Auto-adjust column widths
            worksheet = writer.sheets['Deep Search Results']
            for idx, col in enumerate(df.columns):
                # Handle NaN values by converting to string first
                col_values = df[col].fillna('').astype(str)
                max_length = max(col_values.apply(len).max(), len(col)) + 2
                worksheet.column_dimensions[chr(65 + idx)].width = min(max_length, 50)

        output.seek(0)

        # Generate filename
        filename = f"deep_search_{search_term}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"

        print(f"   ✅ Export complete: {filename}")

        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        print(f"❌ Error exporting deep search: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/workspace/content-search/<workspace_id>')
@login_required
def search_workspace_content(workspace_id):
    """Unified content search: searches across visual titles, measure names, table names, and column names"""
    try:
        from scanner_connector import PowerBIScanner
        import re

        search_term = request.args.get('query', '').strip()

        if not search_term:
            return jsonify({'success': False, 'error': 'Search query is required'}), 400

        # Use case-insensitive matching (SQL LIKE equivalent)
        search_term_lower = search_term.lower()

        print(f"\n📄 ===============================================")
        print(f"📄 CONTENT SEARCH: '{search_term}' in workspace {workspace_id}")
        print(f"📄 ===============================================")

        # Initialize Scanner API
        scanner = PowerBIScanner()
        scanner.access_token = scanner.get_access_token()

        # Run scan to get comprehensive metadata
        print("   📊 Running Scanner API scan...")
        scan_data = scanner.run_scan(workspace_id=workspace_id)

        if not scan_data or "workspaces" not in scan_data:
            return jsonify({'success': False, 'error': 'Failed to scan workspace'}), 500

        # Collect all matches
        all_matches = []
        reports_with_matches = set()

        for ws in scan_data["workspaces"]:
            # SECTION 1: Search in VISUAL TITLES
            print("\n   🎨 Searching visual titles...")
            for report in ws.get("reports", []):
                report_id = report.get("id")
                report_name = report.get("name", "Unknown Report")

                for page in report.get("pages", []):
                    page_name = page.get("displayName", "Unknown Page")

                    for visual in page.get("visuals", []):
                        visual_title = visual.get("title", "")

                        # Search in visual title using LIKE matching
                        if visual_title and search_term_lower in visual_title.lower():
                            reports_with_matches.add(report_id)
                            all_matches.append({
                                'match_type': 'visual_title',
                                'matched_text': visual_title,
                                'report_id': report_id,
                                'report_name': report_name,
                                'page_name': page_name,
                                'dataset_name': None,
                                'context': f"Visual on page '{page_name}'"
                            })
                            print(f"      ✅ VISUAL MATCH: '{visual_title}' in report '{report_name}' / page '{page_name}'")

            # SECTION 2: Search in MEASURE NAMES
            print("\n   📐 Searching measure names...")
            for dataset in ws.get("datasets", []):
                dataset_id = dataset.get("id")
                dataset_name = dataset.get("name", "Unknown Dataset")

                # Find reports using this dataset
                using_reports = []
                for report in ws.get("reports", []):
                    if report.get("datasetId") == dataset_id:
                        using_reports.append({
                            'id': report.get('id'),
                            'name': report.get('name', 'Unknown Report')
                        })

                for table in dataset.get("tables", []):
                    for measure in table.get("measures", []):
                        measure_name = measure.get("name", "")

                        # Search in measure name using LIKE matching
                        if measure_name and search_term_lower in measure_name.lower():
                            # Add a match for each report that uses this dataset
                            if using_reports:
                                for report in using_reports:
                                    reports_with_matches.add(report['id'])
                                    all_matches.append({
                                        'match_type': 'measure',
                                        'matched_text': measure_name,
                                        'report_id': report['id'],
                                        'report_name': report['name'],
                                        'page_name': None,
                                        'dataset_name': dataset_name,
                                        'context': f"DAX measure in table '{table.get('name', 'Unknown')}'"
                                    })
                                    print(f"      ✅ MEASURE MATCH: '{measure_name}' in dataset '{dataset_name}' used by '{report['name']}'")
                            else:
                                # Dataset not used by any report
                                all_matches.append({
                                    'match_type': 'measure',
                                    'matched_text': measure_name,
                                    'report_id': None,
                                    'report_name': 'No reports using this dataset',
                                    'page_name': None,
                                    'dataset_name': dataset_name,
                                    'context': f"DAX measure in table '{table.get('name', 'Unknown')}' (dataset not used in any report)"
                                })
                                print(f"      ✅ MEASURE MATCH: '{measure_name}' in dataset '{dataset_name}' (no reports)")

            # SECTION 3: Search in TABLE NAMES
            print("\n   🗂️ Searching table names...")
            for dataset in ws.get("datasets", []):
                dataset_id = dataset.get("id")
                dataset_name = dataset.get("name", "Unknown Dataset")

                # Find reports using this dataset
                using_reports = []
                for report in ws.get("reports", []):
                    if report.get("datasetId") == dataset_id:
                        using_reports.append({
                            'id': report.get('id'),
                            'name': report.get('name', 'Unknown Report')
                        })

                for table in dataset.get("tables", []):
                    table_display_name = table.get("name", "")

                    # Extract native source name (reuse logic from table-search)
                    native_source_name = None
                    source = table.get("source", [])
                    if source and len(source) > 0:
                        for src in source:
                            expr = src.get("expression", "")
                            if expr:
                                # Use simplified extraction (just SQL pattern)
                                sql_curly_match = re.search(r'\{?\[Schema="([^"]+)",\s*Item="([^"]+)"\]\}?', expr)
                                if sql_curly_match:
                                    schema = sql_curly_match.group(1)
                                    item = sql_curly_match.group(2)
                                    if '[' in schema and ']' in schema:
                                        schema = re.sub(r'[\[\]]', '', schema)
                                    native_source_name = f"[{schema}].[{item}]"
                                    break

                    if not native_source_name:
                        native_source_name = table_display_name

                    # Normalize for search
                    def normalize_for_search(text):
                        text = re.sub(r'[\[\]]', '', text)
                        text = re.sub(r'[._]', ' ', text)
                        return text.lower()

                    native_normalized = normalize_for_search(native_source_name)
                    display_normalized = normalize_for_search(table_display_name)

                    # Check if matches using LIKE matching
                    search_in_native = (search_term_lower in native_source_name.lower() or
                                       search_term_lower in native_normalized)
                    search_in_display = (search_term_lower in table_display_name.lower() or
                                        search_term_lower in display_normalized)

                    if search_in_native or search_in_display:
                        # Add a match for each report that uses this dataset
                        if using_reports:
                            for report in using_reports:
                                reports_with_matches.add(report['id'])
                                all_matches.append({
                                    'match_type': 'table',
                                    'matched_text': f"{native_source_name}" if native_source_name != table_display_name else table_display_name,
                                    'report_id': report['id'],
                                    'report_name': report['name'],
                                    'page_name': None,
                                    'dataset_name': dataset_name,
                                    'context': f"Table in dataset (display name: {table_display_name})" if native_source_name != table_display_name else f"Table in dataset"
                                })
                                print(f"      ✅ TABLE MATCH: '{native_source_name}' in dataset '{dataset_name}' used by '{report['name']}'")
                        else:
                            # Dataset not used by any report
                            all_matches.append({
                                'match_type': 'table',
                                'matched_text': f"{native_source_name}" if native_source_name != table_display_name else table_display_name,
                                'report_id': None,
                                'report_name': 'No reports using this dataset',
                                'page_name': None,
                                'dataset_name': dataset_name,
                                'context': f"Table in dataset (display name: {table_display_name})" if native_source_name != table_display_name else f"Table in dataset"
                            })
                            print(f"      ✅ TABLE MATCH: '{native_source_name}' in dataset '{dataset_name}' (no reports)")

            # SECTION 4: Search in COLUMN NAMES
            print("\n   📝 Searching column names...")
            total_columns_searched = 0
            for dataset in ws.get("datasets", []):
                dataset_id = dataset.get("id")
                dataset_name = dataset.get("name", "Unknown Dataset")

                # Find reports using this dataset
                using_reports = []
                for report in ws.get("reports", []):
                    if report.get("datasetId") == dataset_id:
                        using_reports.append({
                            'id': report.get('id'),
                            'name': report.get('name', 'Unknown Report')
                        })

                print(f"      📁 Dataset: {dataset_name} - used by {len(using_reports)} report(s)")

                for table in dataset.get("tables", []):
                    table_name = table.get("name", "")
                    columns = table.get("columns", [])
                    total_columns_searched += len(columns)

                    # Debug: Print table and column info if search term might match
                    if search_term in table_name.lower():
                        print(f"         🔎 Table '{table_name}' name contains search term")

                    for col in table.get("columns", []):
                        col_name = col.get('name', '')
                        col_source = col.get('sourceColumn', col_name)

                        # Debug: Log ALL column names if searching for HTS (for diagnostic)
                        if 'hts' in search_term_lower and 'hts' in col_name.lower():
                            print(f"         📝 Found column with 'hts': name='{col_name}', source='{col_source}'")

                        # Normalize column names for search
                        def normalize_for_search(text):
                            # Only remove brackets and convert dots/underscores to spaces
                            # DO NOT remove other special characters like #, $, etc.
                            text = re.sub(r'[\[\]]', '', text)
                            text = re.sub(r'[._]', ' ', text)
                            return text.lower()

                        col_name_normalized = normalize_for_search(col_name)
                        col_source_normalized = normalize_for_search(col_source)

                        # Enhanced matching: Check both exact (with special chars) and normalized
                        # This ensures "HTS#" matches "HTS#" even with the # character
                        # Using LIKE matching (case-insensitive substring match)
                        is_col_name_match = (search_term_lower in col_name.lower() or
                                            search_term_lower in col_name_normalized)
                        is_col_source_match = (search_term_lower in col_source.lower() or
                                              search_term_lower in col_source_normalized)

                        # Check if search term matches
                        if is_col_name_match or is_col_source_match:
                            # Debug: Log what matched
                            match_reason = []
                            if search_term_lower in col_name.lower():
                                match_reason.append(f"exact match in col_name '{col_name}'")
                            if search_term_lower in col_name_normalized:
                                match_reason.append(f"normalized match in col_name '{col_name_normalized}'")
                            if search_term_lower in col_source.lower():
                                match_reason.append(f"exact match in col_source '{col_source}'")
                            if search_term_lower in col_source_normalized:
                                match_reason.append(f"normalized match in col_source '{col_source_normalized}'")

                            print(f"         🎯 COLUMN MATCH in table '{table_name}': {col_name} | Reason: {', '.join(match_reason)}")

                            # Add a match for each report that uses this dataset
                            column_display = f"{col_source} (as: {col_name})" if col_source != col_name else col_name

                            if using_reports:
                                for report in using_reports:
                                    reports_with_matches.add(report['id'])
                                    all_matches.append({
                                        'match_type': 'column',
                                        'matched_text': column_display,
                                        'report_id': report['id'],
                                        'report_name': report['name'],
                                        'page_name': None,
                                        'dataset_name': dataset_name,
                                        'context': f"Column in table '{table_name}'"
                                    })
                                    print(f"      ✅ Added match for report: '{report['name']}'")
                            else:
                                # Dataset not used by any report
                                all_matches.append({
                                    'match_type': 'column',
                                    'matched_text': column_display,
                                    'report_id': None,
                                    'report_name': 'No reports using this dataset',
                                    'page_name': None,
                                    'dataset_name': dataset_name,
                                    'context': f"Column in table '{table_name}' (dataset not used in any report)"
                                })
                                print(f"      ✅ Added match (dataset not used in reports)")

        # Sort results: Visual matches first, then measures, then tables, then columns
        match_type_order = {'visual_title': 1, 'measure': 2, 'table': 3, 'column': 4}
        all_matches.sort(key=lambda x: (match_type_order.get(x['match_type'], 5), x['matched_text']))

        print(f"\n📄 ===============================================")
        print(f"📄 CONTENT SEARCH SUMMARY")
        print(f"📄 ===============================================")
        print(f"   🔍 Search term: '{search_term}'")
        print(f"   ✅ Total matches found: {len(all_matches)}")
        print(f"   📊 Reports with matches: {len(reports_with_matches)}")
        print(f"   📋 Breakdown:")
        print(f"      - Visual Titles: {len([m for m in all_matches if m['match_type'] == 'visual_title'])}")
        print(f"      - DAX Measures: {len([m for m in all_matches if m['match_type'] == 'measure'])}")
        print(f"      - Tables: {len([m for m in all_matches if m['match_type'] == 'table'])}")
        print(f"      - Columns: {len([m for m in all_matches if m['match_type'] == 'column'])}")
        print(f"📄 ===============================================\n")

        return jsonify({
            'success': True,
            'workspace_id': workspace_id,
            'search_term': search_term,
            'result_count': len(all_matches),
            'reports_count': len(reports_with_matches),
            'matches': all_matches
        })

    except Exception as e:
        print(f"❌ Error in content search: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500


# Global progress tracker for deep search
deep_search_progress = {}

@app.route('/api/workspace/deep-search-progress/<workspace_id>')
@login_required
def get_deep_search_progress(workspace_id):
    """Get current progress of deep search operation"""
    progress = deep_search_progress.get(workspace_id, {}).get('message', 'Processing...')
    return jsonify({'progress': progress})

@app.route('/api/workspace/deep-content-search/<workspace_id>')
@login_required
def deep_content_search(workspace_id):
    """
    OPTIMIZED Playwright Visual Extraction with:
    - Parallel processing
    - Metadata-driven pre-filtering
    - Smart caching with modifiedDateTime validation
    - Progress tracking for large workspaces
    - Folder-level scoping for large workspaces
    """
    try:
        from scanner_connector import PowerBIScanner
        from visual_metadata_extractor import VisualMetadataExtractor
        import re
        import asyncio
        import os
        import pickle
        from datetime import datetime, timedelta
        from concurrent.futures import ThreadPoolExecutor, as_completed
        import threading

        search_term = request.args.get('query', '').strip()
        folder_id = request.args.get('folder_id')  # Optional folder filter
        report_id = request.args.get('report_id')  # Optional single report filter

        if not search_term:
            return jsonify({'success': False, 'error': 'Search query is required'}), 400

        search_term_lower = search_term.lower()

        # Build scope description
        scope_info = []
        if folder_id:
            scope_info.append(f"Folder: {folder_id[:8]}...")
        if report_id:
            scope_info.append(f"Report: {report_id[:8]}...")
        if not scope_info:
            scope_info.append("All Reports")

        scope_str = " | ".join(scope_info)
        print(f"\n🔬 ===============================================")
        print(f"🔬 OPTIMIZED DEEP SEARCH: '{search_term}' in workspace {workspace_id}")
        print(f"🔬 Scope: {scope_str}")
        print(f"🔬 ===============================================")

        # Initialize progress tracker
        deep_search_progress[workspace_id] = {'message': 'Initializing...', 'lock': threading.Lock()}

        # Get user token
        user_token = session.get('access_token')
        if not user_token:
            return jsonify({'success': False, 'error': 'Not authenticated'}), 401

        # Create cache directory
        cache_dir = os.path.join(os.getcwd(), '.visual_cache')
        os.makedirs(cache_dir, exist_ok=True)

        # Step 1: Get Scanner API data for dataset metadata
        deep_search_progress[workspace_id]['message'] = 'Scanning workspace metadata...'
        print("\n   📊 Step 1: Running Scanner API scan for dataset metadata...")
        scanner = PowerBIScanner()
        scanner.access_token = scanner.get_access_token()
        scan_data = scanner.run_scan(workspace_id=workspace_id)

        if not scan_data or "workspaces" not in scan_data:
            return jsonify({'success': False, 'error': 'Failed to scan workspace'}), 500

        # Build dataset metadata index
        dataset_metadata = {}
        for ws in scan_data["workspaces"]:
            for dataset in ws.get("datasets", []):
                dataset_id = dataset.get("id")
                dataset_name = dataset.get("name", "Unknown Dataset")

                # Get datasources
                datasources = dataset.get("datasources", [])
                datasource_info = []
                for ds in datasources:
                    connection_details = ds.get("connectionDetails", {})
                    datasource_info.append({
                        'datasourceType': ds.get("datasourceType", "Unknown"),
                        'server': connection_details.get("server", "N/A"),
                        'database': connection_details.get("database", "N/A")
                    })

                # Build table/column index
                tables = {}
                for table in dataset.get("tables", []):
                    table_name = table.get("name", "")

                    # Extract native source name
                    native_source_name = None
                    source = table.get("source", [])
                    if source and len(source) > 0:
                        for src in source:
                            expr = src.get("expression", "")
                            if expr:
                                sql_curly_match = re.search(r'\{?\[Schema="([^"]+)",\s*Item="([^"]+)"\]\}?', expr)
                                if sql_curly_match:
                                    schema = sql_curly_match.group(1)
                                    item = sql_curly_match.group(2)
                                    if '[' in schema and ']' in schema:
                                        schema = re.sub(r'[\[\]]', '', schema)
                                    native_source_name = f"[{schema}].[{item}]"
                                    break

                    if not native_source_name:
                        native_source_name = table_name

                    # Build column index
                    columns = {}
                    for col in table.get("columns", []):
                        col_name = col.get('name', '')
                        col_source = col.get('sourceColumn', col_name)
                        columns[col_name.lower()] = {
                            'name': col_name,
                            'source_name': col_source,
                            'dataType': col.get('dataType', 'Unknown'),
                            'isHidden': col.get('isHidden', False)
                        }

                    tables[table_name] = {
                        'native_source_name': native_source_name,
                        'columns': columns
                    }

                dataset_metadata[dataset_id] = {
                    'name': dataset_name,
                    'tables': tables,
                    'datasources': datasource_info
                }

        print(f"   ✅ Indexed {len(dataset_metadata)} dataset(s)")

        # Step 1.5: PRE-FILTERING - Identify datasets containing search term
        deep_search_progress[workspace_id]['message'] = 'Pre-filtering datasets...'
        print("\n   🔍 Step 1.5: Pre-filtering datasets by search term...")
        search_normalized = re.sub(r'[._-]', ' ', search_term_lower)
        relevant_dataset_ids = set()

        for dataset_id, dataset_meta in dataset_metadata.items():
            tables = dataset_meta.get('tables', {})
            for table_name, table_data in tables.items():
                # Check table name
                if search_term_lower in table_name.lower():
                    relevant_dataset_ids.add(dataset_id)
                    print(f"      ✓ Dataset '{dataset_meta['name']}' contains table '{table_name}'")
                    break
                # Check columns
                for col_name in table_data.get('columns', {}).keys():
                    col_normalized = re.sub(r'[._-]', ' ', col_name.lower())
                    if search_term_lower in col_name.lower() or search_normalized in col_normalized:
                        relevant_dataset_ids.add(dataset_id)
                        print(f"      ✓ Dataset '{dataset_meta['name']}' contains column '{col_name}'")
                        break
                if dataset_id in relevant_dataset_ids:
                    break

        print(f"   ✅ Found {len(relevant_dataset_ids)} dataset(s) potentially containing '{search_term}'")
        if not relevant_dataset_ids:
            print(f"   ⚠️  No datasets contain the search term in schema - will search visual titles and aliases only")

        # Step 2: Get all reports in workspace (using Scanner API to get folder info)
        deep_search_progress[workspace_id]['message'] = 'Loading reports...'
        print("\n   📄 Step 2: Getting reports from workspace (with folder info)...")

        # Use Scanner API scan data which already has folder information
        reports = []
        for ws in scan_data.get("workspaces", []):
            if ws.get("id") == workspace_id:
                reports = ws.get("reports", [])
                break

        total_before_filter = len(reports)
        print(f"   ✅ Found {total_before_filter} report(s) from Scanner API")

        # NEW: Filter by folder if specified (including nested subfolders)
        if folder_id:
            print(f"\n   🗂️  Step 2.5: Filtering reports by folder {folder_id[:8]}...")

            # Get folder hierarchy from Fabric API
            try:
                folders_url = f"https://api.fabric.microsoft.com/v1/workspaces/{workspace_id}/folders?recursive=true"
                folders_response = requests.get(folders_url, headers=headers)

                folder_hierarchy = {}
                if folders_response.status_code == 200:
                    folders_data = folders_response.json().get('value', [])
                    for folder in folders_data:
                        fid = folder.get('id')
                        if fid:
                            folder_hierarchy[fid] = {
                                'id': fid,
                                'name': folder.get('displayName', ''),
                                'parentFolderId': folder.get('parentFolderId')
                            }

                    # Helper to get all child folder IDs recursively
                    def get_child_folders(parent_id, hierarchy):
                        children = {parent_id}
                        for fid, finfo in hierarchy.items():
                            if finfo.get('parentFolderId') == parent_id:
                                children.add(fid)
                                children.update(get_child_folders(fid, hierarchy))
                        return children

                    all_folder_ids = get_child_folders(folder_id, folder_hierarchy)
                    print(f"   📁 Including {len(all_folder_ids)} folder(s) (selected + nested subfolders)")

                    # Filter reports by folder hierarchy (Scanner API uses 'folderId')
                    reports = [r for r in reports if r.get('folderId') in all_folder_ids]
                else:
                    # Fallback to single folder only if Fabric API fails
                    reports = [r for r in reports if r.get('folderId') == folder_id]
            except Exception as e:
                print(f"   ⚠️ Error fetching folder hierarchy: {e}, filtering single folder only")
                reports = [r for r in reports if r.get('folderId') == folder_id]

            print(f"   ✅ Filtered from {total_before_filter} to {len(reports)} reports in selected folder")
        else:
            print(f"   ✅ Found {len(reports)} report(s) in workspace (all folders)")

        # NEW: Filter by specific report if report_id is provided
        if report_id:
            reports_before_report_filter = len(reports)
            reports = [r for r in reports if r.get('id') == report_id]
            print(f"   🎯 Filtered to specific report: {reports_before_report_filter} → {len(reports)} report(s)")

        # Step 3: Extract visual metadata with smart caching
        deep_search_progress[workspace_id]['message'] = 'Extracting visual metadata...'
        print(f"\n   🎨 Step 3: Extracting visual metadata with smart caching...")

        all_matches = []
        reports_processed = 0
        reports_with_matches = set()
        reports_from_cache = 0
        reports_extracted = 0
        reports_skipped = 0
        reports_filtered = 0

        # Initialize Visual Extractor
        extractor = VisualMetadataExtractor(user_token=user_token)

        total_reports = len(reports)
        for idx, report in enumerate(reports, 1):
            report_id = report.get('id')
            report_name = report.get('name', 'Unknown Report')
            dataset_id = report.get('datasetId')

            # Update progress
            with deep_search_progress[workspace_id]['lock']:
                deep_search_progress[workspace_id]['message'] = f'Processing {idx}/{total_reports} reports...'

            # Skip platform / system reports (usage metrics, capacity metrics, [App] shells)
            if (
                _is_excluded_report_name(report_name)
                or 'Fabric Capacity Metrics' in report_name
            ):
                print(f"\n      ⏭️  Skipping system report: {report_name}")
                reports_skipped += 1
                continue

            # PRE-FILTER DISABLED: We need to search ALL reports because the term might exist in:
            # - Visual titles, field aliases, DAX measures, or visual-level renamed fields
            # The pre-filtering was TOO AGGRESSIVE and caused false negatives
            # if relevant_dataset_ids and dataset_id not in relevant_dataset_ids:
            #     print(f"\n      🚫 Filtered out: {report_name} (dataset doesn't contain search term)")
            #     reports_filtered += 1
            #     continue

            print(f"\n      📊 Processing report: {report_name}")

            # Check cache first (SMART CACHE with modifiedDateTime validation)
            cache_file = os.path.join(cache_dir, f"{report_id}.pkl")
            visual_result = None
            use_cache = False

            if os.path.exists(cache_file):
                try:
                    with open(cache_file, 'rb') as f:
                        cached_data = pickle.load(f)

                    # Smart cache validation using report modifiedDateTime
                    report_modified = report.get('modifiedDateTime')
                    cache_timestamp = cached_data.get('cached_at')

                    if report_modified and cache_timestamp:
                        try:
                            from dateutil import parser
                            report_modified_dt = parser.parse(report_modified)
                            cache_dt = datetime.fromisoformat(cache_timestamp) if isinstance(cache_timestamp, str) else cache_timestamp

                            if report_modified_dt <= cache_dt:
                                print(f"         💾 Using cache (report unchanged since {cache_dt.strftime('%Y-%m-%d %H:%M')})")
                                visual_result = cached_data
                                use_cache = True
                                reports_from_cache += 1
                            else:
                                print(f"         🔄 Report modified on {report_modified_dt.strftime('%Y-%m-%d %H:%M')}, re-extracting...")
                        except Exception as e:
                            print(f"         ⚠️  Date parsing failed: {e}, using fallback cache logic")
                            # Fallback to 24h TTL
                            cache_age = datetime.now() - datetime.fromtimestamp(os.path.getmtime(cache_file))
                            if cache_age < timedelta(hours=24):
                                visual_result = cached_data
                                use_cache = True
                                reports_from_cache += 1
                                print(f"         💾 Using cache (age: {cache_age.seconds // 3600}h)")
                    else:
                        # Fallback to 24h TTL if no modifiedDateTime
                        cache_age = datetime.now() - datetime.fromtimestamp(os.path.getmtime(cache_file))
                        if cache_age < timedelta(hours=24):
                            visual_result = cached_data
                            use_cache = True
                            reports_from_cache += 1
                            print(f"         💾 Using cache (age: {cache_age.seconds // 3600}h)")
                        else:
                            print(f"         🕐 Cache expired (age: {cache_age.days}d)")
                except Exception as e:
                    print(f"         ⚠️ Cache validation failed: {e}, will extract fresh")
                    visual_result = None

            # Extract if not in cache or cache invalid
            if not use_cache:
                print(f"         🌐 Extracting visual metadata via Playwright...")
                try:
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)
                    visual_result = loop.run_until_complete(
                        extractor.extract_visuals(workspace_id, report_id, timeout=60)
                    )
                    loop.close()

                    if visual_result.get('success'):
                        # Cache the result with timestamp
                        visual_result['cached_at'] = datetime.now().isoformat()
                        with open(cache_file, 'wb') as f:
                            pickle.dump(visual_result, f)
                        print(f"         💾 Cached visual metadata")
                        reports_extracted += 1
                    else:
                        print(f"         ⚠️ Extraction failed: {visual_result.get('error')}")
                        continue
                except Exception as e:
                    print(f"         ❌ Error extracting visuals: {str(e)}")
                    continue

            if not visual_result or not visual_result.get('success'):
                continue

            pages = visual_result.get('pages', [])
            print(f"         ✅ Loaded {len(pages)} page(s)")

            # Get dataset metadata
            dataset_meta = dataset_metadata.get(dataset_id, {})
            dataset_name = dataset_meta.get('name', 'Unknown Dataset')
            dataset_tables = dataset_meta.get('tables', {})
            dataset_datasources = dataset_meta.get('datasources', [])

            # Search through visual titles AND field bindings
            for page in pages:
                page_name = page.get('displayName', 'Unknown Page')

                for visual in page.get('visuals', []):
                    visual_name = visual.get('name', 'Unknown')
                    visual_type = visual.get('type', 'unknown')
                    visual_title = visual.get('title', '')
                    visual_fields = visual.get('fields', [])

                    # FILTER: Skip ONLY non-data decorative visuals
                    # Keep slicers, cards in case user wants to search for fields used there
                    non_data_visuals = ['actionButton', 'textbox', 'image', 'shape', 'rectangle']
                    if visual_type in non_data_visuals:
                        continue  # Skip these visual types entirely

                    # DEBUG: Show all fields for data visuals only
                    if visual_fields:
                        print(f"\n         🔍 DEBUG: Visual '{visual_title or visual_name}' (type: {visual_type}) on page '{page_name}' has {len(visual_fields)} fields:")
                        for fld in visual_fields[:15]:  # Show first 15 fields
                            print(f"            - name: '{fld.get('name')}', displayName: '{fld.get('displayName')}', table: '{fld.get('table', 'N/A')}'")

                    # Check for matches in:
                    # 1. Visual Title
                    # 2. Field Names (columns/measures used in the visual)
                    match_found = False
                    match_reason = None
                    matched_fields = []

                    # Search in visual title
                    if visual_title and search_term_lower in visual_title.lower():
                        match_found = True
                        match_reason = 'visual_title'
                        print(f"         🎯 MATCH in TITLE: '{visual_title}' (type: {visual_type}) on page '{page_name}'")

                    # Normalize function for field name matching (handle spaces, underscores, hyphens)
                    def normalize_field_name(text):
                        if not text:
                            return ''
                        # Replace underscores, hyphens, dots with spaces for matching
                        text = re.sub(r'[._-]', ' ', text)
                        return text.lower().strip()

                    search_normalized = normalize_field_name(search_term)

                    # Search in field bindings (METHOD 1: Direct field match)
                    if not match_found and visual_fields:
                        for field in visual_fields:
                            field_name = field.get('name', '')
                            field_display = field.get('displayName', '')

                            # Normalize both field name and display name
                            field_name_normalized = normalize_field_name(field_name)
                            field_display_normalized = normalize_field_name(field_display)

                            # Check if search term matches either the exact field name or normalized version
                            if (search_term_lower in field_name.lower() or
                                search_term_lower in field_display.lower() or
                                search_normalized in field_name_normalized or
                                search_normalized in field_display_normalized):
                                match_found = True
                                match_reason = 'field_binding'
                                matched_fields.append(field.get('displayName') or field.get('name'))

                        if match_found:
                            fields_str = ', '.join(matched_fields[:3])  # Show first 3 fields
                            if len(matched_fields) > 3:
                                fields_str += f" (+{len(matched_fields)-3} more)"
                            title_display = f"'{visual_title}'" if visual_title else "(no title)"
                            print(f"         🎯 MATCH in FIELDS: {title_display} (type: {visual_type}) on page '{page_name}' - Fields: {fields_str}")

                    # ENHANCED SEARCH (METHOD 2: Dataset column match via lineage)
                    # Even if fields weren't captured, check if the visual uses tables with matching columns
                    if not match_found:
                        # Build list of tables this visual uses (from extracted fields)
                        visual_tables = set()
                        for field in visual_fields:
                            table = field.get('table', '')
                            if table:
                                visual_tables.add(table.lower())

                        # If no tables found, search ALL tables (for visuals with incomplete field extraction)
                        # Otherwise, only search tables the visual is confirmed to use
                        tables_to_search = visual_tables if visual_tables else [t.lower() for t in dataset_tables.keys()]

                        # Search dataset tables for columns matching search term
                        matching_tables = {}  # {table_name: [matching_columns]}
                        for table_name, table_data in dataset_tables.items():
                            # Skip if this visual doesn't use this table (if we have that info)
                            if visual_tables and table_name.lower() not in visual_tables:
                                continue

                            table_columns = table_data.get('columns', {})
                            table_matches = []
                            for col_name_key, col_info in table_columns.items():
                                col_name = col_info.get('name', '')
                                col_normalized = normalize_field_name(col_name)

                                # Check if column name matches search term
                                if (search_term_lower in col_name.lower() or
                                    search_normalized in col_normalized):
                                    table_matches.append(col_name)
                                    matched_fields.append(col_name)

                            if table_matches:
                                matching_tables[table_name] = table_matches

                        if matched_fields:
                            match_found = True
                            match_reason = 'dataset_column_lineage'
                            fields_str = ', '.join(matched_fields[:5])  # Show first 5 matches
                            if len(matched_fields) > 5:
                                fields_str += f" (+{len(matched_fields)-5} more)"
                            title_display = f"'{visual_title}'" if visual_title else "(no title)"
                            confidence = "HIGH" if visual_tables else "POSSIBLE"
                            print(f"         🎯 {confidence} MATCH via LINEAGE: {title_display} (type: {visual_type}) on page '{page_name}' - Columns: {fields_str}")

                    if match_found:

                        # Try to find matching column in dataset
                        matching_column_info = None
                        matching_table_info = None

                        # Try to match based on search term or matched fields
                        search_fields = matched_fields if matched_fields else [search_term]

                        for search_field in search_fields:
                            search_field_lower = search_field.lower().strip()

                            for table_name, table_data in dataset_tables.items():
                                table_columns = table_data.get('columns', {})

                                if search_field_lower in table_columns:
                                    col_info = table_columns[search_field_lower]
                                    matching_column_info = {
                                        'column_name': col_info.get('name'),
                                        'source_name': col_info.get('source_name'),
                                        'dataType': col_info.get('dataType')
                                    }
                                    matching_table_info = {
                                        'table_name': table_name,
                                        'native_source_name': table_data.get('native_source_name')
                                    }
                                    break

                            if matching_column_info:
                                break

                        # Build datasource text
                        datasource_text = "N/A"
                        if dataset_datasources:
                            ds = dataset_datasources[0]
                            if ds['server'] != 'N/A':
                                datasource_text = f"{ds['server']} / {ds['database']}"

                        # Add match
                        all_matches.append({
                            'match_type': match_reason,
                            'report_id': report_id,
                            'report_name': report_name,
                            'page_name': page_name,
                            'visual_type': visual_type,
                            'visual_title': visual_title or '(no title)',
                            'visual_name': visual_name,
                            'matched_fields': matched_fields if matched_fields else None,
                            'dataset_name': dataset_name,
                            'table_name': matching_table_info.get('table_name') if matching_table_info else None,
                            'table_native_name': matching_table_info.get('native_source_name') if matching_table_info else None,
                            'column_name': matching_column_info.get('column_name') if matching_column_info else None,
                            'column_source_name': matching_column_info.get('source_name') if matching_column_info else None,
                            'column_dataType': matching_column_info.get('dataType') if matching_column_info else None,
                            'datasource': datasource_text,
                            'is_alias': visual_title != matching_column_info.get('source_name') if matching_column_info else False
                        })

                        reports_with_matches.add(report_id)

            reports_processed += 1

        print(f"\n🔬 ===============================================")
        print(f"🔬 OPTIMIZED DEEP SEARCH SUMMARY")
        print(f"🔬 ===============================================")
        print(f"   🔍 Search term: '{search_term}'")
        print(f"   📊 Total reports: {len(reports)}")
        print(f"   ✅ Processed: {reports_processed}")
        print(f"   ⏭️  Skipped: {reports_skipped} (system reports)")
        print(f"   🚫 Filtered: {reports_filtered} (dataset pre-filter)")
        print(f"   💾 From cache: {reports_from_cache}")
        print(f"   🌐 Newly extracted: {reports_extracted}")
        print(f"   🎯 Matches found: {len(all_matches)}")
        print(f"   📊 Reports with matches: {len(reports_with_matches)}")
        print(f"🔬 ===============================================\n")

        # Clear progress
        deep_search_progress[workspace_id]['message'] = 'Complete'

        return jsonify({
            'success': True,
            'workspace_id': workspace_id,
            'search_term': search_term,
            'result_count': len(all_matches),
            'reports_count': len(reports_with_matches),
            'reports_total': len(reports),
            'reports_processed': reports_processed,
            'reports_skipped': reports_skipped,
            'reports_filtered': reports_filtered,
            'reports_from_cache': reports_from_cache,
            'reports_extracted': reports_extracted,
            'matches': all_matches
        })

    except Exception as e:
        print(f"❌ Error in optimized deep search: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/clear-cache', methods=['POST'])
@login_required
def clear_cache():
    """
    ⚡ PERFORMANCE OPTIMIZATION: Clear all caches for current user
    This endpoint allows users to force a refresh of all cached data
    """
    try:
        user_id = session.get('user', {}).get('oid')
        if not user_id:
            return jsonify({'success': False, 'error': 'Not authenticated'}), 401

        # Clear user-specific caches
        cleared_count = 0

        # Clear workspaces cache
        workspace_key = f"workspaces_{user_id}"
        if workspace_key in workspaces_cache:
            del workspaces_cache[workspace_key]
            cleared_count += 1

        # Clear all workspace summaries for this user
        keys_to_delete = [k for k in workspace_cache.keys() if k.endswith(f"_{user_id}")]
        for key in keys_to_delete:
            del workspace_cache[key]
            cleared_count += 1

        # Clear reports cache for this user
        report_keys_to_delete = [k for k in reports_cache.keys() if f"_{user_id}_" in k]
        for key in report_keys_to_delete:
            del reports_cache[key]
            cleared_count += 1

        print(f"🗑️  Cleared {cleared_count} cache entries for user: {user_id}")

        return jsonify({
            'success': True,
            'message': f'Cache cleared successfully ({cleared_count} entries)',
            'cleared_count': cleared_count
        })

    except Exception as e:
        print(f"❌ Error clearing cache: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/workspaces')
@login_required
def get_workspaces():
    """API endpoint to get workspaces the logged-in user has access to (user-delegated permissions)"""
    try:
        import time
        import requests
        current_time = time.time()

        # Get user info for cache key - CRITICAL: Validate user is properly authenticated
        user_id = session.get('user', {}).get('oid')
        user_email = session.get('user', {}).get('preferred_username', 'Unknown')

        # If no valid user ID, force re-authentication
        if not user_id:
            print("\n⚠️ NO USER ID IN SESSION - Forcing re-authentication")
            session.clear()
            return jsonify({
                'success': False,
                'error': 'Session invalid. Please log in again.',
                'redirect': '/login'
            }), 401

        cache_key = f"workspaces_{user_id}"
        print(f"\n📊 Fetching workspaces for user: {user_email} (ID: {user_id})")

        # Check cache (per-user caching)
        if cache_key in workspaces_cache and workspaces_cache[cache_key].get('data') and \
           (current_time - workspaces_cache[cache_key].get('timestamp', 0)) < CACHE_DURATION:
            print(f"   ✅ Returning cached workspaces ({len(workspaces_cache[cache_key]['data'])} items)")
            return jsonify({
                'success': True,
                'workspaces': workspaces_cache[cache_key]['data']
            })

        # Fetch fresh data using USER'S delegated token (not service principal)
        base_url = "https://api.powerbi.com/v1.0/myorg"
        url = f"{base_url}/groups"

        try:
            # First check if we have a valid token
            token = get_user_powerbi_token()
            if not token:
                print("\n❌ NO VALID TOKEN - User needs to re-authenticate")
                # Clear session and redirect to login
                session.pop("access_token", None)
                return jsonify({
                    'success': False,
                    'error': 'Session expired. Please log in again.',
                    'redirect': '/login'
                }), 401

            headers = get_user_powerbi_headers()
            response = requests.get(url, headers=headers)
            response.raise_for_status()

            workspaces = response.json().get('value', [])

            # Log authentication method and user info
            user_info = session.get('user', {})
            user_name = user_info.get('name', 'Unknown')
            user_email = user_info.get('preferred_username', 'Unknown')

            print("\n" + "="*80)
            print("🔐 SSO-BASED WORKSPACE ACCESS CONTROL TEST")
            print("="*80)
            print(f"✅ Authentication Method: USER DELEGATED TOKEN (SSO)")
            print(f"👤 Logged-in User: {user_name} ({user_email})")
            print(f"📊 Total Workspaces User Has Access To: {len(workspaces)}")

            if workspaces:
                print("\n📁 User's Accessible Workspaces:")
                for idx, workspace in enumerate(workspaces, 1):
                    print(f"   {idx}. {workspace['name']} (ID: {workspace['id']})")
            else:
                print("\n⚠️  User has no workspace access")
            print("="*80 + "\n")

        except requests.exceptions.RequestException as e:
            print("\n" + "="*80)
            print("❌ USER TOKEN FAILED")
            print("="*80)
            print(f"Error: {str(e)}")
            if hasattr(e, 'response') and e.response is not None:
                print(f"Status Code: {e.response.status_code}")
                print(f"Response: {e.response.text[:500]}")

                # Check if token expired
                response_text = e.response.text.lower()
                if 'tokenexpired' in response_text or 'token has expired' in response_text:
                    print("\n⚠️ TOKEN EXPIRED - User needs to re-authenticate")
                    print("="*80 + "\n")
                    # Clear the expired token from session
                    session.pop("access_token", None)
                    # Return 401 to trigger re-login
                    return jsonify({
                        'success': False,
                        'error': 'Session expired. Please log in again.',
                        'redirect': '/login'
                    }), 401

                # If 403 Forbidden (not token expired), try fallback to service principal
                if e.response.status_code == 403:
                    print("\n⚠️  User delegated permissions not granted. Falling back to service principal...")
                    print("   To fix this permanently, grant admin consent for Power BI API permissions in Azure Portal.")
                    print("\n🔄 Authentication Method: SERVICE PRINCIPAL (FALLBACK)")

                    try:
                        # Fallback to service principal
                        workspaces = powerbi.get_workspaces()
                        print(f"✅ Service Principal returned {len(workspaces)} workspaces")
                        print("⚠️  WARNING: User is seeing ALL workspaces (not filtered by user permissions)")
                        print("="*80 + "\n")
                    except Exception as fallback_error:
                        print(f"❌ Fallback authentication failed: {str(fallback_error)}")
                        print("="*80 + "\n")
                        return jsonify({
                            'success': False,
                            'error': 'Unable to fetch workspaces. Please contact your administrator to grant Power BI API permissions.',
                            'details': str(e)
                        }), 500
                else:
                    # Other HTTP error - return error response
                    print("="*80 + "\n")
                    return jsonify({
                        'success': False,
                        'error': f'Failed to fetch workspaces: {str(e)}',
                        'status_code': e.response.status_code if hasattr(e, 'response') and e.response else 500
                    }), 500
            else:
                # Network error or other non-HTTP exception - try service principal fallback
                print("\n⚠️  Network error accessing Power BI API. Falling back to service principal...")
                print("="*80 + "\n")

                try:
                    # Fallback to service principal
                    workspaces = powerbi.get_workspaces()
                    print(f"✅ Service Principal returned {len(workspaces)} workspaces")
                    print("⚠️  WARNING: User is seeing ALL workspaces (not filtered by user permissions)")
                    print("="*80 + "\n")
                except Exception as fallback_error:
                    print(f"❌ Fallback authentication failed: {str(fallback_error)}")
                    print("="*80 + "\n")
                    return jsonify({
                        'success': False,
                        'error': f'Network error and fallback failed: {str(e)}',
                        'details': str(fallback_error)
                    }), 500

        # Show all workspaces the user has access to (no filtering)
        workspace_list = [
            {
                'id': ws['id'],
                'name': ws['name']
            }
            for ws in workspaces
        ]

        # Optional: Uncomment below to filter to show only "Fabric Admin Governance" workspace
        # workspace_list = [ws for ws in workspace_list if ws['name'] == 'Fabric Admin Governance']
        # if not workspace_list:
        #     workspace_list = [{'id': ws['id'], 'name': ws['name']} for ws in workspaces]

        print(f"📤 Returning {len(workspace_list)} workspace(s) to frontend")
        print()

        # Update cache (per-user)
        if cache_key not in workspaces_cache:
            workspaces_cache[cache_key] = {'data': None, 'timestamp': 0}

        workspaces_cache[cache_key]['data'] = workspace_list
        workspaces_cache[cache_key]['timestamp'] = current_time

        return jsonify({
            'success': True,
            'workspaces': workspace_list
        })
    except Exception as e:
        print(f"❌ Error fetching workspaces: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/datasets')
@login_required
def get_datasets():
    """API endpoint to get datasets for a workspace"""
    try:
        workspace_id = request.args.get('workspaceId')

        if not workspace_id:
            return jsonify({
                'success': False,
                'error': 'workspaceId parameter is required'
            }), 400

        # Get user's access token
        token = get_user_powerbi_token()
        if not token:
            return jsonify({
                'success': False,
                'error': 'Session expired. Please log in again.',
                'redirect': '/login'
            }), 401

        # Fetch datasets from Power BI API
        base_url = "https://api.powerbi.com/v1.0/myorg"
        url = f"{base_url}/groups/{workspace_id}/datasets"

        headers = get_user_powerbi_headers()
        response = requests.get(url, headers=headers)
        response.raise_for_status()

        datasets = response.json().get('value', [])

        # Filter and format datasets
        dataset_list = [
            {
                'id': ds['id'],
                'name': ds['name'],
                'configuredBy': ds.get('configuredBy', 'Unknown'),
                'isRefreshable': ds.get('isRefreshable', False),
                'isOnPremGatewayRequired': ds.get('isOnPremGatewayRequired', False)
            }
            for ds in datasets
        ]

        print(f"📊 Fetched {len(dataset_list)} dataset(s) from workspace {workspace_id}")

        return jsonify({
            'success': True,
            'datasets': dataset_list
        })

    except requests.exceptions.RequestException as e:
        print(f"❌ Error fetching datasets: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/semantic-models')
@login_required
def get_semantic_models():
    """
    List semantic models for a workspace.

    Fast path: precomputed workspace_catalog (+ ops refresh fields).
    Fallback: one Scanner scan (NO live refresh fan-out — that was causing multi-minute hangs).
    """
    try:
        workspace_id = request.args.get('workspace_id')
        if not workspace_id:
            return jsonify({
                'success': False,
                'error': 'workspace_id parameter is required'
            }), 400

        # Optional: only when explicitly requested (and never the default)
        live_refresh = str(request.args.get('live_refresh', '0')).lower() in ('1', 'true', 'yes')

        def _measures_from_tables(tables):
            measures = []
            for table in tables or []:
                for measure in table.get('measures') or []:
                    expr = measure.get('expression') or ''
                    measures.append({
                        'name': measure.get('name'),
                        'table': table.get('name'),
                        'expression': expr[:100] if isinstance(expr, str) else '',
                    })
            return measures

        def _model_from_dataset(dataset, ops=None):
            tables = dataset.get('tables') or []
            relationships = dataset.get('relationships') or []
            measures = _measures_from_tables(tables)
            last_refresh = None
            refresh_status = 'Unknown'
            refresh_type = None

            ops = ops or {}
            # Catalog ops fields (preferred)
            last_refresh = (
                ops.get('last_refreshed')
                or dataset.get('last_refreshed')
                or dataset.get('lastRefresh')
                or dataset.get('lastRefreshTime')
            )
            refresh_status = (
                ops.get('last_refresh_status')
                or dataset.get('last_refresh_status')
                or dataset.get('refreshStatus')
                or ('Unknown' if not last_refresh else 'Completed')
            )
            refresh_type = ops.get('refresh_type') or dataset.get('refresh_type')

            return {
                'id': dataset.get('id'),
                'name': dataset.get('name'),
                'tables': tables,
                'relationships': relationships,
                'measures': measures,
                'tableCount': len(tables),
                'relationshipCount': len(relationships),
                'measureCount': len(measures),
                'lastRefresh': last_refresh,
                'refreshStatus': refresh_status,
                'refreshType': refresh_type,
                'configuredBy': dataset.get('configuredBy') or dataset.get('configuredByUser') or 'Unknown',
                'source': 'catalog' if ops is not None or dataset.get('_from_catalog') else 'scanner',
            }

        # ---------- Catalog fast path ----------
        if CATALOG_AVAILABLE and catalog_service is not None and catalog_service.is_available():
            try:
                cat = catalog_service.get_workspace_catalog()
                ws = None
                if cat:
                    ws = next(
                        (w for w in (cat.get('workspaces') or []) if w.get('id') == workspace_id),
                        None,
                    )
                if ws is not None:
                    datasets = list(ws.get('datasets') or [])
                    # Some catalogs only embed dataset ids on reports — pull from top-level map
                    if not datasets and isinstance(cat.get('datasets'), dict):
                        for did, d in cat['datasets'].items():
                            if (d or {}).get('workspaceId') == workspace_id:
                                dd = dict(d)
                                dd.setdefault('id', did)
                                datasets.append(dd)

                    # Preload ops refresh snapshot once (dataset id → refresh fields)
                    snap_map = {}
                    try:
                        snap = catalog_service.get_json('refresh_snapshot.json') or {}
                        snap_map = snap.get('datasets') or {}
                    except Exception:
                        snap_map = {}

                    # Infer light stats from reports when dataset schema not embedded
                    report_ds_stats = {}
                    for r in (ws.get('reports') or []):
                        did = r.get('datasetId')
                        if not did:
                            continue
                        st = report_ds_stats.setdefault(did, {
                            'reportCount': 0,
                            'last_refreshed': r.get('last_refreshed'),
                            'last_refresh_status': r.get('last_refresh_status'),
                            'refresh_type': r.get('refresh_type'),
                        })
                        st['reportCount'] += 1
                        # prefer report ops if present
                        if r.get('last_refreshed') and not st.get('last_refreshed'):
                            st['last_refreshed'] = r.get('last_refreshed')
                        if r.get('last_refresh_status') and not st.get('last_refresh_status'):
                            st['last_refresh_status'] = r.get('last_refresh_status')

                    # Full schema lives on top-level cat['datasets'] map (not always on ws.datasets)
                    dmap = cat.get('datasets') if isinstance(cat.get('datasets'), dict) else {}

                    models = []
                    for ds in datasets:
                        ds = dict(ds or {})
                        ds['_from_catalog'] = True
                        did = ds.get('id') or ''
                        rich = dict(dmap.get(did) or {}) if did else {}
                        info = snap_map.get(did) or {}
                        rstat = report_ds_stats.get(did) or {}
                        ops = {
                            'last_refreshed': (
                                ds.get('last_refreshed')
                                or rich.get('last_refreshed')
                                or info.get('last_refreshed')
                                or rstat.get('last_refreshed')
                            ),
                            'last_refresh_status': (
                                ds.get('last_refresh_status')
                                or rich.get('last_refresh_status')
                                or info.get('last_refresh_status')
                                or rstat.get('last_refresh_status')
                            ),
                            'refresh_type': (
                                ds.get('refresh_type')
                                or rich.get('refresh_type')
                                or info.get('refresh_type')
                                or rstat.get('refresh_type')
                            ),
                            'refresh_schedule': (
                                ds.get('refresh_schedule')
                                or rich.get('refresh_schedule')
                                or info.get('refresh_schedule')
                            ),
                        }
                        model = _model_from_dataset(ds, ops=ops)

                        # Counts from rich schema when list entry is thin
                        rich_tables = rich.get('tables') or []
                        table_count = (
                            ds.get('tableCount')
                            or rich.get('tableCount')
                            or len(rich_tables)
                            or 0
                        )
                        measure_count = 0
                        for t in rich_tables:
                            measure_count += len(t.get('measures') or [])
                            if t.get('measureCount'):
                                # prefer explicit count if measures array empty in extract
                                if not t.get('measures'):
                                    measure_count += int(t.get('measureCount') or 0)
                        rel_count = len(rich.get('relationships') or ds.get('relationships') or [])

                        model['tableCount'] = table_count
                        model['measureCount'] = measure_count
                        model['relationshipCount'] = rel_count
                        model['reportCount'] = rstat.get('reportCount') or 0
                        # List payload stays light — full schema on Details
                        model['tables'] = []
                        model['relationships'] = []
                        model['measures'] = []
                        models.append(model)

                    models.sort(key=lambda m: (m.get('name') or '').lower())
                    print(
                        f"⚡ SEMANTIC MODELS from catalog: ws={workspace_id[:8]}… "
                        f"models={len(models)} opsHit={sum(1 for m in models if m.get('lastRefresh'))} "
                        f"opsEnrichedAt={cat.get('opsEnrichedAt')}"
                    )
                    return jsonify({
                        'success': True,
                        'models': models,
                        'source': 'catalog',
                        'opsEnrichedAt': cat.get('opsEnrichedAt'),
                        'generatedAt': cat.get('generatedAt'),
                    })
            except Exception as cat_err:
                print(f"⚠️ Semantic models catalog path failed, falling back to scanner: {cat_err}")

        # ---------- Scanner fallback (list only; no per-dataset live refresh) ----------
        from scanner_connector import PowerBIScanner

        print(f"📊 Fetching semantic models via Scanner for workspace {workspace_id}...")
        scanner = PowerBIScanner()
        scan_result = scanner.run_scan(workspace_id=workspace_id)
        if not scan_result:
            return jsonify({'success': False, 'error': 'Failed to scan workspace'}), 500

        workspaces = scan_result.get('workspaces', [])
        if not workspaces:
            return jsonify({'success': True, 'models': [], 'source': 'scanner'})

        datasets = workspaces[0].get('datasets', []) or []
        print(f"✅ Scanner returned {len(datasets)} semantic model(s)")

        # Optional live refresh (opt-in only) — never default; kills the UI spinner for minutes
        refresh_by_id = {}
        if live_refresh and datasets:
            print(f"⚠️ live_refresh=1 — resolving refresh for {len(datasets)} models (slow)")
            try:
                from powerbi_connector import resolve_dataset_refresh_info
                headers = get_user_powerbi_headers()
                for dataset in datasets:
                    did = dataset.get('id')
                    if not did:
                        continue
                    try:
                        refresh_by_id[did] = resolve_dataset_refresh_info(
                            headers=headers,
                            workspace_id=workspace_id,
                            dataset_id=did,
                            dataset_workspace_id=workspace_id,
                            history_top=3,
                            timeout=6,
                        )
                    except Exception:
                        continue
            except Exception as e:
                print(f"⚠️ live refresh batch failed: {e}")

        models = []
        for dataset in datasets:
            ops = None
            info = refresh_by_id.get(dataset.get('id') or '')
            if info:
                ops = {
                    'last_refreshed': info.get('last_refreshed'),
                    'last_refresh_status': info.get('last_refresh_status'),
                    'refresh_type': info.get('refresh_type'),
                }
            models.append(_model_from_dataset(dataset, ops=ops))

        models.sort(key=lambda m: (m.get('name') or '').lower())
        print(f"✅ Returning {len(models)} semantic model(s) (source=scanner, live_refresh={live_refresh})")
        return jsonify({
            'success': True,
            'models': models,
            'source': 'scanner',
            'liveRefresh': live_refresh,
        })

    except Exception as e:
        print(f"❌ Error fetching semantic models: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


# Workspace-level Scanner schema cache for Semantic Models Details
# (one scan fills measures/relationships for all models in that workspace)
_semantic_scan_cache = {}
_semantic_scan_lock = threading.Lock()
_SEMANTIC_SCAN_TTL_SEC = 45 * 60  # 45 minutes


def _get_workspace_scan_cached(workspace_id: str) -> dict:
    """Run Admin Scanner once per workspace and cache datasets by id."""
    now = time.time()
    with _semantic_scan_lock:
        hit = _semantic_scan_cache.get(workspace_id)
        if hit and (now - hit.get('ts', 0)) < _SEMANTIC_SCAN_TTL_SEC:
            return hit.get('by_id') or {}

    from scanner_connector import PowerBIScanner
    print(f"📡 Scanner schema scan for semantic details ws={workspace_id[:8]}…")
    scanner = PowerBIScanner()
    scan_result = scanner.run_scan(workspace_id=workspace_id) or {}
    workspaces = scan_result.get('workspaces') or []
    by_id = {}
    if workspaces:
        for ds in workspaces[0].get('datasets') or []:
            did = ds.get('id')
            if did:
                by_id[did] = ds
    print(f"   cached {len(by_id)} dataset schemas for workspace {workspace_id[:8]}…")

    with _semantic_scan_lock:
        _semantic_scan_cache[workspace_id] = {'ts': time.time(), 'by_id': by_id}
    return by_id


def _get_scanned_dataset_cached(workspace_id: str, dataset_id: str):
    by_id = _get_workspace_scan_cached(workspace_id)
    return by_id.get(dataset_id)


def _merge_dataset_schema(base: dict, scan_ds: dict) -> dict:
    """Merge scanner measures/relationships/columns onto catalog dataset."""
    out = dict(base or {})
    if not scan_ds:
        return out

    if not out.get('relationships') and scan_ds.get('relationships'):
        out['relationships'] = scan_ds.get('relationships')
    if scan_ds.get('configuredBy') and not out.get('configuredBy'):
        out['configuredBy'] = scan_ds.get('configuredBy')
    if scan_ds.get('createdDate') and not out.get('createdDate'):
        out['createdDate'] = scan_ds.get('createdDate')

    scan_tables = {
        (t.get('name') or ''): t for t in (scan_ds.get('tables') or [])
    }
    merged = []
    for t in out.get('tables') or []:
        tt = dict(t)
        st = scan_tables.get(tt.get('name') or '')
        if st:
            if not tt.get('measures') and st.get('measures'):
                tt['measures'] = st.get('measures')
                tt['measureCount'] = len(st.get('measures') or [])
            if not tt.get('columns') and st.get('columns'):
                tt['columns'] = st.get('columns')
                tt['columnCount'] = len(st.get('columns') or [])
            # fill measure expressions if catalog only had empty measures
            if tt.get('measures') and st.get('measures'):
                by_name = {
                    (m.get('name') if isinstance(m, dict) else str(m)): m
                    for m in (st.get('measures') or [])
                    if isinstance(m, dict) or isinstance(m, str)
                }
                fixed = []
                for m in tt.get('measures') or []:
                    if isinstance(m, dict) and not m.get('expression'):
                        sm = by_name.get(m.get('name') or '')
                        if isinstance(sm, dict) and sm.get('expression'):
                            mm = dict(m)
                            mm['expression'] = sm.get('expression')
                            fixed.append(mm)
                            continue
                    fixed.append(m)
                tt['measures'] = fixed
        merged.append(tt)

    have = {t.get('name') for t in merged}
    for name, st in scan_tables.items():
        if name and name not in have:
            merged.append(st)
    out['tables'] = merged
    if not out.get('relationshipCount'):
        out['relationshipCount'] = len(out.get('relationships') or [])
    return out


def _normalize_semantic_tables(tables):
    """Return UI-friendly table objects with explicit column lists."""
    out = []
    for table in tables or []:
        cols_in = table.get('columns') or []
        columns = []
        for c in cols_in:
            if isinstance(c, str):
                columns.append({'name': c, 'dataType': '', 'isHidden': False})
            elif isinstance(c, dict):
                columns.append({
                    'name': c.get('name') or c.get('columnName') or '',
                    'dataType': c.get('dataType') or c.get('dataTypeName') or c.get('type') or '',
                    'isHidden': bool(c.get('isHidden') or c.get('isHiddenColumn')),
                    'usedInReport': c.get('usedInReport'),
                })
        measures_in = table.get('measures') or []
        measures = []
        for m in measures_in:
            if isinstance(m, str):
                measures.append({'name': m, 'expression': ''})
            elif isinstance(m, dict):
                measures.append({
                    'name': m.get('name') or '',
                    'expression': m.get('expression') or '',
                })
        out.append({
            'name': table.get('name') or 'Unknown',
            'isHidden': bool(table.get('isHidden')),
            'columnCount': table.get('columnCount') if table.get('columnCount') is not None else len(columns),
            'measureCount': table.get('measureCount') if table.get('measureCount') is not None else len(measures),
            'columns': columns,
            'measures': measures,
            'sourceTypeLabel': table.get('sourceTypeLabel'),
            'serverName': table.get('serverName'),
            'sqlSourceTables': table.get('sqlSourceTables') or [],
            'sqlQuery': table.get('sqlQuery') or '',
            'fileName': table.get('fileName') or '',
            'sourceUrl': table.get('sourceUrl') or '',
            'sourceExpression': table.get('sourceExpression') or '',
        })
    return out


def _rel_field(rel, *keys):
    """Read a relationship endpoint from common Scanner / catalog key shapes."""
    if not isinstance(rel, dict):
        return ''
    for k in keys:
        v = rel.get(k)
        if v is None or v == '':
            continue
        if isinstance(v, dict):
            # nested { table, column } / { name }
            return (
                v.get('table')
                or v.get('tableName')
                or v.get('column')
                or v.get('columnName')
                or v.get('name')
                or ''
            )
        return v
    return ''


def _extract_measures_and_relationships(dataset):
    tables = dataset.get('tables') or []
    all_measures = []
    for table in tables:
        tname = table.get('name') or 'Unknown'
        for measure in table.get('measures') or []:
            if isinstance(measure, str):
                all_measures.append({'name': measure, 'table': tname, 'expression': ''})
            elif isinstance(measure, dict):
                all_measures.append({
                    'name': measure.get('name'),
                    'table': tname,
                    'expression': measure.get('expression') or '',
                })

    # Some extracts put measures only under dataset.expressions
    if not all_measures:
        for expr in dataset.get('expressions') or []:
            if not isinstance(expr, dict):
                continue
            name = expr.get('name') or ''
            body = expr.get('expression') or expr.get('query') or ''
            if name and body:
                all_measures.append({
                    'name': name,
                    'table': expr.get('table') or 'Expression',
                    'expression': body if isinstance(body, str) else str(body),
                })

    raw_rels = (
        dataset.get('relationships')
        or dataset.get('modelRelationships')
        or dataset.get('datasetRelationships')
        or []
    )
    relationships = []
    for rel in raw_rels:
        if not isinstance(rel, dict):
            continue
        from_table = _rel_field(
            rel, 'fromTable', 'sourceTable', 'fromTableName', 'from', 'FromTable',
        )
        from_col = _rel_field(
            rel, 'fromColumn', 'sourceColumn', 'fromColumnName', 'FromColumn',
        )
        to_table = _rel_field(
            rel, 'toTable', 'targetTable', 'toTableName', 'to', 'ToTable',
        )
        to_col = _rel_field(
            rel, 'toColumn', 'targetColumn', 'toColumnName', 'ToColumn',
        )
        # nested from/to objects: { from: { table, column }, to: {...} }
        if isinstance(rel.get('from'), dict) and not from_table:
            fr = rel['from']
            from_table = fr.get('table') or fr.get('tableName') or ''
            from_col = from_col or fr.get('column') or fr.get('columnName') or ''
        if isinstance(rel.get('to'), dict) and not to_table:
            to = rel['to']
            to_table = to.get('table') or to.get('tableName') or ''
            to_col = to_col or to.get('column') or to.get('columnName') or ''
        if not (from_table or to_table or from_col or to_col):
            continue
        relationships.append({
            'fromTable': from_table or '',
            'fromColumn': from_col or '',
            'toTable': to_table or '',
            'toColumn': to_col or '',
            'cardinality': (
                rel.get('cardinality')
                or rel.get('Cardinality')
                or rel.get('crossFilteringBehavior')
                or rel.get('relationshipType')
                or ''
            ),
            'isActive': rel.get('isActive', rel.get('IsActive', True)),
            'crossFilteringBehavior': rel.get('crossFilteringBehavior') or '',
        })
    return all_measures, relationships


@app.route('/api/semantic-model-details')
@login_required
def get_semantic_model_details():
    """
    Detailed semantic model schema for the Details modal.
    Prefer catalog datasets map (has columns) — Scanner only if missing.
    """
    try:
        workspace_id = request.args.get('workspace_id')
        dataset_id = request.args.get('dataset_id')

        if not workspace_id or not dataset_id:
            return jsonify({
                'success': False,
                'error': 'workspace_id and dataset_id are required'
            }), 400

        dataset = None
        source = None

        # 1) Catalog top-level datasets map (full schema when extract included it)
        if CATALOG_AVAILABLE and catalog_service is not None and catalog_service.is_available():
            try:
                cat = catalog_service.get_workspace_catalog()
                dmap = (cat or {}).get('datasets') or {}
                if isinstance(dmap, dict) and dataset_id in dmap:
                    dataset = dict(dmap.get(dataset_id) or {})
                    dataset.setdefault('id', dataset_id)
                    source = 'catalog'
                # Also try workspace.datasets if richer
                if dataset is not None and not (dataset.get('tables') or []):
                    ws = next(
                        (w for w in (cat.get('workspaces') or []) if w.get('id') == workspace_id),
                        None,
                    )
                    if ws:
                        for ds in ws.get('datasets') or []:
                            if ds.get('id') == dataset_id and (ds.get('tables') or []):
                                dataset = dict(ds)
                                source = 'catalog-ws'
                                break
            except Exception as cat_err:
                print(f"⚠️ semantic-model-details catalog miss: {cat_err}")

        # 2) Fill missing measures / relationships / columns via Scanner (cached per workspace).
        # IMPORTANT: having measures does NOT mean relationships are present.
        # Old logic treated (has measures) as complete and skipped Scanner → Relationships (0).
        def _schema_incomplete(ds_obj):
            if not ds_obj:
                return True
            tables = ds_obj.get('tables') or []
            if not tables:
                return True
            has_cols = any((t.get('columns') or t.get('columnCount')) for t in tables)
            has_meas = (
                any((t.get('measures') or t.get('measureCount')) for t in tables)
                or bool(ds_obj.get('measureCount'))
                or bool(ds_obj.get('expressions'))
            )
            rels = (
                ds_obj.get('relationships')
                or ds_obj.get('modelRelationships')
                or []
            )
            rel_n = len(rels) if isinstance(rels, list) else 0
            if not rel_n:
                try:
                    rel_n = int(ds_obj.get('relationshipCount') or 0)
                except Exception:
                    rel_n = 0
            has_rel = rel_n > 0
            # Multi-table models almost always have relationships in Desktop —
            # if catalog has 2+ tables and 0 rels, force Scanner enrich.
            multi_table = len(tables) >= 2
            missing_rels = (not has_rel) and multi_table
            return (not has_cols) or (not has_meas) or missing_rels

        # opt-out: enrich=0 skips Scanner (catalog-only)
        allow_enrich = str(request.args.get('enrich', '1')).lower() not in ('0', 'false', 'no')
        need_scan = _schema_incomplete(dataset)

        if allow_enrich and need_scan:
            try:
                print(
                    f"🔎 semantic-model-details enriching via Scanner "
                    f"(ws={workspace_id[:8]}… ds={dataset_id[:8]}… "
                    f"had_tables={len((dataset or {}).get('tables') or [])} "
                    f"had_rels={len((dataset or {}).get('relationships') or [])})"
                )
                scan_ds = _get_scanned_dataset_cached(workspace_id, dataset_id)
                if scan_ds:
                    scan_rel_n = len(scan_ds.get('relationships') or [])
                    print(f"   scanner dataset keys sample rels={scan_rel_n}")
                    if dataset is None or not (dataset.get('tables') or []):
                        dataset = scan_ds
                        source = 'scanner'
                    else:
                        dataset = _merge_dataset_schema(dataset, scan_ds)
                        source = 'catalog+scanner'
            except Exception as scan_err:
                print(f"⚠️ semantic-model-details scanner enrich failed: {scan_err}")
                if dataset is None:
                    return jsonify({'success': False, 'error': f'Failed to load model schema: {scan_err}'}), 500

        if not dataset:
            return jsonify({
                'success': False,
                'error': f'Dataset {dataset_id} not found in workspace'
            }), 404

        raw_tables = dataset.get('tables') or []
        tables = _normalize_semantic_tables(raw_tables)
        all_measures, relationships = _extract_measures_and_relationships(dataset)

        # If still no relationships after merge, try one more direct read of scan keys
        if not relationships and allow_enrich:
            try:
                scan_ds = _get_scanned_dataset_cached(workspace_id, dataset_id)
                if scan_ds and (scan_ds.get('relationships') or []):
                    _, relationships = _extract_measures_and_relationships(scan_ds)
                    if relationships and source and 'scanner' not in str(source):
                        source = f'{source}+rels'
            except Exception:
                pass

        owner = dataset.get('configuredBy') or dataset.get('configuredByUser') or 'Unknown'
        created_date = dataset.get('createdDate') or dataset.get('createdDateTime') or 'Unknown'

        # Attach ops refresh if present
        last_refresh = dataset.get('last_refreshed') or dataset.get('lastRefresh')
        refresh_status = dataset.get('last_refresh_status') or dataset.get('refreshStatus')
        if CATALOG_AVAILABLE and catalog_service is not None and (not last_refresh or not refresh_status):
            try:
                snap = catalog_service.get_json('refresh_snapshot.json') or {}
                info = (snap.get('datasets') or {}).get(dataset_id) or {}
                last_refresh = last_refresh or info.get('last_refreshed')
                refresh_status = refresh_status or info.get('last_refresh_status')
            except Exception:
                pass

        print(
            f"✅ semantic-model-details source={source} tables={len(tables)} "
            f"cols={sum(len(t.get('columns') or []) for t in tables)} "
            f"measures={len(all_measures)} rels={len(relationships)}"
        )

        return jsonify({
            'success': True,
            'source': source or 'unknown',
            'tables': tables,
            'relationships': relationships,
            'measures': all_measures,
            'owner': owner,
            'configuredBy': owner,
            'modifiedBy': dataset.get('modifiedBy') or 'Not available for datasets',
            'createdDate': created_date,
            'modifiedDate': dataset.get('modifiedDateTime') or dataset.get('modifiedDate') or 'Not available for datasets',
            'name': dataset.get('name', 'Unknown'),
            'lastRefresh': last_refresh,
            'refreshStatus': refresh_status,
            'tableCount': len(tables),
            'columnCount': sum(len(t.get('columns') or []) for t in tables),
            'measureCount': len(all_measures),
            'relationshipCount': len(relationships),
        })

    except Exception as e:
        print(f"❌ Error getting semantic model details: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/semantic-model-health-check', methods=['POST'])
@login_required
def semantic_model_health_check():
    """Run health check on a semantic model to detect issues"""
    try:
        data = request.get_json()
        workspace_id = data.get('workspace_id')
        dataset_id = data.get('dataset_id')

        if not workspace_id or not dataset_id:
            return jsonify({
                'success': False,
                'error': 'workspace_id and dataset_id are required'
            }), 400

        from scanner_connector import PowerBIScanner

        # Get scanner instance
        scanner = PowerBIScanner()
        scan_result = scanner.run_scan(workspace_id=workspace_id)

        if not scan_result:
            return jsonify({
                'success': False,
                'error': 'Failed to scan workspace'
            }), 500

        # Extract datasets
        workspaces = scan_result.get('workspaces', [])
        if not workspaces:
            return jsonify({'success': False, 'error': 'No workspace data'}), 404

        datasets = workspaces[0].get('datasets', [])
        dataset = None
        for ds in datasets:
            if ds.get('id') == dataset_id:
                dataset = ds
                break

        if not dataset:
            return jsonify({'success': False, 'error': 'Dataset not found'}), 404

        # Perform health checks
        issues = []
        warnings = []
        health_score = 100

        tables = dataset.get('tables', [])
        relationships = dataset.get('relationships', [])

        # Check 1: Empty tables (no columns)
        for table in tables:
            if not table.get('columns') or len(table.get('columns', [])) == 0:
                issues.append({
                    'severity': 'warning',
                    'message': f"Table '{table.get('name')}' has no columns"
                })
                health_score -= 5

        # Check 2: Tables with no relationships
        table_names = [t.get('name') for t in tables]
        tables_in_relationships = set()
        for rel in relationships:
            tables_in_relationships.add(rel.get('fromTable'))
            tables_in_relationships.add(rel.get('toTable'))

        orphaned_tables = [t for t in table_names if t not in tables_in_relationships]
        if len(orphaned_tables) > 0 and len(tables) > 1:
            warnings.append(f"{len(orphaned_tables)} table(s) have no relationships: {', '.join(orphaned_tables[:3])}")
            health_score -= 10

        # Check 3: Circular relationships (simplified check)
        if len(relationships) > len(tables):
            warnings.append("Model has many relationships - check for potential circular dependencies")

        # Check 4: Tables with many columns (performance concern)
        for table in tables:
            column_count = len(table.get('columns', []))
            if column_count > 100:
                warnings.append(f"Table '{table.get('name')}' has {column_count} columns (performance concern)")
                health_score -= 5

        # Check 5: Measures without expressions
        for table in tables:
            for measure in table.get('measures', []):
                if not measure.get('expression'):
                    issues.append({
                        'severity': 'danger',
                        'message': f"Measure '{measure.get('name')}' in table '{table.get('name')}' has no expression"
                    })
                    health_score -= 10

        # Ensure score doesn't go below 0
        health_score = max(0, health_score)

        return jsonify({
            'success': True,
            'health_score': health_score,
            'issues': issues,
            'warnings': warnings,
            'summary': {
                'total_tables': len(tables),
                'total_relationships': len(relationships),
                'total_measures': sum(len(t.get('measures', [])) for t in tables),
                'orphaned_tables': len(orphaned_tables)
            }
        })

    except Exception as e:
        print(f"❌ Error running health check: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/semantic-model-documentation')
@login_required
def semantic_model_documentation():
    """Generate and download documentation for a semantic model"""
    try:
        workspace_id = request.args.get('workspace_id')
        dataset_id = request.args.get('dataset_id')

        if not workspace_id or not dataset_id:
            return jsonify({
                'success': False,
                'error': 'workspace_id and dataset_id are required'
            }), 400

        from scanner_connector import PowerBIScanner
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io

        # Get scanner instance
        scanner = PowerBIScanner()
        scan_result = scanner.run_scan(workspace_id=workspace_id)

        if not scan_result:
            return jsonify({'success': False, 'error': 'Failed to scan'}), 500

        # Extract dataset
        workspaces = scan_result.get('workspaces', [])
        if not workspaces:
            return jsonify({'success': False, 'error': 'No workspace data'}), 404

        datasets = workspaces[0].get('datasets', [])
        dataset = None
        for ds in datasets:
            if ds.get('id') == dataset_id:
                dataset = ds
                break

        if not dataset:
            return jsonify({'success': False, 'error': 'Dataset not found'}), 404

        # Create Word document
        doc = Document()

        # Title
        title = doc.add_heading('Semantic Model Documentation', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Model name
        dataset_name = dataset.get('name', 'Unknown Model')
        doc.add_heading(f"Model: {dataset_name}", 1)

        # Summary
        doc.add_heading('Summary', 2)
        tables = dataset.get('tables', [])
        relationships = dataset.get('relationships', [])

        doc.add_paragraph(f"Total Tables: {len(tables)}")
        doc.add_paragraph(f"Total Relationships: {len(relationships)}")

        total_measures = sum(len(t.get('measures', [])) for t in tables)
        doc.add_paragraph(f"Total Measures: {total_measures}")

        # Tables section
        doc.add_heading('Tables', 2)

        for table in tables:
            doc.add_heading(table.get('name', 'Unknown'), 3)

            # Columns
            columns = table.get('columns', [])
            if columns:
                doc.add_paragraph('Columns:', style='Heading 4')
                for col in columns:
                    col_text = f"  • {col.get('name', 'Unknown')} ({col.get('dataType', 'Unknown')})"
                    doc.add_paragraph(col_text)

            # Measures
            measures = table.get('measures', [])
            if measures:
                doc.add_paragraph('Measures:', style='Heading 4')
                for measure in measures:
                    doc.add_paragraph(f"  • {measure.get('name', 'Unknown')}")
                    if measure.get('expression'):
                        doc.add_paragraph(f"    Expression: {measure['expression'][:200]}")

        # Relationships section
        if relationships:
            doc.add_heading('Relationships', 2)

            for rel in relationships:
                rel_text = f"{rel.get('fromTable')}.{rel.get('fromColumn')} → {rel.get('toTable')}.{rel.get('toColumn')}"
                doc.add_paragraph(f"  • {rel_text}")
                if rel.get('crossFilteringBehavior'):
                    doc.add_paragraph(f"    Filtering: {rel['crossFilteringBehavior']}")

        # Save to BytesIO
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        return send_file(
            file_stream,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'{dataset_name}_SemanticModel_Documentation.docx'
        )

    except Exception as e:
        print(f"❌ Error generating documentation: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


# Helper function to get all child folder IDs recursively
def get_all_child_folders_recursive(folder_id, folder_map):
    """Get all nested child folder IDs recursively

    Args:
        folder_id: The parent folder ID to start from
        folder_map: Dict mapping folder_id -> folder_info (with 'parentFolderId' key)

    Returns:
        Set of all child folder IDs (including the parent folder itself)
    """
    child_ids = {folder_id}
    for fid, finfo in folder_map.items():
        parent_id = finfo.get('parentFolderId') if isinstance(finfo, dict) else finfo
        if parent_id == folder_id:
            child_ids.add(fid)
            child_ids.update(get_all_child_folders_recursive(fid, folder_map))
    return child_ids


def _build_folder_display_path(folder_id, folder_map):
    """Hierarchical path like 'Parent / Child' for dropdown labels."""
    if not folder_id or folder_id not in folder_map:
        return None
    path_parts = []
    current_id = folder_id
    visited = set()
    while current_id and current_id not in visited:
        visited.add(current_id)
        info = folder_map.get(current_id) or {}
        name = info.get('name') or info.get('displayName') or ''
        if name:
            path_parts.insert(0, name)
        current_id = info.get('parentFolderId')
    return ' / '.join(path_parts) if path_parts else None


def _fetch_workspace_folder_meta(workspace_id, timeout=12):
    """
    Fetch workspace folder tree + report→folder membership for Report Catalog.

    Returns:
      {
        'folders': [{id, name, parentFolderId, report_count, hasChildren}, ...],
        'report_folder_map': {reportId: folderId|None},  # None = root
        'folder_names_map': {folderId: {id, name, parentFolderId}},
      }

    Folder names come from Fabric Folders API.
    Membership prefers Fabric Items (report/parentFolderId) — cheap vs Admin Scanner.
    Cached per workspace for WORKSPACE_FOLDERS_CACHE_DURATION seconds.
    """
    import time as _time
    import requests as _requests

    empty = {'folders': [], 'report_folder_map': {}, 'folder_names_map': {}}
    if not workspace_id:
        return empty

    cached = workspace_folders_cache.get(workspace_id)
    if cached and (_time.time() - float(cached.get('timestamp') or 0)) < WORKSPACE_FOLDERS_CACHE_DURATION:
        return {
            'folders': list(cached.get('folders') or []),
            'report_folder_map': dict(cached.get('report_folder_map') or {}),
            'folder_names_map': dict(cached.get('folder_names_map') or {}),
        }

    folder_names_map = {}
    report_folder_map = {}  # reportId -> folderId (omit root / None not stored; treated as root)

    def _ingest_folder_rows(rows):
        for folder in rows or []:
            fid = folder.get('id')
            if not fid:
                continue
            fname = (
                folder.get('displayName')
                or folder.get('name')
                or f'Folder {str(fid)[:8]}'
            )
            folder_names_map[fid] = {
                'id': fid,
                'name': fname,
                'parentFolderId': folder.get('parentFolderId'),
            }

    def _auth_headers(token):
        return {'Authorization': f'Bearer {token}', 'Content-Type': 'application/json'}

    def _try_fabric_folders(token, label):
        if not token:
            return False
        try:
            url = f"https://api.fabric.microsoft.com/v1/workspaces/{workspace_id}/folders?recursive=true"
            resp = _requests.get(url, headers=_auth_headers(token), timeout=timeout)
            if resp.status_code == 200:
                _ingest_folder_rows(resp.json().get('value', []))
                print(f"   📁 Folders via Fabric ({label}): {len(folder_names_map)}")
                return True
            print(f"   ⚠️ Fabric folders ({label}) HTTP {resp.status_code}")
        except Exception as ex:
            print(f"   ⚠️ Fabric folders ({label}) error: {ex}")
        return False

    def _paginate_fabric_items(token, params_base, label):
        """Yield items from Fabric workspace items API with continuation."""
        url = f"https://api.fabric.microsoft.com/v1/workspaces/{workspace_id}/items"
        headers = _auth_headers(token)
        continuation = None
        pages = 0
        while pages < 30:
            pages += 1
            params = dict(params_base or {})
            if continuation:
                params['continuationToken'] = continuation
            resp = _requests.get(url, headers=headers, params=params, timeout=timeout)
            if resp.status_code != 200:
                print(f"   ⚠️ Fabric items ({label}) HTTP {resp.status_code} params={params_base}")
                return
            body = resp.json() or {}
            for item in body.get('value') or []:
                yield item
            continuation = body.get('continuationToken')
            if not continuation:
                break

    def _is_report_item(item):
        itype = (item.get('type') or '').replace(' ', '').lower()
        return itype in (
            'report',
            'paginatedreport',
            'pagedreport',
            'powerbireport',
            'rdl',
        )

    def _item_folder_id(item):
        return (
            item.get('folderId')
            or item.get('parentFolderId')
            or item.get('folderObjectId')
            or None
        )

    def _try_fabric_items_membership(token, label):
        """
        ONE workspace items list (paginated) → reportId → folderId.
        Avoids N+1 per-folder calls that made workspace select slow.
        """
        nonlocal report_folder_map
        if not token:
            return False
        try:
            rootish = 0
            in_folder = 0
            for item in _paginate_fabric_items(token, {'recursive': 'true'}, f"{label}/items"):
                if not _is_report_item(item):
                    continue
                rid = item.get('id')
                if not rid:
                    continue
                rid_s = str(rid)
                parent = _item_folder_id(item)
                if parent:
                    report_folder_map[rid_s] = str(parent)
                    in_folder += 1
                else:
                    report_folder_map[rid_s] = None
                    rootish += 1

            print(
                f"   📁 Report→folder map ({label}): "
                f"total={len(report_folder_map)} in_folder={in_folder} root={rootish}"
            )
            return len(report_folder_map) > 0
        except Exception as ex:
            print(f"   ⚠️ Fabric items membership ({label}) error: {ex}")
        return False

    # Tokens: try Fabric user token FIRST only. Do not fan-out PBI+SP+Scanner
    # unless Fabric fails — that was forcing new tokens every workspace select.
    tok = None
    label = 'user-fabric'
    try:
        tok = get_user_fabric_token()
    except Exception as ex:
        print(f"   ⚠️ get_user_fabric_token failed: {ex}")

    if tok:
        _try_fabric_folders(tok, label)
        _try_fabric_items_membership(tok, label)

    # Fallback 1: user Power BI token (sometimes accepted on Fabric endpoints)
    if not folder_names_map or not report_folder_map:
        try:
            pbi_tok = get_user_powerbi_token()
            if pbi_tok:
                if not folder_names_map:
                    _try_fabric_folders(pbi_tok, 'user-pbi')
                if not report_folder_map:
                    _try_fabric_items_membership(pbi_tok, 'user-pbi')
        except Exception as ex:
            print(f"   ⚠️ user-pbi folder fallback failed: {ex}")

    # Fallback 2: service principal — only if still empty (slow token + possible scan)
    if not folder_names_map or not report_folder_map:
        try:
            from scanner_connector import PowerBIScanner
            sp = PowerBIScanner()
            sp_token = None
            if hasattr(sp, 'get_access_token'):
                try:
                    try:
                        sp_token = sp.get_access_token(
                            scope='https://api.fabric.microsoft.com/.default'
                        )
                    except TypeError:
                        sp_token = sp.get_access_token()
                except Exception:
                    sp_token = None
            if sp_token:
                if not folder_names_map:
                    _try_fabric_folders(sp_token, 'service-principal')
                if not report_folder_map:
                    _try_fabric_items_membership(sp_token, 'service-principal')
            # Admin Scanner only as last resort (expensive)
            if not folder_names_map or not report_folder_map:
                print("   ⚠️ Folder meta incomplete — Admin Scanner last resort…")
                scanner = PowerBIScanner()
                scanner.access_token = sp_token or scanner.get_access_token()
                try:
                    scan_data = scanner.run_scan(
                        workspace_id=workspace_id,
                        dataset_schema=False,
                        dataset_expressions=False,
                        lineage=False,
                    )
                except TypeError:
                    scan_data = scanner.run_scan(workspace_id=workspace_id)
                if scan_data and scan_data.get('workspaces'):
                    for ws in scan_data['workspaces']:
                        if ws.get('id') != workspace_id:
                            continue
                        if not folder_names_map:
                            _ingest_folder_rows(ws.get('folders') or [])
                        for rep in ws.get('reports') or []:
                            rid = rep.get('id')
                            if not rid:
                                continue
                            fid = rep.get('folderObjectId') or rep.get('folderId')
                            if str(rid) not in report_folder_map:
                                report_folder_map[str(rid)] = str(fid) if fid else None
                            if fid and fid not in folder_names_map:
                                folder_names_map[fid] = {
                                    'id': fid,
                                    'name': rep.get('folderName') or f'Folder {str(fid)[:8]}',
                                    'parentFolderId': None,
                                }
                        print(
                            f"   📁 Scanner fallback folders={len(folder_names_map)} "
                            f"membership={len(report_folder_map)}"
                        )
                        break
        except Exception as ex:
            print(f"   ⚠️ SP/Scanner folder fallback failed: {ex}")

    # report_count per folder from membership
    counts = {}
    root_count = 0
    for _rid, fid in report_folder_map.items():
        if fid:
            counts[fid] = counts.get(fid, 0) + 1
        else:
            root_count += 1

    folders_list = []
    for fid, finfo in folder_names_map.items():
        leaf = finfo.get('name') or fid
        parent_id = finfo.get('parentFolderId')
        folders_list.append({
            'id': fid,
            'name': leaf,
            'parentFolderId': parent_id,
            'report_count': int(counts.get(fid) or 0),
            'hasChildren': False,
        })

    by_id = {f['id']: f for f in folders_list}
    for f in folders_list:
        pid = f.get('parentFolderId')
        if pid and pid in by_id:
            by_id[pid]['hasChildren'] = True

    # Virtual root entry when there are uncategorized reports
    if root_count > 0:
        folders_list.append({
            'id': '__ROOT__',
            'name': 'Root Directory (Uncategorized)',
            'parentFolderId': None,
            'report_count': root_count,
            'hasChildren': False,
        })

    folders_list = sorted(folders_list, key=lambda f: (f.get('name') or '').lower())

    workspace_folders_cache[workspace_id] = {
        'folders': folders_list,
        'report_folder_map': report_folder_map,
        'folder_names_map': folder_names_map,
        'timestamp': _time.time(),
    }
    print(
        f"   ✅ Workspace folder meta ready: folders={len(folders_list)} "
        f"mapped_reports={len(report_folder_map)} for {workspace_id[:8]}…"
    )
    return {
        'folders': list(folders_list),
        'report_folder_map': dict(report_folder_map),
        'folder_names_map': dict(folder_names_map),
    }


def _fetch_workspace_folders_list(workspace_id, timeout=12):
    """Backward-compatible: folder dropdown rows only."""
    return _fetch_workspace_folder_meta(workspace_id, timeout=timeout).get('folders') or []


def _filter_reports_by_folder(report_list, folder_id, folder_meta):
    """
    Filter catalog/live report shells by folder using membership map + hierarchy.
    folder_id='__ROOT__' → reports with no folder.
    """
    if not folder_id:
        return report_list

    folder_names_map = (folder_meta or {}).get('folder_names_map') or {}
    report_folder_map = (folder_meta or {}).get('report_folder_map') or {}

    def _report_folder(rep):
        rid = str(rep.get('id') or '')
        if rid and rid in report_folder_map:
            return report_folder_map.get(rid)
        return rep.get('folderId') or rep.get('folderObjectId') or None

    if folder_id == '__ROOT__':
        return [r for r in report_list if not _report_folder(r)]

    # Selected folder + nested children
    if folder_names_map:
        allowed = get_all_child_folders_recursive(folder_id, folder_names_map)
    else:
        allowed = {folder_id}

    out = []
    for r in report_list:
        rf = _report_folder(r)
        if rf and rf in allowed:
            out.append(r)
    return out


def _attach_folder_ids(report_list, report_folder_map):
    """Stamp folderId onto report shells for UI client-side filtering."""
    if not report_folder_map:
        return report_list
    out = []
    for r in report_list or []:
        if not isinstance(r, dict):
            out.append(r)
            continue
        rr = dict(r)
        rid = str(rr.get('id') or '')
        if rid and rid in report_folder_map:
            fid = report_folder_map.get(rid)
            rr['folderId'] = fid
            rr['folderObjectId'] = fid
        out.append(rr)
    return out


def _catalog_reports_shell(catalog_reports):
    """
    Instant catalog rows. Prefer precomputed ops fields (refresh + views)
    when present in SharePoint catalog; otherwise leave null/pending.
    """
    report_list = []
    for r in catalog_reports or []:
        has_ops = bool(
            r.get('ops_from_catalog')
            or r.get('last_refreshed')
            or r.get('last_refresh_status')
            or r.get('refresh_schedule')
            or r.get('refresh_type')
            or r.get('view_count') is not None
        )
        # Optional owner fields from catalog (after extract preserves them).
        # UI still runs /api/reports-metadata for live fill; seeding avoids N/A flash.
        # Seed owner fields when catalog has them, but never mark metadata_loaded here.
        # UI waits for ONE batch /api/reports-metadata so ALL rows stop spinning together.
        created_by = _pick_person(r.get('createdBy'), r.get('created_by'))
        modified_by = _pick_person(r.get('modifiedBy'), r.get('modified_by'))
        created_dt = _pick_datetime(
            r.get('createdDateTime'), r.get('created_date_time'), r.get('created_date')
        )
        modified_dt = _pick_datetime(
            r.get('modifiedDateTime'), r.get('modified_date_time'), r.get('modified_date')
        )

        last_ref = r.get('last_refreshed')
        last_status = r.get('last_refresh_status')
        refresh_note = r.get('refresh_note')
        refresh_source = r.get('refresh_source')
        days = r.get('days_since_refresh')
        # Fast paint: if ops has no scheduled timestamp, fall back to report modified
        # (OneDrive-tab history is not in REST; modified is the closest public signal)
        if not last_ref and modified_dt:
            try:
                from powerbi_connector import (
                    refresh_info_from_content_modified as _cm,
                    days_since_refresh as _dsr,
                )
                cm = _cm(None, {'modifiedDateTime': modified_dt}, dataset_workspace_id=None)
                if cm and cm.get('last_refreshed'):
                    last_ref = cm.get('last_refreshed')
                    last_status = cm.get('last_refresh_status') or last_status or 'Completed'
                    refresh_note = cm.get('refresh_note') or refresh_note
                    refresh_source = cm.get('refresh_source') or 'content_modified'
                    days = _dsr(last_ref)
            except Exception:
                pass
        if days is None and last_ref:
            try:
                from powerbi_connector import days_since_refresh as _dsr
                days = _dsr(last_ref)
            except Exception:
                days = None

        report_list.append({
            'id': r.get('id'),
            'name': r.get('name'),
            'datasetId': r.get('datasetId') or '',
            'folderId': r.get('folderId'),
            'folderName': r.get('folderName'),
            'refresh_schedule': r.get('refresh_schedule'),
            'last_refreshed': last_ref,
            'last_refresh_status': last_status,
            'refresh_type': r.get('refresh_type'),
            'refresh_note': refresh_note,
            'refresh_source': refresh_source,
            'days_since_refresh': days,
            'dataset_workspace_id': r.get('dataset_workspace_id'),
            'dataset_owner': r.get('dataset_owner'),
            'view_count': r.get('view_count'),
            'last_viewed': r.get('last_viewed'),
            'usage_loaded': r.get('view_count') is not None,
            'created_by': created_by or None,
            'modified_by': modified_by or None,
            'created_date_time': created_dt or None,
            'modified_date_time': modified_dt or None,
            'metadata_loaded': False,
            'last_accessed': None,
            'last_accessed_by': None,
            'last_modified': None,
            'last_generated': None,
            'has_documentation': False,
            'needs_update': False,
            'webUrl': r.get('webUrl') or '',
            'from_catalog': True,
            'ops_from_catalog': has_ops,
            'datasetName': r.get('datasetName'),
            'reportType': r.get('reportType'),
            # Only pending when ops snapshot not yet in catalog
            'refresh_pending': not has_ops,
        })
    return report_list


def _enrich_catalog_reports_with_refresh(catalog_reports, workspace_id):
    """
    Take catalog report shells and overlay live refresh data (robust resolver).
    Structure from catalog (fast); refresh/status still live (accurate).
    """
    from concurrent.futures import ThreadPoolExecutor, as_completed
    from powerbi_connector import (
        resolve_dataset_refresh_info,
        _empty_refresh_info,
        merge_refresh_candidates,
        days_since_refresh,
        refresh_info_from_content_modified,
    )

    report_list = []
    dataset_ids = set()
    for r in catalog_reports:
        ds = r.get('datasetId')
        if ds:
            dataset_ids.add(ds)

    dataset_refresh_map = {}
    if dataset_ids:
        try:
            headers = get_user_powerbi_headers()
        except Exception:
            headers = {}
        base_url = "https://api.powerbi.com/v1.0/myorg"

        def _one(ds_id):
            try:
                return ds_id, resolve_dataset_refresh_info(
                    headers=headers,
                    workspace_id=workspace_id,
                    dataset_id=ds_id,
                    dataset_workspace_id=workspace_id,
                    history_top=3,   # enough for endTime fallback; keeps bulk path lighter
                    base_url=base_url,
                    timeout=5,      # fail fast under load; UI already painted
                )
            except Exception as exc:
                return ds_id, _empty_refresh_info(
                    refresh_schedule='Error',
                    last_refresh_status='Error',
                    refresh_type='error',
                    refresh_note=str(exc),
                    dataset_workspace_id=workspace_id,
                )

        # Cap concurrency to reduce Power BI throttling / cascading timeouts
        workers = min(6, max(2, len(dataset_ids)))
        print(f"   ⚡ Catalog path: live refresh for {len(dataset_ids)} dataset(s) (workers={workers})…")
        with ThreadPoolExecutor(max_workers=workers) as ex:
            futs = {ex.submit(_one, ds): ds for ds in dataset_ids}
            for fut in as_completed(futs):
                ds = futs[fut]
                try:
                    rid, info = fut.result(timeout=60)
                    dataset_refresh_map[rid] = info
                except Exception as exc:
                    dataset_refresh_map[ds] = _empty_refresh_info(
                        refresh_schedule='Error',
                        last_refresh_status='Error',
                        refresh_type='error',
                        refresh_note=str(exc),
                    )

    for r in catalog_reports:
        ds_id = r.get('datasetId') or ''
        live = dataset_refresh_map.get(ds_id, {}) or {}
        # Catalog / SharePoint ops snapshot (scheduled + admin + prior fallbacks)
        catalog_ops = {
            'last_refreshed': r.get('last_refreshed'),
            'last_refresh_status': r.get('last_refresh_status'),
            'refresh_schedule': r.get('refresh_schedule'),
            'refresh_type': r.get('refresh_type'),
            'refresh_note': r.get('refresh_note'),
            'refresh_source': r.get('refresh_source'),
            'history_refresh_type': r.get('history_refresh_type'),
            'dataset_owner': r.get('dataset_owner'),
            'dataset_workspace_id': r.get('dataset_workspace_id') or workspace_id,
        }
        created_by = _pick_person(r.get('createdBy'), r.get('created_by'))
        modified_by = _pick_person(r.get('modifiedBy'), r.get('modified_by'))
        created_dt = _pick_datetime(
            r.get('createdDateTime'), r.get('created_date_time'), r.get('created_date')
        )
        modified_dt = _pick_datetime(
            r.get('modifiedDateTime'), r.get('modified_date_time'), r.get('modified_date')
        )
        # Content-modified fallback when scheduled + catalog ops both lack a timestamp
        # (common for OneDrive-only models — portal OneDrive tab is not in REST APIs)
        content_side = None
        if not live.get('last_refreshed') and not catalog_ops.get('last_refreshed') and modified_dt:
            content_side = refresh_info_from_content_modified(
                None,
                {
                    'modifiedDateTime': modified_dt,
                    'name': r.get('name'),
                },
                dataset_workspace_id=workspace_id,
            )
        # Latest timestamp wins across live scheduled, catalog ops, content-modified
        info = merge_refresh_candidates(live, catalog_ops, content_side or {})
        has_people = bool(created_by or modified_by or created_dt or modified_dt)
        days = info.get('days_since_refresh')
        if days is None:
            days = days_since_refresh(info.get('last_refreshed'))
        report_list.append({
            'id': r.get('id'),
            'name': r.get('name'),
            'datasetId': ds_id,
            'folderId': r.get('folderId'),
            'folderName': r.get('folderName'),
            'refresh_schedule': info.get('refresh_schedule'),
            'last_refreshed': info.get('last_refreshed'),
            'last_refresh_status': info.get('last_refresh_status'),
            'refresh_type': info.get('refresh_type'),
            'refresh_note': info.get('refresh_note'),
            'refresh_source': info.get('refresh_source'),
            'days_since_refresh': days,
            'dataset_workspace_id': info.get('dataset_workspace_id'),
            'dataset_owner': info.get('dataset_owner') or r.get('dataset_owner'),
            'view_count': r.get('view_count'),
            'last_viewed': r.get('last_viewed'),
            'created_by': created_by or None,
            'modified_by': modified_by or None,
            'created_date_time': created_dt or None,
            'modified_date_time': modified_dt or None,
            # Always false — UI batch-loads owner columns together
            'metadata_loaded': False,
            'last_accessed': None,
            'last_accessed_by': None,
            'last_modified': None,
            'last_generated': None,
            'has_documentation': False,
            'needs_update': False,
            'webUrl': r.get('webUrl') or '',
            'from_catalog': True,
            'datasetName': r.get('datasetName'),
            'reportType': r.get('reportType'),
        })
    return report_list


def _filter_reports_by_refresh_status(report_list, status_filter):
    """Active/inactive filter shared by catalog + live paths."""
    from datetime import datetime, timezone as tz

    def _is_live(rep):
        rtype = (rep.get('refresh_type') or '').lower()
        status = (rep.get('last_refresh_status') or '').lower()
        schedule = (rep.get('refresh_schedule') or '').lower()
        return (
            rtype in ('directquery', 'live')
            or 'directquery' in status or 'live' in status
            or 'directquery' in schedule or 'live' in schedule
        )

    filtered = []
    for report in report_list:
        last_refreshed = report.get('last_refreshed')
        is_live = _is_live(report)
        if status_filter == 'inactive':
            if is_live:
                continue
            if not last_refreshed:
                filtered.append(report)
            else:
                try:
                    s = last_refreshed
                    if s.endswith('Z'):
                        s = s[:-1] + '+00:00'
                    rt = datetime.fromisoformat(s)
                    if rt.tzinfo is None:
                        rt = rt.replace(tzinfo=tz.utc)
                    if (datetime.now(tz.utc) - rt).days >= 30:
                        filtered.append(report)
                except Exception:
                    filtered.append(report)
        elif status_filter == 'active':
            if is_live:
                filtered.append(report)
                continue
            if last_refreshed:
                try:
                    s = last_refreshed
                    if s.endswith('Z'):
                        s = s[:-1] + '+00:00'
                    rt = datetime.fromisoformat(s)
                    if rt.tzinfo is None:
                        rt = rt.replace(tzinfo=tz.utc)
                    if (datetime.now(tz.utc) - rt).days < 30:
                        filtered.append(report)
                except Exception:
                    pass
    return filtered


@app.route('/api/reports')
@login_required
def get_reports():
    """API endpoint to get reports for a workspace using user-delegated permissions - FAST initial load
    Supports optional folder_id parameter for folder-level scoping
    Supports optional status parameter for active/inactive filtering
    """
    try:
        import time
        import requests
        workspace_id = request.args.get('workspace_id')
        folder_id = request.args.get('folder_id')  # NEW: Optional folder filter
        status_filter = request.args.get('status')  # NEW: Optional status filter (active/inactive)

        print(f"\n🎯 get_reports() called:")
        print(f"   workspace_id = {workspace_id}")
        print(f"   folder_id (from URL) = {folder_id if folder_id else 'NULL (All Folders)'}")
        print(f"   Full request URL: {request.url}")

        if not workspace_id:
            return jsonify({
                'success': False,
                'error': 'workspace_id is required'
            }), 400

        # Get user info for cache key (include folder_id in cache key)
        user_id = session.get('user', {}).get('oid', 'unknown')
        cache_key = f"{workspace_id}_{user_id}_{folder_id or 'all'}"
        print(f"   Cache key: {cache_key}")

        # Check cache first (per-user, per-workspace, per-folder)
        current_time = time.time()
        force_live = request.args.get('source') in ('live', 'api') or request.args.get('force_live') in ('1', 'true', 'yes')
        if cache_key in reports_cache and not force_live:
            cache_entry = reports_cache[cache_key]
            if cache_entry['data'] and (current_time - cache_entry['timestamp']) < CACHE_DURATION:
                print(f"📦 Using cached reports (folder: {folder_id or 'all'})")
                return jsonify({
                    'success': True,
                    'reports': cache_entry['data'],
                    'folders': cache_entry.get('folders', []),
                    'reportFolderMap': cache_entry.get('reportFolderMap') or {},
                    'source': cache_entry.get('source', 'cache'),
                    'catalogGeneratedAt': cache_entry.get('catalogGeneratedAt'),
                    'opsFromCatalog': cache_entry.get('opsFromCatalog'),
                    'refreshIncluded': cache_entry.get('refreshIncluded'),
                    'opsEnrichedAt': cache_entry.get('opsEnrichedAt'),
                })

        # =====================================================================
        # FAST PATH: precomputed catalog (SharePoint / local)
        # Default: return list immediately (no live refresh) — UI paints fast.
        # Optional: include_refresh=1 overlays live refresh (slower; use after paint).
        # Skip Scanner when catalog has this workspace. Fall through on miss.
        # =====================================================================
        include_refresh = request.args.get('include_refresh', '0').lower() in ('1', 'true', 'yes')
        # Active/inactive filter needs refresh timestamps
        if status_filter in ('active', 'inactive'):
            include_refresh = True

        if (
            not force_live
            and CATALOG_AVAILABLE
            and catalog_service is not None
            and catalog_service.is_available()
        ):
            try:
                allowed = _user_allowed_workspace_ids()
                if workspace_id not in (allowed or set()):
                    print(f"   ⚠️ Catalog fast-path: user cannot access workspace {workspace_id[:8]}…")
                else:
                    catalog_payload = catalog_service.get_workspace_reports(
                        workspace_id, allowed_workspace_ids=allowed
                    )
                    if catalog_payload and catalog_payload.get('reports') is not None:
                        n_reports = len(catalog_payload['reports'])
                        ops_cov = catalog_payload.get('opsCoverage') or {}
                        ops_hits = int(ops_cov.get('reportsWithOps') or 0)
                        # Prefer precomputed ops in catalog (6h batch). Live refresh only
                        # when explicitly requested AND catalog lacks ops coverage.
                        shell = _catalog_reports_shell(catalog_payload['reports'])
                        catalog_has_ops = ops_hits > 0 or any(
                            r.get('ops_from_catalog') for r in shell
                        )
                        need_live = include_refresh and not catalog_has_ops
                        print(
                            f"⚡ CATALOG FAST PATH for workspace {workspace_id[:8]}… "
                            f"({n_reports} reports, ops={ops_hits}/{n_reports}, "
                            f"generatedAt={catalog_payload.get('generatedAt')}, "
                            f"opsEnrichedAt={catalog_payload.get('opsEnrichedAt')}, "
                            f"live_refresh={need_live}, folder={folder_id or 'all'})"
                        )
                        if need_live:
                            report_list = _enrich_catalog_reports_with_refresh(
                                catalog_payload['reports'], workspace_id
                            )
                            source_label = 'catalog+live-refresh'
                            refresh_included = True
                        else:
                            report_list = shell
                            source_label = 'catalog+ops' if catalog_has_ops else 'catalog'
                            refresh_included = catalog_has_ops

                        if status_filter in ('active', 'inactive'):
                            report_list = _filter_reports_by_refresh_status(
                                report_list, status_filter
                            )

                        # Folders + report membership (Fabric items, cached).
                        # Used for dropdown AND for fast folder filter without Scanner.
                        folder_meta = {
                            'folders': [],
                            'report_folder_map': {},
                            'folder_names_map': {},
                        }
                        try:
                            folder_meta = _fetch_workspace_folder_meta(workspace_id)
                        except Exception as folders_exc:
                            print(f"   ⚠️ Folder meta on catalog path failed: {folders_exc}")
                            folder_meta = {
                                'folders': [],
                                'report_folder_map': {},
                                'folder_names_map': {},
                            }
                        folders_list = folder_meta.get('folders') or []
                        report_list = _attach_folder_ids(
                            report_list, folder_meta.get('report_folder_map') or {}
                        )

                        if folder_id:
                            before = len(report_list)
                            report_list = _filter_reports_by_folder(
                                report_list, folder_id, folder_meta
                            )
                            print(
                                f"   📂 Catalog folder filter '{folder_id[:8]}…': "
                                f"{before} → {len(report_list)} reports"
                            )
                            source_label = f"{source_label}+folder"

                        report_folder_map_out = folder_meta.get('report_folder_map') or {}
                        # Cache instant/ops responses (safe — no per-click live fan-out)
                        reports_cache[cache_key] = {
                            'data': report_list,
                            'folders': folders_list,
                            'reportFolderMap': report_folder_map_out,
                            'timestamp': current_time,
                            'source': source_label,
                            'catalogGeneratedAt': catalog_payload.get('generatedAt'),
                            'opsEnrichedAt': catalog_payload.get('opsEnrichedAt'),
                            'opsFromCatalog': catalog_has_ops,
                            'refreshIncluded': refresh_included,
                        }
                        return jsonify({
                            'success': True,
                            'reports': report_list,
                            'folders': folders_list,
                            'reportFolderMap': report_folder_map_out,
                            'source': source_label,
                            'refreshIncluded': refresh_included,
                            'opsFromCatalog': catalog_has_ops,
                            'catalogGeneratedAt': catalog_payload.get('generatedAt'),
                            'opsEnrichedAt': catalog_payload.get('opsEnrichedAt'),
                            'ops': catalog_payload.get('ops'),
                            'workspaceName': (catalog_payload.get('workspace') or {}).get('name'),
                        })
                    else:
                        print(f"   ℹ️ Catalog miss for workspace {workspace_id[:8]}… — live Scanner fallback")
            except Exception as catalog_exc:
                print(f"   ⚠️ Catalog fast-path error (falling back to live): {catalog_exc}")
                import traceback
                traceback.print_exc()

        # Get reports using Scanner API to get folder information
        # Regular API doesn't return folderId/folderObjectId, so we need Scanner API
        from scanner_connector import PowerBIScanner

        scan_data = None
        reports = []
        try:
            print(f"\n📊 Getting reports with folder info from Scanner API...")
            scanner = PowerBIScanner()
            scanner.access_token = scanner.get_access_token()
            scan_data = scanner.run_scan(workspace_id=workspace_id)

            reports = []
            if scan_data and "workspaces" in scan_data:
                for ws in scan_data["workspaces"]:
                    if ws.get("id") == workspace_id:
                        reports = ws.get("reports", [])
                        break
                print(f"   ✅ Found {len(reports)} reports from Scanner API (including [App] reports)")
            else:
                # Fallback to regular API if Scanner fails (no folder info)
                print("⚠️ Scanner API failed, falling back to regular API (no folder info)")
                base_url = "https://api.powerbi.com/v1.0/myorg"
                url = f"{base_url}/groups/{workspace_id}/reports"
                headers = get_user_powerbi_headers()
                response = requests.get(url, headers=headers)
                response.raise_for_status()
                reports = response.json().get('value', [])

            if reports:
                print("\nReports in workspace:")
                for idx, report in enumerate(reports, 1):
                    print(f"   {idx}. {report['name']}")

        except requests.exceptions.RequestException as e:
            print(f"❌ Error fetching reports with user token: {str(e)}")
            if hasattr(e, 'response') and e.response is not None:
                print(f"   Status Code: {e.response.status_code}")
                print(f"   Response: {e.response.text[:500]}")

                # If 403 Forbidden, try fallback to service principal
                if e.response.status_code == 403:
                    print("\n⚠️  User delegated permissions not granted. Falling back to service principal...")
                    print("   To fix this permanently, grant admin consent for Power BI API permissions in Azure Portal.")

                    try:
                        # Fallback to service principal
                        reports = powerbi.get_all_reports(workspace_id)
                    except Exception as fallback_error:
                        print(f"Error: Fallback authentication failed: {str(fallback_error)}")
                        return jsonify({
                            'success': False,
                            'error': 'Unable to fetch reports. Please contact your administrator to grant Power BI API permissions.',
                            'details': str(e)
                        }), 500
                else:
                    raise

        # Fetch folder names via shared helper (Fabric token + SP + Scanner fallbacks)
        folder_names_map = {}
        fabric_folders_list = []  # Store full folder hierarchy

        try:
            print("\n📁 Fetching folder names for live /api/reports path...")
            shared_folders = _fetch_workspace_folders_list(workspace_id)
            for f in shared_folders or []:
                fid = f.get('id')
                if not fid or fid == '__ROOT__':
                    continue
                # Prefer leaf displayName if path was already built into name
                folder_names_map[fid] = {
                    'id': fid,
                    'name': f.get('name') or fid,
                    'parentFolderId': f.get('parentFolderId'),
                }
                fabric_folders_list.append({
                    'id': fid,
                    'name': f.get('name') or fid,
                    'parentFolderId': f.get('parentFolderId'),
                })
            print(f"   ✅ Folder map size: {len(folder_names_map)}")
        except Exception as e:
            print(f"   ⚠️ Error fetching folders: {e}, using folder IDs as names")

        # Build report list with minimal processing
        report_list = []
        folders_map = {}  # Track unique folders
        output_dir = 'output'
        app_reports_excluded = 0  # Count of [App] reports excluded

        # ⚡ Fetch refresh data from Power BI REST API (Scanner doesn't include refresh history).
        # Robust resolver handles: null endTime (in-progress), DirectQuery/Live, cross-workspace datasets.
        from powerbi_connector import resolve_dataset_refresh_info, _empty_refresh_info
        from concurrent.futures import ThreadPoolExecutor, as_completed

        dataset_refresh_map = {}

        print(f"\n🔥 STARTING REFRESH DATA FETCH - Total reports from Scanner: {len(reports)}")

        # Map dataset_id -> preferred dataset home workspace (from Scanner lineage when present)
        dataset_workspace_hints = {}
        dataset_ids = set()

        # Scanner datasets in this workspace (local datasets)
        scan_datasets_in_ws = set()
        if scan_data and "workspaces" in scan_data:
            for ws in scan_data["workspaces"]:
                if ws.get("id") == workspace_id:
                    for ds in ws.get("datasets", []) or []:
                        if ds.get("id"):
                            scan_datasets_in_ws.add(ds["id"])
                            dataset_workspace_hints[ds["id"]] = workspace_id
                    break

        for report in reports:
            if _is_excluded_report_name(report.get('name')):
                continue
            dataset_id = report.get('datasetId')
            if not dataset_id:
                continue
            dataset_ids.add(dataset_id)

            # Prefer explicit dataset workspace from Scanner report object when available
            ds_ws = (
                report.get('datasetWorkspaceId')
                or report.get('datasetWorkspaceObjectId')
                or report.get('workspaceId')
            )
            # If dataset is not among local workspace datasets, keep any cross-ws hint
            if ds_ws and ds_ws != workspace_id:
                dataset_workspace_hints[dataset_id] = ds_ws
            elif dataset_id not in scan_datasets_in_ws and dataset_id not in dataset_workspace_hints:
                # Unknown home workspace — resolver will try report WS then /datasets/{id}
                dataset_workspace_hints[dataset_id] = None
            elif dataset_id in scan_datasets_in_ws:
                dataset_workspace_hints[dataset_id] = workspace_id

        print(f"📊 Found {len(dataset_ids)} unique datasets after filtering [App] reports")
        if not dataset_ids:
            print("⚠️ WARNING: No dataset IDs found! Check if reports have 'datasetId' field.")

        if dataset_ids:
            print(f"\n📊 Fetching refresh data for {len(dataset_ids)} datasets...")
            print(
                f"   Dataset IDs: {list(dataset_ids)[:3]}..."
                if len(dataset_ids) > 3
                else f"   Dataset IDs: {list(dataset_ids)}"
            )

            # Headers must be captured before threads (Flask request context not available there)
            headers = get_user_powerbi_headers()
            base_url = "https://api.powerbi.com/v1.0/myorg"

            def fetch_dataset_refresh(ds_id, headers_param, base_url_param, report_ws_id, ds_ws_hint):
                """Thread-safe refresh fetch via shared robust resolver."""
                try:
                    print(f"   🔍 Fetching refresh for dataset: {ds_id[:8]}...")
                    info = resolve_dataset_refresh_info(
                        headers=headers_param,
                        workspace_id=report_ws_id,
                        dataset_id=ds_id,
                        dataset_workspace_id=ds_ws_hint,
                        history_top=5,
                        base_url=base_url_param,
                        timeout=8,
                    )
                    return ds_id, info
                except Exception as e:
                    print(f"      ❌ Error fetching dataset {ds_id[:8]}: {str(e)}")
                    return ds_id, _empty_refresh_info(
                        refresh_schedule='Error',
                        last_refresh_status='Error',
                        refresh_type='error',
                        refresh_note=str(e),
                        dataset_workspace_id=ds_ws_hint or report_ws_id,
                    )

            with ThreadPoolExecutor(max_workers=10) as executor:
                futures = {
                    executor.submit(
                        fetch_dataset_refresh,
                        ds_id,
                        headers,
                        base_url,
                        workspace_id,
                        dataset_workspace_hints.get(ds_id),
                    ): ds_id
                    for ds_id in dataset_ids
                }

                for future in as_completed(futures):
                    ds_id = futures[future]
                    try:
                        # Per-dataset budget: resolver itself uses 8s timeouts
                        resolved_id, refresh_info = future.result(timeout=45)
                        dataset_refresh_map[resolved_id] = refresh_info
                    except Exception as e:
                        print(f"   ⚠️ Error/timeout fetching refresh for dataset {ds_id[:8]}: {e}")
                        dataset_refresh_map[ds_id] = _empty_refresh_info(
                            refresh_schedule='Error',
                            last_refresh_status='Error',
                            refresh_type='error',
                            refresh_note=f'Timeout or worker error: {e}',
                            dataset_workspace_id=dataset_workspace_hints.get(ds_id) or workspace_id,
                        )

            print(f"   ✅ Fetched refresh data for {len(dataset_refresh_map)} datasets")

            if dataset_refresh_map:
                sample_id = list(dataset_refresh_map.keys())[0]
                sample_data = dataset_refresh_map[sample_id]
                print(f"   📋 Sample refresh data for {sample_id[:8]}...: {sample_data}")
            else:
                print(f"   ⚠️ WARNING: dataset_refresh_map is EMPTY!")

        # Helper function to build full folder path for display
        def build_folder_path(folder_id, folder_map):
            """Build hierarchical path like 'Parent / Child / Grandchild'"""
            if not folder_id or folder_id not in folder_map:
                return None

            path_parts = []
            current_id = folder_id
            visited = set()  # Prevent infinite loops

            while current_id and current_id not in visited:
                visited.add(current_id)
                folder_info = folder_map.get(current_id)
                if not folder_info:
                    break
                path_parts.insert(0, folder_info['name'])
                current_id = folder_info.get('parentFolderId')

            return ' / '.join(path_parts) if path_parts else None

        # Use global helper function for getting child folders

        # Track root-level reports (no folder assigned)
        root_reports_count = 0

        for report in reports:
            # SKIP [App] shells + platform Usage Metrics reports
            report_name = report.get('name', '')
            if _is_excluded_report_name(report_name):
                app_reports_excluded += 1
                continue

            dataset_id = report.get('datasetId', '')
            report_folder_id = report.get('folderObjectId') or report.get('folderId')

            # Debug first 5 reports' folder IDs
            if reports.index(report) < 5:
                print(f"\n🔍 DEBUG Report #{reports.index(report) + 1} '{report.get('name')}':")
                print(f"   folder ID = {report_folder_id if report_folder_id else 'NULL (root-level)'}")
                if report_folder_id:
                    print(f"   Exists in folder_names_map: {report_folder_id in folder_names_map}")
                    if report_folder_id in folder_names_map:
                        print(f"   Folder name: {folder_names_map[report_folder_id]['name']}")

            # Count root-level reports
            if not report_folder_id:
                root_reports_count += 1

            # NEW: Track folders for dropdown population with proper names
            if report_folder_id:
                # Use folder name from API if available, otherwise use report's folderName, or fallback to ID
                if report_folder_id in folder_names_map:
                    folder_info = folder_names_map[report_folder_id]
                    folder_name = build_folder_path(report_folder_id, folder_names_map) or folder_info['name']
                    parent_folder_id = folder_info.get('parentFolderId')
                else:
                    folder_name = report.get('folderName', f'Folder {report_folder_id[:8]}')
                    parent_folder_id = None

                if report_folder_id not in folders_map:
                    folders_map[report_folder_id] = {
                        'id': report_folder_id,
                        'name': folder_name,
                        'parentFolderId': parent_folder_id,
                        'report_count': 0,
                        'hasChildren': False  # Will be updated later
                    }
                folders_map[report_folder_id]['report_count'] += 1

            # Quick file check only
            doc_filename = f"{report['name']}_Documentation.docx"
            doc_path = os.path.join(output_dir, doc_filename)
            has_documentation = os.path.exists(doc_path)

            # Get file timestamp if exists
            last_generated = None
            if has_documentation:
                try:
                    doc_mtime = os.path.getmtime(doc_path)
                    from datetime import datetime
                    last_generated = datetime.fromtimestamp(doc_mtime).isoformat()
                except:
                    pass

            # Get dataset refresh info if available
            dataset_refresh_info = dataset_refresh_map.get(dataset_id, {})

            # Debug: Log if we found refresh info for this dataset
            if dataset_id and dataset_refresh_info:
                print(
                    f"   🔍 Report '{report['name'][:30]}' - Dataset {dataset_id[:8]}: "
                    f"last={dataset_refresh_info.get('last_refreshed')} "
                    f"status={dataset_refresh_info.get('last_refresh_status')} "
                    f"type={dataset_refresh_info.get('refresh_type')}"
                )
            elif dataset_id:
                print(f"   ⚠️ Report '{report['name'][:30]}' - Dataset {dataset_id[:8]}: NO REFRESH INFO FOUND")

            # Build minimal report object for fast initial load
            report_obj = {
                'id': report['id'],
                'name': report['name'],
                'datasetId': dataset_id,
                'folderId': report_folder_id,  # NEW: Include folder info
                'folderName': report.get('folderName'),
                'refresh_schedule': dataset_refresh_info.get('refresh_schedule'),
                'last_refreshed': dataset_refresh_info.get('last_refreshed'),
                'last_refresh_status': dataset_refresh_info.get('last_refresh_status'),
                'refresh_type': dataset_refresh_info.get('refresh_type'),
                'refresh_note': dataset_refresh_info.get('refresh_note'),
                'dataset_workspace_id': dataset_refresh_info.get('dataset_workspace_id'),
                'last_accessed': None,  # Not available in Scanner API
                'last_accessed_by': None,  # Not available in Scanner API
                'last_modified': report.get('modifiedDateTime'),
                'last_generated': last_generated,
                'has_documentation': has_documentation,
                'needs_update': False,
                'webUrl': report.get('webUrl', '')
            }

            report_list.append(report_obj)

        # Log exclusion summary
        if app_reports_excluded > 0:
            print(f"\n🚫 Excluded {app_reports_excluded} [App] reports (published copies)")
            print(f"✅ Returning {len(report_list)} original reports")

        # CRITICAL FIX: Add ALL folders from Fabric API to folders_map (not just folders with reports)
        # This ensures empty folders are visible in the dropdown
        if folder_names_map:
            print(f"\n📁 Merging ALL Fabric API folders into folders_map...")
            for fabric_folder_id, fabric_folder_info in folder_names_map.items():
                if fabric_folder_id not in folders_map:
                    # Add folder with 0 reports
                    folder_name = build_folder_path(fabric_folder_id, folder_names_map) or fabric_folder_info['name']
                    folders_map[fabric_folder_id] = {
                        'id': fabric_folder_id,
                        'name': folder_name,
                        'parentFolderId': fabric_folder_info.get('parentFolderId'),
                        'report_count': 0,  # No reports in this folder
                        'hasChildren': False
                    }
                    print(f"   + Added empty folder: {fabric_folder_info['name']} (0 reports)")

        # Add virtual "Root Directory" folder for reports without folder assignment
        if root_reports_count > 0:
            print(f"\n📁 Found {root_reports_count} root-level reports (no folder assigned)")
            folders_map['__ROOT__'] = {
                'id': '__ROOT__',
                'name': '📂 Root Directory (Uncategorized)',
                'parentFolderId': None,
                'report_count': root_reports_count,
                'hasChildren': False
            }

        # Mark folders that have children
        for folder_id_key, folder_info in folders_map.items():
            parent_id = folder_info.get('parentFolderId')
            if parent_id and parent_id in folders_map:
                folders_map[parent_id]['hasChildren'] = True

        # HIERARCHICAL FOLDER FILTERING: Filter by folder if specified (including nested subfolders)
        if folder_id:
            print(f"\n🗂️  Hierarchical Filtering: folder_id = {folder_id}")

            # Special handling for Root Directory
            if folder_id == '__ROOT__':
                print(f"   ✅ Filtering for ROOT DIRECTORY (uncategorized reports)")
                original_count = len(report_list)
                report_list = [r for r in report_list if not (r.get('folderObjectId') or r.get('folderId'))]
                print(f"   📊 Reports filtered: {original_count} → {len(report_list)} root-level reports")
            else:
                # SAFETY CHECK: Validate folder ID exists in workspace
                folder_exists = folder_id in folder_names_map or folder_id in folders_map

                if not folder_exists:
                    print(f"   ⚠️  INVALID FOLDER ID: '{folder_id[:8]}...' not found in workspace!")
                    print(f"   ✅ FALLBACK: Showing all {len(report_list)} reports (ignoring invalid filter)")
                    # Don't filter - show all reports
                else:
                    # Get the folder name for logging
                    folder_name = folder_names_map.get(folder_id, folders_map.get(folder_id, {})).get('name', 'Unknown')
                    print(f"   ✅ Valid folder: '{folder_name}' (ID: {folder_id[:8]}...)")

                    # Get all child folders recursively (folder + all descendants)
                    all_folder_ids = get_all_child_folders_recursive(folder_id, folder_names_map if folder_names_map else folders_map)

                    print(f"   📁 Folder hierarchy: {len(all_folder_ids)} folder(s) in tree")
                    print(f"   📁 Folder IDs to match: {[fid[:8] + '...' for fid in list(all_folder_ids)[:3]]}")

                    # Filter reports to only those in the selected folder or its children
                    original_count = len(report_list)

                    # DEBUG: Show what we're matching against
                    filtered_reports = []
                    for r in report_list:
                        r_folder_id = r.get('folderObjectId') or r.get('folderId')
                        if r_folder_id in all_folder_ids:
                            filtered_reports.append(r)
                            # Debug first match
                            if len(filtered_reports) == 1:
                                print(f"   🔍 First matching report: '{r.get('name')}' (folder: {r_folder_id[:8]}...)")

                    report_list = filtered_reports

                    print(f"   📊 Reports filtered: {original_count} → {len(report_list)} reports")

                    if len(report_list) == 0:
                        print(f"   ⚠️  WARNING: Zero reports matched! Possible ID mismatch.")
                        print(f"   💡 TIP: Check if Scanner API folderObjectId matches Fabric API folder id")

                    # Keep ALL folders in dropdown (don't filter folders_map)
                    # This allows users to navigate to sibling folders without reloading workspace

        # Apply status filter (active/inactive) if requested
        if status_filter in ['active', 'inactive']:
            from datetime import datetime, timezone

            print(f"\n🔍 Applying status filter: {status_filter}")
            original_count = len(report_list)
            filtered_reports = []

            def _is_live_connection(rep):
                rtype = (rep.get('refresh_type') or '').lower()
                status = (rep.get('last_refresh_status') or '').lower()
                schedule = (rep.get('refresh_schedule') or '').lower()
                return (
                    rtype in ('directquery', 'live')
                    or 'directquery' in status
                    or 'live' in status
                    or 'directquery' in schedule
                    or 'live' in schedule
                )

            for report in report_list:
                last_refreshed = report.get('last_refreshed')
                is_live = _is_live_connection(report)

                if status_filter == 'inactive':
                    # DirectQuery/Live is not "stale import refresh" — exclude from inactive
                    if is_live:
                        continue
                    # Inactive: no last_refreshed OR last_refreshed > 30 days ago
                    if not last_refreshed:
                        filtered_reports.append(report)
                    else:
                        try:
                            # Parse the refresh time
                            refresh_time_str = last_refreshed
                            if refresh_time_str.endswith('Z'):
                                refresh_time_str = refresh_time_str[:-1] + '+00:00'

                            refresh_time = datetime.fromisoformat(refresh_time_str)
                            if refresh_time.tzinfo is None:
                                refresh_time = refresh_time.replace(tzinfo=timezone.utc)

                            now = datetime.now(timezone.utc)
                            days_since_refresh = (now - refresh_time).days

                            if days_since_refresh >= 30:
                                filtered_reports.append(report)
                        except Exception:
                            # Can't parse date - treat as inactive
                            filtered_reports.append(report)

                elif status_filter == 'active':
                    # Live connections count as active (always current)
                    if is_live:
                        filtered_reports.append(report)
                        continue
                    # Active: has last_refreshed AND last_refreshed <= 30 days ago
                    if last_refreshed:
                        try:
                            # Parse the refresh time
                            refresh_time_str = last_refreshed
                            if refresh_time_str.endswith('Z'):
                                refresh_time_str = refresh_time_str[:-1] + '+00:00'

                            refresh_time = datetime.fromisoformat(refresh_time_str)
                            if refresh_time.tzinfo is None:
                                refresh_time = refresh_time.replace(tzinfo=timezone.utc)

                            now = datetime.now(timezone.utc)
                            days_since_refresh = (now - refresh_time).days

                            if days_since_refresh < 30:
                                filtered_reports.append(report)
                        except Exception:
                            # Can't parse date - not active
                            pass

            report_list = filtered_reports
            print(f"   📊 Status filtered: {original_count} → {len(report_list)} reports ({status_filter})")

        # Convert folders to list and sort by hierarchy (root folders first, then children)
        def get_folder_sort_key(folder):
            """Sort by hierarchy: root folders first, then by name"""
            # Build full path for sorting
            path = build_folder_path(folder['id'], folder_names_map if folder_names_map else {f['id']: f for f in [folder]})
            return path if path else folder['name']

        folders_list = sorted(folders_map.values(), key=get_folder_sort_key)

        # Build report→folder map for client-side filtering
        live_report_folder_map = {}
        for r in report_list or []:
            rid = r.get('id')
            if not rid:
                continue
            live_report_folder_map[str(rid)] = r.get('folderId') or r.get('folderObjectId')

        # Update cache (per-user, per-workspace, per-folder)
        reports_cache[cache_key] = {
            'data': report_list,
            'folders': folders_list,
            'reportFolderMap': live_report_folder_map,
            'timestamp': current_time,
        }

        print(f"✅ Returning {len(report_list)} reports, {len(folders_list)} folders")

        # Debug: Check if first report has refresh data
        if report_list and len(report_list) > 0:
            first_report = report_list[0]
            print(f"\n🔍 FINAL CHECK - First report '{first_report.get('name', 'Unknown')[:30]}':")
            print(f"   refresh_schedule: {first_report.get('refresh_schedule')}")
            print(f"   last_refreshed: {first_report.get('last_refreshed')}")
            print(f"   last_refresh_status: {first_report.get('last_refresh_status')}")

        return jsonify({
            'success': True,
            'reports': report_list,
            'folders': folders_list,
            'reportFolderMap': live_report_folder_map,
        })

    except Exception as e:
        print(f"❌ Error fetching reports: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/export-inactive-reports/<workspace_id>')
@login_required
def export_inactive_reports(workspace_id):
    """
    Generate decommissioning Excel for a workspace.

    A report is a decommission candidate if ANY of:
      1) days_since_refresh > 90
      2) no import refresh history (for models that should refresh)
      3) views in last 30 days == 0

    DirectQuery / Live connections are NOT flagged solely for "no refresh history"
    (they have no import refresh cycle) — they still flag on 0 views / stale if applicable.

    Prefers catalog ops fields; falls back to usage_cache + Scanner when needed.
    """
    try:
        from datetime import datetime, timezone
        import io
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from flask import send_file

        print(f"\n📋 GENERATING DECOMMISSION REPORT FOR WORKSPACE: {workspace_id}")
        REFRESH_STALE_DAYS = 90
        VIEW_LOOKBACK_LABEL = '60 days'
        VIEW_LOOKBACK_DAYS = 60

        def _is_live_or_dq(report_or_ds):
            rtype = str(
                report_or_ds.get('refresh_type')
                or report_or_ds.get('refreshType')
                or ''
            ).lower()
            status = str(
                report_or_ds.get('last_refresh_status')
                or report_or_ds.get('refresh_schedule')
                or ''
            ).lower()
            note = str(report_or_ds.get('refresh_note') or '').lower()
            if rtype in {'directquery', 'live', 'push', 'streaming'}:
                return True
            if 'directquery' in status or 'live' in status or 'direct query' in status:
                return True
            if 'live connection' in note or 'directquery' in note:
                return True
            return False

        def _parse_dt(value):
            if value in (None, '', 'Unknown', '—', '-'):
                return None
            try:
                if isinstance(value, (int, float)):
                    ts = float(value)
                    if ts > 1e12:
                        ts /= 1000.0
                    return datetime.fromtimestamp(ts, tz=timezone.utc)
                s = str(value).strip()
                if s.endswith('Z'):
                    s = s[:-1] + '+00:00'
                dt = datetime.fromisoformat(s)
                if dt.tzinfo is None:
                    dt = dt.replace(tzinfo=timezone.utc)
                return dt
            except Exception:
                return None

        def _days_since(value):
            dt = _parse_dt(value)
            if not dt:
                return None
            return max(0, (datetime.now(timezone.utc) - dt).days)

        def _clean_person(*vals):
            for v in vals:
                if v is None:
                    continue
                if isinstance(v, dict):
                    v = (
                        v.get('userPrincipalName')
                        or v.get('emailAddress')
                        or v.get('displayName')
                        or ''
                    )
                s = str(v).strip()
                if s and s.lower() not in {'unknown', 'n/a', 'none', 'null', '-', '—'}:
                    return s
            return ''

        # ---------- 1) Catalog-first report list + ops ----------
        workspace_name = 'Unknown Workspace'
        reports_in = []
        source = 'catalog'
        report_views = {}
        last_viewed_map = {}

        pack = None
        if CATALOG_AVAILABLE:
            try:
                pack = catalog_service.get_workspace_reports(workspace_id)
            except Exception as e:
                print(f"   ⚠️ catalog get_workspace_reports failed: {e}")
                pack = None

        if pack and (pack.get('reports') or []):
            workspace_name = (
                (pack.get('workspace') or {}).get('name')
                or pack.get('workspace_name')
                or workspace_name
            )
            # resolve workspace name from catalog if needed
            if workspace_name == 'Unknown Workspace':
                try:
                    cat = catalog_service.get_workspace_catalog() or {}
                    ws = next(
                        (w for w in (cat.get('workspaces') or []) if w.get('id') == workspace_id),
                        None,
                    )
                    if ws:
                        workspace_name = ws.get('name') or workspace_name
                except Exception:
                    pass
            reports_in = list(pack.get('reports') or [])
            print(f"   📦 Catalog reports: {len(reports_in)}")
        else:
            # Scanner fallback for report list
            source = 'scanner'
            from scanner_connector import PowerBIScanner
            scanner = PowerBIScanner()
            scan_result = scanner.run_scan(workspace_id=workspace_id) or {}
            workspace_data = None
            for ws in scan_result.get('workspaces') or []:
                if ws.get('id') == workspace_id:
                    workspace_data = ws
                    break
            if not workspace_data and (scan_result.get('workspaces') or []):
                workspace_data = scan_result['workspaces'][0]
            if not workspace_data:
                return jsonify({'error': 'Could not retrieve workspace data'}), 500
            workspace_name = workspace_data.get('name', workspace_name)
            reports_in = list(workspace_data.get('reports') or [])
            print(f"   🛰️ Scanner reports: {len(reports_in)}")

        # ---------- 2) Views (60-day window preferred) ----------
        # a) catalog report.view_count
        for r in reports_in:
            rid = r.get('id')
            if rid and r.get('view_count') is not None:
                try:
                    report_views[rid] = int(r.get('view_count') or 0)
                except Exception:
                    report_views[rid] = 0
            lv = r.get('last_viewed')
            if rid and lv:
                last_viewed_map[rid] = lv if isinstance(lv, dict) else {'timestamp': lv}

        # b) in-memory usage_cache overlay
        cache_key = f"usage_{workspace_id}"
        if cache_key in usage_cache:
            cached = usage_cache[cache_key].get('data') or {}
            for rid, cnt in (cached.get('report_views') or {}).items():
                report_views[rid] = int(cnt or 0)
            for rid, info in (cached.get('last_viewed') or {}).items():
                last_viewed_map[rid] = info

        # c) usage_snapshot.json overlay (catalog ops)
        data_as_of_raw = None  # when usage / ops data is current through
        if CATALOG_AVAILABLE:
            try:
                usage_snap = catalog_service.get_json('usage_snapshot.json') or {}
                data_as_of_raw = (
                    usage_snap.get('generatedAt')
                    or usage_snap.get('opsEnrichedAt')
                    or usage_snap.get('asOf')
                    or usage_snap.get('endDate')
                )
                for rid, cnt in (usage_snap.get('report_views') or {}).items():
                    # don't wipe known values with empty snapshot
                    if rid not in report_views or report_views.get(rid) in (None,):
                        report_views[rid] = int(cnt or 0)
                    else:
                        # prefer higher fidelity if snapshot has value
                        report_views[rid] = int(cnt or report_views.get(rid) or 0)
                for rid, info in (usage_snap.get('last_viewed') or {}).items():
                    last_viewed_map.setdefault(rid, info)
            except Exception as e:
                print(f"   ⚠️ usage_snapshot overlay failed: {e}")

        # ops / catalog timestamps for "as of" day math when usage has none
        if not data_as_of_raw and pack:
            data_as_of_raw = (
                (pack.get('catalog_meta') or {}).get('opsEnrichedAt')
                or pack.get('opsEnrichedAt')
                or (pack.get('catalog_meta') or {}).get('generatedAt')
                or pack.get('generatedAt')
            )
        if not data_as_of_raw and CATALOG_AVAILABLE:
            try:
                summary = catalog_service.get_summary() or {}
                data_as_of_raw = summary.get('opsEnrichedAt') or summary.get('generatedAt')
            except Exception:
                pass

        data_as_of_dt = _parse_dt(data_as_of_raw) or datetime.now(timezone.utc)
        # Calendar date label for column header (local US style MM-DD-YYYY)
        data_as_of_label = data_as_of_dt.strftime('%m-%d-%Y')

        # Folder path map (best-effort) for Sub Folder column
        folder_path_by_report = {}
        try:
            folder_meta = _fetch_workspace_folder_meta(workspace_id) or {}
            report_folder_map = folder_meta.get('report_folder_map') or {}
            folder_names_map = folder_meta.get('folder_names_map') or {}

            def _build_folder_path(folder_id):
                if not folder_id or folder_id not in folder_names_map:
                    return None
                parts = []
                cur = folder_id
                seen = set()
                while cur and cur not in seen:
                    seen.add(cur)
                    info = folder_names_map.get(cur) or {}
                    nm = info.get('name')
                    if nm:
                        parts.insert(0, nm)
                    cur = info.get('parentFolderId')
                return ' / '.join(parts) if parts else None

            for rid, fid in report_folder_map.items():
                path = _build_folder_path(fid)
                if path:
                    folder_path_by_report[rid] = path
        except Exception as folder_err:
            print(f"   ⚠️ folder map for decomm export skipped: {folder_err}")
            folder_path_by_report = {}

        def _days_until_as_of(value):
            """Days from last refresh to catalog/usage data-as-of date."""
            dt = _parse_dt(value)
            if not dt:
                return None
            return max(0, (data_as_of_dt - dt).days)

        # ---------- 3) Evaluate each report ----------
        def evaluate_report(report):
            name = report.get('name') or 'Unknown'
            if _is_excluded_report_name(name):
                return None  # skip [App] shells + platform usage metrics

            rid = report.get('id') or ''
            live_dq = _is_live_or_dq(report)

            # Refresh signals
            days = report.get('days_since_refresh')
            try:
                days = int(days) if days is not None and days != '' else None
            except Exception:
                days = None
            last_ref = report.get('last_refreshed') or report.get('lastRefresh') or ''
            # Prefer days relative to data-as-of (not wall-clock now) for the export column
            days_as_of = _days_until_as_of(last_ref) if last_ref else None
            if days_as_of is not None:
                days = days_as_of
            elif days is None and last_ref:
                days = _days_since(last_ref)

            refresh_status = (
                report.get('last_refresh_status')
                or report.get('refreshStatus')
                or ''
            )
            refresh_type = report.get('refresh_type') or report.get('refreshType') or ''

            # Views (60-day window when snapshot provides it)
            if rid in report_views:
                views = int(report_views.get(rid) or 0)
                views_known = True
            elif report.get('view_count') is not None:
                views = int(report.get('view_count') or 0)
                views_known = True
            else:
                views = 0
                views_known = False

            reasons = []
            # Rule 1: stale refresh > 90 days
            if days is not None and days > REFRESH_STALE_DAYS:
                reasons.append(f'Days since refresh > {REFRESH_STALE_DAYS} ({days}d)')

            # Rule 2: no refresh history (import models only)
            no_history = False
            if not live_dq:
                if (not last_ref) and (days is None):
                    no_history = True
                elif str(refresh_status).lower() in {'error', 'failed'} and not last_ref:
                    no_history = True
            if no_history:
                reasons.append('No refresh history')

            # Rule 3: last 60 days views == 0
            zero_views = False
            if (views_known or report.get('view_count') is not None) and int(views or 0) == 0:
                zero_views = True
                reasons.append(f'0 views in last {VIEW_LOOKBACK_LABEL}')

            # Rule 4: Verify in PBI (content-modified / Unverified — not confirmed refresh history)
            refresh_src = str(
                report.get('refresh_source') or report.get('refreshSource') or ''
            ).lower().replace('-', '_').replace(' ', '')
            status_l = str(refresh_status or '').lower()
            note_l = str(report.get('refresh_note') or report.get('refreshNote') or '').lower()
            needs_verify = (
                status_l in {'unverified', 'estimated'}
                or refresh_src in {
                    'content_modified', 'content_created', 'contentcreated', 'created',
                }
                or 'verify in power bi' in note_l
                or 'verify in pbi' in note_l
            )
            if needs_verify and not live_dq:
                reasons.append('Verify in PBI (unconfirmed refresh)')

            is_candidate = len(reasons) > 0

            owner = _clean_person(
                report.get('dataset_owner'),
                report.get('configuredBy'),
                report.get('createdBy'),
                report.get('created_by'),
                report.get('modifiedBy'),
                report.get('modified_by'),
            )

            # Sub folder: root = NA
            sub_folder = (
                report.get('folderPath')
                or report.get('folder_path')
                or report.get('folderName')
                or report.get('folder_name')
                or folder_path_by_report.get(rid)
                or ''
            )
            sub_folder = str(sub_folder).strip() if sub_folder else ''
            if not sub_folder or sub_folder.lower() in {'root', 'workspace', 'none', 'n/a', 'na', '—'}:
                sub_folder = 'NA'

            model_name = (
                report.get('datasetName')
                or report.get('dataset_name')
                or ''
            ).strip() or 'NA'

            last_ref_disp = ''
            if last_ref and str(last_ref) not in {'—', '-', 'Unknown', 'None'}:
                last_ref_disp = str(last_ref)[:19].replace('T', ' ')

            # Comments: verify flag first, then zero views
            comment_parts = []
            if needs_verify and not live_dq:
                comment_parts.append('Verify in PBI — last refresh unconfirmed')
            if zero_views:
                comment_parts.append(f'Zero views in last {VIEW_LOOKBACK_LABEL}')
            comments = '; '.join(comment_parts)

            return {
                'report_id': rid,
                'workspace_name': workspace_name,
                'sub_folder': sub_folder,
                'report_name': name,
                'model_name': model_name,
                'last_refreshed_on': last_ref_disp or 'NA',
                'days_since_refresh': days if days is not None else '',
                'contact_owner': owner or 'NA',
                'archived_status': '',
                'comments': comments,
                'can_decommission': '',
                'review_comments': '',
                'is_candidate': is_candidate,
                'reasons': reasons,
                'zero_views': zero_views,
                'needs_verify': bool(needs_verify and not live_dq),
            }

        evaluated = []
        for r in reports_in:
            row = evaluate_report(r)
            if row:
                evaluated.append(row)

        candidates = [r for r in evaluated if r['is_candidate']]
        candidates.sort(key=lambda x: (
            (x.get('workspace_name') or '').lower(),
            (x.get('sub_folder') or '').lower(),
            (x.get('report_name') or '').lower(),
        ))

        print(
            f"   ✅ total={len(evaluated)} candidates={len(candidates)} "
            f"source={source} views_keys={len(report_views)} as_of={data_as_of_label}"
        )

        # ---------- 4) Excel (single sheet — requested columns only) ----------
        wb = Workbook()
        ws_out = wb.active
        ws_out.title = 'Decommission List'

        header_fill = PatternFill(start_color='2B6CB0', end_color='2B6CB0', fill_type='solid')
        header_font = Font(color='FFFFFF', bold=True, size=11)
        center_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
        left_align = Alignment(horizontal='left', vertical='center', wrap_text=True)
        # Verify-in-PBI rows: light red so reviewers know to check before acting
        verify_fill = PatternFill(start_color='FFC7CE', end_color='FFC7CE', fill_type='solid')
        verify_font = Font(color='9C0006')

        days_header = f'LastRefresh in days till {data_as_of_label}'
        headers = [
            'Workspace Name',
            'Sub Folder',
            'Report Name',
            'Model Name',
            'LastRefreshedOn',
            days_header,
            'Contact/Owner',
            'Archived Status',
            'Comments',
            'Can Decommission',
            'Review Comments',
        ]

        ws_out.append(headers)
        for col_num, _h in enumerate(headers, 1):
            cell = ws_out.cell(row=1, column=col_num)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = center_align
        # AutoFilter + freeze
        last_col_letter = chr(ord('A') + len(headers) - 1) if len(headers) <= 26 else 'K'
        ws_out.auto_filter.ref = f'A1:{last_col_letter}1'
        ws_out.freeze_panes = 'A2'
        # Note in header comment on days column
        ws_out.cell(row=1, column=6).comment = None
        try:
            from openpyxl.comments import Comment
            ws_out.cell(row=1, column=6).comment = Comment(
                f'Days from LastRefreshedOn through data-as-of date ({data_as_of_label}). '
                f'Source: catalog/ops usage extract.',
                'Power BI Control Center',
            )
        except Exception:
            pass

        for r in candidates:
            ws_out.append([
                r['workspace_name'],
                r['sub_folder'],
                r['report_name'],
                r['model_name'],
                r['last_refreshed_on'],
                r['days_since_refresh'] if r['days_since_refresh'] != '' else 'NA',
                r['contact_owner'],
                r['archived_status'],  # empty for reviewers
                r['comments'],
                r['can_decommission'],  # empty
                r['review_comments'],  # empty
            ])
            # Left-align text columns; red highlight = Verify in PBI
            rn = ws_out.max_row
            for c in range(1, 12):
                cell = ws_out.cell(row=rn, column=c)
                cell.alignment = left_align if c != 6 else center_align
                if r.get('needs_verify'):
                    cell.fill = verify_fill
                    cell.font = verify_font

        if not candidates:
            ws_out.append(
                [workspace_name, 'NA', 'No decommission candidates under current rules',
                 'NA', 'NA', 'NA', 'NA', '', '', '', '']
            )

        # Column widths
        widths = {
            'A': 28, 'B': 22, 'C': 36, 'D': 28, 'E': 18, 'F': 28,
            'G': 28, 'H': 16, 'I': 32, 'J': 16, 'K': 22,
        }
        for letter, w in widths.items():
            ws_out.column_dimensions[letter].width = w

        # Thin Rules sheet (optional reference — not the main list)
        ws_rules = wb.create_sheet('Rules')
        ws_rules.append(['Decommission list export'])
        ws_rules.append(['Workspace', workspace_name])
        ws_rules.append(['Generated (UTC)', datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M')])
        ws_rules.append(['Data as of', data_as_of_label])
        ws_rules.append(['Data source', source])
        ws_rules.append(['Candidates', len(candidates)])
        ws_rules.append([])
        ws_rules.append(['Candidate if ANY:'])
        ws_rules.append(['1', f'Days since last refresh > {REFRESH_STALE_DAYS} (through {data_as_of_label})'])
        ws_rules.append(['2', 'No import refresh history (not applied alone to DirectQuery/Live)'])
        ws_rules.append(['3', f'Views in last {VIEW_LOOKBACK_LABEL} == 0'])
        ws_rules.append(['4', 'Verify in PBI — Unverified / content-modified fallback (not confirmed Scheduled refresh)'])
        ws_rules.append([])
        ws_rules.append(['Comments column', f'Pre-filled with verify note and/or "Zero views in last {VIEW_LOOKBACK_LABEL}" when applicable'])
        ws_rules.append(['Red rows', 'Verify in PBI — last refresh unconfirmed; review in Power BI before acting. You may change the fill color after review.'])
        ws_rules.append(['Archived Status / Can Decommission / Review Comments', 'Left blank for reviewers'])
        ws_rules.append(['Excluded', 'Usage Metrics Report, Report Usage Metrics Report, Dashboard Usage Metrics Report, [App] copies'])
        ws_rules.column_dimensions['A'].width = 40
        ws_rules.column_dimensions['B'].width = 80
        ws_rules.cell(row=1, column=1).font = Font(bold=True, size=13, color='FFFFFF')
        ws_rules.cell(row=1, column=1).fill = header_fill

        output = io.BytesIO()
        wb.save(output)
        output.seek(0)

        safe_ws = ''.join(c if c.isalnum() or c in ('-', '_') else '_' for c in workspace_name)[:60]
        filename = f'Decommission_List_{safe_ws}_{datetime.now(timezone.utc).strftime("%Y%m%d")}.xlsx'
        print(f"   📁 {filename} candidates={len(candidates)}/{len(evaluated)} as_of={data_as_of_label}")

        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=filename,
        )

    except Exception as e:
        print(f"❌ Error generating decommission report: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/decommission-list-candidates/<workspace_id>')
@login_required
def decommission_list_candidates(workspace_id):
    """
    JSON list of Report Catalog decommission-list candidates for bulk SharePoint archive.

    Same candidate rules as /api/export-inactive-reports. Does not export files —
    the client calls POST /api/reports/archive-to-sharepoint once per candidate
    (identical to the per-row Download button).

    Requires archive allow-list (same as individual Download).
    """
    try:
        from datetime import datetime, timezone

        from features.report_archive_service import user_can_archive

        email = session.get('user', {}).get('preferred_username') or ''
        if not user_can_archive(email):
            return jsonify({
                'success': False,
                'error': 'Not authorized to archive reports to SharePoint.',
            }), 403

        workspace_id = (workspace_id or '').strip()
        if not workspace_id:
            return jsonify({'success': False, 'error': 'workspace_id required'}), 400

        REFRESH_STALE_DAYS = 90
        VIEW_LOOKBACK_LABEL = '60 days'

        def _is_live_or_dq(report_or_ds):
            rtype = str(
                report_or_ds.get('refresh_type')
                or report_or_ds.get('refreshType')
                or ''
            ).lower()
            status = str(
                report_or_ds.get('last_refresh_status')
                or report_or_ds.get('refresh_schedule')
                or ''
            ).lower()
            note = str(report_or_ds.get('refresh_note') or '').lower()
            if rtype in {'directquery', 'live', 'push', 'streaming'}:
                return True
            if 'directquery' in status or 'live' in status or 'direct query' in status:
                return True
            if 'live connection' in note or 'directquery' in note:
                return True
            return False

        def _parse_dt(value):
            if value in (None, '', 'Unknown', '—', '-'):
                return None
            try:
                if isinstance(value, (int, float)):
                    ts = float(value)
                    if ts > 1e12:
                        ts /= 1000.0
                    return datetime.fromtimestamp(ts, tz=timezone.utc)
                s = str(value).strip()
                if s.endswith('Z'):
                    s = s[:-1] + '+00:00'
                dt = datetime.fromisoformat(s)
                if dt.tzinfo is None:
                    dt = dt.replace(tzinfo=timezone.utc)
                return dt
            except Exception:
                return None

        def _days_since(value, as_of=None):
            dt = _parse_dt(value)
            if not dt:
                return None
            end = as_of or datetime.now(timezone.utc)
            return max(0, (end - dt).days)

        workspace_name = 'Unknown Workspace'
        reports_in = []
        report_views = {}

        pack = None
        if CATALOG_AVAILABLE:
            try:
                pack = catalog_service.get_workspace_reports(workspace_id)
            except Exception as e:
                print(f"   ⚠️ decomm candidates catalog failed: {e}")
                pack = None

        if pack and (pack.get('reports') or []):
            workspace_name = (
                (pack.get('workspace') or {}).get('name')
                or pack.get('workspace_name')
                or workspace_name
            )
            if workspace_name == 'Unknown Workspace' and CATALOG_AVAILABLE:
                try:
                    cat = catalog_service.get_workspace_catalog() or {}
                    ws = next(
                        (w for w in (cat.get('workspaces') or []) if w.get('id') == workspace_id),
                        None,
                    )
                    if ws:
                        workspace_name = ws.get('name') or workspace_name
                except Exception:
                    pass
            reports_in = list(pack.get('reports') or [])
        else:
            return jsonify({
                'success': False,
                'error': 'No reports found in catalog for this workspace.',
            }), 404

        for r in reports_in:
            rid = r.get('id')
            if rid and r.get('view_count') is not None:
                try:
                    report_views[rid] = int(r.get('view_count') or 0)
                except Exception:
                    report_views[rid] = 0

        data_as_of_raw = None
        if CATALOG_AVAILABLE:
            try:
                usage_snap = catalog_service.get_json('usage_snapshot.json') or {}
                data_as_of_raw = (
                    usage_snap.get('generatedAt')
                    or usage_snap.get('opsEnrichedAt')
                    or usage_snap.get('asOf')
                )
                for rid, cnt in (usage_snap.get('report_views') or {}).items():
                    report_views[rid] = int(cnt or report_views.get(rid) or 0)
            except Exception:
                pass
        if not data_as_of_raw and pack:
            data_as_of_raw = (
                (pack.get('catalog_meta') or {}).get('opsEnrichedAt')
                or pack.get('opsEnrichedAt')
                or pack.get('generatedAt')
            )
        data_as_of_dt = _parse_dt(data_as_of_raw) or datetime.now(timezone.utc)

        # Folder map for archive path (same idea as single Download)
        folder_path_by_report = {}
        report_folder_map = {}
        folder_names_map = {}
        try:
            folder_meta = _fetch_workspace_folder_meta(workspace_id) or {}
            report_folder_map = folder_meta.get('report_folder_map') or {}
            folder_names_map = folder_meta.get('folder_names_map') or {}
            for rid, fid in report_folder_map.items():
                info = folder_names_map.get(fid) or {}
                nm = info.get('name')
                if nm:
                    folder_path_by_report[rid] = nm
        except Exception as folder_err:
            print(f"   ⚠️ folder map for decomm candidates skipped: {folder_err}")

        candidates = []
        for report in reports_in:
            name = report.get('name') or 'Unknown'
            if _is_excluded_report_name(name):
                continue
            rid = report.get('id') or ''
            if not rid:
                continue
            live_dq = _is_live_or_dq(report)
            days = report.get('days_since_refresh')
            try:
                days = int(days) if days is not None and days != '' else None
            except Exception:
                days = None
            last_ref = report.get('last_refreshed') or report.get('lastRefresh') or ''
            days_as_of = _days_since(last_ref, data_as_of_dt) if last_ref else None
            if days_as_of is not None:
                days = days_as_of
            elif days is None and last_ref:
                days = _days_since(last_ref)

            refresh_status = report.get('last_refresh_status') or report.get('refreshStatus') or ''
            if rid in report_views:
                views = int(report_views.get(rid) or 0)
                views_known = True
            elif report.get('view_count') is not None:
                views = int(report.get('view_count') or 0)
                views_known = True
            else:
                views = 0
                views_known = False

            reasons = []
            if days is not None and days > REFRESH_STALE_DAYS:
                reasons.append(f'Days since refresh > {REFRESH_STALE_DAYS} ({days}d)')
            no_history = False
            if not live_dq:
                if (not last_ref) and (days is None):
                    no_history = True
                elif str(refresh_status).lower() in {'error', 'failed'} and not last_ref:
                    no_history = True
            if no_history:
                reasons.append('No refresh history')
            if (views_known or report.get('view_count') is not None) and int(views or 0) == 0:
                reasons.append(f'0 views in last {VIEW_LOOKBACK_LABEL}')

            # Bulk SharePoint archive: original rules only (stale / no history / 0 views).
            # Verify-in-PBI is Excel decommission-list only — not auto-archived here.

            if not reasons:
                continue

            folder_id = (
                report.get('folderId')
                or report.get('folderObjectId')
                or report_folder_map.get(rid)
                or None
            )
            if folder_id in ('__ROOT__', '', None):
                folder_id = None
            folder_name = (
                report.get('folderName')
                or report.get('folder_name')
                or folder_path_by_report.get(rid)
                or None
            )
            if not folder_id:
                folder_name = None

            candidates.append({
                'report_id': rid,
                'report_name': name,
                'folder_id': folder_id,
                'folder_name': folder_name,
                'reasons': reasons,
            })

        candidates.sort(key=lambda x: (x.get('report_name') or '').lower())

        print(
            f"📋 decomm candidates ws={workspace_name!r} "
            f"count={len(candidates)}/{len(reports_in)}"
        )
        return jsonify({
            'success': True,
            'workspace_id': workspace_id,
            'workspace_name': workspace_name,
            'count': len(candidates),
            'candidates': candidates,
        })

    except Exception as e:
        print(f"❌ Error listing decommission candidates: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500


def _usage_cache_path(workspace_id: str) -> str:
    """Disk path for per-workspace Activity Events day cache."""
    root = (os.getenv("USAGE_CACHE_DIR") or "").strip()
    if not root:
        # Prefer durable App Service path; fall back to repo-local .usage_cache
        if os.getenv("WEBSITE_HOSTNAME"):
            root = os.path.join("/home", "data", "usage_cache")
        else:
            root = os.path.join(os.getcwd(), ".usage_cache")
    return os.path.join(root, f"usage_{workspace_id}.json")


def _aggregate_usage_days(daily_data: dict):
    """Sum report_views / last_viewed across day buckets."""
    report_views = {}
    last_viewed = {}
    for _date_key, day_data in (daily_data or {}).items():
        if not isinstance(day_data, dict):
            continue
        for report_id, count in (day_data.get("report_views") or {}).items():
            try:
                report_views[report_id] = report_views.get(report_id, 0) + int(count or 0)
            except (TypeError, ValueError):
                continue
        for report_id, view_info in (day_data.get("last_viewed") or {}).items():
            if not isinstance(view_info, dict):
                continue
            ts = view_info.get("timestamp") or ""
            prev = last_viewed.get(report_id)
            if not prev or ts > (prev.get("timestamp") or ""):
                last_viewed[report_id] = view_info
    return report_views, last_viewed


def _usage_from_catalog_snapshot(workspace_id: str = None):
    """
    Fast path: precomputed usage_snapshot.json (ops extract).
    Tenant-wide report_views / last_viewed — optionally filtered to workspace
    report ids when known. Returns None when snapshot unavailable.
    """
    if not CATALOG_AVAILABLE or catalog_service is None:
        return None
    try:
        snap = catalog_service.get_json("usage_snapshot.json")
    except Exception as exc:
        print(f"   ⚠️ Catalog usage_snapshot unavailable: {exc}")
        return None
    if not isinstance(snap, dict):
        return None
    report_views = snap.get("report_views") if isinstance(snap.get("report_views"), dict) else {}
    last_viewed = snap.get("last_viewed") if isinstance(snap.get("last_viewed"), dict) else {}
    if not report_views and not last_viewed:
        return None

    # Optionally narrow to reports present in this workspace (catalog or thin pack)
    ws_report_ids = None
    if workspace_id:
        try:
            pack = catalog_service.get_workspace_reports(workspace_id)
            if pack and isinstance(pack.get("reports"), list):
                ws_report_ids = {
                    str(r.get("id") or r.get("reportId") or "").lower()
                    for r in pack["reports"]
                    if (r.get("id") or r.get("reportId"))
                }
                ws_report_ids.discard("")
        except Exception:
            ws_report_ids = None

    if ws_report_ids:
        # Keep original casing from snapshot keys
        report_views = {
            rid: cnt for rid, cnt in report_views.items()
            if str(rid).lower() in ws_report_ids
        }
        last_viewed = {
            rid: info for rid, info in last_viewed.items()
            if str(rid).lower() in ws_report_ids
        }

    days = int(snap.get("lookbackDays") or os.getenv("USAGE_LOOKBACK_DAYS", "60") or 60)
    return {
        "success": True,
        "workspace_id": workspace_id,
        "days_analyzed": days,
        "report_views": report_views,
        "last_viewed": last_viewed,
        "source": "catalog_usage_snapshot",
        "generatedAt": snap.get("generatedAt"),
        "note": (
            f"Catalog usage snapshot"
            f"{(' @ ' + str(snap.get('generatedAt'))) if snap.get('generatedAt') else ''}"
            f" ({len(report_views)} reports with views)"
        ),
    }


@app.route('/api/report-usage/<workspace_id>')
@login_required
def get_report_usage(workspace_id):
    """
    Get report usage metrics (view counts) for the last 60 days.

    Priority:
      1) Catalog usage_snapshot.json (ops extract — fast, reliable)
      2) Non-empty per-workspace Activity Events day cache (< 24h)
      3) Live Admin Activity Events fetch (service principal)
      4) Stale / empty cache last resort
    """
    try:
        from datetime import datetime, timedelta, timezone
        import requests
        from scanner_connector import PowerBIScanner

        print(f"\n📊 FETCHING REPORT USAGE METRICS FOR WORKSPACE: {workspace_id}")
        print(f"   ⚠️  NOTE: Live Activity Events requires Service Principal with Power BI Admin role")

        days_back = int(os.getenv("USAGE_LOOKBACK_DAYS", "60"))  # default 60-day views
        max_workers = 30
        force_live = str(request.args.get("force") or request.args.get("refresh") or "").strip().lower() in (
            "1", "true", "yes", "live"
        )

        # --- 1) Catalog snapshot first (instant, pre-aggregated by ops job) ---
        if not force_live:
            cat_usage = _usage_from_catalog_snapshot(workspace_id)
            if cat_usage and (
                cat_usage.get("report_views") or cat_usage.get("last_viewed")
            ):
                print(
                    f"   ✅ Using catalog usage_snapshot "
                    f"({len(cat_usage.get('report_views') or {})} reports, "
                    f"generatedAt={cat_usage.get('generatedAt')})"
                )
                return jsonify(cat_usage)

        # === PERSISTENT CACHE (Activity Events incremental) ===
        cache_file = _usage_cache_path(workspace_id)
        persistent_cache = {'daily_data': {}, 'last_updated': None}

        if os.path.exists(cache_file):
            try:
                with open(cache_file, 'r', encoding='utf-8') as f:
                    persistent_cache = json.load(f)
                    print(f"   📂 Loaded persistent cache from disk: {cache_file}")
            except Exception as e:
                print(f"   ⚠️ Cache load error: {e}")
        else:
            print(f"   🔍 No persistent cache found at {cache_file}")

        cached_dates = set((persistent_cache.get('daily_data') or {}).keys())
        all_dates_needed = [
            (datetime.now(timezone.utc) - timedelta(days=i)).strftime('%Y-%m-%d')
            for i in range(days_back)
        ]
        missing_dates = [d for d in all_dates_needed if d not in cached_dates]

        cached_views, cached_last = _aggregate_usage_days(persistent_cache.get('daily_data') or {})
        cached_total_views = sum(cached_views.values()) if cached_views else 0
        print(
            f"   📅 Period: {days_back} days | Cached days: {len(cached_dates)} | "
            f"Missing: {len(missing_dates)} | Cached views: {cached_total_views}"
        )

        # Only short-circuit on cache when it actually has view data (avoid empty poison).
        last_updated = persistent_cache.get('last_updated')
        if (
            not force_live
            and last_updated
            and len(missing_dates) == 0
            and cached_total_views > 0
        ):
            try:
                last_updated_dt = datetime.fromisoformat(last_updated)
                if last_updated_dt.tzinfo is None:
                    last_updated_dt = last_updated_dt.replace(tzinfo=timezone.utc)
                cache_age_hours = (datetime.now(timezone.utc) - last_updated_dt).total_seconds() / 3600
            except Exception:
                cache_age_hours = 999

            if cache_age_hours < 24:
                print(
                    f"   ✅ Using non-empty Activity cache "
                    f"(age: {cache_age_hours:.1f}h, views={cached_total_views})"
                )
                return jsonify({
                    'success': True,
                    'workspace_id': workspace_id,
                    'days_analyzed': days_back,
                    'report_views': cached_views,
                    'last_viewed': cached_last,
                    'source': 'activity_cache',
                    'note': f'Loaded from cache (age: {cache_age_hours:.1f}h, views={cached_total_views})'
                })

        # Empty full cache is useless — drop it so we re-fetch rather than serving zeros forever
        if len(missing_dates) == 0 and cached_total_views == 0 and cached_dates:
            print("   🗑️ Discarding empty Activity day-cache (0 total views) — will re-fetch")
            missing_dates = list(all_dates_needed)
            persistent_cache = {'daily_data': {}, 'last_updated': None}
            cached_dates = set()
            try:
                if os.path.exists(cache_file):
                    os.remove(cache_file)
            except Exception:
                pass

        # Only fetch missing days (incremental update)
        print(f"   🚀 Fetching {len(missing_dates)} missing Activity days...")

        report_views = {}  # Dictionary to store view counts per report_id

        # Use service principal token (Scanner API connector) instead of user token
        # Activity Events API requires service principal authentication
        scanner = PowerBIScanner()
        service_principal_token = scanner.get_access_token()

        if not service_principal_token:
            # Fall back to whatever non-empty cache / catalog we still have
            if cached_total_views > 0:
                return jsonify({
                    'success': True,
                    'workspace_id': workspace_id,
                    'days_analyzed': days_back,
                    'report_views': cached_views,
                    'last_viewed': cached_last,
                    'source': 'activity_cache_stale',
                    'note': 'Service principal token unavailable; serving Activity cache',
                })
            cat_usage = _usage_from_catalog_snapshot(workspace_id)
            if cat_usage:
                return jsonify(cat_usage)
            return jsonify({
                'success': False,
                'error': 'Unable to obtain service principal token for Activity Events API'
            }), 500

        headers = {
            'Authorization': f'Bearer {service_principal_token}',
            'Content-Type': 'application/json'
        }
        base_url = "https://api.powerbi.com/v1.0/myorg/admin/activityevents"

        # FIX FOR SYSTEM DATE ISSUE: Get real-world current date from Power BI API
        # System date is 2026, but we need the actual current date
        # Make a lightweight API call to get the server's current date from response headers
        from datetime import timezone
        import dateutil.parser

        try:
            # Make a simple GET request to Power BI to get server date from headers
            test_url = f"https://api.powerbi.com/v1.0/myorg/groups/{workspace_id}"
            test_response = requests.get(test_url, headers=headers, timeout=10)
            server_date_str = test_response.headers.get('Date')

            if server_date_str:
                # Parse the HTTP Date header to get current server time
                server_date = dateutil.parser.parse(server_date_str)
                # Use yesterday as end date (Activity Events API only has complete days)
                end_date = (server_date - timedelta(days=1)).replace(hour=0, minute=0, second=0, microsecond=0, tzinfo=timezone.utc)
                print(f"   ✅ Retrieved real-world date from Power BI server: {server_date_str}")
                print(f"   📅 Using end date: {end_date.strftime('%Y-%m-%d')} (yesterday from server time)")
            else:
                # Fallback: use utcnow() - timedelta(days=1) if header not available
                end_date = (datetime.utcnow() - timedelta(days=1)).replace(hour=0, minute=0, second=0, microsecond=0, tzinfo=timezone.utc)
                print(f"   ⚠️  Could not get server date, using UTC now: {end_date.strftime('%Y-%m-%d')}")
        except Exception as e:
            # Final fallback
            end_date = (datetime.utcnow() - timedelta(days=1)).replace(hour=0, minute=0, second=0, microsecond=0, tzinfo=timezone.utc)
            print(f"   ⚠️  Error getting server date ({str(e)}), using UTC now: {end_date.strftime('%Y-%m-%d')}")

        print(f"   ⚡ MAXIMUM SPEED MODE: {max_workers} parallel workers with intelligent retry logic for {days_back} days...")

        # Import parallel processing tools
        from concurrent.futures import ThreadPoolExecutor, as_completed
        from threading import Lock

        # Thread-safe lock for updating shared dictionary
        views_lock = Lock()
        last_viewed = {}  # Track last viewed info per report

        # Define hex checker ONCE outside the loop (major performance gain)
        def is_hex_identifier(value):
            """
            Check if value is a hex ID or GUID rather than email/username
            Returns True if value looks like a user ID (hex or GUID) that should be rejected
            Returns False if value looks like a human-readable identifier (email, UPN, name)
            """
            if not value:
                return True  # Empty = invalid, treat as ID

            # Valid human identifiers contain @ or spaces
            if '@' in value or ' ' in value:
                return False  # Has @ or space = email/name, keep it

            # GUID format: 8-4-4-4-12 characters with dashes (e.g., 02453143-dc86-4125-a211-64162fc93005)
            if '-' in value and len(value) >= 32:
                # Check if it matches GUID pattern
                parts = value.split('-')
                if len(parts) == 5 and all(len(p) in [8, 4, 4, 4, 12] for p in parts):
                    return True  # GUID format = reject

            # Hex ID (no dashes, all hex chars, long string like 100FFF92C7717B)
            if len(value) >= 12 and '.' not in value:
                hex_chars = set('0123456789ABCDEFabcdef')
                if all(c in hex_chars for c in value):
                    return True  # Pure hex = reject

            # Default: keep it (might be username or other readable format)
            return False

        def _resolve_viewer(activity):
            """Human-readable viewer identity (prefer email/UPN over hex UserKey)."""
            user_id = activity.get('UserId')
            if user_id and not is_hex_identifier(str(user_id)):
                return str(user_id)
            return (
                activity.get('UserPrincipalName')
                or activity.get('UserEmail')
                or activity.get('User')
                or 'Unknown User'
            )

        def _ingest_view_activity(activity, day_views, day_last_viewed):
            """Count a ViewReport for this workspace (crash-test compatible)."""
            act = (activity.get('Activity') or '').strip()
            if act and act != 'ViewReport':
                return
            ws = activity.get('WorkspaceId') or activity.get('WorkSpaceId')
            if ws and str(ws).lower() != str(workspace_id).lower():
                return
            report_id = activity.get('ReportId') or activity.get('ArtifactId')
            if not report_id:
                return
            report_id = str(report_id)
            activity_time = activity.get('CreationTime') or ''
            user_identifier = _resolve_viewer(activity)
            # Skip known service/admin identity from counts (existing product rule)
            if str(user_identifier).lower() == 'admin-rsteinke@ashleyfurniture.com':
                return
            day_views[report_id] = day_views.get(report_id, 0) + 1
            if activity_time:
                prev = day_last_viewed.get(report_id)
                if not prev or activity_time > (prev.get('timestamp') or ''):
                    day_last_viewed[report_id] = {
                        'timestamp': activity_time,
                        'user': user_identifier,
                    }

        def fetch_day_activities(day_offset):
            """
            Fetch ViewReport activities for a single day.
            Aligned with PowerBI-Crash-Test (working in tenant):
              - plain ISO times without trailing Z
              - $filter=Activity eq 'ViewReport'
              - pagination via continuationUri (full URL)
            """
            import time
            import random

            current_date = end_date - timedelta(days=day_offset)
            day_key = current_date.strftime('%Y-%m-%d')
            # No .000Z — matches working crash-test Activity Events calls
            start_datetime = f"{day_key}T00:00:00"
            end_datetime = f"{day_key}T23:59:59"
            url = (
                f"{base_url}"
                f"?startDateTime='{start_datetime}'&endDateTime='{end_datetime}'"
                f"&$filter=Activity eq 'ViewReport'"
            )

            day_views = {}
            day_last_viewed = {}
            pages = 0
            max_pages = 100
            max_retries = 3
            base_delay = 0.5

            while url and pages < max_pages:
                response = None
                for attempt in range(max_retries):
                    try:
                        if attempt == 0 and pages == 0:
                            time.sleep(random.uniform(0.01, 0.02))
                        response = requests.get(url, headers=headers, timeout=60)
                        if response.status_code == 400:
                            return (day_offset, {}, {}, 0, "HTTP 400 (no data)")
                        if response.status_code == 429:
                            if attempt < max_retries - 1:
                                max_delay = min(10, base_delay * (2 ** attempt))
                                time.sleep(random.uniform(base_delay, max_delay))
                                continue
                            return (day_offset, {}, {}, 0, "HTTP 429")
                        break
                    except requests.Timeout:
                        if attempt < max_retries - 1:
                            time.sleep(random.uniform(0.5, 1.0))
                            continue
                        return (day_offset, {}, {}, 0, "Timeout")
                    except Exception as e:
                        if attempt < max_retries - 1 and ('Connection' in str(e) or '10054' in str(e)):
                            time.sleep(random.uniform(1.0, 2.0))
                            continue
                        return (day_offset, {}, {}, 0, f"{str(e)[:50]}")

                if response is None or response.status_code != 200:
                    status = getattr(response, 'status_code', 'n/a')
                    if pages == 0:
                        return (day_offset, {}, {}, 0, f"HTTP {status}")
                    break

                pages += 1
                try:
                    data = response.json() or {}
                except Exception:
                    break

                for activity in data.get('activityEventEntities') or []:
                    _ingest_view_activity(activity, day_views, day_last_viewed)

                cont_uri = data.get('continuationUri')
                cont_tok = data.get('continuationToken')
                if cont_uri:
                    url = cont_uri
                elif cont_tok:
                    tok = str(cont_tok).strip().strip("'")
                    url = f"{base_url}?continuationToken='{tok}'"
                else:
                    url = None

            # (day_offset, day_views, day_last_viewed, event_count, error)
            event_count = sum(day_views.values())
            return (day_offset, day_views, day_last_viewed, event_count, None)

        # Fetch only missing days in parallel (incremental fetching)
        # Convert missing date strings to day offsets
        date_to_offset = {}
        for date_str in missing_dates:
            target_date = datetime.strptime(date_str, '%Y-%m-%d').replace(tzinfo=timezone.utc)
            day_offset = (end_date - target_date).days
            date_to_offset[day_offset] = date_str

        new_daily_data = {}  # Store newly fetched data

        if len(missing_dates) > 0:
            with ThreadPoolExecutor(max_workers=min(len(missing_dates), max_workers)) as executor:
                futures = [executor.submit(fetch_day_activities, day_offset) for day_offset in date_to_offset.keys()]

                for future in as_completed(futures):
                    day_offset, day_views, day_last_viewed, activity_count, error = future.result()
                    current_date = end_date - timedelta(days=day_offset)
                    date_str = current_date.strftime('%Y-%m-%d')

                    if error:
                        # Do NOT cache failed days — empty error days poison the 24h short-circuit
                        if "429" in str(error):
                            print(f"      ⚠️  {date_str}: Rate limit - {error}")
                        elif "400" not in str(error) and error != "HTTP 200":
                            print(f"      ⚠️  {date_str}: {error}")
                        continue

                    if day_views:
                        print(f"      ✅ {date_str}: {sum(day_views.values())} views")

                    # Successful API day (may legitimately have 0 views)
                    with views_lock:
                        new_daily_data[date_str] = {
                            'report_views': day_views or {},
                            'last_viewed': day_last_viewed or {},
                        }

        # Merge only successfully fetched days into cache
        if 'daily_data' not in persistent_cache or not isinstance(persistent_cache.get('daily_data'), dict):
            persistent_cache['daily_data'] = {}
        persistent_cache['daily_data'].update(new_daily_data)

        # Aggregate from ALL data (cached + new)
        report_views, last_viewed = _aggregate_usage_days(persistent_cache.get('daily_data') or {})

        print(f"\n✅ Report usage summary:")
        print(f"   Total reports with views: {len(report_views)}")
        print(f"   Total views across all reports: {sum(report_views.values()) if report_views else 0}")

        # Debug: Show sample of last viewed users to verify email resolution
        if last_viewed:
            sample_users = list(set([info.get('user', '') for info in last_viewed.values() if isinstance(info, dict)]))[:5]
            print(f"   📧 Sample users: {', '.join(str(u) for u in sample_users if u)}")

        total_views = sum(report_views.values()) if report_views else 0

        # If live Activity Events produced nothing, fall back to catalog snapshot
        if total_views == 0:
            cat_usage = _usage_from_catalog_snapshot(workspace_id)
            if cat_usage and (cat_usage.get("report_views") or cat_usage.get("last_viewed")):
                print(
                    f"   ↩️ Live Activity returned 0 views — serving catalog usage_snapshot "
                    f"({len(cat_usage.get('report_views') or {})} reports)"
                )
                return jsonify(cat_usage)

        result = {
            'success': True,
            'workspace_id': workspace_id,
            'days_analyzed': days_back,
            'report_views': report_views,
            'last_viewed': last_viewed,
            'source': 'activity_events',
            'note': (
                f'Activity data: {days_back} days '
                f'(fetched {len(new_daily_data)} new, '
                f'had {len(cached_dates)} cached days)'
            ),
        }

        # Save cache only when we have successful day data (avoid empty poison files)
        if persistent_cache.get('daily_data'):
            persistent_cache['last_updated'] = datetime.now(timezone.utc).isoformat()
            persistent_cache['workspace_id'] = workspace_id
            try:
                os.makedirs(os.path.dirname(cache_file) or ".", exist_ok=True)
                with open(cache_file, 'w', encoding='utf-8') as f:
                    json.dump(persistent_cache, f, indent=2)
                print(f"   💾 Persistent cache saved: {cache_file}")
            except Exception as e:
                print(f"   ⚠️ Cache save error: {e}")
        else:
            print("   ℹ️ Skipping cache save (no successful Activity day buckets)")

        return jsonify(result)

    except Exception as e:
        print(f"❌ Error fetching report usage: {str(e)}")
        import traceback
        traceback.print_exc()
        # Last-chance catalog fallback on unexpected errors
        try:
            cat_usage = _usage_from_catalog_snapshot(workspace_id)
            if cat_usage:
                cat_usage["note"] = (cat_usage.get("note") or "") + f" (fallback after error: {e})"
                return jsonify(cat_usage)
        except Exception:
            pass
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/orphaned-reports/<workspace_id>')
@login_required
def get_orphaned_reports(workspace_id):
    """
    Identify ownership gaps for reports in a workspace.

    Why fields used to show "Unknown":
      - Old path only read Scanner keys createdBy/modifiedBy/modifiedDateTime.
      - Tenant Scanner often omits those on reports (especially without user graph data),
        so every row became "Unknown" and almost everything looked orphaned.
      - Catalog normalizer also dropped creator/modifier; it keeps dataset_owner only.

    Fix:
      1) Prefer Fast catalog list (names + dataset_owner / configuredBy).
      2) Overlay Groups REST /reports which usually has createdBy, modifiedBy, dates.
      3) Optional lightweight Scanner fallback for remaining gaps.
      4) Treat as orphaned only when NO usable owner signal remains
         (createdBy / modifiedBy / dataset owner / configuredBy).
    """
    try:
        from datetime import datetime, timezone
        import requests

        print(f"\n👻 DETECTING ORPHANED REPORTS FOR WORKSPACE: {workspace_id}")

        def _clean_person(value):
            if value is None:
                return ""
            if isinstance(value, dict):
                value = (
                    value.get("userPrincipalName")
                    or value.get("emailAddress")
                    or value.get("displayName")
                    or value.get("identifier")
                    or value.get("id")
                    or ""
                )
            s = str(value).strip()
            if not s or s.lower() in {"unknown", "n/a", "none", "null", "-"}:
                return ""
            return s

        def _pick(*vals):
            for v in vals:
                c = _clean_person(v)
                if c:
                    return c
            return ""

        def _parse_dt(value):
            if not value or value in ("Unknown", "N/A", "-", None):
                return None
            try:
                if isinstance(value, (int, float)):
                    # epoch seconds / ms
                    ts = float(value)
                    if ts > 1e12:
                        ts /= 1000.0
                    return datetime.fromtimestamp(ts, tz=timezone.utc)
                s = str(value).strip()
                if s.endswith("Z"):
                    s = s[:-1] + "+00:00"
                return datetime.fromisoformat(s)
            except Exception:
                return None

        def _age_days_from(*date_vals):
            for raw in date_vals:
                dt = _parse_dt(raw)
                if dt is None:
                    continue
                if dt.tzinfo is None:
                    dt = dt.replace(tzinfo=timezone.utc)
                return max(0, (datetime.now(timezone.utc) - dt).days)
            return None

        # ---- 1) Catalog (fast structural + dataset owner) ----
        by_id = {}
        catalog_meta = None
        pack = None
        try:
            if CATALOG_AVAILABLE:
                pack = catalog_service.get_workspace_reports(workspace_id)
        except Exception as e:
            print(f"   ⚠️ catalog get_workspace_reports failed: {e}")
            pack = None

        if pack:
            catalog_meta = pack.get("catalog_meta") or {
                "generatedAt": pack.get("generatedAt"),
                "source": pack.get("source") or "catalog",
            }
            datasets_map = pack.get("datasets_map") or {}
            for r in pack.get("reports") or []:
                rid = r.get("id")
                if not rid:
                    continue
                name = r.get("name") or "Unknown"
                if _is_excluded_report_name(name):
                    continue
                ds_id = r.get("datasetId") or ""
                ds = datasets_map.get(ds_id) or {}
                dataset_owner = _pick(
                    r.get("dataset_owner"),
                    r.get("configuredBy"),
                    ds.get("configuredBy"),
                    ds.get("configured_by"),
                    ds.get("owner"),
                )
                by_id[rid] = {
                    "report_id": rid,
                    "report_name": name,
                    "dataset_id": ds_id,
                    "dataset_name": r.get("datasetName") or ds.get("name") or "",
                    "created_by": _pick(r.get("createdBy"), r.get("created_by")),
                    "modified_by": _pick(r.get("modifiedBy"), r.get("modified_by")),
                    "created_date": r.get("createdDateTime") or r.get("created_date") or "",
                    "modified_date": r.get("modifiedDateTime") or r.get("modified_date") or "",
                    "dataset_owner": dataset_owner,
                    "last_refreshed": r.get("last_refreshed") or ds.get("last_refreshed") or "",
                    "source": "catalog",
                }
            print(f"   📦 Catalog reports: {len(by_id)}")

        # ---- 2) Groups REST reports (best source for created/modified + dates) ----
        rest_count = 0
        try:
            headers = get_user_powerbi_headers()
            url = f"https://api.powerbi.com/v1.0/myorg/groups/{workspace_id}/reports"
            resp = requests.get(url, headers=headers, timeout=60)
            if resp.status_code == 200:
                for r in resp.json().get("value") or []:
                    rid = r.get("id")
                    name = r.get("name") or "Unknown"
                    if not rid or _is_excluded_report_name(name):
                        continue
                    row = by_id.get(rid) or {
                        "report_id": rid,
                        "report_name": name,
                        "dataset_id": r.get("datasetId") or "",
                        "dataset_name": "",
                        "created_by": "",
                        "modified_by": "",
                        "created_date": "",
                        "modified_date": "",
                        "dataset_owner": "",
                        "last_refreshed": "",
                        "source": "rest",
                    }
                    row["report_name"] = name
                    row["dataset_id"] = r.get("datasetId") or row.get("dataset_id") or ""
                    row["created_by"] = _pick(
                        r.get("createdBy"),
                        r.get("createdByUser"),
                        row.get("created_by"),
                    )
                    row["modified_by"] = _pick(
                        r.get("modifiedBy"),
                        r.get("modifiedByUser"),
                        row.get("modified_by"),
                    )
                    row["created_date"] = (
                        r.get("createdDateTime")
                        or r.get("createdDate")
                        or row.get("created_date")
                        or ""
                    )
                    row["modified_date"] = (
                        r.get("modifiedDateTime")
                        or r.get("modifiedDate")
                        or row.get("modified_date")
                        or ""
                    )
                    row["source"] = (row.get("source") or "") + "+rest"
                    by_id[rid] = row
                    rest_count += 1
                print(f"   🌐 REST reports overlaid: {rest_count}")
            else:
                print(f"   ⚠️ REST reports HTTP {resp.status_code}: {resp.text[:200]}")
        except Exception as e:
            print(f"   ⚠️ REST reports overlay failed: {e}")

        # ---- 3) Scanner fallback only if we still lack people/dates for many rows ----
        missing_meta = sum(
            1
            for row in by_id.values()
            if not row.get("created_by") and not row.get("modified_by") and not row.get("modified_date")
        )
        if missing_meta and (missing_meta >= max(3, len(by_id) // 3) or not by_id):
            try:
                from scanner_connector import PowerBIScanner

                print(f"   🛰️ Scanner fallback ({missing_meta} rows missing people/dates)…")
                scanner = PowerBIScanner()
                scan_result = scanner.run_scan(workspace_id=workspace_id) or {}
                for ws in scan_result.get("workspaces") or []:
                    if ws.get("id") != workspace_id:
                        continue
                    ds_owner_by_id = {}
                    for d in ws.get("datasets") or []:
                        did = d.get("id")
                        if did:
                            ds_owner_by_id[did] = _pick(
                                d.get("configuredBy"),
                                d.get("configured_by"),
                                d.get("owner"),
                            )
                    for r in ws.get("reports") or []:
                        rid = r.get("id")
                        name = r.get("name") or "Unknown"
                        if not rid or _is_excluded_report_name(name):
                            continue
                        row = by_id.get(rid) or {
                            "report_id": rid,
                            "report_name": name,
                            "dataset_id": r.get("datasetId") or "",
                            "dataset_name": "",
                            "created_by": "",
                            "modified_by": "",
                            "created_date": "",
                            "modified_date": "",
                            "dataset_owner": "",
                            "last_refreshed": "",
                            "source": "scanner",
                        }
                        row["created_by"] = _pick(
                            row.get("created_by"),
                            r.get("createdBy"),
                            r.get("createdByUserPrincipalName"),
                        )
                        row["modified_by"] = _pick(
                            row.get("modified_by"),
                            r.get("modifiedBy"),
                            r.get("modifiedByUserPrincipalName"),
                        )
                        row["created_date"] = (
                            row.get("created_date")
                            or r.get("createdDateTime")
                            or r.get("createdDate")
                            or ""
                        )
                        row["modified_date"] = (
                            row.get("modified_date")
                            or r.get("modifiedDateTime")
                            or r.get("modifiedDate")
                            or ""
                        )
                        ds_id = row.get("dataset_id") or r.get("datasetId") or ""
                        row["dataset_id"] = ds_id
                        if not row.get("dataset_owner"):
                            row["dataset_owner"] = ds_owner_by_id.get(ds_id, "")
                        row["source"] = (row.get("source") or "") + "+scanner"
                        by_id[rid] = row
                    break
            except Exception as e:
                print(f"   ⚠️ Scanner fallback failed: {e}")

        if not by_id:
            return jsonify({
                "success": False,
                "error": "No reports found for workspace (catalog + REST + scanner empty)",
            }), 404

        # ---- Classify orphaned ----
        orphaned_reports = []
        total_reports = 0
        for row in sorted(by_id.values(), key=lambda x: (x.get("report_name") or "").lower()):
            total_reports += 1
            created_by = _clean_person(row.get("created_by"))
            modified_by = _clean_person(row.get("modified_by"))
            dataset_owner = _clean_person(row.get("dataset_owner"))
            created_date = row.get("created_date") or ""
            modified_date = row.get("modified_date") or ""

            age_days = _age_days_from(
                modified_date,
                created_date,
                row.get("last_refreshed"),
            )

            orphan_reason = []
            if not created_by:
                orphan_reason.append("No report creator")
            if not modified_by:
                orphan_reason.append("No report modifier")
            if not dataset_owner:
                orphan_reason.append("No dataset owner")

            # True orphan only when there is NO ownership signal at all
            has_any_owner = bool(created_by or modified_by or dataset_owner)
            is_orphaned = not has_any_owner
            if is_orphaned:
                orphan_reason = ["No creator, modifier, or dataset owner"]

            risk_level = "Low"
            if is_orphaned:
                if age_days is not None and age_days > 365:
                    risk_level = "Critical"
                elif age_days is not None and age_days > 180:
                    risk_level = "High"
                elif age_days is not None and age_days > 90:
                    risk_level = "Medium"
                else:
                    risk_level = "High" if age_days is None else "Medium"
            elif not created_by and not modified_by:
                # has dataset owner only — ownership gap, not full orphan
                risk_level = "Low"
                if age_days is not None and age_days > 365:
                    risk_level = "Medium"

            # Surface orphans + ownership gaps (missing creator/modifier)
            if is_orphaned or not created_by or not modified_by:
                orphaned_reports.append({
                    "report_id": row.get("report_id"),
                    "report_name": row.get("report_name") or "Unknown",
                    "created_by": created_by or "—",
                    "modified_by": modified_by or "—",
                    "dataset_owner": dataset_owner or "—",
                    "created_date": created_date or "—",
                    "modified_date": modified_date or "—",
                    "age_days": age_days,
                    "is_orphaned": is_orphaned,
                    "orphan_reasons": orphan_reason,
                    "risk_level": risk_level,
                    "source": row.get("source") or "",
                })

        orphaned_count = len([r for r in orphaned_reports if r["is_orphaned"]])
        # Keep UI KPI meaningful: "orphaned" = no owner at all; table may also list gaps
        gap_count = len(orphaned_reports)
        orphaned_percentage = round((orphaned_count / total_reports) * 100) if total_reports else 0

        critical_count = len([r for r in orphaned_reports if r["risk_level"] == "Critical"])
        high_count = len([r for r in orphaned_reports if r["risk_level"] == "High"])
        medium_count = len([r for r in orphaned_reports if r["risk_level"] == "Medium"])

        print(
            f"   ✅ total={total_reports} true_orphans={orphaned_count} "
            f"rows_listed={gap_count} critical={critical_count}"
        )

        return jsonify({
            "success": True,
            "workspace_id": workspace_id,
            "total_reports": total_reports,
            "orphaned_count": orphaned_count,
            "orphaned_percentage": orphaned_percentage,
            "listed_count": gap_count,
            "reports": orphaned_reports,
            "risk_summary": {
                "critical": critical_count,
                "high": high_count,
                "medium": medium_count,
                "low": len(orphaned_reports) - critical_count - high_count - medium_count,
            },
            "catalog_meta": catalog_meta,
            "note": (
                "Unowned = no creator, modifier, or dataset owner signal. "
                "Table also lists ownership gaps (missing creator/modifier)."
            ),
        })

    except Exception as e:
        print(f"Error detecting unowned reports: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/similarity-analysis')
@login_required
def get_similarity_analysis():
    """
    Multi-type similarity analysis.
    Query params:
      - workspace_id (required)
      - type: reports | visuals | measures | tables
      - threshold: 0.0-1.0 (default 0.3)

    Data sources (same family as Semantic Models):
      - tables / report structure: catalog (columns fast)
      - DAX measures: same workspace Scanner enrich cache used by
        /api/semantic-model-details?enrich=1  (_get_workspace_scan_cached)
      - visuals: full Scanner workspace payload (pages/visuals) when available
    """
    try:
        workspace_id = request.args.get('workspace_id')
        analysis_type = (request.args.get('type') or 'reports').strip().lower()
        try:
            threshold = float(request.args.get('threshold', 0.3))
        except (TypeError, ValueError):
            threshold = 0.3
        threshold = max(0.0, min(1.0, threshold))

        if not workspace_id:
            return jsonify({'success': False, 'error': 'workspace_id is required'}), 400
        if analysis_type not in {'reports', 'visuals', 'measures', 'tables'}:
            return jsonify({
                'success': False,
                'error': f"Unsupported type '{analysis_type}'. Use reports|visuals|measures|tables"
            }), 400

        print(f"\n🔍 SIMILARITY ANALYSIS: type={analysis_type} ws={workspace_id} thr={threshold}")

        # Catalog-first payload for reports/tables
        catalog_ws, catalog_datasets = _similarity_catalog_workspace(workspace_id)

        # Semantic-models style dataset schema cache (measures + tables with expressions)
        # Shared with /api/semantic-model-details enrich — one scan serves both UIs.
        schema_by_id = {}
        scan_ws = None
        scan_note = None
        schema_source = None

        def _load_schema_cache():
            nonlocal schema_by_id, schema_source, scan_note
            try:
                schema_by_id = _get_workspace_scan_cached(workspace_id) or {}
                schema_source = 'semantic_scan_cache'
                print(f"   📦 schema cache datasets={len(schema_by_id)}")
            except Exception as e:
                scan_note = f'Schema enrich scan failed: {e}'
                print(f"   ⚠️ {scan_note}")
                schema_by_id = {}

        def _load_full_scan_ws():
            """Full workspace scan (needed for report pages/visuals)."""
            nonlocal scan_ws, scan_note
            try:
                from scanner_connector import PowerBIScanner
                scanner = PowerBIScanner()
                scan_result = scanner.run_scan(workspace_id=workspace_id) or {}
                for ws in scan_result.get('workspaces') or []:
                    if ws.get('id') == workspace_id:
                        scan_ws = ws
                        break
                if not scan_ws and (scan_result.get('workspaces') or []):
                    scan_ws = (scan_result.get('workspaces') or [None])[0]
                if not scan_ws:
                    scan_note = (scan_note + '; ' if scan_note else '') + 'Scanner returned no workspace payload'
                else:
                    # Also seed schema cache from this full scan when empty
                    if not schema_by_id:
                        by_id = {}
                        for ds in scan_ws.get('datasets') or []:
                            did = ds.get('id')
                            if did:
                                by_id[did] = ds
                        if by_id:
                            schema_by_id.update(by_id)
                            with _semantic_scan_lock:
                                _semantic_scan_cache[workspace_id] = {
                                    'ts': time.time(),
                                    'by_id': dict(by_id),
                                }
            except Exception as e:
                scan_note = f'Scanner failed: {e}'
                print(f"   ⚠️ {scan_note}")

        notes = []
        data_src = {'catalog': bool(catalog_ws), 'scanner': False, 'schema_cache': False}

        if analysis_type == 'tables':
            # Catalog columns are enough; optional schema enrich not required
            matches = _sim_analyze_tables(catalog_ws, catalog_datasets, None, threshold)

        elif analysis_type == 'reports':
            # Prefer catalog structure; enrich measure names from schema cache if available
            # (same source Semantic Models Details uses for DAX)
            if not catalog_ws:
                _load_schema_cache()
            else:
                # cheap: use cache if already warm; don't force scan for report mode
                with _semantic_scan_lock:
                    hit = _semantic_scan_cache.get(workspace_id)
                if hit and hit.get('by_id'):
                    schema_by_id = hit.get('by_id') or {}
                    schema_source = 'semantic_scan_cache_warm'
            scan_ws_stub = {'datasets': list(schema_by_id.values())} if schema_by_id else None
            matches = _sim_analyze_reports(catalog_ws, catalog_datasets, scan_ws_stub, threshold)
            data_src['schema_cache'] = bool(schema_by_id)

        elif analysis_type == 'measures':
            # Same DAX source as Semantic Models Details (workspace scan cache)
            _load_schema_cache()
            data_src['schema_cache'] = bool(schema_by_id)
            data_src['scanner'] = bool(schema_by_id)
            matches, mnote = _sim_analyze_measures(
                catalog_ws,
                catalog_datasets,
                schema_by_id,
                threshold,
                schema_source=schema_source or 'scanner',
            )
            if mnote:
                notes.append(mnote)

        else:  # visuals — still need full report pages payload
            _load_full_scan_ws()
            data_src['scanner'] = bool(scan_ws)
            matches, vnote = _sim_analyze_visuals(scan_ws, threshold)
            if vnote:
                notes.append(vnote)

        if scan_note and analysis_type in {'visuals', 'measures'} and not matches:
            notes.append(scan_note)

        # Cap huge pairwise result sets for UI responsiveness
        max_rows = 500
        truncated = len(matches) > max_rows
        if truncated:
            matches = matches[:max_rows]
            notes.append(f'Showing top {max_rows} matches (truncated)')

        print(f"   ✅ {analysis_type}: {len(matches)} matches")

        return jsonify({
            'success': True,
            'workspace_id': workspace_id,
            'analysis_type': analysis_type,
            'threshold': threshold,
            'matches': matches,
            'total_matches': len(matches),
            'truncated': truncated,
            'notes': notes,
            'data_source': data_src,
        })

    except Exception as e:
        print(f"❌ Error in similarity analysis: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500


def _similarity_catalog_workspace(workspace_id):
    """Return (workspace_dict_or_None, datasets_by_id) from catalog."""
    if not CATALOG_AVAILABLE:
        return None, {}
    try:
        cat = catalog_service.get_workspace_catalog() or {}
    except Exception as e:
        print(f"   ⚠️ catalog load failed for similarity: {e}")
        return None, {}
    datasets = cat.get('datasets') or {}
    ws = next((w for w in (cat.get('workspaces') or []) if w.get('id') == workspace_id), None)
    return ws, datasets


def _sim_dataset_tables(dataset_obj):
    """Normalize Scanner/catalog dataset → list of table dicts with columns/measures."""
    if not dataset_obj:
        return []
    # Catalog & modern Scanner: tables at top-level. Legacy some payloads nested under model.
    tables = dataset_obj.get('tables')
    if tables is None:
        tables = (dataset_obj.get('model') or {}).get('tables') or []
    return tables or []


def _sim_col_name(col):
    if isinstance(col, dict):
        return (col.get('name') or col.get('column') or '').strip()
    return str(col or '').strip()


def _sim_measure_expr(measure):
    if not isinstance(measure, dict):
        return '', ''
    name = (measure.get('name') or '').strip()
    expr = (
        measure.get('expression')
        or measure.get('Expression')
        or measure.get('dax')
        or ''
    )
    return name, str(expr).strip()


def _sim_extract_visual_fields(visual):
    """Best-effort field extraction across Scanner visual shapes."""
    fields = set()
    if not isinstance(visual, dict):
        return fields

    def add_field(val):
        if val is None:
            return
        if isinstance(val, dict):
            name = (
                val.get('name')
                or val.get('column')
                or val.get('measure')
                or val.get('field')
                or val.get('queryRef')
                or val.get('nativeQueryRef')
                or ''
            )
            if name:
                fields.add(str(name).split('.')[-1].strip('[]'))
            return
        s = str(val).strip()
        if s:
            fields.add(s.split('.')[-1].strip('[]'))

    # Common buckets
    for key in (
        'columns', 'measures', 'values', 'categories', 'rows', 'series',
        'fields', 'projections', 'dataRoles', 'selects',
    ):
        bucket = visual.get(key)
        if isinstance(bucket, list):
            for item in bucket:
                add_field(item)
        elif isinstance(bucket, dict):
            for item in bucket.values():
                if isinstance(item, list):
                    for x in item:
                        add_field(x)
                else:
                    add_field(item)

    # Nested payload used by some scanner dumps
    for nest_key in ('query', 'config', 'prototypeQuery', 'visual'):
        nest = visual.get(nest_key)
        if isinstance(nest, dict):
            for k, v in nest.items():
                if k.lower() in {
                    'columns', 'measures', 'values', 'categories', 'rows', 'fields', 'select'
                }:
                    if isinstance(v, list):
                        for item in v:
                            add_field(item)

    return {f for f in fields if f}


def _sim_analyze_reports(catalog_ws, catalog_datasets, scan_ws, threshold):
    """Pairwise report similarity via shared table/measure names on bound datasets."""
    datasets_map = {}
    reports = []

    if catalog_ws:
        reports = list(catalog_ws.get('reports') or [])
        # Prefer global catalog datasets map; fall back to workspace-local datasets
        datasets_map = dict(catalog_datasets or {})
        for d in catalog_ws.get('datasets') or []:
            if d.get('id') and d['id'] not in datasets_map:
                datasets_map[d['id']] = d

    if scan_ws:
        if not reports:
            reports = list(scan_ws.get('reports') or [])
        for d in scan_ws.get('datasets') or []:
            did = d.get('id')
            if did:
                # Scanner often has richer measure expressions — prefer if present
                existing = datasets_map.get(did)
                scan_tables = _sim_dataset_tables(d)
                if not existing:
                    datasets_map[did] = d
                else:
                    # merge: if catalog tables lack measures, keep scanner tables
                    cat_tables = _sim_dataset_tables(existing)
                    cat_has_meas = any((t.get('measures') or []) for t in cat_tables)
                    scan_has_meas = any((t.get('measures') or []) for t in scan_tables)
                    if scan_has_meas and not cat_has_meas:
                        datasets_map[did] = d

    profiles = []
    for report in reports:
        name = report.get('name') or ''
        if _is_excluded_report_name(name):
            continue
        ds_id = report.get('datasetId') or ''
        ds = datasets_map.get(ds_id) or {}
        tables = set()
        measures = set()
        for t in _sim_dataset_tables(ds):
            tname = (t.get('name') or '').strip()
            if tname:
                tables.add(tname)
            for m in t.get('measures') or []:
                if isinstance(m, dict):
                    mn = (m.get('name') or '').strip()
                else:
                    mn = str(m or '').strip()
                if mn:
                    measures.add(mn)
            # some catalogs store measureCount only — still useful via columns for table sim
        if not tables and not measures:
            # still include with empty sets only if dataset known? skip empty noise
            continue
        profiles.append({
            'id': report.get('id'),
            'name': name,
            'dataset_id': ds_id,
            'dataset_name': ds.get('name') or report.get('datasetName') or '',
            'tables': tables,
            'measures': measures,
        })

    matches = []
    for i in range(len(profiles)):
        for j in range(i + 1, len(profiles)):
            r1, r2 = profiles[i], profiles[j]
            # Same dataset means effectively identical model — still valid similarity
            common_tables = r1['tables'] & r2['tables']
            common_measures = r1['measures'] & r2['measures']
            all_tables = r1['tables'] | r2['tables']
            all_measures = r1['measures'] | r2['measures']
            if not all_tables and not all_measures:
                continue
            table_sim = (len(common_tables) / len(all_tables)) if all_tables else 0.0
            measure_sim = (len(common_measures) / len(all_measures)) if all_measures else 0.0
            if all_tables and all_measures:
                score = (table_sim + measure_sim) / 2.0
            elif all_tables:
                score = table_sim
            else:
                score = measure_sim
            if score >= threshold:
                matches.append({
                    'report1_id': r1['id'],
                    'report1_name': r1['name'],
                    'report2_id': r2['id'],
                    'report2_name': r2['name'],
                    'score': round(score, 3),
                    'common_tables': sorted(common_tables)[:40],
                    'common_tables_count': len(common_tables),
                    'common_measures': sorted(common_measures)[:40],
                    'common_measures_count': len(common_measures),
                    'dataset1_name': r1['dataset_name'],
                    'dataset2_name': r2['dataset_name'],
                    'same_dataset': bool(r1['dataset_id'] and r1['dataset_id'] == r2['dataset_id']),
                })
    matches.sort(key=lambda x: x['score'], reverse=True)
    return matches


def _sim_analyze_tables(catalog_ws, catalog_datasets, scan_ws, threshold):
    """Pairwise table similarity by shared column names."""
    datasets = []
    seen = set()

    def add_ds(d):
        if not d:
            return
        did = d.get('id') or id(d)
        if did in seen:
            return
        seen.add(did)
        datasets.append(d)

    if catalog_ws:
        # Prefer global dataset map entries for this workspace
        ws_id = catalog_ws.get('id')
        for d in (catalog_datasets or {}).values():
            if d.get('workspaceId') == ws_id or d.get('workspace_id') == ws_id:
                add_ds(d)
        for d in catalog_ws.get('datasets') or []:
            full = (catalog_datasets or {}).get(d.get('id')) or d
            add_ds(full)

    if scan_ws:
        for d in scan_ws.get('datasets') or []:
            add_ds(d)

    all_tables = []
    for dataset in datasets:
        ds_name = dataset.get('name') or ''
        for table in _sim_dataset_tables(dataset):
            cols = set()
            for c in table.get('columns') or []:
                cn = _sim_col_name(c)
                if cn:
                    cols.add(cn)
            if not cols:
                continue
            source_type = (
                table.get('sourceTypeLabel')
                or table.get('source_type')
                or 'Unknown'
            )
            if source_type == 'Unknown':
                for partition in table.get('partitions') or []:
                    source = partition.get('source') or {}
                    if isinstance(source, dict):
                        if source.get('type') == 'M':
                            source_type = 'M Query'
                        elif source.get('expression'):
                            source_type = 'DAX / expression'
            all_tables.append({
                'name': table.get('name') or '',
                'dataset': ds_name,
                'columns': cols,
                'source_type': source_type,
            })

    matches = []
    n = len(all_tables)
    # Limit O(n^2) explosion
    hard_cap = 400
    if n > hard_cap:
        all_tables = all_tables[:hard_cap]

    for i in range(len(all_tables)):
        for j in range(i + 1, len(all_tables)):
            t1, t2 = all_tables[i], all_tables[j]
            # skip exact same table name in same dataset
            if t1['dataset'] == t2['dataset'] and t1['name'] == t2['name']:
                continue
            common = t1['columns'] & t2['columns']
            union = t1['columns'] | t2['columns']
            if not union:
                continue
            score = len(common) / len(union)
            if score >= threshold:
                matches.append({
                    'table1_name': t1['name'],
                    'table2_name': t2['name'],
                    'dataset1': t1['dataset'],
                    'dataset2': t2['dataset'],
                    'score': round(score, 3),
                    'common_columns': sorted(common)[:50],
                    'common_columns_count': len(common),
                    'source_type': t1['source_type'],
                })
    matches.sort(key=lambda x: x['score'], reverse=True)
    return matches


def _sim_analyze_measures(catalog_ws, catalog_datasets, schema_source_obj, threshold, schema_source='scanner'):
    """
    Pairwise measure similarity using normalized DAX.

    schema_source_obj can be:
      - dict of dataset_id -> dataset  (from _get_workspace_scan_cached — same as Semantic Models Details)
      - workspace scan dict with .datasets list
      - None
    """
    from difflib import SequenceMatcher

    all_measures = []
    sources_tried = []

    def collect_from_datasets(datasets, label):
        count = 0
        for dataset in datasets or []:
            if not dataset:
                continue
            ds_name = dataset.get('name') or ''
            ds_id = dataset.get('id') or ''
            before = count
            # Same flatten helper as /api/semantic-model-details
            flat_measures, _rels = _extract_measures_and_relationships(dataset)
            for m in flat_measures or []:
                name = (m.get('name') or '').strip()
                dax = str(m.get('expression') or '').strip()
                if not dax:
                    continue
                all_measures.append({
                    'name': name or '(unnamed)',
                    'dax': dax,
                    'table': m.get('table') or '',
                    'dataset': ds_name,
                    'dataset_id': ds_id,
                    'normalized_dax': normalize_dax(dax),
                })
                count += 1
            # Per-dataset fallback if flatten had no expressions
            if count == before:
                for table in _sim_dataset_tables(dataset):
                    tname = table.get('name') or ''
                    for measure in table.get('measures') or []:
                        name, dax = _sim_measure_expr(measure)
                        if not dax:
                            continue
                        all_measures.append({
                            'name': name or '(unnamed)',
                            'dax': dax,
                            'table': tname,
                            'dataset': ds_name,
                            'dataset_id': ds_id,
                            'normalized_dax': normalize_dax(dax),
                        })
                        count += 1
        sources_tried.append(f'{label}:{count}')
        return count

    # 1) Preferred: Semantic Models schema cache (dataset id → dataset with measures/expressions)
    if isinstance(schema_source_obj, dict) and schema_source_obj:
        # Heuristic: cache map keys are dataset GUIDs; workspace object has 'datasets' list
        if 'datasets' in schema_source_obj and isinstance(schema_source_obj.get('datasets'), list):
            collect_from_datasets(schema_source_obj.get('datasets') or [], schema_source or 'scanner_ws')
        else:
            # Treat as by_id map (possibly mixed with non-dataset keys — only keep dict values that look like datasets)
            ds_list = []
            for k, v in schema_source_obj.items():
                if isinstance(v, dict) and (v.get('tables') is not None or v.get('name') or v.get('id')):
                    # skip accidental nested structures
                    if k in {'datasets', 'reports', 'dashboards', 'workspaces'}:
                        continue
                    ds_list.append(v)
            collect_from_datasets(ds_list, schema_source or 'semantic_scan_cache')

    # 2) Catalog fallback (usually has columns, rarely DAX expressions today)
    if not all_measures and catalog_ws:
        ws_id = catalog_ws.get('id')
        ds_list = []
        for d in (catalog_datasets or {}).values():
            if d.get('workspaceId') == ws_id or d.get('workspace_id') == ws_id:
                ds_list.append(d)
        if not ds_list:
            for d in catalog_ws.get('datasets') or []:
                ds_list.append((catalog_datasets or {}).get(d.get('id')) or d)
        collect_from_datasets(ds_list, 'catalog')

    note = None
    if not all_measures:
        note = (
            'No DAX measure expressions found for this workspace. '
            'Semantic Models Details uses the same Scanner enrich cache — open a model Details once '
            'to warm the cache, or ensure Admin Scanner returns datasetSchema + measures. '
            f"Tried: {', '.join(sources_tried) or 'none'}"
        )
        return [], note

    if len(all_measures) > 600:
        all_measures = all_measures[:600]
        note = (
            f'Using {schema_source or "scanner"} DAX (same source as Semantic Models). '
            'Compared first 600 measures with expressions (capped).'
        )
    else:
        note = (
            f'Using {schema_source or "scanner"} DAX expressions '
            f'({len(all_measures)} measures — same enrich path as Semantic Models Details).'
        )

    matches = []
    for i in range(len(all_measures)):
        for j in range(i + 1, len(all_measures)):
            m1, m2 = all_measures[i], all_measures[j]
            if (
                m1.get('dataset_id') and m1.get('dataset_id') == m2.get('dataset_id')
                and m1['name'] == m2['name'] and m1['table'] == m2['table']
            ):
                continue
            if m1['dataset'] == m2['dataset'] and m1['name'] == m2['name'] and m1['table'] == m2['table']:
                continue
            ratio = SequenceMatcher(None, m1['normalized_dax'], m2['normalized_dax']).ratio()
            tokens1 = set(m1['normalized_dax'].split())
            tokens2 = set(m2['normalized_dax'].split())
            union = tokens1 | tokens2
            token_score = (len(tokens1 & tokens2) / len(union)) if union else 0.0
            score = max(ratio, token_score)
            if score >= threshold:
                matches.append({
                    'measure1_name': m1['name'],
                    'measure2_name': m2['name'],
                    'table1_name': m1['table'],
                    'table2_name': m2['table'],
                    'dataset1': m1['dataset'],
                    'dataset2': m2['dataset'],
                    'score': round(score, 3),
                    'dax1': m1['dax'][:500],
                    'dax2': m2['dax'][:500],
                    'dax_pattern': extract_dax_pattern(m1['dax']),
                })
    matches.sort(key=lambda x: x['score'], reverse=True)
    return matches, note


def _sim_analyze_visuals(scan_ws, threshold):
    """Pairwise visual similarity by type + shared fields (Scanner pages/visuals)."""
    if not scan_ws:
        return [], (
            'Visual similarity requires Admin Scanner report pages/visuals. '
            'Scan returned no workspace data for this id.'
        )

    all_visuals = []
    reports_with_pages = 0
    for report in scan_ws.get('reports') or []:
        name = report.get('name') or ''
        if _is_excluded_report_name(name):
            continue
        pages = report.get('pages') or []
        if pages:
            reports_with_pages += 1
        for page in pages:
            page_name = page.get('displayName') or page.get('name') or ''
            for visual in page.get('visuals') or []:
                fields = _sim_extract_visual_fields(visual)
                vtype = (
                    visual.get('type')
                    or visual.get('visualType')
                    or visual.get('visualTypeName')
                    or 'Unknown'
                )
                title = (
                    visual.get('title')
                    or visual.get('name')
                    or visual.get('displayName')
                    or 'Untitled'
                )
                all_visuals.append({
                    'report_name': name,
                    'page_name': page_name,
                    'visual_title': title,
                    'visual_type': vtype,
                    'fields': fields,
                })

    if not all_visuals:
        if reports_with_pages == 0:
            return [], (
                'No report pages/visuals in Scanner payload. '
                'Tenant Scanner often omits visual metadata — report/table/measure modes still work.'
            )
        return [], 'Pages found but no visuals extracted from Scanner payload.'

    # Cap
    if len(all_visuals) > 800:
        all_visuals = all_visuals[:800]

    matches = []
    for i in range(len(all_visuals)):
        for j in range(i + 1, len(all_visuals)):
            v1, v2 = all_visuals[i], all_visuals[j]
            if v1['visual_type'] != v2['visual_type']:
                continue
            # avoid comparing a visual to itself on same report/page/title with no fields
            common = v1['fields'] & v2['fields']
            union = v1['fields'] | v2['fields']
            if not union:
                # Same type only → weak score; skip empty field pairs
                continue
            score = len(common) / len(union)
            if score >= threshold:
                matches.append({
                    'visual1_title': v1['visual_title'],
                    'visual2_title': v2['visual_title'],
                    'visual_type': v1['visual_type'],
                    'report1_name': v1['report_name'],
                    'report2_name': v2['report_name'],
                    'page1': v1['page_name'],
                    'page2': v2['page_name'],
                    'score': round(score, 3),
                    'common_fields': sorted(common)[:40],
                    'common_fields_count': len(common),
                })
    matches.sort(key=lambda x: x['score'], reverse=True)
    return matches, None


def normalize_dax(dax_expr):
    """Normalize DAX expression for comparison"""
    import re
    # Remove whitespace, convert to lowercase, remove comments
    normalized = re.sub(r'--.*', '', dax_expr)  # Remove comments
    normalized = re.sub(r'/\*.*?\*/', '', normalized, flags=re.DOTALL)  # Remove block comments
    normalized = normalized.lower()
    normalized = re.sub(r'\s+', ' ', normalized)  # Normalize whitespace
    return normalized.strip()


def extract_dax_pattern(dax_expr):
    """Extract common DAX pattern (e.g., SUM, CALCULATE, etc.)"""
    import re
    # Find main DAX function
    match = re.search(r'(\w+)\s*\(', dax_expr.strip())
    if match:
        return match.group(1).upper()
    return 'Custom'


@app.route('/api/workspace-summary/<workspace_id>')
@login_required
def get_workspace_summary(workspace_id):
    """
    Get workspace summary including total reports and inactive reports count
    Inactive reports = reports with datasets not refreshed in 30+ days
    Results are cached for 15 minutes per workspace per user
    """
    from datetime import datetime, timezone, timedelta
    from concurrent.futures import ThreadPoolExecutor, as_completed

    try:
        # Get user info for cache key
        user_id = session.get('user', {}).get('oid', 'unknown')
        cache_key = f"workspace_summary_{workspace_id}_{user_id}"

        # Check cache first (15 minute cache for better performance)
        current_time = time.time()
        SUMMARY_CACHE_DURATION = 900  # 15 minutes
        if cache_key in workspace_cache:
            cache_entry = workspace_cache[cache_key]
            if (current_time - cache_entry['timestamp']) < SUMMARY_CACHE_DURATION:
                print(f"✅ Cache hit for workspace summary: {workspace_id}")
                return jsonify(cache_entry['data'])

        print(f"📊 Fetching workspace summary for: {workspace_id}")

        # Get all reports in workspace (fast API call)
        base_url = "https://api.powerbi.com/v1.0/myorg"
        reports_url = f"{base_url}/groups/{workspace_id}/reports"

        headers = get_user_powerbi_headers()
        reports_response = requests.get(reports_url, headers=headers)
        reports_response.raise_for_status()

        reports = reports_response.json().get('value', [])
        total_reports = len(reports)

        print(f"   Found {total_reports} reports")

        # Group reports by dataset ID
        dataset_to_reports = {}
        for report in reports:
            dataset_id = report.get('datasetId')
            if dataset_id:
                if dataset_id not in dataset_to_reports:
                    dataset_to_reports[dataset_id] = []
                dataset_to_reports[dataset_id].append(report['id'])

        print(f"   Found {len(dataset_to_reports)} unique datasets")

        # Create a session for connection pooling (reuse TCP connections)
        session_obj = requests.Session()
        session_obj.headers.update(headers)

        # Function to check if a dataset is inactive
        def check_dataset_inactive(dataset_id):
            """
            Check if dataset is inactive based on:
            - ONLY datasets with >30 days since last successful refresh

            Returns False (not inactive) for:
            - Datasets with no refresh history (likely DirectQuery/Live)
            - Datasets refreshed within 30 days
            - Live connection datasets (API 415)
            - API errors (403, 404, etc.)
            """
            try:
                from powerbi_connector import pick_best_refresh_from_history

                # Pull a few history rows so in-progress (null endTime) doesn't hide last completed
                refresh_url = f"{base_url}/groups/{workspace_id}/datasets/{dataset_id}/refreshes?$top=5"
                refresh_response = session_obj.get(refresh_url, timeout=5)

                # API 415 means non-model dataset (live/DirectQuery)
                if refresh_response.status_code == 415:
                    print(f"      ℹ️  Dataset {dataset_id[:8]}... is live/DirectQuery (API 415) - NOT inactive")
                    return False, dataset_id

                # Permission/access errors - don't count as inactive
                if refresh_response.status_code in [403, 404]:
                    print(f"      ℹ️  Dataset {dataset_id[:8]}... access error (Status: {refresh_response.status_code}) - NOT inactive")
                    return False, dataset_id

                if refresh_response.status_code == 200:
                    refresh_history = refresh_response.json().get('value', [])
                    picked = pick_best_refresh_from_history(refresh_history)
                    end_time_str = picked.get('last_refreshed')

                    if end_time_str:
                        try:
                            # Handle ISO 8601 format with Z suffix
                            if end_time_str.endswith('Z'):
                                end_time_str = end_time_str[:-1] + '+00:00'

                            end_time = datetime.fromisoformat(end_time_str)

                            # Ensure timezone-aware comparison
                            if end_time.tzinfo is None:
                                end_time = end_time.replace(tzinfo=timezone.utc)

                            now = datetime.now(timezone.utc)
                            days_since_refresh = (now - end_time).days

                            # ONLY count as inactive if >30 days since last completed refresh
                            if days_since_refresh >= 30:
                                print(f"      ⚠️  Dataset {dataset_id[:8]}... is INACTIVE - {days_since_refresh} days since last refresh")
                                return True, dataset_id
                            else:
                                return False, dataset_id
                        except Exception:
                            # Can't parse date - don't count as inactive
                            print(f"      ℹ️  Dataset {dataset_id[:8]}... date parse error - treating as active")
                            return False, dataset_id
                    else:
                        # No usable timestamp - don't count as inactive (DQ/Live or never refreshed)
                        print(f"      ℹ️  Dataset {dataset_id[:8]}... no usable refresh timestamp - NOT inactive")
                        return False, dataset_id
                else:
                    # Other API errors - don't count as inactive
                    print(f"      ℹ️  Dataset {dataset_id[:8]}... API error (Status: {refresh_response.status_code}) - NOT inactive")
                    return False, dataset_id

            except Exception as e:
                # On error - don't count as inactive
                print(f"      ℹ️  Dataset {dataset_id[:8]}... exception: {str(e)} - treating as active")
                return False, dataset_id

            return False, dataset_id

        # Check datasets in parallel (max 20 concurrent workers for faster processing)
        inactive_dataset_ids = set()

        if dataset_to_reports:
            print(f"   Checking {len(dataset_to_reports)} datasets for inactivity...")

            # ⚡ PERFORMANCE OPTIMIZATION: Parallel processing for refresh history checks
            # Use ThreadPoolExecutor with more workers and process in smaller batches
            with ThreadPoolExecutor(max_workers=20) as executor:
                # Submit all dataset checks at once
                future_to_dataset = {
                    executor.submit(check_dataset_inactive, dataset_id): dataset_id
                    for dataset_id in dataset_to_reports.keys()
                }

                # Collect results as they complete (don't wait for all)
                for future in as_completed(future_to_dataset):
                    try:
                        is_inactive, dataset_id = future.result(timeout=10)
                        if is_inactive:
                            inactive_dataset_ids.add(dataset_id)
                    except Exception as e:
                        # Silently mark as inactive on error to avoid blocking
                        dataset_id = future_to_dataset.get(future)
                        if dataset_id:
                            inactive_dataset_ids.add(dataset_id)

        # Close the session to free up connections
        session_obj.close()

        # Count inactive reports (all reports linked to inactive datasets)
        inactive_reports_count = 0
        for dataset_id in inactive_dataset_ids:
            inactive_reports_count += len(dataset_to_reports.get(dataset_id, []))

        print(f"   ✅ Summary: {total_reports} total, {inactive_reports_count} inactive")

        # Build response
        summary_data = {
            'success': True,
            'workspace_id': workspace_id,
            'total_reports': total_reports,
            'inactive_reports': inactive_reports_count
        }

        # Cache the result (15 minutes)
        workspace_cache[cache_key] = {
            'data': summary_data,
            'timestamp': current_time
        }

        return jsonify(summary_data)

    except requests.exceptions.RequestException as e:
        print(f"❌ API Error in workspace summary: {str(e)}")
        if hasattr(e, 'response') and e.response is not None:
            print(f"   Status: {e.response.status_code}")
            print(f"   Response: {e.response.text[:200]}")
        return jsonify({
            'success': False,
            'error': str(e),
            'workspace_id': workspace_id,
            'total_reports': 0,
            'inactive_reports': 'N/A'
        }), 500

    except Exception as e:
        print(f"❌ Error in workspace summary: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e),
            'workspace_id': workspace_id,
            'total_reports': 0,
            'inactive_reports': 'N/A'
        }), 500


@app.route('/api/me/capabilities')
@login_required
def api_me_capabilities():
    """UI feature flags for the signed-in user (e.g. archive download button)."""
    try:
        from features.report_archive_service import user_can_archive
        email = session.get('user', {}).get('preferred_username') or ''
        return jsonify({
            'success': True,
            'canArchiveReports': user_can_archive(email),
            'email': email,
        })
    except Exception as e:
        return jsonify({'success': False, 'canArchiveReports': False, 'error': str(e)}), 500


@app.route('/api/reports/archive-to-sharepoint', methods=['POST'])
@login_required
def api_archive_report_to_sharepoint():
    """
    Export a Power BI report (.pbix/.rdl) and upload to SharePoint
    Report Decommission Activity / <latest dated folder> / Workspace / [Folder] /.

    Restricted to Central Analytics team UPNs (see report_archive_service).
    Does not alter catalog fast-path, crash test, or generate flows.
    """
    try:
        from features.report_archive_service import (
            archive_report_to_sharepoint,
            user_can_archive,
        )

        email = session.get('user', {}).get('preferred_username') or ''
        if not user_can_archive(email):
            return jsonify({
                'success': False,
                'error': 'Not authorized to archive reports to SharePoint.',
            }), 403

        data = request.get_json(silent=True) or {}
        workspace_id = (data.get('workspace_id') or request.args.get('workspace_id') or '').strip()
        report_id = (data.get('report_id') or request.args.get('report_id') or '').strip()
        report_name = (data.get('report_name') or data.get('name') or 'Report').strip()
        workspace_name = (data.get('workspace_name') or data.get('workspaceName') or '').strip()
        folder_name = (data.get('folder_name') or data.get('folderName') or '').strip() or None
        folder_id = (data.get('folder_id') or data.get('folderId') or '').strip() or None

        if not workspace_id or not report_id:
            return jsonify({
                'success': False,
                'error': 'workspace_id and report_id are required',
            }), 400

        # Resolve workspace display name if missing
        if not workspace_name:
            try:
                if CATALOG_AVAILABLE and catalog_service is not None:
                    allowed = _user_allowed_workspace_ids()
                    pack = catalog_service.get_workspace_reports(
                        workspace_id, allowed_workspace_ids=allowed
                    )
                    if pack:
                        workspace_name = (
                            (pack.get('workspace') or {}).get('name')
                            or pack.get('workspace_name')
                            or ''
                        )
            except Exception:
                pass
        if not workspace_name:
            workspace_name = workspace_id[:8]

        # Prefer user delegated token (same access as UI). Also acquire SP token as
        # fallback — REST Export sometimes 500s on large PBIX with one principal
        # and succeeds with the other (Service UI download uses a different path).
        token = None
        sp_token = None
        try:
            token = get_user_powerbi_token()
        except Exception:
            token = None
        try:
            from scanner_connector import PowerBIScanner
            sc = PowerBIScanner()
            sp_token = sc.get_access_token()
        except Exception as ex:
            print(f"   ⚠️ SP token for export fallback unavailable: {ex}")
            sp_token = None
        if not token:
            token = sp_token
        if not token:
            return jsonify({
                'success': False,
                'error': 'Unable to obtain Power BI token for Export',
            }), 401

        print(f"\n📦 ARCHIVE TO SHAREPOINT")
        print(f"   User: {email}")
        print(f"   Workspace: {workspace_name} ({workspace_id[:8]}…)")
        print(f"   Report: {report_name} ({report_id[:8]}…)")
        print(f"   PBI folder: {folder_name or '(root)'}")
        print(f"   Tokens: user={'yes' if token and token != sp_token else 'no'} sp={'yes' if sp_token else 'no'}")

        result = archive_report_to_sharepoint(
            access_token=token,
            workspace_id=workspace_id,
            workspace_name=workspace_name,
            report_id=report_id,
            report_name=report_name,
            folder_name=folder_name,
            folder_id=folder_id,
            fallback_token=sp_token if (sp_token and sp_token != token) else None,
        )
        if result.get('success'):
            print(f"   ✅ ARCHIVE OK → {result.get('remotePath')}")
            return jsonify(result), 200

        err = result.get('error') or 'Archive failed'
        stage = result.get('stage') or 'unknown'
        print(f"   ❌ ARCHIVE FAILED stage={stage}: {err}")
        status = 400
        sc = result.get('status_code')
        if sc in (401, 403):
            status = int(sc)
        elif sc == 404:
            status = 404
        return jsonify(result), status
    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"   ❌ ARCHIVE EXCEPTION: {e}")
        return jsonify({'success': False, 'error': str(e), 'stage': 'exception'}), 500


@app.route('/api/reports/crash-test/<report_id>', methods=['GET', 'POST'])
@login_required
def crash_test_report(report_id):
    """
    Run crash test analysis on a specific report with optional visual field bindings

    Accepts both GET (legacy) and POST (with visual field bindings from frontend)
    Returns health score and detected issues with deep root cause analysis
    """
    try:
        # Handle both GET and POST requests
        if request.method == 'POST':
            data = request.get_json() or {}
            workspace_id = data.get('workspace_id')
            dataset_id = data.get('dataset_id')
            mode = data.get('mode', 'standard')
            visual_field_bindings = data.get('visual_field_bindings')  # ⭐ NEW!
        else:
            # Legacy GET support
            workspace_id = request.args.get('workspace_id')
            dataset_id = request.args.get('dataset_id')
            mode = request.args.get('mode', 'standard')
            visual_field_bindings = None

        if not workspace_id or not dataset_id:
            return jsonify({
                'success': False,
                'error': 'Missing required parameters: workspace_id and dataset_id'
            }), 400

        print(f"\n🔬 HYBRID CRASH TEST REQUEST")
        print(f"   Report: {report_id}")
        print(f"   Workspace: {workspace_id}")
        print(f"   Dataset: {dataset_id}")
        print(f"   Mode: {mode}")
        if visual_field_bindings:
            print(f"   ⭐ Visual Field Bindings: {len(visual_field_bindings)} visuals extracted from frontend")
        else:
            print(f"   ℹ️  No visual field bindings provided (legacy mode)")

        # Import crash test analyzer
        from crash_test_analyzer import CrashTestAnalyzer
        import os

        # Get service principal credentials for Enhanced Mode
        client_id = os.getenv('CLIENT_ID')
        client_secret = os.getenv('CLIENT_SECRET')
        tenant_id = os.getenv('TENANT_ID')

        # Fetch report metadata (modifiedBy, modifiedDateTime) from Scanner API
        print(f"   📋 Fetching report metadata from Scanner API...")
        from scanner_connector import PowerBIScanner
        report_metadata = {'modified_by': 'N/A', 'modified_date': None}

        try:
            scanner = PowerBIScanner()
            scan_result = scanner.run_scan(workspace_id=workspace_id)

            if scan_result and "workspaces" in scan_result:
                for workspace in scan_result.get('workspaces', []):
                    if workspace.get('id') == workspace_id:
                        for report in workspace.get('reports', []):
                            if report.get('id') == report_id:
                                # Scanner API returns modifiedBy as email/UPN directly
                                modified_by = report.get('modifiedBy', '')
                                modified_date = report.get('modifiedDateTime')

                                print(f"      🔍 Scanner API report found")
                                print(f"      📋 modifiedBy: '{modified_by}'")
                                print(f"      📋 modifiedDateTime: '{modified_date}'")

                                # If modifiedBy is empty, try createdBy as fallback
                                if not modified_by or modified_by == '':
                                    modified_by = report.get('createdBy', 'N/A')
                                    print(f"      ⚠️  modifiedBy empty, using createdBy: '{modified_by}'")

                                report_metadata = {
                                    'modified_by': modified_by,
                                    'modified_date': modified_date
                                }
                                break
                        break
                print(f"      ✅ Report metadata extracted: {report_metadata}")
            else:
                print(f"      ⚠️  Scanner API returned no data")
        except Exception as e:
            print(f"      ⚠️  Could not fetch report metadata from Scanner API: {e}")

        # Initialize analyzer with credentials for Enhanced Mode
        analyzer = CrashTestAnalyzer(
            workspace_id=workspace_id,
            report_id=report_id,
            dataset_id=dataset_id,
            access_token=None,  # Will use service principal
            client_id=client_id,
            client_secret=client_secret,
            tenant_id=tenant_id
        )

        # Set report metadata on analyzer
        analyzer.report_metadata = report_metadata
        print(f"   ✅ Report metadata set: {report_metadata}")

        # ⭐ NEW: Pass visual field bindings to analyzer if available
        if visual_field_bindings:
            analyzer.set_visual_field_bindings(visual_field_bindings)

        # Run crash test (with visual analysis, lineage, and version history if Enhanced Mode)
        include_visual_analysis = (mode == 'enhanced')
        include_lineage_analysis = (mode == 'enhanced')
        include_version_history = (mode == 'enhanced')

        # ⭐ Enable XMLA schema analysis if we have visual bindings
        use_xmla_schema = (visual_field_bindings is not None and len(visual_field_bindings) > 0)

        print(f"   🚀 Running HYBRID deep-dive crash test...")
        print(f"      Visual Analysis: {include_visual_analysis}")
        print(f"      Lineage Analysis: {include_lineage_analysis}")
        print(f"      Version History: {include_version_history}")
        print(f"      XMLA Schema Analysis: {use_xmla_schema}")

        results = analyzer.run_crash_test(
            include_visual_analysis=include_visual_analysis,
            include_lineage_analysis=include_lineage_analysis,
            include_version_history=include_version_history,
            use_xmla_schema=use_xmla_schema  # ⭐ NEW!
        )

        # Format response
        health_score = results.get('health_score', 0)
        issues = results.get('issues', [])
        warnings = results.get('warnings', [])
        lineage_analysis = results.get('lineage_analysis', {})
        root_cause_analysis = results.get('root_cause_analysis', [])
        change_impact_summary = results.get('change_impact_summary', {})

        # Get refresh history (robust: walk past in-progress null endTime)
        print(f"   📊 Fetching refresh history...")
        from powerbi_connector import resolve_dataset_refresh_info, pick_best_refresh_from_history
        powerbi = PowerBIConnector(user_token=session.get('access_token'))
        refresh_history = powerbi.get_refresh_history(workspace_id, dataset_id, top=5)

        refresh_info = None
        if refresh_history and len(refresh_history) > 0:
            picked = pick_best_refresh_from_history(refresh_history)
            # Prefer the raw history row used for timestamp; fall back to newest
            source_idx = picked.get('source_index') if picked.get('source_index') is not None else 0
            source_idx = max(0, min(source_idx, len(refresh_history) - 1))
            last_refresh = refresh_history[source_idx] or refresh_history[0]
            refresh_info = {
                'status': picked.get('last_refresh_status') or last_refresh.get('status', 'Unknown'),
                'refreshType': last_refresh.get('refreshType', 'Unknown'),
                'startTime': last_refresh.get('startTime', ''),
                'endTime': picked.get('last_refreshed') or last_refresh.get('endTime', ''),
                'serviceExceptionJson': last_refresh.get('serviceExceptionJson', ''),
                'note': picked.get('refresh_note'),
            }
            print(f"      ✓ Last refresh: {refresh_info['status']} at {refresh_info['endTime']}")
        else:
            # Try full resolver for DirectQuery/Live labeling
            try:
                resolved = resolve_dataset_refresh_info(
                    headers=get_user_powerbi_headers(),
                    workspace_id=workspace_id,
                    dataset_id=dataset_id,
                    history_top=5,
                )
                if resolved.get('refresh_type') in ('directquery', 'live') or (
                    resolved.get('last_refresh_status') and 'directquery' in str(resolved.get('last_refresh_status')).lower()
                ):
                    refresh_info = {
                        'status': 'DirectQuery/Live',
                        'refreshType': 'DirectQuery/Live',
                        'startTime': '',
                        'endTime': '',
                        'serviceExceptionJson': '',
                        'note': resolved.get('refresh_note'),
                    }
                    print(f"      ℹ️ DirectQuery/Live dataset")
                else:
                    print(f"      ⚠ No refresh history available")
            except Exception:
                print(f"      ⚠ No refresh history available")

        # Calculate score breakdown
        critical_count = len([i for i in issues if i.get('severity') == 'Critical'])
        high_count = len([i for i in issues if i.get('severity') == 'High'])
        medium_count = len([i for i in issues if i.get('severity') == 'Medium'])
        warning_count = len(warnings)

        score_breakdown = {
            'starting_score': 100,
            'critical_issues': critical_count,
            'critical_deduction': critical_count * 20,
            'high_issues': high_count,
            'high_deduction': high_count * 10,
            'medium_issues': medium_count,
            'medium_deduction': medium_count * 5,
            'warnings': warning_count,
            'warning_deduction': warning_count * 2,
            'final_score': health_score
        }

        print(f"   ✅ Deep-dive crash test complete: {health_score}/100")
        print(f"   Issues: {len(issues)}, Warnings: {len(warnings)}")
        print(f"   Root Causes: {len(root_cause_analysis)}")
        if lineage_analysis:
            print(f"   Lineage: {lineage_analysis.get('affected_tables_count', 0)} tables analyzed")
        if change_impact_summary:
            breaking_changes = change_impact_summary.get('breaking_changes', [])
            print(f"   Breaking Changes: {len(breaking_changes)}")

        return jsonify({
            'success': True,
            'health_score': health_score,
            'status': 'excellent' if health_score >= 90 else 'good' if health_score >= 70 else 'fair' if health_score >= 50 else 'poor',
            'issues': issues,
            'warnings': warnings,
            'mode': mode,
            'timestamp': time.strftime('%Y-%m-%d %H:%M:%S'),
            'refresh_info': refresh_info,
            'score_breakdown': score_breakdown,
            'lineage_analysis': lineage_analysis,
            'root_cause_analysis': root_cause_analysis,  # NEW!
            'change_impact_summary': change_impact_summary  # NEW!
        })

    except Exception as e:
        print(f"   ❌ Crash test error: {e}")
        print(f"   Error type: {type(e).__name__}")
        import traceback
        traceback.print_exc()

        # Get full traceback as string for debugging
        import sys
        exc_info = sys.exc_info()
        tb_lines = traceback.format_exception(*exc_info)
        full_traceback = ''.join(tb_lines)

        print(f"\n{'='*80}")
        print(f"FULL TRACEBACK:")
        print(full_traceback)
        print(f"{'='*80}\n")

        return jsonify({
            'success': False,
            'error': str(e),
            'error_type': type(e).__name__,
            'traceback': full_traceback if os.getenv('FLASK_ENV') == 'development' else 'Enable debug mode to see traceback'
        }), 500


@app.route('/api/report-details/<report_id>')
@login_required
def get_report_details(report_id):
    """API endpoint to get detailed information for a specific report (lazy loading)"""
    try:
        workspace_id = request.args.get('workspace_id')
        dataset_id = request.args.get('dataset_id')

        if not workspace_id or not dataset_id:
            return jsonify({
                'success': False,
                'error': 'workspace_id and dataset_id are required'
            }), 400

        print(f"\n🔍 Fetching report details for report {report_id} in workspace {workspace_id}")
        print(f"   Dataset ID: {dataset_id}")

        # ✅ FIX: Use user-delegated token instead of service principal
        user_token = session.get('access_token')
        if user_token:
            powerbi.set_user_token(user_token)
            print(f"   🔑 Using user-delegated token for API calls")
        else:
            print(f"   ⚠️ No user token found, falling back to service principal")

        details = {}

        try:
            # Robust refresh resolution (history fallback, DirectQuery/Live, schedule)
            print(f"   → Resolving dataset refresh info...")
            refresh_info = powerbi.resolve_dataset_refresh(
                workspace_id=workspace_id,
                dataset_id=dataset_id,
                history_top=5,
            )

            details['refresh_type'] = refresh_info.get('refresh_type')
            details['refresh_note'] = refresh_info.get('refresh_note')
            details['last_refreshed'] = refresh_info.get('last_refreshed')
            details['last_refresh_status'] = refresh_info.get('last_refresh_status')

            schedule_text = refresh_info.get('refresh_schedule')
            if refresh_info.get('schedule_days') or refresh_info.get('schedule_times'):
                details['refresh_schedule'] = {
                    'enabled': bool(refresh_info.get('schedule_days') or refresh_info.get('schedule_times')),
                    'days': refresh_info.get('schedule_days') or [],
                    'times': refresh_info.get('schedule_times') or [],
                    'display': schedule_text,
                }
            elif schedule_text:
                details['refresh_schedule'] = schedule_text

            print(
                f"      ✓ Last refresh: {details.get('last_refreshed')} "
                f"({details.get('last_refresh_status')}) "
                f"type={details.get('refresh_type')}"
            )

            # Get dataset details
            print(f"   → Fetching dataset info...")
            dataset_info = powerbi.get_dataset_info(workspace_id, dataset_id)
            if dataset_info:
                configured_by = dataset_info.get('configuredBy', '')
                if configured_by:
                    details['last_accessed_by'] = configured_by

                created_date = dataset_info.get('createdDate', '')
                if created_date:
                    details['last_accessed'] = created_date
                print(f"      ✓ Dataset info retrieved")
            else:
                print(f"      ⚠ No dataset info available")

        except Exception as e:
            print(f"❌ Error getting report details: {str(e)}")
            import traceback
            traceback.print_exc()

        print(f"✅ Report details fetch complete. Found {len(details)} detail fields\n")
        return jsonify({
            'success': True,
            'report_id': report_id,
            'details': details
        })

    except Exception as e:
        print(f"❌ Error fetching report details: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


def scan_report_visual_columns(workspace_id, report_id, access_token, fabric_token=None):
    """
    Scan the report definition to extract all columns used in visuals across all pages.

    Uses two approaches:
    1. Fabric API getDefinition (PBIR-Legacy report.json) - most comprehensive (requires Fabric token)
    2. Fallback: Power BI Export API to download .pbix and parse report layout (uses Power BI token)

    Args:
        workspace_id: Power BI workspace ID
        report_id: Power BI report ID
        access_token: Bearer token for Power BI API (api.powerbi.com)
        fabric_token: Bearer token for Fabric API (api.fabric.microsoft.com), optional

    Returns:
        dict: {table_name: set(column_names)} of columns used in report visuals
    """
    # --- COMMENTED OUT: Visual scan methods disabled (returning 403/timeout) ---
    # All 4 methods (Fabric getDefinition, PBI token getDefinition, Export API, Pages API)
    # are currently returning 403 InsufficientScopes or timing out.
    # Returning empty dict; column usage is handled by the measure dependency fallback.
    return {}

    # import json
    # import base64
    #
    # used_columns = {}
    #
    # print(f"\n{'='*70}")
    # print(f"📊 SCANNING REPORT VISUALS FOR COLUMN USAGE")
    # print(f"   Report ID: {report_id}")
    # print(f"{'='*70}\n")
    #
    # # METHOD 1: Try Fabric API getDefinition (requires Fabric-scoped token)
    # try:
    #     if not fabric_token:
    #         print(f"   🔍 METHOD 1: Skipping Fabric API getDefinition (no Fabric token available)")
    #     else:
    #         print(f"   🔍 METHOD 1: Trying Fabric API getDefinition...")
    #
    #         fabric_headers = {
    #             'Authorization': f'Bearer {fabric_token}',
    #             'Content-Type': 'application/json'
    #         }
    #
    #         definition_url = f"https://api.fabric.microsoft.com/v1/workspaces/{workspace_id}/reports/{report_id}/getDefinition"
    #
    #         response = requests.post(definition_url, headers=fabric_headers, json={}, timeout=30)
    #
    #         if response.status_code == 200:
    #             definition = response.json()
    #             parts = definition.get('definition', {}).get('parts', [])
    #
    #             print(f"      ✓ Got report definition with {len(parts)} parts")
    #
    #             for part in parts:
    #                 path = part.get('path', '')
    #                 payload = part.get('payload', '')
    #
    #                 # PBIR-Legacy: report.json contains all visuals
    #                 # PBIR: visual.json files under definition/pages/*/visuals/*/
    #                 if path == 'report.json' or 'visual.json' in path or path.endswith('.json'):
    #                     try:
    #                         decoded = base64.b64decode(payload).decode('utf-8')
    #                         part_json = json.loads(decoded)
    #                         if path == 'report.json':
    #                             _extract_columns_from_report_json(part_json, used_columns)
    #                         elif 'visual.json' in path:
    #                             # PBIR visual file — extract directly
    #                             _extract_columns_from_visual_config(part_json, used_columns, path)
    #                         else:
    #                             # Other JSON — try both approaches
    #                             _extract_columns_from_report_json(part_json, used_columns)
    #                     except Exception as e:
    #                         print(f"      ⚠️  Error parsing {path}: {e}")
    #
    #             if used_columns:
    #                 print(f"      ✅ Fabric API: Found columns in {len(used_columns)} tables")
    #                 return used_columns
    #         elif response.status_code == 202:
    #             # Long-running operation - try to follow it
    #             print(f"      ⏳ LRO triggered, checking operation status...")
    #             operation_url = response.headers.get('Location', '')
    #             retry_after = int(response.headers.get('Retry-After', '5'))
    #
    #             import time
    #             for attempt in range(3):
    #                 time.sleep(retry_after)
    #                 op_response = requests.get(operation_url, headers=fabric_headers, timeout=30)
    #                 if op_response.status_code == 200:
    #                     definition = op_response.json()
    #                     parts = definition.get('definition', {}).get('parts', [])
    #                     for part in parts:
    #                         path = part.get('path', '')
    #                         payload = part.get('payload', '')
    #                         if path == 'report.json' or 'visual.json' in path or path.endswith('.json'):
    #                             try:
    #                                 decoded = base64.b64decode(payload).decode('utf-8')
    #                                 part_json = json.loads(decoded)
    #                                 if path == 'report.json':
    #                                     _extract_columns_from_report_json(part_json, used_columns)
    #                                 elif 'visual.json' in path:
    #                                     _extract_columns_from_visual_config(part_json, used_columns, path)
    #                                 else:
    #                                     _extract_columns_from_report_json(part_json, used_columns)
    #                             except Exception as e:
    #                                 print(f"      ⚠️  Error parsing {path}: {e}")
    #                     break
    #
    #             if used_columns:
    #                 print(f"      ✅ Fabric API (LRO): Found columns in {len(used_columns)} tables")
    #                 return used_columns
    #         else:
    #             print(f"      ⚠️  Fabric API returned {response.status_code}: {response.text[:200]}")
    # except Exception as e:
    #     print(f"      ⚠️  Fabric API error: {e}")

    # # METHOD 2: Try Fabric getDefinition with Power BI token (often has broader access)
    # if not used_columns:
    #     try:
    #         print(f"\n   🔍 METHOD 2: Trying Fabric getDefinition with Power BI token...")
    #
    #         pbi_fabric_headers = {
    #             'Authorization': f'Bearer {access_token}',
    #             'Content-Type': 'application/json'
    #         }
    #
    #         # Try with PBIR-Legacy format explicitly
    #         for fmt_param in ['?format=PBIR-Legacy', '']:
    #             def_url = f"https://api.fabric.microsoft.com/v1/workspaces/{workspace_id}/reports/{report_id}/getDefinition{fmt_param}"
    #             response = requests.post(def_url, headers=pbi_fabric_headers, json={}, timeout=30)
    #
    #             if response.status_code == 200:
    #                 definition = response.json()
    #                 parts = definition.get('definition', {}).get('parts', [])
    #                 print(f"      ✓ Got report definition with {len(parts)} parts (PBI token, fmt='{fmt_param}')")
    #
    #                 for part in parts:
    #                     path = part.get('path', '')
    #                     payload = part.get('payload', '')
    #                     if 'report' in path.lower() or path.endswith('.json'):
    #                         try:
    #                             decoded = base64.b64decode(payload).decode('utf-8')
    #                             report_json = json.loads(decoded)
    #                             _extract_columns_from_report_json(report_json, used_columns)
    #                         except Exception as e:
    #                             print(f"      ⚠️  Error parsing {path}: {e}")
    #
    #                 if used_columns:
    #                     print(f"      ✅ Fabric API (PBI token): Found columns in {len(used_columns)} tables")
    #                     return used_columns
    #             elif response.status_code == 202:
    #                 print(f"      ⏳ LRO triggered with PBI token, checking operation status...")
    #                 operation_url = response.headers.get('Location', '')
    #                 retry_after = int(response.headers.get('Retry-After', '5'))
    #                 import time
    #                 for _ in range(3):
    #                     time.sleep(retry_after)
    #                     op_response = requests.get(operation_url, headers=pbi_fabric_headers, timeout=30)
    #                     if op_response.status_code == 200:
    #                         definition = op_response.json()
    #                         parts = definition.get('definition', {}).get('parts', [])
    #                         for part in parts:
    #                             path = part.get('path', '')
    #                             payload = part.get('payload', '')
    #                             if 'report' in path.lower() or path.endswith('.json'):
    #                                 try:
    #                                     decoded = base64.b64decode(payload).decode('utf-8')
    #                                     report_json = json.loads(decoded)
    #                                     _extract_columns_from_report_json(report_json, used_columns)
    #                                 except Exception as e:
    #                                     print(f"      ⚠️  Error parsing {path}: {e}")
    #                         break
    #                 if used_columns:
    #                     print(f"      ✅ Fabric API (PBI token, LRO): Found columns in {len(used_columns)} tables")
    #                     return used_columns
    #             else:
    #                 print(f"      ⚠️  Fabric API (PBI token, fmt='{fmt_param}') returned {response.status_code}")
    #     except Exception as e:
    #         print(f"      ⚠️  Fabric API (PBI token) error: {e}")
    #
    # # METHOD 3: Try Power BI Export API to download .pbix
    # if not used_columns:
    #     try:
    #         print(f"\n   🔍 METHOD 3: Trying Power BI Export/Download report...")
    #
    #         export_url = f"https://api.powerbi.com/v1.0/myorg/groups/{workspace_id}/reports/{report_id}/Export"
    #         response = requests.get(export_url, headers={
    #             'Authorization': f'Bearer {access_token}'
    #         }, timeout=60)
    #
    #         if response.status_code == 200:
    #             print(f"      ✓ Downloaded report ({len(response.content)} bytes)")
    #
    #             import zipfile
    #             import io
    #
    #             try:
    #                 with zipfile.ZipFile(io.BytesIO(response.content)) as zf:
    #                     # Look for Report/Layout or report.json
    #                     for name in zf.namelist():
    #                         if 'Layout' in name or name == 'Report/Layout':
    #                             with zf.open(name) as f:
    #                                 layout_data = f.read()
    #                                 try:
    #                                     layout_json = json.loads(layout_data.decode('utf-16-le'))
    #                                 except:
    #                                     layout_json = json.loads(layout_data.decode('utf-8'))
    #                                 _extract_columns_from_report_json(layout_json, used_columns)
    #                                 print(f"      ✓ Parsed Layout from .pbix")
    #                                 break
    #                         elif name == 'report.json':
    #                             with zf.open(name) as f:
    #                                 report_json = json.loads(f.read().decode('utf-8'))
    #                                 _extract_columns_from_report_json(report_json, used_columns)
    #                                 print(f"      ✓ Parsed report.json from .pbix")
    #                                 break
    #             except zipfile.BadZipFile:
    #                 print(f"      ⚠️  Downloaded file is not a valid zip/pbix")
    #         else:
    #             print(f"      ⚠️  Export API returned {response.status_code}")
    #     except Exception as e:
    #         print(f"      ⚠️  Export API error: {e}")
    #
    # # METHOD 4: Use Power BI ExportToFile async API (works when direct Export is blocked)
    # if not used_columns:
    #     try:
    #         print(f"\n   🔍 METHOD 4: Trying Power BI Pages API for report structure...")
    #
    #         # Try getting pages which IS available with standard token
    #         pages_url = f"https://api.powerbi.com/v1.0/myorg/groups/{workspace_id}/reports/{report_id}/pages"
    #         pages_resp = requests.get(pages_url, headers={
    #             'Authorization': f'Bearer {access_token}'
    #         }, timeout=15)
    #
    #         if pages_resp.status_code == 200:
    #             pages = pages_resp.json().get('value', [])
    #             print(f"      ✓ Report has {len(pages)} pages: {[p.get('displayName', p.get('name')) for p in pages]}")
    #             # Pages API confirms the report is accessible but doesn't give visual field info
    #             # Log this for debugging — the pages themselves confirm connectivity
    #         else:
    #             print(f"      ⚠️  Pages API returned {pages_resp.status_code}")
    #     except Exception as e:
    #         print(f"      ⚠️  Pages API error: {e}")
    #
    # total_cols = sum(len(v) for v in used_columns.values())
    # print(f"\n   📊 Report visual scan complete: {len(used_columns)} tables, {total_cols} columns")
    # for table, cols in list(used_columns.items())[:5]:
    #     print(f"      📋 {table}: {list(cols)[:5]}{'...' if len(cols) > 5 else ''}")
    #
    # return used_columns


def _extract_columns_from_report_json(report_json, used_columns):
    """
    Parse a Power BI report.json / Layout JSON to extract all columns used in visuals.

    Extracts from:
    - Visual projections (queryRef: "TableName.ColumnName")
    - prototypeQuery Select (Entity + Property)
    - Filters (column references)
    - Sort expressions
    """
    import re
    import json

    # Convert to string for regex-based extraction as well
    json_str = json.dumps(report_json) if isinstance(report_json, dict) else str(report_json)

    # Pattern 1: queryRef values like "TableName.ColumnName"
    query_ref_pattern = r'"queryRef"\s*:\s*"([^"]+)"'
    for match in re.finditer(query_ref_pattern, json_str):
        ref = match.group(1)
        if '.' in ref:
            parts = ref.split('.', 1)
            table_name = parts[0]
            column_name = parts[1]
            if table_name not in used_columns:
                used_columns[table_name] = set()
            used_columns[table_name].add(column_name)

    # Pattern 2: Entity + Property pairs from prototypeQuery Select
    # "Entity":"TableName" ... "Property":"ColumnName"
    entity_prop_pattern = r'"Entity"\s*:\s*"([^"]+)"[^}]*?"Property"\s*:\s*"([^"]+)"'
    for match in re.finditer(entity_prop_pattern, json_str):
        table_name = match.group(1)
        column_name = match.group(2)
        if table_name not in used_columns:
            used_columns[table_name] = set()
        used_columns[table_name].add(column_name)

    # Pattern 3: NativeReferenceName patterns like "TableName.ColumnName"
    native_ref_pattern = r'"NativeReferenceName"\s*:\s*"([^"]+)"'
    for match in re.finditer(native_ref_pattern, json_str):
        ref = match.group(1)
        # Some NativeReferenceName are just column names, skip those
        # Only process if it looks like Table.Column
        if '.' in ref:
            parts = ref.split('.', 1)
            table_name = parts[0]
            column_name = parts[1]
            if table_name not in used_columns:
                used_columns[table_name] = set()
            used_columns[table_name].add(column_name)

    # Pattern 4: Filter column references
    # "Column":{"Expression":{"SourceRef":{"Entity":"Table"}},"Property":"Column"}
    filter_col_pattern = r'"SourceRef"\s*:\s*\{\s*"Entity"\s*:\s*"([^"]+)"\s*\}[^}]*?"Property"\s*:\s*"([^"]+)"'
    for match in re.finditer(filter_col_pattern, json_str):
        table_name = match.group(1)
        column_name = match.group(2)
        if table_name not in used_columns:
            used_columns[table_name] = set()
        used_columns[table_name].add(column_name)

    # Pattern 5: HierarchyLevel references (matrix visuals, drilldowns)
    # "Level":"ColumnName" near "Entity":"TableName"
    level_pattern = r'"Level"\s*:\s*"([^"]+)"'
    for match in re.finditer(level_pattern, json_str):
        level_name = match.group(1)
        # Look backwards in the string for the nearest Entity reference
        start_pos = max(0, match.start() - 500)
        context = json_str[start_pos:match.start()]
        entity_match = re.search(r'"Entity"\s*:\s*"([^"]+)"', context)
        if entity_match:
            table_name = entity_match.group(1)
            if table_name not in used_columns:
                used_columns[table_name] = set()
            used_columns[table_name].add(level_name)

    # Pattern 6: DAX expression references in visual configs (Table[Column] pattern)
    dax_col_pattern = r"(?:'([^']+)'|([A-Za-z_][A-Za-z0-9_]*))\[([^\]]+)\]"
    # Only search in DAX-like contexts (measure expressions, calculated fields)
    dax_contexts = re.findall(r'"expression"\s*:\s*"([^"]*\[.*?\][^"]*)"', json_str, re.IGNORECASE)
    for dax_expr in dax_contexts:
        matches = re.findall(dax_col_pattern, dax_expr)
        for m in matches:
            tbl = m[0] if m[0] else m[1]
            col = m[2]
            if tbl and col:
                if tbl not in used_columns:
                    used_columns[tbl] = set()
                used_columns[tbl].add(col)

    # Also try to parse the structured sections if available
    try:
        sections = report_json.get('sections', [])
        for section in sections:
            page_name = section.get('displayName', section.get('name', 'Unknown'))
            containers = section.get('visualContainers', [])

            for container in containers:
                config_str = container.get('config', '{}')
                try:
                    config = json.loads(config_str) if isinstance(config_str, str) else config_str
                    _extract_columns_from_visual_config(config, used_columns, page_name)
                except json.JSONDecodeError:
                    pass

                # Also check filters on the container
                filters_str = container.get('filters', '[]')
                try:
                    filters = json.loads(filters_str) if isinstance(filters_str, str) else filters_str
                    _extract_columns_from_filters(filters, used_columns)
                except json.JSONDecodeError:
                    pass

            # Page-level filters
            page_filters_str = section.get('filters', '[]')
            try:
                page_filters = json.loads(page_filters_str) if isinstance(page_filters_str, str) else page_filters_str
                _extract_columns_from_filters(page_filters, used_columns)
            except json.JSONDecodeError:
                pass
    except (AttributeError, TypeError):
        pass  # report_json may not have structured sections

    # Report-level filters
    try:
        report_filters_str = report_json.get('filters', '[]')
        report_filters = json.loads(report_filters_str) if isinstance(report_filters_str, str) else report_filters_str
        _extract_columns_from_filters(report_filters, used_columns)
    except (AttributeError, TypeError, json.JSONDecodeError):
        pass


def _extract_columns_from_visual_config(config, used_columns, page_name=""):
    """Extract column references from a visual's config JSON"""

    single_visual = config.get('singleVisual', {})
    if not single_visual:
        return

    visual_type = single_visual.get('visualType', 'unknown')

    # Extract from projections
    projections = single_visual.get('projections', {})
    for role, fields in projections.items():
        if isinstance(fields, list):
            for field in fields:
                query_ref = field.get('queryRef', '')
                if '.' in query_ref:
                    parts = query_ref.split('.', 1)
                    table_name = parts[0]
                    column_name = parts[1]
                    if table_name not in used_columns:
                        used_columns[table_name] = set()
                    used_columns[table_name].add(column_name)
                    print(f"         ✓ Page '{page_name}' visual '{visual_type}' → {table_name}[{column_name}]")

    # Extract from prototypeQuery
    proto_query = single_visual.get('prototypeQuery', {})
    selects = proto_query.get('Select', [])
    from_clauses = proto_query.get('From', [])

    # Build alias-to-entity mapping
    alias_map = {}
    for frm in from_clauses:
        alias = frm.get('Name', '')
        entity = frm.get('Entity', '')
        if alias and entity:
            alias_map[alias] = entity

    for sel in selects:
        # Column references
        col_ref = sel.get('Column', {})
        if col_ref:
            expr = col_ref.get('Expression', {})
            source_ref = expr.get('SourceRef', {})
            source_alias = source_ref.get('Source', '')
            prop = col_ref.get('Property', '')
            table_name = alias_map.get(source_alias, source_alias)
            if table_name and prop:
                if table_name not in used_columns:
                    used_columns[table_name] = set()
                used_columns[table_name].add(prop)

        # Measure references
        measure_ref = sel.get('Measure', {})
        if measure_ref:
            expr = measure_ref.get('Expression', {})
            source_ref = expr.get('SourceRef', {})
            source_alias = source_ref.get('Source', '')
            prop = measure_ref.get('Property', '')
            table_name = alias_map.get(source_alias, source_alias)
            if table_name and prop:
                if table_name not in used_columns:
                    used_columns[table_name] = set()
                used_columns[table_name].add(prop)

        # Aggregation references
        agg_ref = sel.get('Aggregation', {})
        if agg_ref:
            agg_expr = agg_ref.get('Expression', {})
            col_inner = agg_expr.get('Column', {})
            if col_inner:
                inner_expr = col_inner.get('Expression', {})
                source_ref = inner_expr.get('SourceRef', {})
                source_alias = source_ref.get('Source', '')
                prop = col_inner.get('Property', '')
                table_name = alias_map.get(source_alias, source_alias)
                if table_name and prop:
                    if table_name not in used_columns:
                        used_columns[table_name] = set()
                    used_columns[table_name].add(prop)

        # HierarchyLevel references (common in matrix visuals)
        # {"HierarchyLevel":{"Expression":{"Hierarchy":{"Expression":{"SourceRef":{"Source":"t"}},"Hierarchy":"HierName"}},"Level":"ColName"}}
        hier_level = sel.get('HierarchyLevel', {})
        if hier_level:
            level_name = hier_level.get('Level', '')
            hier_expr = hier_level.get('Expression', {})
            hier_ref = hier_expr.get('Hierarchy', {})
            if hier_ref:
                inner_expr = hier_ref.get('Expression', {})
                source_ref = inner_expr.get('SourceRef', {})
                source_alias = source_ref.get('Source', '')
                table_name = alias_map.get(source_alias, source_alias)
                if table_name and level_name:
                    if table_name not in used_columns:
                        used_columns[table_name] = set()
                    used_columns[table_name].add(level_name)

        # Hierarchy references (the hierarchy itself — mark the hierarchy name as a reference)
        hier_ref = sel.get('Hierarchy', {})
        if hier_ref:
            inner_expr = hier_ref.get('Expression', {})
            source_ref = inner_expr.get('SourceRef', {})
            source_alias = source_ref.get('Source', '')
            hier_name = hier_ref.get('Hierarchy', '')
            table_name = alias_map.get(source_alias, source_alias)
            if table_name and hier_name:
                if table_name not in used_columns:
                    used_columns[table_name] = set()
                used_columns[table_name].add(hier_name)


def _extract_columns_from_filters(filters, used_columns):
    """Extract column references from filter definitions, including hierarchy and nested filters"""
    if not isinstance(filters, list):
        return

    def _extract_col_from_expr(expr):
        """Recursively extract column refs from a filter expression dict"""
        if not isinstance(expr, dict):
            return
        # Column reference
        col = expr.get('Column', {})
        if col:
            entity_expr = col.get('Expression', {})
            source_ref = entity_expr.get('SourceRef', {})
            table_name = source_ref.get('Entity', '')
            column_name = col.get('Property', '')
            if table_name and column_name:
                if table_name not in used_columns:
                    used_columns[table_name] = set()
                used_columns[table_name].add(column_name)

        # HierarchyLevel filter
        hier_level = expr.get('HierarchyLevel', {})
        if hier_level:
            level_name = hier_level.get('Level', '')
            hier_expr = hier_level.get('Expression', {})
            hier_ref = hier_expr.get('Hierarchy', {})
            if hier_ref:
                inner_expr = hier_ref.get('Expression', {})
                source_ref = inner_expr.get('SourceRef', {})
                table_name = source_ref.get('Entity', '')
                if table_name and level_name:
                    if table_name not in used_columns:
                        used_columns[table_name] = set()
                    used_columns[table_name].add(level_name)

        # Measure reference in filters
        measure = expr.get('Measure', {})
        if measure:
            m_expr = measure.get('Expression', {})
            source_ref = m_expr.get('SourceRef', {})
            table_name = source_ref.get('Entity', '')
            measure_name = measure.get('Property', '')
            if table_name and measure_name:
                if table_name not in used_columns:
                    used_columns[table_name] = set()
                used_columns[table_name].add(measure_name)

        # Recurse into nested expressions (And, Or, Not, Comparison, etc.)
        for key in ('Left', 'Right', 'Expression', 'Condition', 'And', 'Or', 'Not'):
            nested = expr.get(key)
            if isinstance(nested, dict):
                _extract_col_from_expr(nested)
            elif isinstance(nested, list):
                for item in nested:
                    _extract_col_from_expr(item)

    for f in filters:
        try:
            expr = f.get('expression', {})
            _extract_col_from_expr(expr)

            # Also check the 'filter' key (some formats use this)
            filter_inner = f.get('filter', {})
            if filter_inner:
                where = filter_inner.get('Where', [])
                for w in (where if isinstance(where, list) else [where]):
                    if isinstance(w, dict):
                        cond = w.get('Condition', {})
                        _extract_col_from_expr(cond)
        except (AttributeError, TypeError):
            pass


@app.route('/api/lineage')
@login_required
def get_lineage():
    """API endpoint to get query and table lineage for a specific report"""
    try:
        import requests
        import re
        workspace_id = request.args.get('workspace_id')
        report_id = request.args.get('report_id')

        if not workspace_id or not report_id:
            return jsonify({
                'success': False,
                'error': 'workspace_id and report_id are required'
            }), 400

        # Get user token from session
        user_token = session.get('access_token')

        if not user_token:
            return jsonify({
                'success': False,
                'error': 'Not authenticated'
            }), 401

        headers = {
            'Authorization': f'Bearer {user_token}',
            'Content-Type': 'application/json'
        }

        # Step 1: Get report details to find dataset ID
        report_url = f'https://api.powerbi.com/v1.0/myorg/groups/{workspace_id}/reports/{report_id}'
        report_response = requests.get(report_url, headers=headers)

        if report_response.status_code != 200:
            return jsonify({
                'success': False,
                'error': f'Failed to fetch report: {report_response.status_code}'
            }), report_response.status_code

        report_data = report_response.json()
        dataset_id = report_data.get('datasetId')

        if not dataset_id:
            return jsonify({
                'success': True,
                'queries': []
            })

        # Step 2: Get all workspaces to create workspace name mapping for Dataflows
        workspace_map = {}
        try:
            # Get workspaces from session cache or API
            workspaces_cache_key = f"workspaces_{session.get('user_id')}"
            cached_workspaces = session.get(workspaces_cache_key)

            if cached_workspaces:
                workspace_map = {ws['id']: ws['name'] for ws in cached_workspaces}
            else:
                # Fetch workspaces if not cached
                workspaces_url = 'https://api.powerbi.com/v1.0/myorg/groups'
                workspaces_response = requests.get(workspaces_url, headers=headers)
                if workspaces_response.status_code == 200:
                    workspaces_data = workspaces_response.json()
                    workspace_map = {ws['id']: ws['name'] for ws in workspaces_data.get('value', [])}
        except Exception as e:
            print(f"   ⚠️ Could not fetch workspace names for Dataflow resolution: {e}")

        # Step 3: Use Scanner API to get dataset expressions and table information
        try:
            import time as time_module
            from concurrent.futures import ThreadPoolExecutor
            from scanner_connector import PowerBIScanner

            lineage_start_time = time_module.time()

            # Check scanner model cache (keyed by dataset_id to avoid redundant scans)
            model_cache_key = f"{workspace_id}_{dataset_id}"
            current_time = time_module.time()
            cached_model = scanner_cache.get(model_cache_key)

            if cached_model and cached_model.get('data') and \
               (current_time - cached_model.get('timestamp', 0)) < SCANNER_CACHE_DURATION:
                print(f"   ⚡ Using cached dataset model (age: {int(current_time - cached_model['timestamp'])}s)")
                model = cached_model['data']
            else:
                scanner = PowerBIScanner()
                print(f"   🔍 Fetching dataset model for dataset {dataset_id}...")
                model = scanner.get_dataset_model(dataset_id, workspace_id=workspace_id)
                # Cache the model result
                scanner_cache[model_cache_key] = {
                    'data': model,
                    'timestamp': time_module.time()
                }
                print(f"   💾 Dataset model cached (key: {model_cache_key})")

            scanner_elapsed = time_module.time() - lineage_start_time
            print(f"   ⏱️ Scanner step took {scanner_elapsed:.1f}s")

            queries = []

            # Get all tables from the model - extract table names (strings) not dict objects
            model_tables = model.get('tables', [])
            all_table_names = set()
            for t in model_tables:
                if isinstance(t, dict):
                    table_name = t.get('name') or t.get('table')
                    if table_name:
                        all_table_names.add(table_name)
                elif isinstance(t, str) and t:
                    all_table_names.add(t)

            # Get columns dictionary from the model
            all_columns = model.get('columns', {})

            print(f"   📊 Dataset contains {len(all_table_names)} total tables/queries")
            print(f"   🔍 Available tables: {list(all_table_names)}")
            print(f"   📋 Column data structure keys: {list(all_columns.keys())[:5] if all_columns else 'None'}")
            if all_columns:
                first_table = list(all_columns.keys())[0] if all_columns else None
                if first_table:
                    print(f"   📋 Sample columns for '{first_table}': {all_columns[first_table][:2] if all_columns[first_table] else 'None'}")
            print(f"   🔍 Processing {len(model.get('expressions', []))} M expressions...")

            # Build column usage map from multiple sources
            # Format: {table_name: {column_name: set_of_sources}}
            # Each source label describes HOW the column is used
            column_usage = {}
            import re
            dax_pattern = r"(?:'([^']+)'|(\w+))\[([^\]]+)\]"

            def _mark_used(tbl, col_nm, source_label):
                """Helper to add a usage source for a column"""
                if tbl and col_nm:
                    if tbl not in column_usage:
                        column_usage[tbl] = {}
                    if col_nm not in column_usage[tbl]:
                        column_usage[tbl][col_nm] = set()
                    column_usage[tbl][col_nm].add(source_label)

            # SOURCE 1: Scanner API's built-in isReferenced flag on columns
            scanner_ref_count = 0
            for table_name, table_cols in all_columns.items():
                for col in table_cols:
                    if col.get('isReferenced') is True:
                        _mark_used(table_name, col.get('name'), 'Scanner API')
                        scanner_ref_count += 1
            print(f"   ✅ Scanner isReferenced: {scanner_ref_count} columns flagged across {len(column_usage)} tables")

            # SOURCE 2: Parse DAX measure expressions for Table[Column] references
            measures = model.get('measures', [])
            measure_col_count = 0
            for measure in measures:
                measure_expr = measure.get('expression', '')
                if measure_expr:
                    matches = re.findall(dax_pattern, measure_expr)
                    for match in matches:
                        table_name = match[0] if match[0] else match[1]
                        column_name = match[2]
                        _mark_used(table_name, column_name, 'DAX Measure')
                        measure_col_count += 1
            print(f"   ✅ Measures: {measure_col_count} column refs")

            # SOURCE 3: Parse calculated column expressions for Table[Column] references
            calc_col_count = 0
            for table_name, table_cols in all_columns.items():
                for col in table_cols:
                    calc_expr = col.get('expression', '')
                    if calc_expr and col.get('columnType') == 'Calculated':
                        matches = re.findall(dax_pattern, calc_expr)
                        for match in matches:
                            ref_table = match[0] if match[0] else match[1]
                            ref_col = match[2]
                            _mark_used(ref_table, ref_col, 'Calculated Column')
                            calc_col_count += 1
                        # Also mark the calculated column itself as used
                        _mark_used(table_name, col.get('name'), 'Calculated Column')
            print(f"   ✅ Calculated columns: {calc_col_count} column refs")

            # SOURCE 4: Extract columns from SQL SELECT clauses in M expressions
            sql_col_count = 0
            for expr_data in model.get('expressions', []):
                expr_table = expr_data.get('table', '')
                expression = expr_data.get('expression', '')
                if expression and ('Sql.Database' in expression or 'Query=' in expression):
                    # Extract the SQL query from the M expression
                    sql_query_pattern = r'Query\s*=\s*"([^"]*)"'
                    sql_matches = re.findall(sql_query_pattern, expression, re.IGNORECASE)
                    for sql_query in sql_matches:
                        sql_query = sql_query.replace('""', '"').replace('#(lf)', '\n').replace('#(cr)', '\r').replace('#(tab)', '\t')
                        sql_columns = extract_sql_column_names(sql_query)
                        for col in sql_columns:
                            _mark_used(expr_table, col, 'SQL Query')
                            sql_col_count += 1
            print(f"   ✅ SQL SELECT columns: {sql_col_count} column refs")

            # SOURCE 5: Query dataset model via DAX for column metadata (SortByColumn, IsHidden)
            dax_meta_count = 0
            try:
                meta_url = f"https://api.powerbi.com/v1.0/myorg/groups/{workspace_id}/datasets/{dataset_id}/executeQueries"
                meta_headers = {
                    'Authorization': f'Bearer {user_token}',
                    'Content-Type': 'application/json'
                }
                # INFO.VIEW.COLUMNS() returns column metadata including IsHidden, SortByColumn
                meta_body = {
                    "queries": [{"query": "EVALUATE INFO.VIEW.COLUMNS()"}],
                    "serializerSettings": {"includeNulls": False}
                }
                meta_resp = requests.post(meta_url, headers=meta_headers, json=meta_body, timeout=15)
                if meta_resp.status_code == 200:
                    meta_result = meta_resp.json()
                    if 'results' in meta_result and meta_result['results']:
                        meta_table = meta_result['results'][0].get('tables', [{}])[0]
                        meta_rows = meta_table.get('rows', [])
                        for row in meta_rows:
                            tbl = row.get('[TableName]', row.get('TableName', ''))
                            col = row.get('[ColumnName]', row.get('ColumnName', ''))
                            is_hidden = row.get('[IsHidden]', row.get('IsHidden', True))
                            sort_by = row.get('[SortByColumn]', row.get('SortByColumn', ''))
                            if not tbl or not col:
                                continue
                            # Skip auto-generated date tables
                            if tbl.startswith('LocalDateTable_') or tbl.startswith('DateTableTemplate_'):
                                continue
                            # Columns that are NOT hidden are visible in the model → likely used
                            if is_hidden is False:
                                _mark_used(tbl, col, 'Model (Visible)')
                                dax_meta_count += 1
                            # Columns used as SortByColumn are definitely used
                            if sort_by:
                                _mark_used(tbl, sort_by, 'Sort By Column')
                                dax_meta_count += 1
                    print(f"   ✅ DAX column metadata: {dax_meta_count} columns from INFO.VIEW.COLUMNS()")
                else:
                    print(f"   ⚠️  INFO.VIEW.COLUMNS() returned {meta_resp.status_code} — skipping")
            except Exception as meta_err:
                print(f"   ⚠️  DAX column metadata error: {meta_err}")

            # SOURCE 6: Parse M expressions for column operations (SelectColumns, RenameColumns, etc.)
            m_col_count = 0
            for expr_data in model.get('expressions', []):
                expr_table = expr_data.get('table', '')
                expression = expr_data.get('expression', '')
                if expression:
                    m_cols = extract_m_expression_columns(expression, expr_table)
                    for col in m_cols:
                        _mark_used(expr_table, col, 'M Expression')
                        m_col_count += 1
            print(f"   ✅ M expression columns: {m_col_count} column refs")

            # Run XMLA column usage and Report visual scan IN PARALLEL for performance
            parallel_start = time_module.time()
            print(f"\n   ⚡ Running XMLA + Visual scan in parallel...")

            # IMPORTANT: Acquire tokens BEFORE spawning threads
            # Flask session/request context is NOT available inside ThreadPoolExecutor threads
            try:
                fabric_token_for_scan = get_user_fabric_token()
            except Exception as e:
                print(f"   ⚠️  Could not acquire Fabric token: {e}")
                fabric_token_for_scan = None

            # Capture values from request context before entering threads
            _ws_id = workspace_id
            _rpt_id = report_id
            _usr_token = user_token

            visual_result = {}

            def _run_visual_scan():
                try:
                    return scan_report_visual_columns(_ws_id, _rpt_id, _usr_token, fabric_token=fabric_token_for_scan)
                except Exception as e:
                    print(f"   ⚠️  Visual scan failed: {e}")
                    import traceback
                    traceback.print_exc()
                    return {}

            with ThreadPoolExecutor(max_workers=1) as executor:
                visual_future = executor.submit(_run_visual_scan)
                try:
                    visual_result = visual_future.result(timeout=120)
                except TimeoutError:
                    print(f"   ⚠️  Visual scan timed out after 120s — continuing without visual column data")
                    visual_result = {}
                except Exception as e:
                    print(f"   ⚠️  Visual scan future error: {e}")
                    visual_result = {}

            parallel_elapsed = time_module.time() - parallel_start
            print(f"   ⏱️ Visual scan took {parallel_elapsed:.1f}s")

            # Merge visual scan results
            if visual_result:
                vis_col_count = sum(len(v) for v in visual_result.values())
                print(f"   ✅ Visual scan found {vis_col_count} columns across {len(visual_result)} tables")
                for table, columns in visual_result.items():
                    for col in columns:
                        _mark_used(table, col, 'Report Visual')
                # Log sample visual scan tables for debugging
                for tbl, cols in list(visual_result.items())[:3]:
                    sample_cols = list(cols)[:5]
                    print(f"      📋 Visual: {tbl} → {sample_cols}{'...' if len(cols) > 5 else ''}")
            else:
                print(f"   ⚠️  Visual scan returned no results — using measure dependency fallback")
                print(f"      Fabric token available: {fabric_token_for_scan is not None}")

                # FALLBACK: When visual scan fails, mark columns that are referenced
                # by measures as "Report (via Measure)" since measures exist to be used
                # in visuals. Also mark all non-hidden columns from the model.
                fallback_count = 0
                measures = model.get('measures', [])
                dax_pattern_fb = r"(?:'([^']+)'|(\w+))\[([^\]]+)\]"
                for measure in measures:
                    m_expr = measure.get('expression', '')
                    m_table = measure.get('table', '')
                    if m_expr:
                        matches = re.findall(dax_pattern_fb, m_expr)
                        for match in matches:
                            ref_table = match[0] if match[0] else match[1]
                            ref_col = match[2]
                            _mark_used(ref_table, ref_col, 'Report (via Measure)')
                            fallback_count += 1
                    # Mark the measure itself as used in its table
                    if m_table and measure.get('name'):
                        _mark_used(m_table, measure.get('name'), 'Report (via Measure)')
                        fallback_count += 1
                print(f"   ✅ Measure fallback added {fallback_count} column refs")

            total_usage = sum(len(v) for v in column_usage.values())
            print(f"   ✅ Column usage map after all sources: {total_usage} columns across {len(column_usage)} tables")

            # Process each expression (M query)
            expr_start = time_module.time()
            for expr_data in model.get('expressions', []):
                query_name = expr_data.get('table', 'Unknown Query')
                expression = expr_data.get('expression', '')

                # Skip tables with dummy/null M expressions (these are measure-only tables)
                if not expression or expression.strip().lower() in ['null', '#"null"', ''] or len(expression.strip()) < 20:
                    print(f"   ⏭️  Skipping table '{query_name}' - has dummy M expression (probably measure-only table)")
                    continue

                # Parse the M expression to find table references
                tables_used = parse_m_expression_for_tables(expression, all_table_names)

                # Exclude self-references
                tables_used = [t for t in tables_used if t != query_name]

                # Determine query type and source type from the M expression
                query_type, source_type = analyze_m_expression(expression)

                # For Expression type, extract just the source line
                display_expression = expression
                if source_type == 'Expression':
                    display_expression = extract_source_line(expression)

                # Extract server/source name
                server_name = extract_server_name(expression, source_type, workspace_map)

                # Build column info from the Power BI model
                query_columns = all_columns.get(query_name, [])
                column_info = [{
                    'name': col.get('name'),
                    'dataType': col.get('dataType', 'Unknown'),
                    'isReferenced': False,
                    'usedIn': '',
                    'columnType': col.get('columnType', 'Data'),
                    'expression': col.get('expression', '') if col.get('columnType') == 'Calculated' else ''
                } for col in query_columns if col.get('name')]

                # Add measures from this table
                table_measures = [m for m in model.get('measures', []) if m.get('table') == query_name]
                measure_info = [{
                    'name': m.get('name'),
                    'dataType': 'Measure',
                    'isReferenced': False,
                    'usedIn': '',
                    'columnType': 'Measure',
                    'expression': m.get('expression', ''),
                    'description': m.get('description', '')
                } for m in table_measures]

                # Build the tablesWithColumns structure
                # Use the Power BI Model table name (query_name), not SQL source table names
                tables_with_columns = [{
                    'tableName': query_name,  # Power BI Model table name
                    'sqlSourceTables': tables_used,  # SQL source tables used by this query
                    'columns': column_info + measure_info  # Include both columns and measures
                }]

                # Extract SQL query from M expression for display
                sql_query = None
                if source_type == 'SQL Server':
                    # Extract SQL query from M expression
                    sql_query_pattern = r'Query\s*=\s*"([^"]*(?:""[^"]*)*)"'
                    sql_matches = re.findall(sql_query_pattern, expression, re.IGNORECASE)
                    if sql_matches:
                        # Clean up the SQL query (remove M escape sequences)
                        sql_query = sql_matches[0].replace('""', '"').replace('#(lf)', '\n').replace('#(cr)', '\r').replace('#(tab)', '\t')

                queries.append({
                    'queryName': query_name,
                    'tables': tables_used,
                    'tablesWithColumns': tables_with_columns,  # NEW: table-to-columns mapping
                    'queryType': query_type,
                    'sourceType': source_type,
                    'expression': display_expression,  # Source line for Expression, full for others
                    'serverName': server_name,  # Server/source identifier
                    'sqlQuery': sql_query  # NEW: Full SQL query for SQL Server sources
                })

            # Also add calculated tables (tables without M expressions)
            # These include DAX-calculated tables (created with CALENDAR, CALENDARAUTO, etc.)
            tables_with_expressions = set([q['queryName'] for q in queries])

            # Build a map of DAX table expressions from the model
            # These come from the Scanner API's partition/source extraction
            dax_table_expressions = {}
            for expr_data in model.get('expressions', []):
                if expr_data.get('expressionType') == 'DAX':
                    table_name = expr_data.get('table', '')
                    expression = expr_data.get('expression', '')
                    if table_name and expression:
                        dax_table_expressions[table_name] = expression

            print(f"   📊 Found DAX table expressions for {len(dax_table_expressions)} calculated table(s)")
            if dax_table_expressions:
                print(f"      Tables with DAX expressions: {list(dax_table_expressions.keys())}")

            print(f"   🔍 Checking for calculated tables...")
            print(f"      Total tables in model: {len(all_table_names)}")
            print(f"      Tables with M expressions: {len(tables_with_expressions)}")
            print(f"      Tables to add as calculated: {len(all_table_names - tables_with_expressions)}")
            if all_table_names - tables_with_expressions:
                print(f"      Calculated table names: {list(all_table_names - tables_with_expressions)[:10]}")  # Show first 10

            for table_name in all_table_names:
                if table_name not in tables_with_expressions:
                    # Check if this table has a DAX expression (e.g., Calendar table)
                    dax_expr = dax_table_expressions.get(table_name, '')

                    # Determine if it's a Calendar/Date table
                    source_type = 'Calculated Table'
                    query_type = 'Calculated Table'
                    if dax_expr:
                        dax_upper = dax_expr.upper()
                        if 'CALENDAR' in dax_upper:
                            query_type = 'Calendar Table (DAX)'
                            source_type = 'DAX Function'
                        elif 'GENERATE' in dax_upper or 'ADDCOLUMNS' in dax_upper:
                            query_type = 'DAX Calculated Table'
                            source_type = 'DAX Expression'

                    # Build column info from the Power BI model
                    query_columns = all_columns.get(table_name, [])
                    column_info = [{
                        'name': col.get('name'),
                        'dataType': col.get('dataType', 'Unknown'),
                        'isReferenced': False,
                        'usedIn': '',
                        'columnType': col.get('columnType', 'Data'),
                        'expression': col.get('expression', '') if col.get('columnType') == 'Calculated' else ''
                    } for col in query_columns if col.get('name')]

                    # Add measures from this table
                    table_measures = [m for m in model.get('measures', []) if m.get('table') == table_name]
                    measure_info = [{
                        'name': m.get('name'),
                        'dataType': 'Measure',
                        'isReferenced': False,
                        'usedIn': '',
                        'columnType': 'Measure',
                        'expression': m.get('expression', ''),
                        'description': m.get('description', '')
                    } for m in table_measures]

                    # Build the tablesWithColumns structure
                    tables_with_columns = [{
                        'tableName': table_name,  # Power BI Model table name
                        'sqlSourceTables': [],  # No SQL sources for calculated tables
                        'columns': column_info + measure_info  # Include both columns and measures
                    }]

                    queries.append({
                        'queryName': table_name,
                        'tables': [],
                        'tablesWithColumns': tables_with_columns,  # NOW INCLUDES COLUMNS!
                        'queryType': query_type,
                        'sourceType': source_type,
                        'expression': dax_expr,  # DAX expression for Calendar/Calculated tables
                        'serverName': 'N/A',
                        'sqlQuery': ''
                    })

            expr_elapsed = time_module.time() - expr_start
            print(f"   ⏱️ Expression processing took {expr_elapsed:.1f}s — {len(queries)} queries/tables, {sum(len(q['tables']) for q in queries)} dependencies")

            # Get relationships from the model
            relationships = model.get('relationships', [])
            print(f"   🔗 Scanner returned {len(relationships)} relationships")

            # ⚡ PERFORMANCE OPTIMIZATION: XMLA only used as fallback when Scanner lacks data
            # If Scanner API didn't return relationships, try XMLA endpoint as fallback
            if not relationships:
                print(f"   ⚡ XMLA fallback for relationships...")
                try:
                    from xmla_connector import XMLAConnector
                    xmla = XMLAConnector(workspace_id, dataset_id, user_token)
                    xmla_result = xmla.get_model_metadata()

                    if xmla_result.get('relationships'):
                        relationships = xmla_result['relationships']
                        print(f"   ✅ XMLA fallback: {len(relationships)} relationships")
                    else:
                        print(f"   ℹ️  XMLA fallback: no relationships found")
                except ImportError:
                    print(f"   ⚠️  XMLA connector not available")
                except Exception as xmla_error:
                    print(f"   ⚠️  XMLA fallback error: {xmla_error}")

            print(f"   📋 Final: {len(relationships)} relationships")

            # Extract table schema for diagram visualization
            table_schemas = {}
            columns_dict = model.get('columns', {})

            for table_name, columns in columns_dict.items():
                if columns:
                    table_schemas[table_name] = {
                        'name': table_name,
                        'columns': [
                            {
                                'name': col.get('name'),
                                'dataType': col.get('dataType', 'Unknown')
                            } for col in columns if col.get('name')
                        ]
                    }

            print(f"   📊 Extracted schema for {len(table_schemas)} tables")

            # Create a mapping from internal table names to user-friendly query names
            # This helps replace technical names like "LocalDateTable_xxx" with actual query names
            table_name_mapping = {}

            # First, add all tables with M expressions
            for expr_data in model.get('expressions', []):
                query_name = expr_data.get('table', '')
                if query_name:
                    # The 'table' field is the internal table name in the model
                    table_name_mapping[query_name] = query_name

            # Also add all tables from the columns dictionary (this includes auto-generated tables)
            for table_name in columns_dict.keys():
                if table_name not in table_name_mapping:
                    # For LocalDateTable_xxx, create a friendly name
                    if table_name.startswith('LocalDateTable_'):
                        # Extract the GUID and create a short friendly name
                        friendly_name = f"Date Table (Auto)"
                        table_name_mapping[table_name] = friendly_name
                    elif table_name.startswith('DateTableTemplate_'):
                        friendly_name = f"Date Table Template"
                        table_name_mapping[table_name] = friendly_name
                    else:
                        # For other tables, use the name as-is
                        table_name_mapping[table_name] = table_name

            print(f"   📝 Created table name mapping for {len(table_name_mapping)} tables (including auto-generated)")

            # Apply the mapping to relationships and filter out auto-generated date tables
            if relationships:
                filtered_relationships = []
                seen_pairs = set()  # Track unique relationship pairs to avoid duplicates

                for rel in relationships:
                    original_from = rel.get('fromTable', '')
                    original_to = rel.get('toTable', '')

                    # Skip relationships involving LocalDateTable or DateTableTemplate
                    if (original_from.startswith('LocalDateTable_') or
                        original_from.startswith('DateTableTemplate_') or
                        original_to.startswith('LocalDateTable_') or
                        original_to.startswith('DateTableTemplate_')):
                        continue

                    # Replace with friendly names if available, otherwise keep original
                    if original_from in table_name_mapping:
                        rel['fromTable'] = table_name_mapping[original_from]
                    if original_to in table_name_mapping:
                        rel['toTable'] = table_name_mapping[original_to]

                    # Create a unique key for this relationship pair (sorted to catch bidirectional duplicates)
                    from_table = rel.get('fromTable', '')
                    to_table = rel.get('toTable', '')
                    from_col = rel.get('fromColumn', '')
                    to_col = rel.get('toColumn', '')

                    # Create normalized pair (alphabetically sorted to catch A→B and B→A as same)
                    pair_key = tuple(sorted([
                        f"{from_table}.{from_col}",
                        f"{to_table}.{to_col}"
                    ]))

                    # Skip if we've already seen this relationship pair
                    if pair_key in seen_pairs:
                        continue

                    seen_pairs.add(pair_key)
                    filtered_relationships.append(rel)

                relationships = filtered_relationships
                print(f"   🔄 Filtered to {len(relationships)} unique relationships")

            # Mark columns used in relationships as "Used"
            if relationships:
                for rel in relationships:
                    from_table = rel.get('fromTable', '')
                    from_col = rel.get('fromColumn', '')
                    to_table = rel.get('toTable', '')
                    to_col = rel.get('toColumn', '')
                    _mark_used(from_table, from_col, 'Relationship')
                    _mark_used(to_table, to_col, 'Relationship')

            # Additional heuristic: Only mark calculated columns not already detected
            heuristic_count = 0

            for table_name, table_columns in all_columns.items():
                for col in table_columns:
                    col_type = col.get('columnType', '')

                    # Calculated columns are always actively used (they were created for a reason)
                    if col_type == 'Calculated':
                        if col.get('name') not in column_usage.get(table_name, {}):
                            _mark_used(table_name, col.get('name'), 'Calculated Column')
                            heuristic_count += 1

            print(f"   ✅ Heuristics added {heuristic_count} calculated columns")

            # ===== FINAL PASS: Apply column_usage to all queries' column info =====
            # Build a case-insensitive lookup for column_usage (merging sets)
            column_usage_lower = {}
            for tbl, cols in column_usage.items():
                tbl_lower = tbl.lower()
                if tbl_lower not in column_usage_lower:
                    column_usage_lower[tbl_lower] = {}
                for col_name_key, sources in cols.items():
                    col_key_lower = col_name_key.lower()
                    if col_key_lower not in column_usage_lower[tbl_lower]:
                        column_usage_lower[tbl_lower][col_key_lower] = set()
                    column_usage_lower[tbl_lower][col_key_lower].update(sources)

            used_count = 0
            total_count = 0
            for query in queries:
                query_name = query.get('queryName', '')
                for twc in query.get('tablesWithColumns', []):
                    twc_table = twc.get('tableName', '')
                    for col in twc.get('columns', []):
                        col_name = col.get('name', '')
                        col_name_lower = col_name.lower()
                        total_count += 1

                        # Collect all usage sources across matching strategies
                        all_sources = set()

                        # 1. Exact match by query_name
                        if query_name in column_usage:
                            sources = column_usage[query_name].get(col_name)
                            if sources:
                                all_sources.update(sources)

                        # 2. Exact match by twc tableName
                        if twc_table and twc_table in column_usage:
                            sources = column_usage[twc_table].get(col_name)
                            if sources:
                                all_sources.update(sources)

                        # 3. Case-insensitive match by query_name
                        tbl_cols = column_usage_lower.get(query_name.lower(), {})
                        sources = tbl_cols.get(col_name_lower)
                        if sources:
                            all_sources.update(sources)

                        # 4. Case-insensitive match by twc tableName
                        if twc_table:
                            tbl_cols = column_usage_lower.get(twc_table.lower(), {})
                            sources = tbl_cols.get(col_name_lower)
                            if sources:
                                all_sources.update(sources)

                        # 5. Direct check against Scanner model's isReferenced flag
                        if not all_sources:
                            model_cols = all_columns.get(query_name, [])
                            for mc in model_cols:
                                if mc.get('name', '').lower() == col_name_lower and mc.get('isReferenced') is True:
                                    all_sources.add('Scanner API')
                                    break

                        # 6. Partial/suffix table name matching (handles schema.table vs table)
                        if not all_sources:
                            query_name_lower = query_name.lower()
                            twc_table_lower = twc_table.lower() if twc_table else ''
                            for usage_tbl, usage_cols in column_usage_lower.items():
                                if col_name_lower in usage_cols:
                                    matched = False
                                    if (query_name_lower.endswith('.' + usage_tbl) or
                                        usage_tbl.endswith('.' + query_name_lower) or
                                        (twc_table_lower and (twc_table_lower.endswith('.' + usage_tbl) or
                                         usage_tbl.endswith('.' + twc_table_lower)))):
                                        matched = True
                                    if not matched:
                                        usage_suffix = usage_tbl.rsplit('.', 1)[-1]
                                        query_suffix = query_name_lower.rsplit('.', 1)[-1]
                                        twc_suffix = twc_table_lower.rsplit('.', 1)[-1] if twc_table_lower else ''
                                        if usage_suffix and (usage_suffix == query_suffix or usage_suffix == twc_suffix):
                                            matched = True
                                    if matched:
                                        all_sources.update(usage_cols[col_name_lower])
                                        break

                        is_used = len(all_sources) > 0
                        col['isReferenced'] = is_used
                        col['usedIn'] = ', '.join(sorted(all_sources)) if all_sources else ''
                        if is_used:
                            used_count += 1

            total_elapsed = time_module.time() - lineage_start_time
            # Count columns with usedIn populated
            used_in_count = sum(1 for q in queries for twc in q.get('tablesWithColumns', []) for c in twc.get('columns', []) if c.get('usedIn'))
            print(f"   ✅ Final: {used_count}/{total_count} columns marked as Used, {used_in_count} with usedIn labels — Total lineage time: {total_elapsed:.1f}s")

            # NEW: Detect and categorize datasets (Primary, Composite, DirectQuery, Live)
            datasets_info = []
            directquery_as_models = []

            try:
                print(f"\n   🔍 Detecting dataset types and composite models...")

                # Ensure we have a scanner instance
                if 'scanner' not in locals():
                    scanner = PowerBIScanner()

                # Get full dataset metadata from Scanner API
                scan_data = scanner.run_scan(workspace_id=workspace_id)

                if scan_data and "workspaces" in scan_data:
                    for ws in scan_data["workspaces"]:
                        # Find all datasets related to this report
                        for dataset in ws.get("datasets", []):
                            if dataset.get("id") == dataset_id:
                                # This is the primary dataset
                                dataset_name = dataset.get('name', 'Unknown Dataset')
                                upstream_datasets = dataset.get('upstreamDatasets', [])
                                datasources = dataset.get('datasources', [])
                                tables_count = len(dataset.get('tables', []))

                                print(f"   📊 Primary Dataset: {dataset_name}")
                                print(f"      - Tables: {tables_count}")
                                print(f"      - Datasources: {len(datasources)}")
                                print(f"      - Upstream Datasets: {len(upstream_datasets)}")

                                # Check for DirectQuery to Analysis Services
                                has_as_live = False
                                for ds in datasources:
                                    ds_type = ds.get('datasourceType', '')
                                    connection_details = ds.get('connectionDetails', {})

                                    # Check for Analysis Services connection
                                    if ds_type in ['AnalysisServices', 'AnalysisServicesDatabase']:
                                        server = connection_details.get('server', 'Unknown Server')
                                        database = connection_details.get('database', 'Unknown Database')

                                        directquery_as_models.append({
                                            'type': 'DirectQuery to Analysis Services',
                                            'server': server,
                                            'database': database,
                                            'connection_mode': ds.get('connectionMode', 'DirectQuery')
                                        })
                                        has_as_live = True
                                        print(f"      🔗 DirectQuery AS: {server}/{database}")

                                # Add primary dataset to list
                                dataset_type = 'Composite Model' if upstream_datasets else 'Import'
                                if has_as_live:
                                    dataset_type = 'DirectQuery/Live'

                                datasets_info.append({
                                    'id': dataset_id,
                                    'name': dataset_name,
                                    'isPrimary': True,
                                    'type': dataset_type,
                                    'tables': tables_count,
                                    'configuredBy': dataset.get('configuredBy', 'Unknown')
                                })

                                # Add upstream datasets (composite model)
                                if upstream_datasets:
                                    print(f"   🔗 Composite Model detected with {len(upstream_datasets)} upstream dataset(s)")
                                    for upstream in upstream_datasets:
                                        upstream_id = upstream.get('targetDatasetId')
                                        upstream_workspace_id = upstream.get('targetWorkspaceId', workspace_id)

                                        # Try to find upstream dataset details
                                        upstream_name = upstream_id  # Fallback to ID
                                        upstream_tables = 0

                                        for upstream_ws in scan_data.get("workspaces", []):
                                            if upstream_ws.get('id') == upstream_workspace_id:
                                                for upstream_ds in upstream_ws.get('datasets', []):
                                                    if upstream_ds.get('id') == upstream_id:
                                                        upstream_name = upstream_ds.get('name', upstream_id)
                                                        upstream_tables = len(upstream_ds.get('tables', []))
                                                        break

                                        datasets_info.append({
                                            'id': upstream_id,
                                            'name': upstream_name,
                                            'isPrimary': False,
                                            'type': 'Composite (Upstream)',
                                            'tables': upstream_tables,
                                            'configuredBy': 'Unknown'
                                        })
                                        print(f"      → Upstream: {upstream_name} ({upstream_tables} tables)")

                                break

                print(f"   ✅ Dataset categorization complete: {len(datasets_info)} dataset(s), {len(directquery_as_models)} DirectQuery AS model(s)")

            except Exception as dataset_error:
                print(f"   ⚠️ Error detecting datasets: {str(dataset_error)}")
                import traceback
                traceback.print_exc()
                # Continue without dataset info - don't fail the entire request

            return jsonify({
                'success': True,
                'queries': queries,
                'relationships': relationships,
                'tableSchemas': table_schemas,
                'datasets': datasets_info,
                'directQueryASModels': directquery_as_models
            })

        except ImportError:
            print("   ⚠️ Scanner connector not available")
            return jsonify({
                'success': False,
                'error': 'Scanner API not available. Admin permissions required.'
            }), 500
        except Exception as scan_error:
            print(f"   ⚠️ Scanner API error: {str(scan_error)}")
            import traceback
            traceback.print_exc()
            return jsonify({
                'success': False,
                'error': f'Scanner API error: {str(scan_error)}'
            }), 500

    except Exception as e:
        print(f"❌ Error fetching lineage: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/report/visual-lineage/<report_id>')
@login_required
def get_visual_lineage(report_id):
    """
    Get visual-level lineage for a report showing Pages -> Visuals -> Data Points -> Tables -> Data Sources

    Returns a hierarchical structure showing:
    - Report pages
    - Visuals on each page
    - Fields (columns/measures) used in each visual
    - Source tables for those fields
    - Upstream data sources
    """
    try:
        from scanner_connector import PowerBIScanner
        import re

        workspace_id = request.args.get('workspace_id')

        if not workspace_id:
            return jsonify({
                'success': False,
                'error': 'workspace_id is required'
            }), 400

        print(f"\n📊 ===============================================")
        print(f"📊 VISUAL LINEAGE REQUEST")
        print(f"📊 ===============================================")
        print(f"   Workspace ID: {workspace_id}")
        print(f"   Report ID: {report_id}")

        # Initialize Scanner API
        scanner = PowerBIScanner()
        scanner.access_token = scanner.get_access_token()

        # Run scan to get visual metadata
        print("   📊 Running Scanner API scan to get visual metadata...")
        scan_data = scanner.run_scan(workspace_id=workspace_id)

        if not scan_data or "workspaces" not in scan_data:
            return jsonify({
                'success': False,
                'error': 'Failed to scan workspace'
            }), 500

        # Find the report in scan results
        report_data = None
        dataset_id = None
        dataset_metadata = {}

        for ws in scan_data["workspaces"]:
            # Build dataset metadata map
            for dataset in ws.get("datasets", []):
                ds_id = dataset.get("id")
                if ds_id:
                    dataset_metadata[ds_id] = {
                        'id': ds_id,
                        'name': dataset.get('name', 'Unknown Dataset'),
                        'tables': dataset.get('tables', []),
                        'datasources': dataset.get('datasources', [])
                    }

            # Find the report
            for report in ws.get("reports", []):
                if report.get("id") == report_id:
                    report_data = report
                    dataset_id = report.get('datasetId')
                    break

            if report_data:
                break

        if not report_data:
            return jsonify({
                'success': False,
                'error': 'Report not found in workspace'
            }), 404

        # Get dataset metadata
        dataset_info = dataset_metadata.get(dataset_id, {})
        tables = dataset_info.get('tables', [])
        datasources = dataset_info.get('datasources', [])

        # Build table-to-datasource mapping AND extract SQL source tables
        table_to_datasource = {}
        table_to_sql_sources = {}  # Map Power Query tables to their SQL source tables
        column_to_sql_source = {}  # NEW: Map individual columns to their exact SQL source table

        for table in tables:
            table_name = table.get('name')
            if table_name:
                # Extract datasource from table source expression
                source = table.get('source', [])
                if source and len(source) > 0:
                    expr = source[0].get('expression', '')
                    # Try to extract datasource info from M query
                    datasource_info = extract_datasource_from_expression(expr, datasources)
                    table_to_datasource[table_name] = datasource_info

                    # NEW: Parse SQL query to extract column-to-table mapping
                    import re

                    # Extract SQL query from M expression
                    # Pattern: Query="SELECT ..."
                    query_match = re.search(r'Query\s*=\s*"([^"]+)"', expr, re.IGNORECASE | re.DOTALL)
                    if query_match:
                        sql_query = query_match.group(1)

                        # Normalize SQL for easier parsing
                        sql_query = sql_query.replace('#(lf)', '\n').replace('#(tab)', '\t')

                        print(f"\n      🔍 Parsing SQL for {table_name}:")
                        print(f"         SQL Preview: {sql_query[:200]}...")

                        # Parse SELECT clause to extract column mappings
                        import re

                        # Normalize SQL for easier parsing
                        sql_normalized = sql_query.replace('#(lf)', ' ').replace('#(tab)', ' ')

                        # Extract table aliases from FROM/JOIN clauses FIRST
                        # Pattern: FROM schema.table alias or JOIN schema.table alias
                        # Also handle: FROM schema.table AS alias
                        alias_patterns = [
                            r'(?:FROM|JOIN)\s+([a-zA-Z_][a-zA-Z0-9_]*\.[a-zA-Z_][a-zA-Z0-9_]*)\s+(?:AS\s+)?([a-zA-Z_][a-zA-Z0-9_]*)',
                            r'(?:FROM|JOIN)\s+([a-zA-Z_][a-zA-Z0-9_]*\.[a-zA-Z_][a-zA-Z0-9_]*)\s*\n'  # No alias case
                        ]

                        # Build alias-to-table mapping
                        alias_to_table = {}
                        sql_tables_found = set()

                        for pattern in alias_patterns:
                            matches = re.findall(pattern, sql_normalized, re.IGNORECASE)
                            for match in matches:
                                if len(match) == 2:
                                    full_table, alias = match
                                    if alias and alias.strip() and alias.upper() not in ['ON', 'WHERE', 'SELECT', 'AND', 'OR']:
                                        alias_to_table[alias.lower()] = full_table
                                        sql_tables_found.add(full_table)
                                        print(f"         🏷️  Alias: {alias} → {full_table}")
                                elif len(match) == 1:
                                    sql_tables_found.add(match[0])

                        table_to_sql_sources[table_name] = list(sql_tables_found)

                        # Now parse column mappings with multiple patterns
                        # Pattern 1: alias.column AS [name]
                        pattern1 = r'([a-zA-Z_][a-zA-Z0-9_]*)\.([a-zA-Z_][a-zA-Z0-9_]*)\s+AS\s+\[([^\]]+)\]'
                        matches1 = re.findall(pattern1, sql_normalized, re.IGNORECASE)

                        for alias_or_table, sql_column, pbi_column in matches1:
                            # Resolve alias to actual table
                            if alias_or_table.lower() in alias_to_table:
                                source_table = alias_to_table[alias_or_table.lower()]
                                key = (table_name, pbi_column)
                                column_to_sql_source[key] = source_table
                                print(f"         ✓ {pbi_column} → {source_table}.{sql_column}")
                            else:
                                # Might be schema.table.column or just table.column - need to check
                                # Try to find in sql_tables_found
                                matching_tables = [t for t in sql_tables_found if t.endswith(f'.{alias_or_table}') or t == alias_or_table]
                                if matching_tables:
                                    key = (table_name, pbi_column)
                                    column_to_sql_source[key] = matching_tables[0]
                                    print(f"         ✓ {pbi_column} → {matching_tables[0]}.{sql_column}")

                        # Pattern 2: schema.table.column AS [name] (less common but possible)
                        pattern2 = r'([a-zA-Z_][a-zA-Z0-9_]*)\.([a-zA-Z_][a-zA-Z0-9_]*)\.([a-zA-Z_][a-zA-Z0-9_]*)\s+AS\s+\[([^\]]+)\]'
                        matches2 = re.findall(pattern2, sql_normalized, re.IGNORECASE)

                        for schema, table, sql_column, pbi_column in matches2:
                            source_table = f"{schema}.{table}"
                            key = (table_name, pbi_column)
                            column_to_sql_source[key] = source_table
                            print(f"         ✓ {pbi_column} → {source_table}.{sql_column}")

                        # Pattern 3: column AS [name] (no table prefix - try to infer)
                        pattern3 = r'(?<![.\w])([a-zA-Z_][a-zA-Z0-9_]*)\s+AS\s+\[([^\]]+)\]'
                        matches3 = re.findall(pattern3, sql_normalized, re.IGNORECASE)

                        for sql_column, pbi_column in matches3:
                            # Only process if not already mapped
                            key = (table_name, pbi_column)
                            if key not in column_to_sql_source:
                                if len(sql_tables_found) == 1:
                                    # Only one table, safe to assume
                                    column_to_sql_source[key] = list(sql_tables_found)[0]
                                    print(f"         ✓ {pbi_column} → {list(sql_tables_found)[0]}.{sql_column} (single table)")
                                else:
                                    # Multiple tables with unprefixed column - cannot determine exact source
                                    # Don't add to column_to_sql_source - let it fall back to showing all tables
                                    print(f"         ⚠️  {pbi_column} → {sql_column} (unprefixed - cannot determine specific table)")

                        print(f"         📊 Mapped {len([k for k in column_to_sql_source if k[0] == table_name])} columns to specific tables")

        # Build column/measure to table mapping
        field_to_table = {}
        for table in tables:
            table_name = table.get('name')

            # Map columns
            for col in table.get('columns', []):
                col_name = col.get('name')
                if col_name:
                    field_to_table[col_name] = {
                        'table': table_name,
                        'type': 'column',
                        'dataType': col.get('dataType', 'Unknown')
                    }

            # Map measures
            for measure in table.get('measures', []):
                measure_name = measure.get('name')
                if measure_name:
                    field_to_table[measure_name] = {
                        'table': table_name,
                        'type': 'measure',
                        'expression': measure.get('expression', '')
                    }

        # Process visual data
        pages_data = []
        pages = report_data.get('pages', [])

        print(f"   ✅ Found {len(pages)} pages in report from Scanner API")

        # FALLBACK: If Scanner API doesn't return visual data, use Playwright extractor
        if not pages or len(pages) == 0:
            print(f"   ⚠️  Scanner API did not return visual metadata for this report")
            print(f"   🔄 FALLBACK: Attempting to extract visuals using Playwright...")

            try:
                from visual_metadata_extractor import VisualMetadataExtractor
                import asyncio
                import pickle
                import os
                from datetime import datetime

                # Check cache first
                cache_dir = '.visual_cache'
                os.makedirs(cache_dir, exist_ok=True)
                cache_file = os.path.join(cache_dir, f"{report_id}.pkl")

                visual_result = None

                # Try cache (valid for 24 hours)
                if os.path.exists(cache_file):
                    try:
                        with open(cache_file, 'rb') as f:
                            cached_data = pickle.load(f)
                            cached_time = datetime.fromisoformat(cached_data.get('cached_at', '2000-01-01'))
                            age_hours = (datetime.now() - cached_time).total_seconds() / 3600

                            if age_hours < 24:
                                print(f"      ✅ Using cached visual data (age: {age_hours:.1f}h)")
                                visual_result = cached_data
                    except Exception as e:
                        print(f"      ⚠️  Cache read failed: {e}")

                # If no cache, extract visuals using Playwright
                if not visual_result:
                    print(f"      🌐 Launching browser to extract visuals...")

                    # Get user's SSO access token
                    user_token = session.get('access_token')

                    print(f"      🔐 Using user's SSO token for extraction...")

                    # Initialize extractor with user's token
                    extractor = VisualMetadataExtractor(user_token=user_token)

                    # Run async extraction
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)
                    visual_result = loop.run_until_complete(
                        extractor.extract_visuals(workspace_id, report_id, timeout=90)
                    )
                    loop.close()

                    # Cache the result
                    if visual_result.get('success'):
                        visual_result['cached_at'] = datetime.now().isoformat()
                        with open(cache_file, 'wb') as f:
                            pickle.dump(visual_result, f)
                        print(f"      ✅ Visual data cached")

                # Process Playwright results
                if visual_result and visual_result.get('success'):
                    pages = visual_result.get('pages', [])
                    print(f"      ✅ Playwright extracted {len(pages)} page(s)")
                else:
                    print(f"      ❌ Playwright extraction failed: {visual_result.get('error') if visual_result else 'Unknown error'}")

            except Exception as e:
                print(f"      ❌ Error during Playwright extraction: {str(e)}")
                import traceback
                traceback.print_exc()

        for page in pages:
            page_name = page.get('displayName', 'Unnamed Page')
            page_ordinal = page.get('ordinal', 0)
            visuals = page.get('visuals', [])

            visuals_data = []

            for visual in visuals:
                visual_name = visual.get('name', 'Unnamed Visual')
                visual_title = visual.get('title', '').strip()  # Remove whitespace
                visual_type = visual.get('visualType') or visual.get('type', 'unknown')

                # If no title is set, use a readable format instead of showing object name
                if not visual_title:
                    # Use visual type as fallback (more readable than "VisualContainer1")
                    visual_title = visual_type.replace('Visual', '').replace('Chart', ' Chart').strip() or 'Unnamed Visual'

                # Extract fields used in the visual
                fields_used = []

                # METHOD 1: Check if visual has 'fields' array (Playwright format)
                playwright_fields = visual.get('fields', [])
                if playwright_fields:
                    print(f"      📊 Processing visual '{visual_title or visual_name}' with {len(playwright_fields)} fields (Playwright format)")

                    for field_obj in playwright_fields:
                        # Handle both formats: string or dict
                        if isinstance(field_obj, str):
                            field_name = field_obj
                            field_display_name = field_obj
                            field_type_hint = 'unknown'
                            playwright_table = None
                        elif isinstance(field_obj, dict):
                            field_name = field_obj.get('name', '')
                            field_display_name = field_obj.get('displayName', field_name)
                            field_type_hint = field_obj.get('type', 'unknown')  # 'Column', 'Measure', etc.
                            playwright_table = field_obj.get('table', '')  # CRITICAL: Get table from Playwright!
                        else:
                            continue

                        # PRIORITY 1: Use table name from Playwright (most accurate!)
                        if playwright_table:
                            table_name = playwright_table
                            datasource = table_to_datasource.get(table_name, {})

                            # Try to get more metadata from dataset if available
                            field_info = field_to_table.get(field_name, {})

                            # Determine field type
                            if field_type_hint in ['Measure', 'Aggregation']:
                                field_type = 'measure'
                            elif field_type_hint == 'Column':
                                field_type = 'column'
                            elif field_info:
                                field_type = field_info.get('type', 'column')
                            else:
                                field_type = 'column'

                            # NEW: Precise column-to-source-table mapping
                            # Check if we have a specific SQL source table for this column
                            column_key = (table_name, field_name)
                            specific_sql_source = column_to_sql_source.get(column_key)

                            if specific_sql_source:
                                # We know the EXACT source table for this column!
                                display_table_name = specific_sql_source
                                print(f"         ✅ Precise mapping: {field_name} → {specific_sql_source}")
                            else:
                                # Fall back to showing all SQL sources for this Power Query table
                                sql_sources = table_to_sql_sources.get(table_name, [])
                                if sql_sources:
                                    if len(sql_sources) == 1:
                                        # Only one source table, use it
                                        display_table_name = sql_sources[0]
                                    else:
                                        # Multiple sources, can't determine which one
                                        display_table_name = ', '.join(sql_sources)
                                else:
                                    # No SQL sources found, use Power Query table name
                                    display_table_name = table_name

                            fields_used.append({
                                'field_name': field_name,
                                'field_type': field_type,
                                'table_name': display_table_name,  # Precise SQL source or Power Query table
                                'model_table_name': table_name,  # Power BI model table (e.g., Query1)
                                'data_type': field_info.get('dataType', 'Unknown') if field_info else 'Unknown',
                                'datasource_type': datasource.get('type', 'Unknown'),
                                'datasource_server': datasource.get('server', 'N/A'),
                                'datasource_database': datasource.get('database', 'N/A'),
                                'expression': field_info.get('expression', '') if field_info and field_type == 'measure' else ''
                            })
                        else:
                            # FALLBACK: Try to find this field in our dataset metadata
                            field_info = field_to_table.get(field_name, {})

                            if field_info:
                                table_name = field_info.get('table')
                                datasource = table_to_datasource.get(table_name, {})
                                field_type = field_info.get('type', 'column')

                                fields_used.append({
                                    'field_name': field_name,
                                    'field_type': field_type,
                                    'table_name': table_name,
                                    'data_type': field_info.get('dataType'),
                                    'datasource_type': datasource.get('type', 'Unknown'),
                                    'datasource_server': datasource.get('server', 'N/A'),
                                    'datasource_database': datasource.get('database', 'N/A'),
                                    'expression': field_info.get('expression', '') if field_type == 'measure' else ''
                                })
                            else:
                                # Field not found in metadata, add with limited info
                                fields_used.append({
                                    'field_name': field_name,
                                    'field_type': 'unknown',
                                    'table_name': 'N/A',
                                    'data_type': 'Unknown',
                                    'datasource_type': 'Unknown',
                                    'datasource_server': 'N/A',
                                    'datasource_database': 'N/A'
                                })

                # METHOD 2: Parse visual config JSON (Scanner API format)
                else:
                    visual_config = visual.get('config', '')

                    if visual_config:
                        # Extract table and field references from visual config JSON
                        # Pattern: "TableName.FieldName" or references in projections
                        field_pattern = r'"(?:Entity|Table)"\s*:\s*"([^"]+)"[^}]*?"(?:Property|Column|Name)"\s*:\s*"([^"]+)"'
                        matches = re.findall(field_pattern, visual_config)

                        for table_ref, field_ref in matches:
                            field_key = field_ref
                            field_info = field_to_table.get(field_key, {})

                            if field_info:
                                table_name = field_info.get('table')
                                datasource = table_to_datasource.get(table_name, {})
                                field_type = field_info.get('type', 'column')

                                fields_used.append({
                                    'field_name': field_ref,
                                    'field_type': field_type,
                                    'table_name': table_name,
                                    'data_type': field_info.get('dataType'),
                                    'datasource_type': datasource.get('type', 'Unknown'),
                                    'datasource_server': datasource.get('server', 'N/A'),
                                    'datasource_database': datasource.get('database', 'N/A'),
                                    'expression': field_info.get('expression', '') if field_type == 'measure' else ''
                                })

                # Remove duplicates
                unique_fields = []
                seen = set()
                for field in fields_used:
                    key = (field['field_name'], field['table_name'])
                    if key not in seen:
                        seen.add(key)
                        unique_fields.append(field)

                visuals_data.append({
                    'name': visual_name,
                    'title': visual_title,  # Now uses actual title or readable type name
                    'type': visual_type,
                    'fields_count': len(unique_fields),
                    'fields': unique_fields
                })

            pages_data.append({
                'name': page_name,
                'ordinal': page_ordinal,
                'visuals_count': len(visuals_data),
                'visuals': visuals_data
            })

        # Sort pages by ordinal
        pages_data.sort(key=lambda x: x['ordinal'])

        print(f"   ✅ Processed {len(pages_data)} pages with visual lineage data")
        print(f"📊 ===============================================\n")

        return jsonify({
            'success': True,
            'report_id': report_id,
            'report_name': report_data.get('name', 'Unknown Report'),
            'dataset_id': dataset_id,
            'dataset_name': dataset_info.get('name', 'Unknown Dataset'),
            'pages_count': len(pages_data),
            'pages': pages_data
        })

    except Exception as e:
        print(f"❌ Error getting visual lineage: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


def extract_datasource_from_expression(expression, datasources):
    """Extract datasource information from M query expression"""
    import re

    datasource_info = {
        'type': 'Unknown',
        'server': 'N/A',
        'database': 'N/A',
        'url': 'N/A'
    }

    if not expression:
        return datasource_info

    # Try to match SQL Server
    sql_match = re.search(r'Sql\.Database\("([^"]+)",\s*"([^"]+)"', expression, re.IGNORECASE)
    if sql_match:
        datasource_info['type'] = 'SQL Server'
        datasource_info['server'] = sql_match.group(1)
        datasource_info['database'] = sql_match.group(2)
        return datasource_info

    # Try to match SharePoint/Web
    web_match = re.search(r'Web\.Contents\("([^"]+)"', expression, re.IGNORECASE)
    if web_match:
        datasource_info['type'] = 'Web/SharePoint'
        datasource_info['url'] = web_match.group(1)
        return datasource_info

    # Try to match Excel file
    excel_match = re.search(r'Excel\.Workbook\(', expression, re.IGNORECASE)
    if excel_match:
        datasource_info['type'] = 'Excel'
        return datasource_info

    # Fallback: check datasources array
    if datasources and len(datasources) > 0:
        ds = datasources[0]
        connection = ds.get('connectionDetails', {})
        datasource_info['type'] = ds.get('datasourceType', 'Unknown')
        datasource_info['server'] = connection.get('server', 'N/A')
        datasource_info['database'] = connection.get('database', 'N/A')
        datasource_info['url'] = connection.get('url', 'N/A')

    return datasource_info


@app.route('/api/lineage/debug')
@login_required
def get_lineage_debug():
    """Debug endpoint to view raw dataset model and expressions"""
    try:
        import requests
        workspace_id = request.args.get('workspace_id')
        report_id = request.args.get('report_id')

        if not workspace_id or not report_id:
            return jsonify({
                'success': False,
                'error': 'workspace_id and report_id are required'
            }), 400

        # Get user token from session
        user_token = session.get('access_token')

        if not user_token:
            return jsonify({
                'success': False,
                'error': 'Not authenticated'
            }), 401

        headers = {
            'Authorization': f'Bearer {user_token}',
            'Content-Type': 'application/json'
        }

        # Get report details to find dataset ID
        report_url = f'https://api.powerbi.com/v1.0/myorg/groups/{workspace_id}/reports/{report_id}'
        report_response = requests.get(report_url, headers=headers)

        if report_response.status_code != 200:
            return jsonify({
                'success': False,
                'error': f'Failed to fetch report: {report_response.status_code}'
            }), report_response.status_code

        report_data = report_response.json()
        dataset_id = report_data.get('datasetId')

        if not dataset_id:
            return jsonify({
                'success': True,
                'message': 'Report has no dataset',
                'model': {}
            })

        # Use Scanner API to get dataset model
        try:
            from scanner_connector import PowerBIScanner
            scanner = PowerBIScanner()

            print(f"   🐛 DEBUG: Fetching dataset model for {dataset_id}...")
            model = scanner.get_dataset_model(dataset_id, workspace_id=workspace_id)

            # Return the raw model for inspection
            return jsonify({
                'success': True,
                'dataset_id': dataset_id,
                'model': {
                    'tables': model.get('tables', []),
                    'expressions_count': len(model.get('expressions', [])),
                    'expressions': model.get('expressions', []),
                    'columns': {k: len(v) for k, v in model.get('columns', {}).items()},
                    'measures_count': len(model.get('measures', [])),
                    'relationships_count': len(model.get('relationships', []))
                }
            })

        except Exception as scan_error:
            return jsonify({
                'success': False,
                'error': f'Scanner API error: {str(scan_error)}'
            }), 500

    except Exception as e:
        print(f"❌ Debug endpoint error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


def extract_sql_table_names(sql_query):
    """
    Extract table names from SQL query text.

    Detects table references in:
    - FROM clauses
    - JOIN clauses (INNER, LEFT, RIGHT, FULL, CROSS)
    - Table-valued functions

    Args:
        sql_query: SQL query string

    Returns:
        Set of table names found in the query
    """
    import re

    if not sql_query:
        return set()

    tables = set()

    # Remove SQL comments
    # -- single line comments
    sql_query = re.sub(r'--[^\n]*', '', sql_query)
    # /* multi-line comments */
    sql_query = re.sub(r'/\*.*?\*/', '', sql_query, flags=re.DOTALL)

    # Pattern for FROM and JOIN clauses
    # Matches: FROM [schema].[table], FROM table, FROM [table]
    # Also: JOIN [schema].[table], INNER JOIN table AS alias

    # Pattern 1: [schema].[table] or [schema].table or schema.[table]
    bracketed_pattern = r'(?:FROM|JOIN|INTO|UPDATE|TABLE)\s+(?:\[([^\]]+)\]\.)?\[?([^\]\s,;)(]+)\]?'
    matches = re.findall(bracketed_pattern, sql_query, re.IGNORECASE)

    for schema, table in matches:
        # Clean up table name
        table = table.strip('[]').strip()
        if table.upper() not in ('SELECT', 'WHERE', 'ON', 'AS', 'AND', 'OR', 'INNER', 'LEFT', 'RIGHT', 'OUTER', 'FULL', 'CROSS'):
            if schema:
                schema = schema.strip('[]').strip()
                tables.add(f"{schema}.{table}")
            else:
                tables.add(table)

    # Pattern 2: Unbracketed schema.table
    # Matches: FROM dbo.TableName, JOIN Sales.FactSales
    dotted_pattern = r'(?:FROM|JOIN|INTO|UPDATE|TABLE)\s+([a-zA-Z_][a-zA-Z0-9_]*\.[a-zA-Z_][a-zA-Z0-9_]*)'
    dotted_matches = re.findall(dotted_pattern, sql_query, re.IGNORECASE)

    for match in dotted_matches:
        # Skip common SQL keywords that might match pattern
        if not any(kw in match.upper() for kw in ['INNER.', 'LEFT.', 'RIGHT.', 'OUTER.', 'FULL.']):
            tables.add(match)

    return tables


def extract_sql_column_names(sql_query):
    """
    Extract column names from SQL SELECT clauses.

    Parses SELECT column lists to identify which columns are being queried.
    These columns are definitively "used" since they're pulled from the database.

    Args:
        sql_query: SQL query string (already cleaned of M-code escapes)

    Returns:
        Set of column names found in SELECT clauses
    """
    import re

    if not sql_query:
        return set()

    columns = set()

    # Remove SQL comments
    sql_query = re.sub(r'--[^\n]*', '', sql_query)
    sql_query = re.sub(r'/\*.*?\*/', '', sql_query, flags=re.DOTALL)

    # Extract the SELECT ... FROM portion(s)
    # Handle multiple SELECT statements (subqueries, CTEs)
    select_blocks = re.findall(
        r'SELECT\s+(?:DISTINCT\s+|TOP\s+\d+\s+)?(.*?)(?:\bFROM\b)',
        sql_query, re.IGNORECASE | re.DOTALL
    )

    for block in select_blocks:
        # Skip SELECT * — doesn't tell us specific columns
        if block.strip() == '*':
            continue

        # Split by comma, handling nested parentheses
        depth = 0
        current = ''
        items = []
        for ch in block:
            if ch == '(':
                depth += 1
                current += ch
            elif ch == ')':
                depth -= 1
                current += ch
            elif ch == ',' and depth == 0:
                items.append(current.strip())
                current = ''
            else:
                current += ch
        if current.strip():
            items.append(current.strip())

        for item in items:
            if not item:
                continue

            # Check for AS alias: extract the source column name, not the alias
            # Pattern: [Table].[Column] AS [Alias] or Column AS Alias
            as_match = re.match(r'(.+?)\s+AS\s+\[?([^\]]+)\]?\s*$', item, re.IGNORECASE)
            if as_match:
                col_expr = as_match.group(1).strip()
            else:
                col_expr = item.strip()

            # Extract bracketed column names: [ColumnName]
            bracket_cols = re.findall(r'\[([^\]]+)\]', col_expr)
            for bc in bracket_cols:
                # Skip schema/table-like names (contain dots or are all uppercase table patterns)
                if bc and not re.match(r'^[A-Z_]+\.[A-Z_]+$', bc):
                    columns.add(bc)

            # Extract unbracketed column references: table.column or just column
            # Match alias.column pattern (e.g., BASE.ITNBR, prod.ManufacturingStatus)
            dot_cols = re.findall(r'(?:^|\s|,)([A-Za-z_]\w*\.)([A-Za-z_]\w*)', col_expr)
            for _, col in dot_cols:
                if col.upper() not in ('AS', 'AND', 'OR', 'NOT', 'NULL', 'CAST', 'CASE',
                                       'WHEN', 'THEN', 'ELSE', 'END', 'FROM', 'WHERE',
                                       'SELECT', 'DISTINCT', 'INT', 'DATE', 'VARCHAR', 'NVARCHAR'):
                    columns.add(col)

    return columns


def extract_m_expression_columns(expression, table_name):
    """
    Parse M/Power Query expression to identify columns referenced in transformations.

    Detects column references from:
    - Table.SelectColumns(source, {"col1", "col2"})
    - Table.RemoveColumns(source, {"col1", "col2"})
    - Table.RenameColumns(source, {{"old", "new"}, ...})
    - Table.TransformColumnTypes(source, {{"col1", type}})
    - Step-level references like [ColumnName]
    - #"Column Name" hash-quoted references in expressions
    - each [ColumnName] row-level references

    Args:
        expression: M expression/Power Query code
        table_name: Name of the table this expression belongs to

    Returns:
        Set of column names found in the expression
    """
    import re

    if not expression:
        return set()

    columns = set()

    # Pattern 1: Table.SelectColumns, Table.RemoveColumns, Table.ReorderColumns
    # These take a list of column names: {"Col1", "Col2", "Col3"}
    select_pattern = r'Table\.(?:SelectColumns|RemoveColumns|ReorderColumns)\s*\([^,]+,\s*\{([^}]+)\}'
    for match in re.finditer(select_pattern, expression, re.IGNORECASE):
        col_list = match.group(1)
        col_names = re.findall(r'"([^"]+)"', col_list)
        columns.update(col_names)

    # Pattern 2: Table.RenameColumns — {{"OldName", "NewName"}, ...}
    rename_pattern = r'Table\.RenameColumns\s*\([^,]+,\s*\{((?:\{[^}]+\}\s*,?\s*)+)\}'
    for match in re.finditer(rename_pattern, expression, re.IGNORECASE):
        pairs = match.group(1)
        pair_cols = re.findall(r'"([^"]+)"', pairs)
        columns.update(pair_cols)  # Both old and new names are relevant

    # Pattern 3: Table.TransformColumnTypes — {{"Col1", type text}, {"Col2", Int64.Type}}
    transform_pattern = r'Table\.TransformColumnTypes\s*\([^,]+,\s*\{((?:\{[^}]+\}\s*,?\s*)+)\}'
    for match in re.finditer(transform_pattern, expression, re.IGNORECASE):
        pairs = match.group(1)
        col_names = re.findall(r'\{\s*"([^"]+)"', pairs)
        columns.update(col_names)

    # Pattern 4: each [ColumnName] — row-level access in M
    each_col_pattern = r'(?:each\s+)?\[([A-Za-z_][A-Za-z0-9_ ]*)\]'
    for match in re.finditer(each_col_pattern, expression):
        col = match.group(1).strip()
        # Filter out M keywords and step names
        if col and len(col) < 80 and col not in ('Source', 'Changes', 'Type', 'Content'):
            columns.add(col)

    # Pattern 5: #"Renamed Columns" step references that contain column names in quotes
    # This catches things like = Table.AddColumn(source, "NewCol", each [ExistingCol])
    add_col_pattern = r'Table\.AddColumn\s*\([^,]+,\s*"([^"]+)"'
    for match in re.finditer(add_col_pattern, expression, re.IGNORECASE):
        columns.add(match.group(1))

    return columns


def parse_m_expression_for_tables(expression, all_tables):
    """
    Parse M expression to identify table references with enhanced pattern matching.

    This function identifies tables referenced within Power Query M expressions by
    detecting various M code patterns where tables are referenced, including:
    - Other Power BI tables/queries in the same dataset
    - SQL tables referenced in embedded SQL queries
    - Dataflow tables referenced in PowerBI.Dataflows
    - Excel tables, SharePoint lists, etc.

    Args:
        expression: The M expression/query code to parse
        all_tables: Set/list of all table names in the dataset model

    Returns:
        List of table names found in the expression (deduplicated)
    """
    import re

    if not expression:
        return []

    tables_found = set()  # Use set for automatic deduplication

    # Convert all_tables to a set for O(1) lookup performance
    all_tables_set = set(all_tables) if not isinstance(all_tables, set) else all_tables

    # ====================
    # PATTERN 1: Hash-quoted references #"TableName" or #"Query Name"
    # This is the MOST COMMON pattern in Power Query
    # Examples:
    #   let Source = #"Sales Data" in Source
    #   Table.Combine({#"Table1", #"Table2"})
    # ====================
    hash_quoted_pattern = r'#"([^"]+)"'
    hash_matches = re.findall(hash_quoted_pattern, expression)

    for match in hash_matches:
        # Only add if it matches a known table name in the model
        if match in all_tables_set:
            tables_found.add(match)

    # ====================
    # PATTERN 2: Direct table name references (unquoted identifiers)
    # Example: let Source = TableName in Source
    # This is less common but can occur with simple table names
    # ====================
    # We need to be careful here to avoid false positives
    # Only check for tables with simple names (no spaces, no special chars)
    for table in all_tables_set:
        # Skip tables that would be quoted (contain spaces or special chars)
        # These would only appear as #"TableName" format
        if ' ' in table or not table.replace('_', '').isalnum():
            continue

        # Create a word boundary pattern to match the table name
        # Use word boundaries to avoid partial matches
        escaped_table = re.escape(table)
        pattern = r'\b' + escaped_table + r'\b'

        # Search for the table name in the expression
        if re.search(pattern, expression):
            tables_found.add(table)

    # ====================
    # PATTERN 3: Table references in curly braces (list syntax)
    # Example: Table.Combine({TableName1, TableName2})
    # Example: {#"Table1", #"Table2"}
    # ====================
    # This pattern is already covered by patterns 1 and 2 above

    # ====================
    # PATTERN 4: Source{[Name="TableName"]} - filtered table references
    # Example: Source{[Name="DimCustomer"]}[Content]
    # ====================
    filtered_pattern = r'\{?\[\s*Name\s*=\s*"([^"]+)"\s*\]\}?'
    filtered_matches = re.findall(filtered_pattern, expression, re.IGNORECASE)

    for match in filtered_matches:
        if match in all_tables_set:
            tables_found.add(match)

    # ====================
    # PATTERN 5: Excel.CurrentWorkbook(){[Name="TableName"]}
    # This is specific to Excel sources with named ranges/tables
    # ====================
    excel_pattern = r'Excel\.CurrentWorkbook\s*\(\s*\)\s*\{?\[\s*Name\s*=\s*"([^"]+)"\s*\]\}?'
    excel_matches = re.findall(excel_pattern, expression, re.IGNORECASE)

    for match in excel_matches:
        if match in all_tables_set:
            tables_found.add(match)

    # ====================
    # PATTERN 6: Nested query references in 'let' statements
    # Example: let Source = #"Query1", Result = #"Query2" in Result
    # This is already covered by Pattern 1
    # ====================

    # ====================
    # PATTERN 7: SQL Table References - Extract from embedded SQL queries
    # Example: Sql.Database("server", "db", [Query="SELECT * FROM dbo.TableName"])
    # Example: FROM [schema].[TableName]
    # ====================
    print(f"         [DEBUG] Checking for SQL patterns...")
    print(f"         [DEBUG] 'Sql.Database' in expression: {'Sql.Database' in expression}")
    print(f"         [DEBUG] 'Query=' in expression: {'Query=' in expression}")

    if 'Sql.Database' in expression or 'Sql.Databases' in expression or 'Query=' in expression:
        print(f"         [DEBUG] SQL pattern detected in expression")
        # Extract the SQL query from the M expression
        # Look for Query="..." or Query='...' parameter
        # The pattern needs to handle #(lf), #(cr), etc. within the quoted string
        sql_query_pattern = r'Query\s*=\s*"([^"]*)"'
        sql_matches = re.findall(sql_query_pattern, expression, re.IGNORECASE)

        print(f"         [DEBUG] Found {len(sql_matches)} SQL query patterns")

        for sql_query in sql_matches:
            print(f"         [DEBUG] SQL Query preview: {sql_query[:100]}")
            # Replace escaped quotes
            sql_query = sql_query.replace('""', '"')
            # Also handle #(lf) line feeds in M code
            sql_query = sql_query.replace('#(lf)', '\n')
            sql_query = sql_query.replace('#(cr)', '\r')
            sql_query = sql_query.replace('#(tab)', '\t')

            print(f"         [DEBUG] After normalization: {sql_query[:150]}")

            # Extract table names from SQL
            sql_tables = extract_sql_table_names(sql_query)
            print(f"         [DEBUG] SQL tables found: {sql_tables}")
            tables_found.update(sql_tables)

    # Also check for direct SQL in Source{[Schema=..., Item=...]} pattern
    schema_item_pattern = r'\[\s*Schema\s*=\s*"([^"]+)"\s*,\s*Item\s*=\s*"([^"]+)"\s*\]'
    schema_items = re.findall(schema_item_pattern, expression, re.IGNORECASE)
    for schema, item in schema_items:
        # Format as schema.table
        table_ref = f"{schema}.{item}"
        tables_found.add(table_ref)

    # ====================
    # PATTERN 8: PowerBI.Dataflows - Extract dataflow table references
    # Example: PowerBI.Dataflows(null){[workspaceId="..."]}[Data]{[dataflowId="..."]}[Data]{[entity="TableName"]}
    # ====================
    if 'PowerBI.Dataflows' in expression:
        # Look for entity= references which indicate dataflow table names
        entity_pattern = r'entity\s*=\s*"([^"]+)"'
        entity_matches = re.findall(entity_pattern, expression, re.IGNORECASE)
        for entity in entity_matches:
            # Add clean entity name without prefix - Source Type column will show "Dataflow"
            tables_found.add(entity)

        # Also look for hash-quoted navigation after dataflows
        # Pattern: #"a47e4573-c455-40af-a9ad-e22c81a07926"[Data]{[entity="WarehouseMaster"]}
        # The hash-quoted GUID is workspace/dataflow ID, not a table

    # ====================
    # EXCLUDE the query's own table name from results
    # A query shouldn't be listed as "using" itself
    # We'll handle this exclusion in the calling function
    # ====================

    # Convert set back to sorted list for consistent output
    return sorted(list(tables_found))



def extract_server_name(expression, source_type, workspace_map=None):
    """
    Extract server/source name from M expression based on source type.

    Args:
        expression: M expression text
        source_type: Type of the data source (SQL Server, Dataflow, etc.)
        workspace_map: Dictionary mapping workspace IDs to workspace names (optional)

    Returns:
        Server name, file path, or source identifier depending on the source type.
    """
    import re

    if not expression:
        return 'N/A'

    if source_type == 'SQL Server':
        # Extract server name from Sql.Database("server.name", "database")
        sql_pattern = r'Sql\.Database\s*\(\s*"([^"]+)"'
        match = re.search(sql_pattern, expression, re.IGNORECASE)
        if match:
            return match.group(1)

    elif source_type == 'Dataflow':
        # Extract workspace ID and resolve to workspace name if available
        # Pattern: PowerBI.Dataflows(null){[workspaceId="..."]}
        workspace_pattern = r'workspaceId\s*=\s*"([^"]+)"'
        match = re.search(workspace_pattern, expression, re.IGNORECASE)
        if match:
            workspace_id = match.group(1)
            # Try to resolve workspace name from the map
            if workspace_map and workspace_id in workspace_map:
                return workspace_map[workspace_id]
            # Fallback to truncated ID if name not available
            return f"Dataflow (Workspace: {workspace_id[:8]}...)"
        return 'Power BI Dataflow'

    elif source_type == 'Excel':
        # Extract file path from Excel.Workbook(File.Contents("path"))
        excel_pattern = r'File\.Contents\s*\(\s*"([^"]+)"'
        match = re.search(excel_pattern, expression, re.IGNORECASE)
        if match:
            file_path = match.group(1)
            # Return just the filename if it's a full path
            if '\\' in file_path or '/' in file_path:
                return file_path.split('\\')[-1].split('/')[-1]
            return file_path
        return 'Local File'

    elif source_type == 'Expression':
        # Internal model references
        return 'Internal Model'

    elif source_type == 'ODBC':
        # Extract DSN or connection string
        odbc_pattern = r'Odbc\.DataSource\s*\(\s*"([^"]+)"'
        match = re.search(odbc_pattern, expression, re.IGNORECASE)
        if match:
            return match.group(1)
        return 'ODBC Source'

    elif source_type == 'Web':
        # Extract URL
        web_pattern = r'Web\.Contents\s*\(\s*"([^"]+)"'
        match = re.search(web_pattern, expression, re.IGNORECASE)
        if match:
            url = match.group(1)
            # Return domain name
            domain_match = re.search(r'https?://([^/]+)', url)
            if domain_match:
                return domain_match.group(1)
            return url
        return 'Web Source'

    elif source_type == 'SharePoint':
        # Extract SharePoint site
        sp_pattern = r'https?://([^/]+)'
        match = re.search(sp_pattern, expression, re.IGNORECASE)
        if match:
            return match.group(1)
        return 'SharePoint'

    elif source_type == 'OData':
        # Extract OData feed URL
        odata_pattern = r'OData\.Feed\s*\(\s*"([^"]+)"'
        match = re.search(odata_pattern, expression, re.IGNORECASE)
        if match:
            url = match.group(1)
            domain_match = re.search(r'https?://([^/]+)', url)
            if domain_match:
                return domain_match.group(1)
            return url
        return 'OData Feed'

    return 'N/A'


def extract_source_line(expression):
    """
    Extract just the source table reference from an M expression.
    For Expression type queries, returns only the table reference without any variable assignment.
    Example: "#\"FactSales\"" (not "Source = #\"FactSales\"")
    """
    import re

    if not expression:
        return expression

    print(f"      🔍 DEBUG extract_source_line - Input: {expression[:100]}")

    # Pattern to find variable = #"TableName" assignments
    # This matches lines like: Source = #"TableName", source = #"Table", BaseTable = #"DimCustomer", etc.
    source_pattern = r'^\s*\w+\s*=\s*(#"[^"]+").*$'

    lines = expression.split('\n')

    # Look for the first line after "let" that assigns a hash-quoted table reference
    in_let_block = False
    for line in lines:
        stripped = line.strip()

        if stripped.lower().startswith('let'):
            in_let_block = True
            continue

        if in_let_block and stripped:
            # Check if this line has a table reference assignment (variable = #"TableName")
            match = re.match(source_pattern, stripped, re.IGNORECASE)
            if match:
                # Return ONLY the table reference part (e.g., #"FactSales")
                # This removes any "Source = ", "source = ", or other variable prefix
                result = match.group(1)
                print(f"      ✅ DEBUG extract_source_line - Extracted from let block: {result}")
                return result

    # If no source line found in let block, search entire expression
    for line in lines:
        stripped = line.strip()
        if stripped and not stripped.lower().startswith('let') and not stripped.lower().startswith('in'):
            # Try to extract just the table reference if it contains a pattern like "Variable = #"Table""
            match = re.match(source_pattern, stripped, re.IGNORECASE)
            if match:
                result = match.group(1)
                print(f"      ✅ DEBUG extract_source_line - Extracted from line scan: {result}")
                return result
            # Otherwise return the whole line
            print(f"      ⚠️ DEBUG extract_source_line - Returning whole line: {stripped}")
            return stripped

    # Fallback: if expression contains the pattern anywhere, extract just the table reference
    final_match = re.search(r'\w+\s*=\s*(#"[^"]+")' , expression, re.IGNORECASE)
    if final_match:
        result = final_match.group(1)
        print(f"      ✅ DEBUG extract_source_line - Extracted from final search: {result}")
        return result

    # Absolute fallback: return the full expression
    print(f"      ⚠️ DEBUG extract_source_line - Returning full expression: {expression[:100]}")
    return expression


def analyze_m_expression(expression):
    """Analyze M expression to determine query type and source type"""
    import re

    if not expression:
        return ('Unknown', 'Unknown')

    expression_lower = expression.lower()

    # List of external source indicators
    external_sources = [
        'sql.database', 'sql server',
        'powerbi.dataflows',
        'odbc.datasource', 'odbc.query',
        'excel.workbook', 'excel.currentworkbook',
        'web.contents', 'web.page',
        'sharepoint',
        'folder.files', 'file.contents',
        'odata.feed',
        'json.document',
        'xml.tables',
        'oracle.database',
        'mysql.database',
        'postgresql.database',
        'azuresql.database'
    ]

    # Check if expression has any external sources
    has_external_source = any(src in expression_lower for src in external_sources)

    # Determine source type
    source_type = 'Unknown'
    if 'sql.database' in expression_lower or 'sql server' in expression_lower:
        source_type = 'SQL Server'
    elif 'powerbi.dataflows' in expression_lower:
        source_type = 'Dataflow'
    elif 'odbc.datasource' in expression_lower or 'odbc.query' in expression_lower:
        source_type = 'ODBC'
    elif 'excel.workbook' in expression_lower or 'excel.currentworkbook' in expression_lower:
        source_type = 'Excel'
    elif 'web.contents' in expression_lower or 'web.page' in expression_lower:
        source_type = 'Web'
    elif 'sharepoint' in expression_lower:
        source_type = 'SharePoint'
    elif 'folder.files' in expression_lower or 'file.contents' in expression_lower:
        source_type = 'File System'
    elif 'odata.feed' in expression_lower:
        source_type = 'OData'
    elif 'json.document' in expression_lower:
        source_type = 'JSON'
    elif 'xml.tables' in expression_lower:
        source_type = 'XML'
    elif '#datetime' in expression_lower or '#date' in expression_lower or 'list.dates' in expression_lower:
        source_type = 'Date Function'
    elif 'table.fromrows' in expression_lower or 'table.fromlist' in expression_lower:
        source_type = 'Manual Table'
    elif not has_external_source:
        # Check if expression contains internal table references (hash-quoted names)
        # Pattern: #"TableName" indicates reference to another table in the model
        hash_quote_pattern = r'#"[^"]+"'
        has_internal_refs = bool(re.search(hash_quote_pattern, expression))

        if has_internal_refs:
            source_type = 'Expression'

    # Determine query type
    query_type = 'M Query'
    if 'value.nativequery' in expression_lower:
        query_type = 'Native Query'
    elif source_type == 'Manual Table' or source_type == 'Date Function':
        query_type = 'Generated Table'
    elif source_type == 'Expression':
        query_type = 'Transformation'

    return (query_type, source_type)


@app.route('/api/generate/progress/<job_id>')
@login_required
def get_generation_progress(job_id):
    """Get the current progress of a documentation generation job"""
    progress_data = generation_progress.get(job_id, {'progress': 0, 'status': 'Not started', 'complete': False})
    return jsonify(progress_data)


@app.route('/api/generate/download/<job_id>')
def download_generated_file(job_id):
    """Download the generated file for a completed job"""
    # Note: Removed @login_required to avoid request context issues
    # The job_id itself acts as a secure token (UUID + timestamp)

    job_data = generation_progress.get(job_id)

    if not job_data:
        return jsonify({'error': 'Job not found'}), 404

    if not job_data.get('complete'):
        return jsonify({'error': 'Generation not complete'}), 400

    if job_data.get('error'):
        return jsonify({'error': job_data['error']}), 500

    file_path = job_data.get('file_path')
    if not file_path or not os.path.exists(file_path):
        return jsonify({'error': 'File not found'}), 404

    # Get the filename from the path
    filename = os.path.basename(file_path)

    return send_file(
        file_path,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=filename
    )


def _generate_documentation_background(workspace_id, report_id, dataset_id, report_name, job_id, user_token):
    """Background task to generate documentation - runs in a separate thread"""
    try:
        print(f"\n🚀 [Job {job_id}] Starting complete documentation generation")
        print(f"   Workspace: {workspace_id}")
        print(f"   Report: {report_id}")
        print(f"   Dataset: {dataset_id}")

        generation_progress[job_id] = {'progress': 5, 'status': 'Initializing...', 'complete': False}

        # ✅ Use user-delegated token passed from main thread (already extracted from session)
        if user_token:
            powerbi.set_user_token(user_token)
            print(f"   🔑 Using user-delegated token for API calls")
        else:
            print(f"   ⚠️ No user token found, falling back to service principal")

        generation_progress[job_id] = {'progress': 10, 'status': 'Authenticating...', 'complete': False}

        # Ensure output directory exists
        output_dir = 'output'
        os.makedirs(output_dir, exist_ok=True)

        # Get dataset_id if not provided
        if not dataset_id:
            print(f"   ⚠️ Dataset ID not provided, fetching from report metadata...")
            generation_progress[job_id] = {'progress': 15, 'status': 'Fetching report info...', 'complete': False}
            # Quick fetch to get dataset_id
            reports = powerbi.get_all_reports(workspace_id)
            report = next((r for r in reports if r['id'] == report_id), None)
            if report:
                dataset_id = report.get('datasetId')
                report_name = report.get('name', report_name)
                print(f"   ✅ Found dataset ID: {dataset_id}")

        if not dataset_id:
            generation_progress[job_id] = {
                'progress': 0,
                'status': 'Error: No dataset ID',
                'complete': True,
                'error': 'Could not determine dataset ID for this report'
            }
            return  # Exit the background thread

        generation_progress[job_id] = {'progress': 20, 'status': 'Fetching metadata...', 'complete': False}

        # Use PowerBIDataFetcher for complete metadata (same as Main.py)
        from ai_generator import PowerBIDataFetcher

        # ✅ FIX: Pass user token to PowerBIDataFetcher
        fetcher = PowerBIDataFetcher(
            config.CLIENT_ID,
            config.CLIENT_SECRET,
            config.TENANT_ID,
            workspace_id,
            user_token=user_token  # Pass user-delegated token
        )

        print(f"   🔑 PowerBIDataFetcher initialized with {'user-delegated' if user_token else 'service principal'} token")

        generation_progress[job_id] = {'progress': 30, 'status': 'Analyzing dataset...', 'complete': False}

        # Get complete metadata including scanner data
        metadata = fetcher.get_complete_metadata(
            dataset_id=dataset_id,
            report_id=report_id,
            history_top=20
        )

        generation_progress[job_id] = {'progress': 45, 'status': 'Metadata collected...', 'complete': False}

        # Add report name to metadata
        metadata['report_name'] = report_name
        metadata['name'] = report_name

        generation_progress[job_id] = {'progress': 50, 'status': 'Generating overview...', 'complete': False}

        # Generate AI documentation sections
        print(f"\n🤖 Generating AI documentation sections...")

        # Create the documentation structure (same as generate_complete_documentation)
        documentation = {
            'overview': ai_generator.generate_overview(metadata),
            'data_sources': None,
            'pages': None,
            'user_guide': None,
            'technical_details': None,
            'migration': None,
            'metadata': metadata
        }

        generation_progress[job_id] = {'progress': 60, 'status': 'Generating data sources...', 'complete': False}
        documentation['data_sources'] = ai_generator.generate_data_sources_doc(metadata.get('data_sources', []))

        generation_progress[job_id] = {'progress': 70, 'status': 'Generating pages...', 'complete': False}
        documentation['pages'] = ai_generator.generate_pages_documentation(metadata.get('pages', []))

        generation_progress[job_id] = {'progress': 75, 'status': 'Generating user guide...', 'complete': False}
        documentation['user_guide'] = ai_generator.generate_user_guide(metadata.get('pages', []), report_name)

        generation_progress[job_id] = {'progress': 80, 'status': 'Generating technical details...', 'complete': False}
        documentation['technical_details'] = ai_generator.generate_technical_details(metadata)

        generation_progress[job_id] = {'progress': 85, 'status': 'Generating migration steps...', 'complete': False}
        documentation['migration'] = ai_generator.generate_migration_steps(report_name)

        generation_progress[job_id] = {'progress': 90, 'status': 'Creating document...', 'complete': False}

        # Create Word document
        doc_creator = PowerBIDocumentCreator()
        doc_filename = f"{report_name}_Documentation.docx"
        doc_path = os.path.join(output_dir, doc_filename)
        print(f"\n📝 Creating Word document: {doc_path}")

        # Create comprehensive document with proper structure
        doc_creator.create_documentation_from_json(
            json_data=documentation,  # Pass the full documentation structure, not just metadata
            output_filename=doc_path,
            author=config.AUTHOR_NAME
        )

        generation_progress[job_id] = {
            'progress': 100,
            'status': 'Complete!',
            'complete': True,
            'file_path': doc_path,
            'filename': doc_filename
        }

        print(f"✅ [Job {job_id}] Documentation generation complete!")

        # Clean up progress after a delay
        def cleanup_progress():
            time.sleep(300)  # Keep progress for 5 minutes
            if job_id in generation_progress:
                print(f"🧹 Cleaning up job {job_id}")
                # Also delete the file
                try:
                    if os.path.exists(doc_path):
                        os.remove(doc_path)
                except:
                    pass
                del generation_progress[job_id]

        threading.Thread(target=cleanup_progress, daemon=True).start()

    except Exception as e:
        print(f"❌ [Job {job_id}] Error generating document: {str(e)}")
        import traceback
        traceback.print_exc()

        # Update progress with error
        generation_progress[job_id] = {
            'progress': 0,
            'status': f'Error: {str(e)[:50]}',
            'complete': True,
            'error': str(e)
        }


@app.route('/api/generate', methods=['POST'])
@login_required
def generate_documentation():
    """Start documentation generation in background and return job ID"""
    data = request.get_json()

    workspace_id = data.get('workspace_id')
    report_id = data.get('report_id')
    dataset_id = data.get('dataset_id')
    report_name = data.get('report_name', 'Report')
    job_id = data.get('job_id', f'{report_id}_{int(time.time())}')

    if not report_id:
        return jsonify({
            'success': False,
            'error': 'Report ID is required'
        }), 400

    # Use provided workspace_id or fall back to config
    if not workspace_id:
        workspace_id = config.WORKSPACE_ID

    # Get user token from session
    user_token = session.get('access_token')

    # Initialize progress tracking
    generation_progress[job_id] = {'progress': 0, 'status': 'Starting...', 'complete': False}

    # Start background thread for generation
    thread = threading.Thread(
        target=_generate_documentation_background,
        args=(workspace_id, report_id, dataset_id, report_name, job_id, user_token),
        daemon=True
    )
    thread.start()

    # Return immediately with job ID
    return jsonify({
        'success': True,
        'job_id': job_id,
        'message': 'Documentation generation started'
    })



def _warm_catalog_async():
    """
    Warm ONLY thin packs (Home + Impact table list + summary).

    Never preload workspace_catalog.json / impact_index.json into workers —
    that was the main App Service OOM path:
      Worker was sent SIGKILL! Perhaps out of memory?
    Full catalog is read on demand from disk/SharePoint and not kept in RAM
    (see CATALOG_KEEP_HEAVY_IN_MEMORY).
    """
    if not CATALOG_AVAILABLE or catalog_service is None:
        return

    def _run():
        try:
            t0 = time.time()
            home = catalog_service.get_json('ui_home_index.json', force_refresh=False)
            tables = catalog_service.get_json('ui_impact_tables.json', force_refresh=False)
            report_dir = catalog_service.get_json('ui_report_directory.json', force_refresh=False)
            summary = catalog_service.get_summary(force_refresh=False)
            try:
                catalog_service.get_json('ops_summary.json', force_refresh=False)
            except Exception:
                pass
            n_home = len((home or {}).get('workspaces') or [])
            n_tables = len((tables or {}).get('rows') or []) if isinstance(tables, dict) else 0
            n_dir = len((report_dir or {}).get('rows') or []) if isinstance(report_dir, dict) else 0
            if not n_tables and hasattr(catalog_service, 'impact_table_rows'):
                try:
                    n_tables = len(catalog_service.impact_table_rows())
                except Exception:
                    pass
            # Ensure we never left a heavy blob from a rebuild path
            try:
                catalog_service.drop_heavy_memory()
            except Exception:
                pass
            print(
                f"⚡ Catalog warm-up (thin only) in {time.time() - t0:.1f}s "
                f"(homeWs={n_home}, impactRows={n_tables}, reportDir={n_dir}, "
                f"summary={bool(summary)}, opsEnrichedAt={(home or summary or {}).get('opsEnrichedAt')})"
            )
            if n_home == 0:
                print(
                    "⚠️ ui_home_index empty/missing after warm-up — Home will show offline "
                    "until SharePoint has the thin pack (extract job). "
                    "Do NOT load full workspace_catalog on this SKU."
                )
        except Exception as exc:
            print(f"⚠️ Catalog warm-up failed: {exc}")

    threading.Thread(target=_run, daemon=True, name='catalog-warm').start()


# Warm catalog on import so Gunicorn/Azure workers also preload thin packs.
# Guard with env if you need to disable on tiny SKUs: CATALOG_WARM_ON_START=false
if os.getenv('CATALOG_WARM_ON_START', 'true').lower() in ('1', 'true', 'yes', 'y'):
    try:
        _warm_catalog_async()
    except Exception as _warm_exc:
        print(f"Catalog warm-on-import skipped: {_warm_exc}")


if __name__ == '__main__':
    # Ensure output directory exists
    os.makedirs('output', exist_ok=True)
    # Use localhost instead of 0.0.0.0 to match Azure AD redirect URI
    app.run(debug=True, host='localhost', port=5000, use_reloader=False)

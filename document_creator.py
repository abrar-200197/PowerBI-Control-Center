# document_creator.py - Enhanced Version with M-to-SQL Conversion
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
import json
import re

# optional AI generator for on‑demand narrative creation
try:
    from ai_generator import AIDocGenerator
except ImportError:
    AIDocGenerator = None

class MQueryToSQLConverter:
    """Convert Power Query M expressions to SQL queries"""
    
    @staticmethod
    def extract_sql_from_m(m_expression):
        """
        Extract SQL query from M expression if present
        
        Args:
            m_expression (str): Power Query M code
            
        Returns:
            dict: {'has_sql': bool, 'sql_query': str, 'server': str, 'database': str}
        """
        result = {
            'has_sql': False,
            'sql_query': None,
            'server': None,
            'database': None,
            'connection_timeout': None
        }
        
        if not m_expression:
            return result
        
        # Extract server and database
        server_match = re.search(r'Sql\.Database\s*\(\s*"([^"]+)"', m_expression)
        if server_match:
            result['server'] = server_match.group(1)
        
        db_match = re.search(r'Sql\.Database\s*\([^,]+,\s*"([^"]+)"', m_expression)
        if db_match:
            result['database'] = db_match.group(1)
        
        # Extract SQL query from Query parameter
        query_match = re.search(r'Query\s*=\s*"([^"]*(?:"[^"]*")*[^"]*)"', m_expression, re.DOTALL)
        if query_match:
            sql_query = query_match.group(1)
            # Clean up the query
            sql_query = sql_query.replace('#(lf)', '\n')
            sql_query = sql_query.replace('#(tab)', '\t')
            sql_query = sql_query.replace('""', '"')
            result['sql_query'] = sql_query.strip()
            result['has_sql'] = True
        
        # Extract timeout if present
        timeout_match = re.search(r'CommandTimeout\s*=\s*#duration\((\d+),\s*(\d+),\s*(\d+),\s*(\d+)\)', m_expression)
        if timeout_match:
            days, hours, mins, secs = map(int, timeout_match.groups())
            total_minutes = (days * 24 * 60) + (hours * 60) + mins + (secs / 60)
            result['connection_timeout'] = f"{int(total_minutes)} minutes"
        
        return result
    
    @staticmethod
    def format_sql_query(sql_query):
        """
        Format SQL query for better readability
        
        Args:
            sql_query (str): Raw SQL query
            
        Returns:
            str: Formatted SQL query
        """
        if not sql_query:
            return ""
        
        # Keywords to format
        keywords = [
            'SELECT', 'FROM', 'WHERE', 'JOIN', 'LEFT JOIN', 'RIGHT JOIN', 
            'INNER JOIN', 'OUTER JOIN', 'GROUP BY', 'ORDER BY', 'HAVING',
            'UNION', 'EXCEPT', 'INTERSECT', 'INSERT', 'UPDATE', 'DELETE',
            'CREATE', 'ALTER', 'DROP', 'EXEC', 'EXECUTE', 'AS', 'ON', 'AND', 'OR'
        ]
        
        formatted = sql_query
        
        # Add newlines before major keywords
        for keyword in ['SELECT', 'FROM', 'WHERE', 'GROUP BY', 'ORDER BY', 'HAVING']:
            formatted = re.sub(f'\\b{keyword}\\b', f'\n{keyword}', formatted, flags=re.IGNORECASE)
        
        # Clean up multiple newlines
        formatted = re.sub(r'\n\s*\n', '\n', formatted)
        
        # Trim whitespace
        lines = [line.strip() for line in formatted.split('\n')]
        formatted = '\n'.join(line for line in lines if line)
        
        return formatted


class PowerBIDocumentCreator:
    """Creates professional Word documents from Power BI documentation"""
    
    def __init__(self):
        self.doc = Document()
        self.setup_styles()
        self.sql_converter = MQueryToSQLConverter()
        
    def setup_styles(self):
        """Configure document styles"""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

    # ---------- Utility ----------

    def safe_str(self, value, default='N/A'):
        """Safely convert any value to string, handling None"""
        if value is None:
            return default
        return str(value)

    # ---------- Cover Page & TOC ----------

    def add_cover_page(self, report_name, author="System Generated", company="Ashley Furniture India"):
        """Create professional cover page"""
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for _ in range(5):
            title.add_run('\n')
        
        title_run = title.add_run(f'{report_name}\n\n')
        title_run.font.size = Pt(32)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 51, 102)
        
        subtitle_run = title.add_run('Technical Documentation\n\n')
        subtitle_run.font.size = Pt(20)
        subtitle_run.font.color.rgb = RGBColor(68, 84, 106)
        
        company_run = title.add_run(f'\n{company}\n')
        company_run.font.size = Pt(14)
        company_run.font.bold = True
        
        dept_run = title.add_run('Quality Department - Business Intelligence')
        dept_run.font.size = Pt(12)
        
        for _ in range(3):
            self.doc.add_paragraph()
        
        table = self.doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'
        
        table.rows[0].cells[0].text = 'Generated Date:'
        table.rows[0].cells[1].text = datetime.now().strftime('%B %d, %Y at %H:%M')
        
        table.rows[1].cells[0].text = 'Author:'
        table.rows[1].cells[1].text = self.safe_str(author)
        
        table.rows[2].cells[0].text = 'Version:'
        table.rows[2].cells[1].text = '1.0 (Auto-Generated)'
        
        table.rows[3].cells[0].text = 'Generator:'
        table.rows[3].cells[1].text = 'AI Documentation System'
        
        for row in table.rows:
            row.cells[0].paragraphs[0].runs[0].font.bold = True
        
        self.doc.add_page_break()
    
    def add_table_of_contents(self):
        """Add table of contents"""
        self.doc.add_heading('Table of Contents', level=1)
        
        toc_items = [
            ('1. Report Overview', 'Business purpose and objectives'),
            ('2. Data Sources & Architecture', 'Source systems and data flow'),
            ('3. Report Pages', 'Page-by-page breakdown'),
            ('4. User Guide', 'How to use the report'),
            ('5. Technical Details', 'Technical specifications'),
            ('6. Migration & Deployment', 'Deployment procedures'),
            ('7. Backend Queries & Data Connections', 'Data connections and query details'),
            ('8. Appendix', 'Additional information')
        ]
        
        for item, description in toc_items:
            p = self.doc.add_paragraph()
            run_item = p.add_run(item)
            run_item.font.bold = True
            run_item.font.size = Pt(11)
            p.add_run(f'\n    {description}')
        
        self.doc.add_page_break()

    # ---------- Markdown Renderer ----------

    def add_markdown_content(self, content):
        """Parse and add markdown-formatted content to document"""
        if not content:
            return
        
        lines = content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                i += 1
                continue
            
            if line.startswith('###'):
                self.doc.add_heading(line.replace('###', '').strip(), level=3)
            elif line.startswith('##'):
                self.doc.add_heading(line.replace('##', '').strip(), level=2)
            elif line.startswith('#'):
                self.doc.add_heading(line.replace('#', '').strip(), level=1)
            elif line.startswith('- ') or line.startswith('* '):
                self.doc.add_paragraph(line[2:].strip(), style='List Bullet')
            elif len(line) > 2 and line[0].isdigit() and line[1] == '.':
                self.doc.add_paragraph(line[2:].strip(), style='List Number')
            elif line.startswith('- [ ]') or line.startswith('- [x]'):
                checked = '☑' if '[x]' in line else '☐'
                text = line.replace('- [ ]', '').replace('- [x]', '').strip()
                self.doc.add_paragraph(f'{checked} {text}', style='List Bullet')
            elif '**' in line:
                p = self.doc.add_paragraph()
                parts = line.split('**')
                for idx, part in enumerate(parts):
                    if part:
                        run = p.add_run(part)
                        if idx % 2 == 1:
                            run.font.bold = True
            elif line.startswith('⚠️') or line.startswith('⚡') or line.startswith('💡'):
                p = self.doc.add_paragraph(line, style='List Bullet')
                p.runs[0].font.color.rgb = RGBColor(255, 140, 0)
            else:
                self.doc.add_paragraph(line)
            
            i += 1

    # ---------- Generic Section Helper ----------

    def add_section(self, title, content, level=1):
        """Add a documentation section"""
        self.doc.add_heading(title, level=level)
        
        if isinstance(content, str):
            self.add_markdown_content(content)
        elif isinstance(content, dict):
            if content:
                table = self.doc.add_table(rows=len(content)+1, cols=2)
                table.style = 'Light Shading Accent 1'
                
                header_cells = table.rows[0].cells
                header_cells[0].text = 'Property'
                header_cells[1].text = 'Value'
                
                for cell in header_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                
                for idx, (key, value) in enumerate(content.items(), 1):
                    row_cells = table.rows[idx].cells
                    row_cells[0].text = self.safe_str(key)
                    row_cells[1].text = self.safe_str(value)
        elif isinstance(content, list):
            for item in content:
                self.doc.add_paragraph(self.safe_str(item), style='List Bullet')
        
        self.doc.add_paragraph()

    # ---------- Data Analysis Methods ----------

    def add_refresh_section(self, metadata):
        """7.x Refresh schedule and history details"""
        # logging for visibility
        print("   🔍 [Document Creator] add_refresh_section called")
        schedule = metadata.get('refresh_details') or metadata.get('refresh_schedule')
        history = metadata.get('refresh_history', [])

        self.doc.add_heading('7.10 Refresh Schedule & History', level=2)

        if schedule:
            self.doc.add_paragraph('**Refresh Schedule Configuration:**')
            tbl = self.doc.add_table(rows=len(schedule) + 1, cols=2)
            tbl.style = 'Light Shading Accent 1'
            hdr = tbl.rows[0].cells
            hdr[0].text = 'Property'
            hdr[1].text = 'Value'
            for cell in hdr:
                cell.paragraphs[0].runs[0].font.bold = True
            for idx, (k, v) in enumerate(schedule.items(), 1):
                row = tbl.rows[idx].cells
                row[0].text = str(k)
                row[1].text = self.safe_str(v)
            self.doc.add_paragraph()
        else:
            self.doc.add_paragraph('No refresh schedule information available.')
            self.doc.add_paragraph()

        if history:
            self.doc.add_paragraph('**Recent Refresh History:**')
            # preferred ordering for common fields, followed by any extras
            preferred = ['startTime', 'endTime', 'status', 'refreshType', 'requestId', 'id', 'serviceExceptionJson', 'refreshAttempts']
            # gather all keys from records
            all_keys = []
            for rec in history:
                for k in rec.keys():
                    if k not in all_keys:
                        all_keys.append(k)
            # sort keys according to preferred list then alphabetically
            ordered_keys = [k for k in preferred if k in all_keys]
            ordered_keys += sorted([k for k in all_keys if k not in ordered_keys])

            cols = len(ordered_keys)
            hist_tbl = self.doc.add_table(rows=len(history) + 1, cols=cols)
            hist_tbl.style = 'Light Grid Accent 1'
            # header row
            for j, key in enumerate(ordered_keys):
                cell = hist_tbl.rows[0].cells[j]
                cell.text = key
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(9)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            # data rows
            for idx, rec in enumerate(history, 1):
                if idx >= len(hist_tbl.rows):
                    break
                cells = hist_tbl.rows[idx].cells
                for j, key in enumerate(ordered_keys):
                    val = rec.get(key, '')
                    cells[j].text = self.safe_str(val)
                    for para in cells[j].paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for run in para.runs:
                            run.font.size = Pt(9)
            self.doc.add_paragraph()
        else:
            self.doc.add_paragraph('No refresh history records available.')
            self.doc.add_paragraph()

    def add_relationship_diagram_section(self, metadata, output_filename=None):
        """Add a Mermaid ER diagram representation of relationships and embed image if possible."""
        rels = metadata.get('relationships', [])
        tables = metadata.get('model_tables', [])
        # always show header
        self.doc.add_heading('7.11 Relationship Diagram', level=2)
        if not rels:
            self.doc.add_paragraph('No relationships detected in the model.')
        self.doc.add_paragraph(
            'Below is a Mermaid `erDiagram` representation of the relationship graph.'
        )
        lines = ['erDiagram']
        # if there are no relationships, still include table nodes
        if not rels and tables:
            for t in tables:
                lines.append(f'    {t}')
        else:
            for rel in rels:
                ft = rel.get('fromTable', 'Unknown')
                tt = rel.get('toTable', 'Unknown')
                card = rel.get('type','').lower() or ''
                if 'manytoone' in card:
                    arrow = '||--o{'
                elif 'onetomany' in card:
                    arrow = '}o--||'
                elif 'manytomany' in card:
                    arrow = '}o--o{'
                else:
                    arrow = '||--||'
                label = rel.get('name','')
                lines.append(f'    {ft} {arrow} {tt} : "{label}"')
        mermaid_text = '\n'.join(lines)

        # save mermaid file
        mmd_path = None
        if output_filename:
            mmd_path = os.path.splitext(output_filename)[0] + '_relationships.mmd'
            try:
                with open(mmd_path, 'w', encoding='utf-8') as f:
                    f.write(mermaid_text)
                print(f"   ✅ Mermaid diagram saved: {mmd_path}")
            except Exception as e:
                print(f"   ⚠️ Could not save mermaid file: {e}")

        # try to render image using mermaid-cli if available
        img_path = None
        try:
            import subprocess, tempfile
            if mmd_path:
                img_path = os.path.splitext(mmd_path)[0] + '.png'
                proc = subprocess.run(['mmdc', '-i', mmd_path, '-o', img_path], capture_output=True)
                if proc.returncode == 0 and os.path.exists(img_path):
                    print(f"   ✅ Mermaid image generated: {img_path}")
                    # embed in document
                    try:
                        self.doc.add_picture(img_path, width=Pt(450))
                    except Exception as e:
                        print(f"   ⚠️ Could not embed image: {e}")
                else:
                    print(f"   ⚠️ Mermaid CLI failed, output: {proc.stderr.decode()}")
        except FileNotFoundError:
            # mmdc not installed
            pass
        except Exception as e:
            print(f"   ⚠️ Error running mermaid-cli: {e}")
        
        # always include code block for reference
        self.doc.add_paragraph(f'```mermaid\n{mermaid_text}\n```')

    def export_refresh_history_csv(self, metadata, output_filename=None):
        """Export refresh history records to a CSV file"""
        history = metadata.get('refresh_history', [])
        if not history or not output_filename:
            return
        csv_file = os.path.splitext(output_filename)[0] + '_refresh_history.csv'
        try:
            import csv
            with open(csv_file, 'w', newline='', encoding='utf-8') as csvfile:
                fieldnames = list({k for rec in history for k in rec.keys()})
                writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
                writer.writeheader()
                for rec in history:
                    writer.writerow(rec)
            print(f"   ✅ Refresh history exported to CSV: {csv_file}")
        except Exception as e:
            print(f"   ⚠️ Error exporting refresh history CSV: {e}")

    def detect_unused_columns_and_tables(self, metadata):
        """
        Detect potentially unused tables and columns in the model
        
        Returns:
            dict: {
                'unused_tables': [],
                'tables_with_unused_columns': {table: [columns]},
                'analysis': str
            }
        """
        result = {
            'unused_tables': [],
            'tables_with_unused_columns': {},
            'analysis': ''
        }
        
        # Get model tables and columns
        model_tables = metadata.get('model_tables', [])
        model_columns = metadata.get('model_columns', {})
        expressions = metadata.get('expressions', [])
        
        # Get tables used in expressions
        tables_in_expressions = set([expr.get('table') for expr in expressions if expr.get('table')])
        
        # Get tables used in relationships
        relationships = metadata.get('relationships', [])
        tables_in_relationships = set()
        for rel in relationships:
            tables_in_relationships.add(rel.get('fromTable'))
            tables_in_relationships.add(rel.get('toTable'))
        
        # Get all used tables
        all_used_tables = tables_in_expressions.union(tables_in_relationships)
        
        # Detect unused tables (tables not in expressions or relationships)
        model_table_set = set([t for t in model_tables if t])
        unused_tables = model_table_set - all_used_tables
        
        if unused_tables:
            result['unused_tables'] = sorted(list(unused_tables))
        
        # Detect columns with no references (very simple heuristic)
        # A column is potentially unused if:
        # - It starts with underscore (convention)
        # - It contains "ID" and there's no matching relationship
        # - It was added as "Calculated" but not referenced
        
        for table_name, columns in model_columns.items():
            unused_in_table = []
            for col in columns:
                col_name = col.get('name', '')
                col_type = col.get('columnType', 'Data')
                
                # Flag potential candidates
                if col_name.startswith('_') or col_name.lower() in ['index', 'row_number']:
                    unused_in_table.append(col_name)
            
            if unused_in_table:
                result['tables_with_unused_columns'][table_name] = unused_in_table
        
        return result

    # ---------- Backend Queries & Connections (7.x) ----------

    def add_queries_and_connections_section(self, metadata):
        """Add backend queries and connections section"""
        self.doc.add_heading('7. Backend Queries & Data Connections', level=1)
        
        intro = """
This section documents the backend queries, data connections, and source configurations that power this report.
Understanding these technical details is essential for troubleshooting, optimization, and maintenance.
"""
        self.doc.add_paragraph(intro)
        
        # 7.1 Data Source Connections
        self.doc.add_heading('7.1 Data Source Connection Details', level=2)
        detailed_sources = metadata.get('detailed_sources', [])
        
        if detailed_sources:
            for idx, source in enumerate(detailed_sources, 1):
                ds_type = self.safe_str(source.get('datasourceType', 'Unknown'))
                self.doc.add_heading(f"Connection {idx}: {ds_type}", level=3)
                
                conn_table = self.doc.add_table(rows=5, cols=2)
                conn_table.style = 'Light Shading Accent 1'
                
                conn_table.rows[0].cells[0].text = 'Property'
                conn_table.rows[0].cells[1].text = 'Value'
                for cell in conn_table.rows[0].cells:
                    cell.paragraphs[0].runs[0].font.bold = True
                
                conn_table.rows[1].cells[0].text = 'Source Type'
                conn_table.rows[1].cells[1].text = ds_type
                
                conn_table.rows[2].cells[0].text = 'Connection String'
                conn_table.rows[2].cells[1].text = self.safe_str(source.get('connectionString', 'N/A'))
                
                conn_table.rows[3].cells[0].text = 'Gateway'
                gateway_status = 'Yes (Gateway Required)' if source.get('gatewayId') else 'No (Direct Cloud Connection)'
                conn_table.rows[3].cells[1].text = gateway_status
                
                if source.get('gatewayId'):
                    conn_table.rows[4].cells[0].text = 'Gateway ID'
                    conn_table.rows[4].cells[1].text = self.safe_str(source.get('gatewayId', 'N/A'))
                else:
                    conn_table.rows[4].cells[0].text = 'Connection Mode'
                    conn_table.rows[4].cells[1].text = 'Direct Cloud Connection'
                
                conn_details = source.get('connectionDetails', {})
                if conn_details:
                    self.doc.add_paragraph()
                    self.doc.add_heading('Connection Parameters:', level=4)
                    for key, value in conn_details.items():
                        if value and value != 'N/A':
                            p = self.doc.add_paragraph(style='List Bullet')
                            p.add_run(f'{key}: ').bold = True
                            p.add_run(self.safe_str(value))
                
                self.doc.add_paragraph()
        else:
            self.doc.add_paragraph("⚠️ Detailed data source information not available.")
        
        # 7.2 Dataset Configuration
        self.doc.add_heading('7.2 Dataset Configuration', level=2)
        refresh_details = metadata.get('refresh_details')
        
        if refresh_details:
            config_table = self.doc.add_table(rows=6, cols=2)
            config_table.style = 'Light List Accent 1'
            
            config_table.rows[0].cells[0].text = 'Dataset Name'
            config_table.rows[0].cells[1].text = self.safe_str(refresh_details.get('name') or 'N/A')
            
            config_table.rows[1].cells[0].text = 'Storage Mode'
            config_table.rows[1].cells[1].text = self.safe_str(refresh_details.get('targetStorageMode') or 'Unknown')
            
            config_table.rows[2].cells[0].text = 'Refreshable'
            config_table.rows[2].cells[1].text = 'Yes' if refresh_details.get('isRefreshable') else 'No'
            
            config_table.rows[3].cells[0].text = 'Content Provider'
            config_table.rows[3].cells[1].text = self.safe_str(refresh_details.get('contentProviderType') or 'N/A')
            
            config_table.rows[4].cells[0].text = 'Created Date'
            created = self.safe_str(refresh_details.get('createdDate') or 'Unknown')
            config_table.rows[4].cells[1].text = created[:10] if len(created) > 10 else created
            
            config_table.rows[5].cells[0].text = 'Configured By'
            config_table.rows[5].cells[1].text = self.safe_str(refresh_details.get('configuredBy') or 'Unknown')
            
            for row in config_table.rows:
                row.cells[0].paragraphs[0].runs[0].font.bold = True
        
        self.doc.add_paragraph()
        
        # 7.3 Dataset Parameters
        self.doc.add_heading('7.3 Dataset Parameters', level=2)
        parameters = metadata.get('parameters', [])
        
        if parameters:
            self.doc.add_paragraph(f"This dataset uses {len(parameters)} parameter(s):")
            self.doc.add_paragraph()
            
            param_table = self.doc.add_table(rows=len(parameters)+1, cols=3)
            param_table.style = 'Light Grid Accent 1'
            
            param_table.rows[0].cells[0].text = 'Parameter Name'
            param_table.rows[0].cells[1].text = 'Type'
            param_table.rows[0].cells[2].text = 'Current Value'
            
            for cell in param_table.rows[0].cells:
                cell.paragraphs[0].runs[0].font.bold = True
            
            for idx, param in enumerate(parameters, 1):
                param_table.rows[idx].cells[0].text = self.safe_str(param.get('name', 'Unknown'))
                param_table.rows[idx].cells[1].text = self.safe_str(param.get('type', 'Unknown'))
                param_table.rows[idx].cells[2].text = self.safe_str(param.get('currentValue', 'N/A'))
        else:
            self.doc.add_paragraph("This dataset does not use parameters.")
        
        self.doc.add_paragraph()
        
        # 7.4 Query Performance
        self.doc.add_heading('7.4 Query Performance & Optimization', level=2)
        storage_mode = self.safe_str(refresh_details.get('targetStorageMode') or 'Unknown') if refresh_details else 'Unknown'
        self.doc.add_heading(f'Current Storage Mode: {storage_mode}', level=3)
        
        perf_notes = [
            "**DirectQuery** - Queries are sent to source in real-time. Slower but always current.",
            "**Import** - Data is cached in Power BI. Faster queries but requires refresh.",
            "**Composite** - Mix of DirectQuery and Import for optimization.",
            "**Query Folding** - Ensure Power Query transformations push to source database.",
            "**Incremental Refresh** - Use for large datasets to improve refresh times.",
            "**Aggregations** - Consider aggregation tables for frequently queried data.",
            "**Gateway Performance** - Monitor gateway health and configure appropriate capacity.",
            "**Parameterization** - Use parameters for dynamic connection strings and filters."
        ]
        
        for note in perf_notes:
            self.doc.add_paragraph(note, style='List Bullet')
        
        # 7.5 Security Notes
        self.doc.add_heading('7.5 Data Security & Access Control', level=2)
        security_notes = [
            "**Row-Level Security (RLS)** - Verify RLS roles are properly configured if needed",
            "**Credential Management** - Ensure data source credentials are securely stored",
            "**Gateway Security** - Gateway should be in secure network zone",
            "**OAuth/SSO** - Use modern authentication where supported",
            "**Encryption** - Data should be encrypted in transit and at rest",
            "**Access Reviews** - Regularly audit who has access to datasets and reports"
        ]
        for note in security_notes:
            self.doc.add_paragraph(note, style='List Bullet')
        
        self.doc.add_paragraph()
        
        # 7.6 Troubleshooting
        self.doc.add_heading('7.6 Troubleshooting Data Connection Issues', level=2)
        troubleshoot_table = self.doc.add_table(rows=5, cols=3)
        troubleshoot_table.style = 'Light List Accent 1'
        
        troubleshoot_table.rows[0].cells[0].text = 'Issue'
        troubleshoot_table.rows[0].cells[1].text = 'Possible Cause'
        troubleshoot_table.rows[0].cells[2].text = 'Solution'
        for cell in troubleshoot_table.rows[0].cells:
            cell.paragraphs[0].runs[0].font.bold = True
        
        issues = [
            {
                'issue': 'Data Not Refreshing',
                'cause': 'Gateway offline, credentials expired',
                'solution': 'Check gateway status, update credentials'
            },
            {
                'issue': 'Slow Query Performance',
                'cause': 'DirectQuery on large tables',
                'solution': 'Consider Import mode or aggregations'
            },
            {
                'issue': 'Connection Timeout',
                'cause': 'Complex query, network latency',
                'solution': 'Optimize source queries, check network'
            },
            {
                'issue': 'Authentication Failed',
                'cause': 'Invalid credentials, expired token',
                'solution': 'Update data source credentials'
            }
        ]
        
        for idx, issue in enumerate(issues, 1):
            troubleshoot_table.rows[idx].cells[0].text = issue['issue']
            troubleshoot_table.rows[idx].cells[1].text = issue['cause']
            troubleshoot_table.rows[idx].cells[2].text = issue['solution']
        
        self.doc.add_paragraph()

    def add_columns_section(self, metadata):
        """7.7 Model Columns (Tables and Columns) - Enhanced with DAX Formulas"""
        columns_dict = metadata.get('model_columns', {})
        if not columns_dict:
            return

        self.doc.add_heading('7.7 Model Columns & Data Types', level=2)
        self.doc.add_paragraph(
            "Summary of all tables and their columns in the data model, including DAX formulas for calculated columns."
        )

        for table_name, columns in columns_dict.items():
            if columns:
                # Count data vs calculated columns
                data_cols = len([c for c in columns if c.get('columnType', 'Data') == 'Data'])
                calc_cols = len([c for c in columns if c.get('columnType', 'Data') == 'Calculated'])

                summary = f'{table_name} ({len(columns)} columns: {data_cols} data, {calc_cols} calculated)'
                self.doc.add_heading(summary, level=3)

                # Create compact column listing with column type indicator
                col_table = self.doc.add_table(rows=len(columns) + 1, cols=3)
                col_table.style = 'Light Shading Accent 1'
                col_table.autofit = False

                headers = col_table.rows[0].cells
                headers[0].text = 'Column'
                headers[1].text = 'Data Type'
                headers[2].text = 'Column Type'

                for cell in headers:
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].runs[0].font.size = Pt(10)

                for idx, col in enumerate(columns, 1):
                    if idx >= len(col_table.rows):
                        break
                    cells = col_table.rows[idx].cells
                    col_name = self.safe_str(col.get('name', 'N/A'))
                    col_data_type = self.safe_str(col.get('dataType', 'N/A'))
                    col_type = self.safe_str(col.get('columnType', 'Data'))

                    cells[0].text = col_name
                    cells[1].text = col_data_type
                    cells[2].text = col_type

                    for cell in cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.size = Pt(9)

                # Show DAX formulas for calculated columns
                calc_columns = [c for c in columns if c.get('columnType') == 'Calculated' and c.get('expression')]
                if calc_columns:
                    self.doc.add_paragraph()
                    self.doc.add_heading(f'Calculated Columns in {table_name}', level=4)

                    for calc_col in calc_columns:
                        col_name = self.safe_str(calc_col.get('name', 'Unknown'))
                        expression = self.safe_str(calc_col.get('expression', 'N/A'))

                        p = self.doc.add_paragraph()
                        run_name = p.add_run(f'{col_name} = ')
                        run_name.font.bold = True
                        run_name.font.size = Pt(10)

                        if expression and expression != 'N/A':
                            p_expr = self.doc.add_paragraph()
                            run_expr = p_expr.add_run(expression[:3000])
                            run_expr.font.name = 'Consolas'
                            run_expr.font.size = Pt(8)
                            run_expr.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue for DAX

                        self.doc.add_paragraph()
            else:
                self.doc.add_paragraph('No columns found.')

        self.doc.add_paragraph()

    def add_dax_measures_section(self, metadata):
        """7.8 DAX Measures - Enhanced with All Formulas"""
        measures = metadata.get('measures', [])

        # DEBUG LOGGING
        print(f"\n   🔍 [Document Creator] add_dax_measures_section called")
        print(f"      - Measures in metadata: {len(measures)}")
        if measures:
            print(f"      - First measure: {measures[0]}")

        if not measures:
            self.doc.add_heading('7.8 DAX Measures', level=2)
            self.doc.add_paragraph('No DAX measures defined in this model.')
            self.doc.add_paragraph()
            return

        self.doc.add_heading('7.8 DAX Measures', level=2)
        self.doc.add_paragraph(
            f'Total Measures: {len(measures)}'
        )

        # Create summary table first
        measure_table = self.doc.add_table(rows=len(measures) + 1, cols=3)
        measure_table.style = 'Light Grid Accent 1'

        headers = measure_table.rows[0].cells
        headers[0].text = 'Table'
        headers[1].text = 'Measure Name'
        headers[2].text = 'Description'

        for cell in headers:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(10)

        for idx, measure in enumerate(measures, 1):
            if idx >= len(measure_table.rows):
                break
            cells = measure_table.rows[idx].cells
            cells[0].text = self.safe_str(measure.get('table', 'Unknown'))
            cells[1].text = self.safe_str(measure.get('name', 'Unknown'))
            desc = self.safe_str(measure.get('description', 'N/A'))
            cells[2].text = desc[:100] if len(desc) > 100 else desc

        self.doc.add_paragraph()

        # Add ALL measure DAX formulas (removed the <=5 limit)
        self.doc.add_heading('Measure DAX Formulas', level=3)
        self.doc.add_paragraph(
            'Below are the DAX formulas for all measures in this data model:'
        )
        self.doc.add_paragraph()

        for measure in measures:
            table_name = self.safe_str(measure.get('table', 'Unknown'))
            measure_name = self.safe_str(measure.get('name', 'Unknown'))
            expression = self.safe_str(measure.get('expression', 'N/A'))
            description = self.safe_str(measure.get('description', ''))

            # Measure header
            p_header = self.doc.add_paragraph()
            run_header = p_header.add_run(f'{table_name}[{measure_name}]')
            run_header.font.bold = True
            run_header.font.size = Pt(11)
            run_header.font.color.rgb = RGBColor(0, 51, 102)

            # Description if available
            if description and description != 'N/A':
                p_desc = self.doc.add_paragraph()
                run_desc = p_desc.add_run(f'Description: {description}')
                run_desc.font.italic = True
                run_desc.font.size = Pt(9)

            # DAX Formula
            if expression and expression != 'N/A':
                p_expr = self.doc.add_paragraph()
                run_expr = p_expr.add_run(expression[:5000])  # Increased limit
                run_expr.font.name = 'Consolas'
                run_expr.font.size = Pt(8)
                run_expr.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue for DAX
            else:
                self.doc.add_paragraph('(No expression available)', style='List Bullet')

            self.doc.add_paragraph()  # Spacing between measures

        self.doc.add_paragraph()

    def add_relationships_section(self, metadata):
        """7.9 Model Relationships - Compact Format"""
        relationships = metadata.get('relationships', [])
        tables = metadata.get('model_tables', [])

        # DEBUG LOGGING
        print(f"\n   🔍 [Document Creator] add_relationships_section called")
        print(f"      - Relationships in metadata: {len(relationships)}")
        if relationships:
            print(f"      - First relationship: {relationships[0]}")

        self.doc.add_heading('7.9 Model Relationships', level=2)

        if not relationships:
            self.doc.add_paragraph('No relationships defined in this model.')
            if tables:
                self.doc.add_paragraph('Tables present:')
                for t in tables:
                    self.doc.add_paragraph(f'- {t}', style='List Bullet')
            self.doc.add_paragraph()
            return
        
        self.doc.add_paragraph(
            f'Total Relationships: {len(relationships)}'
        )
        
        # lookup all keys to display
        extra_keys = set()
        for rel in relationships:
            extra_keys.update(rel.keys())
        # fields we always show in table
        base_cols = ['fromTable','fromColumn','toTable','toColumn','type','joinType','isActive']
        # determine columns dynamically
        cols = base_cols + [k for k in extra_keys if k not in base_cols]
        rel_table = self.doc.add_table(rows=len(relationships) + 1, cols=len(cols))
        rel_table.style = 'Light Grid Accent 1'
        
        # header row
        for j, colname in enumerate(cols):
            cell = rel_table.rows[0].cells[j]
            cell.text = colname
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)
        
        # data rows
        for idx, rel in enumerate(relationships, 1):
            if idx >= len(rel_table.rows):
                break
            cells = rel_table.rows[idx].cells
            for j, colname in enumerate(cols):
                val = rel.get(colname)
                # boolean formatting
                if isinstance(val, bool):
                    val = 'Yes' if val else 'No'
                cells[j].text = self.safe_str(val)
                for para in cells[j].paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
            
            # after table, provide descriptions for any extras
            for rel in relationships:
                extra_info = {k:v for k,v in rel.items() if k not in base_cols}
                if extra_info:
                    self.doc.add_paragraph(f"\nRelationship '{rel.get('name', '')}' additional details:", style='List Bullet')
                    for k,v in extra_info.items():
                        self.doc.add_paragraph(f"{k}: {v}", style='List Bullet')
        
        self.doc.add_paragraph()
        
        self.doc.add_paragraph()

    def add_unused_elements_section(self, metadata):
        """7.8 Unused Tables and Columns Analysis"""
        analysis = self.detect_unused_columns_and_tables(metadata)
        
        if not analysis['unused_tables'] and not analysis['tables_with_unused_columns']:
            return
        
        self.doc.add_heading('7.8 Unused Elements & Optimization Opportunities', level=2)
        
        # Unused Tables
        if analysis['unused_tables']:
            self.doc.add_heading('Unused Tables', level=3)
            self.doc.add_paragraph(
                'The following tables are defined in the model but not used in any queries or relationships:'
            )
            for table in analysis['unused_tables']:
                self.doc.add_paragraph(table, style='List Bullet')
            self.doc.add_paragraph()
        
        # Columns with potential issues
        if analysis['tables_with_unused_columns']:
            self.doc.add_heading('Columns Flagged for Review', level=3)
            self.doc.add_paragraph(
                'The following columns may be unused or can be optimized:'
            )
            
            for table_name, cols in analysis['tables_with_unused_columns'].items():
                self.doc.add_paragraph(f'{table_name}:', style='List Bullet')
                for col in cols:
                    self.doc.add_paragraph(col, style='List Bullet 2')
            
            self.doc.add_paragraph()
        
        # Optimization tips
        self.doc.add_heading('Optimization Recommendations', level=3)
        tips = [
            'Remove/archive unused tables to reduce model size and complexity',
            'Consolidate duplicate data sources when possible',
            'Index frequently filtered columns for better performance',
            'Review relationships to ensure they support all reports',
            'Consider summarizing or aggregating large fact tables'
        ]
        for tip in tips:
            self.doc.add_paragraph(tip, style='List Bullet')

        self.doc.add_paragraph()

    def add_backend_expressions_section(self, metadata):
        """7.10 Backend SQL & M Query Expressions with SQL Extraction"""
        expressions = metadata.get('expressions', [])
        if not expressions:
            return

        self.doc.add_heading('7.10 Backend SQL & M Query Expressions', level=2)
        self.doc.add_paragraph(
            "The following expressions were extracted using the Power BI Admin Scanner API. "
            "For queries containing SQL, both the M code and extracted SQL are shown."
        )

        for idx, expr in enumerate(expressions, 1):
            table_name = self.safe_str(expr.get('table', f'Expression {idx}'))
            m_code = self.safe_str(expr.get('expression', ''))

            self.doc.add_heading(f'Table: {table_name}', level=3)
            
            # Convert M to SQL
            sql_info = self.sql_converter.extract_sql_from_m(m_code)
            
            if sql_info['has_sql']:
                # Show connection details
                self.doc.add_heading('Connection Details', level=4)
                conn_table = self.doc.add_table(rows=4, cols=2)
                conn_table.style = 'Light Shading Accent 1'
                
                conn_table.rows[0].cells[0].text = 'Server'
                conn_table.rows[0].cells[1].text = self.safe_str(sql_info.get('server', 'N/A'))
                
                conn_table.rows[1].cells[0].text = 'Database'
                conn_table.rows[1].cells[1].text = self.safe_str(sql_info.get('database', 'N/A'))
                
                conn_table.rows[2].cells[0].text = 'Query Type'
                conn_table.rows[2].cells[1].text = 'SQL Server Query'
                
                conn_table.rows[3].cells[0].text = 'Timeout'
                conn_table.rows[3].cells[1].text = self.safe_str(sql_info.get('connection_timeout', 'Default'))
                
                for row in conn_table.rows:
                    row.cells[0].paragraphs[0].runs[0].font.bold = True
                
                # Show SQL Query
                self.doc.add_paragraph()
                self.doc.add_heading('SQL Query (Extracted & Formatted)', level=4)
                
                formatted_sql = self.sql_converter.format_sql_query(sql_info['sql_query'])
                
                p = self.doc.add_paragraph()
                run = p.add_run(formatted_sql[:5000])
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0, 100, 0)  # Green for SQL
                
                # Show original M code in collapsed form
                self.doc.add_paragraph()
                self.doc.add_heading('Original M Query Expression', level=4)
                p = self.doc.add_paragraph()
                run = p.add_run(m_code[:2000])
                run.font.name = 'Consolas'
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(100, 100, 100)  # Gray for M
                
            else:
                # No SQL found, show M code normally
                self.doc.add_heading('M Query Expression', level=4)
                p = self.doc.add_paragraph()
                run = p.add_run(m_code[:4000])
                run.font.name = 'Consolas'
                run.font.size = Pt(8)
            
            self.doc.add_paragraph()

    def export_sql_queries_to_file(self, metadata, output_filename):
        """
        Export all extracted SQL queries to a separate .sql file
        
        Args:
            metadata (dict): Report metadata containing expressions
            output_filename (str): Base filename for SQL export
        """
        expressions = metadata.get('expressions', [])
        if not expressions:
            return None
        
        sql_filename = output_filename.replace('.docx', '_SQL_Queries.sql')
        
        with open(sql_filename, 'w', encoding='utf-8') as f:
            f.write("-- =====================================================\n")
            f.write("-- POWER BI REPORT SQL QUERIES\n")
            f.write(f"-- Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write("-- =====================================================\n\n")
            
            sql_count = 0
            
            for expr in expressions:
                table_name = self.safe_str(expr.get('table', 'Unknown'))
                m_code = self.safe_str(expr.get('expression', ''))
                
                sql_info = self.sql_converter.extract_sql_from_m(m_code)
                
                if sql_info['has_sql']:
                    sql_count += 1
                    
                    f.write(f"-- =====================================================\n")
                    f.write(f"-- TABLE: {table_name}\n")
                    f.write(f"-- Server: {sql_info.get('server', 'N/A')}\n")
                    f.write(f"-- Database: {sql_info.get('database', 'N/A')}\n")
                    if sql_info.get('connection_timeout'):
                        f.write(f"-- Timeout: {sql_info['connection_timeout']}\n")
                    f.write(f"-- =====================================================\n\n")
                    
                    formatted_sql = self.sql_converter.format_sql_query(sql_info['sql_query'])
                    f.write(formatted_sql)
                    f.write("\n\n")
                    f.write("GO\n\n")
            
            f.write(f"-- =====================================================\n")
            f.write(f"-- Total SQL Queries Extracted: {sql_count}\n")
            f.write(f"-- =====================================================\n")
        
        if sql_count > 0:
            print(f"   📝 Exported {sql_count} SQL queries to: {sql_filename}")
            return sql_filename
        else:
            os.remove(sql_filename)
            return None

    # ---------- Main Creation Method ----------

    def create_documentation_from_json(self, json_data, output_filename, author="System Generated"):
        """Create complete documentation from JSON data"""
        print(f"\n📄 Creating Word document...")
        
        metadata = json_data.get('metadata', {})
        report_name = self.safe_str(metadata.get('report_name', 'Power BI Report'))
        
        print("   ✍️  Adding cover page...")
        self.add_cover_page(report_name, author)
        
        print("   ✍️  Adding table of contents...")
        self.add_table_of_contents()
        
        if 'overview' in json_data:
            print("   ✍️  Adding overview section...")
            self.add_section('1. Report Overview', json_data['overview'], level=1)
        
        if 'data_sources' in json_data:
            print("   ✍️  Adding data sources section...")
            self.add_section('2. Data Sources & Architecture', json_data['data_sources'], level=1)
        
        if 'pages' in json_data:
            print("   ✍️  Adding pages section...")
            self.add_section('3. Report Pages', json_data['pages'], level=1)
        
        if 'user_guide' in json_data:
            print("   ✍️  Adding user guide...")
            self.add_section('4. User Guide', json_data['user_guide'], level=1)
        
        if 'technical_details' in json_data:
            print("   ✍️  Adding technical details...")
            self.add_section('5. Technical Details', json_data['technical_details'], level=1)
        
        if 'migration' in json_data:
            print("   ✍️  Adding migration procedure...")
            self.add_section('6. Migration & Deployment Procedure', json_data['migration'], level=1)
        
        # Backend queries & connections (7.x)
        print("   ✍️  Adding data model documentation...")
        self.doc.add_heading('7. Data Model Overview', level=1)
        
        # 7.1 Model Tables Overview
        model_tables = metadata.get('model_tables', [])
        if model_tables:
            self.doc.add_heading('7.1 Model Tables', level=2)
            self.doc.add_paragraph(
                f'This dataset contains {len(model_tables)} table(s):'
            )
            for table in model_tables:
                self.doc.add_paragraph(table, style='List Bullet')
            self.doc.add_paragraph()
        
        self.add_columns_section(metadata)
        self.add_dax_measures_section(metadata)
        self.add_relationships_section(metadata)
        self.add_unused_elements_section(metadata)

        # new: refresh schedule & history
        print("   ✍️  Adding refresh schedule & history section...")
        self.add_refresh_section(metadata)

        # export CSV if history exists and the caller requested it
        if getattr(self, 'export_csv', False):
            self.export_refresh_history_csv(metadata, output_filename)

        # generate mermaid diagram and save file
        self.add_relationship_diagram_section(metadata, output_filename)
        
        # Only add detailed queries section if there are expressions
        expressions = metadata.get('expressions', [])
        if expressions:
            print("   ✍️  Adding backend SQL & M query expressions...")
            self.add_backend_expressions_section(metadata)
            self.export_sql_queries_to_file(metadata, output_filename)
        
        # Data sources section (simplified)
        detailed_sources = metadata.get('detailed_sources', [])
        if detailed_sources:
            self.doc.add_heading('7. Data Connections', level=2)
            for idx, source in enumerate(detailed_sources, 1):
                ds_type = self.safe_str(source.get('datasourceType', 'Unknown'))
                self.doc.add_heading(f"Source {idx}: {ds_type}", level=3)
                
                conn_table = self.doc.add_table(rows=3, cols=2)
                conn_table.style = 'Light Shading Accent 1'
                
                conn_table.rows[0].cells[0].text = 'Type'
                conn_table.rows[0].cells[1].text = ds_type
                
                conn_table.rows[1].cells[0].text = 'Server'
                conn_table.rows[1].cells[1].text = self.safe_str(source.get('connectionDetails', {}).get('server', 'N/A'))
                
                conn_table.rows[2].cells[0].text = 'Database'
                conn_table.rows[2].cells[1].text = self.safe_str(source.get('connectionDetails', {}).get('database', 'N/A'))
                
                for row in conn_table.rows:
                    row.cells[0].paragraphs[0].runs[0].font.bold = True
        
        self.doc.add_paragraph()
        
        # Appendix
        print("   ✍️  Adding appendix...")
        self.doc.add_heading('8. Appendix', level=1)
        
        if metadata:
            self.doc.add_heading('A. Report Metadata Summary', level=2)
            appendix_data = {
                'Report Name': self.safe_str(metadata.get('report_name', 'N/A')),
                'Report ID': self.safe_str(metadata.get('report_id', 'N/A')),
                'Dataset Name': self.safe_str(metadata.get('dataset_name', 'N/A')),
                'Tables': str(len(metadata.get('model_tables', []))),
                'Data Sources': str(len(metadata.get('detailed_sources', []))),
                'Storage Mode': self.safe_str(metadata.get('refresh_details', {}).get('targetStorageMode', 'N/A')),
                'Generated': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            }
            self.add_section('', appendix_data, level=3)
        
        print(f"   💾 Saving document...")
        output_dir = os.path.dirname(output_filename)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # --- FIX: SAFE SAVE LOGIC ---
        if os.path.exists(output_filename):
            try:
                os.remove(output_filename)
            except PermissionError:
                base, ext = os.path.splitext(output_filename)
                ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_filename = f"{base}_{ts}{ext}"
                print(f"   ⚠️ Existing file locked. Saving as: {output_filename}")

        self.doc.save(output_filename)
        print(f"   ✅ Document saved: {output_filename}")
        
        return output_filename

# ---------- Top-level Helper ----------

def create_document_from_json_file(json_filename, output_filename=None, author="System Generated", use_ai=False, export_csv=False):
    """Create Word document from JSON documentation file

    By default the JSON is treated as complete.  When ``use_ai`` is True the
    metadata portion of the JSON is passed through the Azure OpenAI generator
    to populate any empty narrative sections before rendering.  This allows the
    Json file itself to remain AI‑free while still leveraging AI when producing
    the Word document.

    The ``export_csv`` flag controls whether a CSV file containing refresh
    history is written.  Older versions of the script always produced a CSV;
    setting this to False disables that behaviour.
    """
    print("\n" + "="*60)
    print("📝 Power BI Document Creator with SQL Extraction")
    print("="*60)
    
    print(f"\n📂 Loading documentation from: {json_filename}")
    try:
        with open(json_filename, 'r', encoding='utf-8') as f:
            json_data = json.load(f)
        print("   ✅ JSON loaded successfully")
    except Exception as e:
        print(f"   ❌ Error loading JSON: {e}")
        return None

    # if requested, enrich blank sections using AI
    if use_ai and AIDocGenerator:
        print("   🤖 Enriching JSON with AI-generated text (document stage)...")
        metadata = json_data.get('metadata', {})
        ai_gen = AIDocGenerator(os.getenv('OPENAI_API_KEY'))
        # populate any missing/empty keys
        if not json_data.get('overview'):
            json_data['overview'] = ai_gen.generate_overview(metadata)
        if not json_data.get('data_sources'):
            json_data['data_sources'] = ai_gen.generate_data_sources_doc(metadata.get('data_sources', []))
        if not json_data.get('pages'):
            json_data['pages'] = ai_gen.generate_pages_documentation(metadata.get('pages', []))
        if not json_data.get('user_guide'):
            json_data['user_guide'] = ai_gen.generate_user_guide(metadata.get('pages', []), metadata.get('report_name', 'Report'))
        if not json_data.get('technical_details'):
            json_data['technical_details'] = ai_gen.generate_technical_details(metadata)
        if not json_data.get('migration'):
            json_data['migration'] = ai_gen.generate_migration_steps(metadata.get('report_name', 'Report'))
        if not json_data.get('semantic_model') and metadata.get('dataset_id'):
            # semantic-only docs may use this
            json_data['semantic_model'] = ai_gen.generate_semantic_model_doc(metadata)
    elif use_ai and not AIDocGenerator:
        print("   ⚠️ AI enrichment requested but AIDocGenerator not available.")
    
    if not output_filename:
        metadata = json_data.get('metadata', {})
        report_name = metadata.get('report_name', 'PowerBI_Report')
        safe_name = "".join(c for c in report_name if c.isalnum() or c in (' ', '-', '_')).strip()
        safe_name = safe_name.replace(' ', '_')
        output_filename = f"{safe_name}_Documentation.docx"
    
    creator = PowerBIDocumentCreator()
    # pass export_csv option through to the creator instance
    creator.export_csv = export_csv
    result = creator.create_documentation_from_json(json_data, output_filename, author)
    
    if result:
        print("\n" + "="*60)
        print("✅ Document creation complete!")
        print(f"📁 Word Document: {os.path.abspath(output_filename)}")
        
        sql_file = output_filename.replace('.docx', '_SQL_Queries.sql')
        if os.path.exists(sql_file):
            print(f"📁 SQL Queries: {os.path.abspath(sql_file)}")
        
        print("="*60 + "\n")
    
    return result

# ---------- Standalone Test ----------

if __name__ == "__main__":
    json_file = 'powerbi_documentation.json'
    if os.path.exists(json_file):
        print(f"Found {json_file}, creating Word document...\n")
        create_document_from_json_file(
            json_filename=json_file,
            output_filename='PowerBI_Report_Documentation.docx',
            author='D Anandkumar'
        )
    else:
        print(f"❌ Error: {json_file} not found!")